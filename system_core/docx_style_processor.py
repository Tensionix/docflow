#!/usr/bin/env python3
"""
Conservative DOCX structural style processor.

Scan mode reports heading/list/caption/TOC/appendix style consistency issues. Fix mode only
assigns existing paragraph styles to text that is already recognizable as a
heading, list item, table caption, figure caption, TOC title, or appendix
title. It does not rewrite text, numbering, tables, images, sections, or page
setup.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from lxml import etree

from _office_common import find_docx_files, md_escape, mirrored_output_path, norm_space, safe_mkdir, truncate, write_json_file
from docx_xml_tools import NS, _etree_from_bytes, _etree_to_bytes, read_zip_map, write_zip_map


W = f"{{{NS['w']}}}"

TABLE_CAPTION_RE = re.compile(r"^\s*(таблица|table)\s+[\dIVXLCDM]+(?:[\.\-–]\d+)*\b", re.IGNORECASE)
FIGURE_CAPTION_RE = re.compile(r"^\s*(рисунок|рис\.?|figure|fig\.?)\s+[\dIVXLCDM]+(?:[\.\-–]\d+)*\b", re.IGNORECASE)
TOC_TITLE_RE = re.compile(r"^\s*(содержание|оглавление|contents|table\s+of\s+contents)\s*$", re.IGNORECASE)
APPENDIX_RE = re.compile(
    r"^\s*приложение\s+[\dIVXLCDMА-ЯA-Z]+(?:[\.\-–—][\dIVXLCDMА-ЯA-Z]+)*\.?(?:\s*[-–—]\s*.*|\s+.*)?$",
    re.IGNORECASE,
)
NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+){1,5})\.?\s+(.+)")
CHAPTER_HEADING_RE = re.compile(r"^\s*глава\s+\d+\.?\s+(.+)", re.IGNORECASE)
LIST_DASH_RE = re.compile(r"^\s*[-–—]\s+\S")
LIST_BULLET_RE = re.compile(r"^\s*[•·●]\s+\S")
LIST_NUMBER_RE = re.compile(r"^\s*(?:\d+|[IVXLCDM]+)[\.)]\s+\S", re.IGNORECASE)
LIST_LETTER_RE = re.compile(r"^\s*[A-Za-zА-Яа-яЁё][\.)]\s+\S")
VALUE_LIKE_RE = re.compile(
    r"^\s*[\d\s.,]+(?:г\.?|год(?:а|ов)?|кв\.?\s*м|км|тонн(?:ы)?|тыс\.?|млн|мест(?:а)?|человек|объект(?:а|ов)?)\s*$",
    re.IGNORECASE,
)


@dataclass
class ParagraphInfo:
    index: int
    text: str
    style_id: str
    style_name: str
    kind: str
    target_style: str
    issue: str


@dataclass
class FileReport:
    path: Path
    paragraphs: int
    tables: int
    images: int
    sections: int
    toc_markers: int
    appendices: int
    headings: int
    lists: int
    caption_tables: int
    caption_figures: int
    issues: list[ParagraphInfo]
    style_counts: Counter[str]
    fixes: int = 0
    output: Path | None = None


@dataclass
class StyleProfile:
    table_caption: str = ""
    figure_caption: str = ""
    toc_title: str = ""
    appendix: str = ""
    list_bullet: str = ""
    list_number: str = ""
    list_paragraph: str = ""
    heading_by_level: dict[int, str] | None = None
    style_names: dict[str, str] | None = None

    def heading(self, level: int) -> str:
        return (self.heading_by_level or {}).get(level, "")

    def display(self, style_id: str) -> str:
        if not style_id:
            return ""
        style_name = (self.style_names or {}).get(style_id, "")
        return f"{style_id} ({style_name})" if style_name and style_name != style_id else style_id


def _kind_label(kind: str) -> str:
    if kind.startswith("heading-"):
        return "Заголовок " + kind.split("-", 1)[1]
    labels = {
        "table-caption": "Название таблицы",
        "figure-caption": "Название рисунка",
        "toc-title": "Заголовок оглавления",
        "appendix": "Приложение",
        "list-bullet": "Маркированный список",
        "list-dash": "Список с тире",
        "list-numbered": "Нумерованный список",
        "list-lettered": "Буквенный список",
        "table-caption-proximity": "Название таблицы рядом",
        "figure-caption-proximity": "Название рисунка рядом",
    }
    return labels.get(kind, kind)


def _style_maps(files: dict[str, bytes]) -> tuple[dict[str, str], dict[str, str]]:
    id_to_name: dict[str, str] = {}
    name_to_id: dict[str, str] = {}
    for part in ("word/styles.xml", "word/stylesWithEffects.xml"):
        if part not in files:
            continue
        root = _etree_from_bytes(files[part]).getroot()
        for style in root.xpath(".//w:style", namespaces=NS):
            style_id = style.get(f"{W}styleId") or ""
            name_el = style.find("w:name", namespaces=NS)
            name = ""
            if name_el is not None:
                name = name_el.get(f"{W}val") or name_el.get("val") or ""
            if style_id:
                id_to_name[style_id] = name or style_id
            if name:
                name_to_id[name.lower()] = style_id
    return id_to_name, name_to_id


def _body_children(tree: etree._ElementTree) -> list[etree._Element]:
    body = tree.getroot().find(".//w:body", namespaces=NS)
    return list(body) if body is not None else []


def _paragraphs(tree: etree._ElementTree) -> list[etree._Element]:
    return tree.getroot().xpath(".//w:body/w:p", namespaces=NS)


def _p_text(p: etree._Element) -> str:
    return norm_space("".join(t.text or "" for t in p.xpath(".//w:t", namespaces=NS)))


def _p_style_id(p: etree._Element) -> str:
    pstyle = p.find("w:pPr/w:pStyle", namespaces=NS)
    if pstyle is None:
        return ""
    return pstyle.get(f"{W}val") or pstyle.get("val") or ""


def _p_num_key(p: etree._Element) -> tuple[str, str] | None:
    numpr = p.find("w:pPr/w:numPr", namespaces=NS)
    if numpr is None:
        return None
    ilvl = numpr.find("w:ilvl", namespaces=NS)
    numid = numpr.find("w:numId", namespaces=NS)
    return (
        ilvl.get(f"{W}val") if ilvl is not None else "",
        numid.get(f"{W}val") if numid is not None else "",
    )


def _set_p_style_id(p: etree._Element, style_id: str) -> None:
    ppr = p.find("w:pPr", namespaces=NS)
    if ppr is None:
        ppr = etree.Element(f"{W}pPr")
        p.insert(0, ppr)
    pstyle = ppr.find("w:pStyle", namespaces=NS)
    if pstyle is None:
        pstyle = etree.Element(f"{W}pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(f"{W}val", style_id)


def _is_heading_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return "heading" in haystack or "заголовок" in haystack


def _is_caption_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return "caption" in haystack or "название" in haystack or "подпись" in haystack


def _is_table_caption_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return (
        "caption" in haystack
        or "подпись" in haystack
        or "название таблиц" in haystack
        or "название таблицы" in haystack
        or "таб наз" in haystack
        or "table caption" in haystack
    )


def _is_figure_caption_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return (
        "caption" in haystack
        or "подпись" in haystack
        or "рисунок" in haystack
        or "figure" in haystack
        or "название рисун" in haystack
    )


def _is_toc_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return "toc" in haystack or "оглавление" in haystack or "содержание" in haystack


def _is_appendix_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return "appendix" in haystack or "приложение" in haystack


def _is_list_style(style_id: str, style_name: str) -> bool:
    haystack = f"{style_id} {style_name}".lower()
    return "list" in haystack or "список" in haystack


def _pick_style(id_to_name: dict[str, str], name_to_id: dict[str, str], candidates: tuple[str, ...]) -> str:
    for name in candidates:
        found = name_to_id.get(name.lower())
        if found:
            return found
    for style_id, style_name in id_to_name.items():
        haystack = f"{style_id} {style_name}".lower()
        for name in candidates:
            if name.lower() in haystack:
                return style_id
    return ""


def _heading_target(level: int, id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(
        id_to_name,
        name_to_id,
        (
            f"Heading{level}",
            f"Heading {level}",
            f"heading {level}",
            f"Заголовок {level}",
        ),
    )


def _table_caption_target(id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(id_to_name, name_to_id, ("_Название таблицы", "Название таблицы", "Таб наз", "Table Caption", "Caption", "Подпись"))


def _figure_caption_target(id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(id_to_name, name_to_id, ("Рисунок", "Название рисунка", "Figure Caption", "Figure", "Caption", "Подпись"))


def _toc_target(id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(id_to_name, name_to_id, ("TOCHeading", "TOC Heading", "Оглавление", "Содержание"))


def _appendix_target(id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(id_to_name, name_to_id, ("Appendix", "Приложение"))


def _list_bullet_target(id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(id_to_name, name_to_id, ("Список (–)", "Список (-)", "List Bullet", "Маркированный список", "List Paragraph", "List"))


def _list_number_target(id_to_name: dict[str, str], name_to_id: dict[str, str]) -> str:
    return _pick_style(id_to_name, name_to_id, ("1) Список", "Список 1)", "List Number", "Numbered List", "Нумерованный список", "List Paragraph", "List"))


def _list_target(kind: str, id_to_name: dict[str, str], name_to_id: dict[str, str], profile: StyleProfile | None = None) -> str:
    if kind in {"list-bullet", "list-dash"}:
        return (profile.list_bullet if profile else "") or _list_bullet_target(id_to_name, name_to_id)
    if kind in {"list-numbered", "list-lettered"}:
        return (profile.list_number if profile else "") or _list_number_target(id_to_name, name_to_id)
    return (profile.list_paragraph if profile else "") or _pick_style(id_to_name, name_to_id, ("List Paragraph", "List", "Список"))


def _classify_paragraph(
    text: str,
    style_id: str,
    style_name: str,
    id_to_name: dict[str, str],
    name_to_id: dict[str, str],
    profile: StyleProfile | None = None,
    num_key: tuple[str, str] | None = None,
) -> tuple[str, str, str]:
    if not text:
        return "", "", ""

    if TABLE_CAPTION_RE.match(text):
        if profile and profile.table_caption and style_id == profile.table_caption:
            return "table-caption", "", ""
        if _is_table_caption_style(style_id, style_name):
            return "table-caption", "", ""
        return "table-caption", (profile.table_caption if profile else "") or _table_caption_target(id_to_name, name_to_id), "Название таблицы оформлено не тем стилем."

    if FIGURE_CAPTION_RE.match(text):
        if profile and profile.figure_caption and style_id == profile.figure_caption:
            return "figure-caption", "", ""
        if _is_figure_caption_style(style_id, style_name):
            return "figure-caption", "", ""
        return "figure-caption", (profile.figure_caption if profile else "") or _figure_caption_target(id_to_name, name_to_id), "Название рисунка оформлено не тем стилем."

    if TOC_TITLE_RE.match(text):
        if profile and profile.toc_title and style_id == profile.toc_title:
            return "toc-title", "", ""
        if _is_toc_style(style_id, style_name) or _is_heading_style(style_id, style_name):
            return "toc-title", "", ""
        return "toc-title", (profile.toc_title if profile else "") or _toc_target(id_to_name, name_to_id) or _heading_target(1, id_to_name, name_to_id), "Заголовок оглавления оформлен не тем стилем."

    if _is_toc_style(style_id, style_name):
        return "", "", ""

    if APPENDIX_RE.match(text):
        if profile and profile.appendix and style_id == profile.appendix:
            return "appendix", "", ""
        if _is_appendix_style(style_id, style_name) or _is_heading_style(style_id, style_name):
            return "appendix", "", ""
        return "appendix", (profile.appendix if profile else "") or _appendix_target(id_to_name, name_to_id) or _heading_target(1, id_to_name, name_to_id), "Заголовок приложения оформлен не тем стилем."

    chapter = CHAPTER_HEADING_RE.match(text)
    if chapter:
        tail = norm_space(chapter.group(1))
        if len(tail) >= 8 and re.search(r"[A-Za-zА-Яа-яЁё]", tail):
            if _is_heading_style(style_id, style_name):
                return "heading-1", "", ""
            return "heading-1", (profile.heading(1) if profile else "") or _heading_target(1, id_to_name, name_to_id), "Похожий на заголовок главы абзац оформлен не стилем заголовка."

    match = NUMBERED_HEADING_RE.match(text)
    if match:
        tail = norm_space(match.group(2))
        if VALUE_LIKE_RE.match(text) or len(tail) < 8 or not re.search(r"[A-Za-zА-Яа-яЁё]", tail):
            return "", "", ""
        level = min(match.group(1).count(".") + 1, 6)
        if _is_heading_style(style_id, style_name):
            return f"heading-{level}", "", ""
        return f"heading-{level}", (profile.heading(level) if profile else "") or _heading_target(level, id_to_name, name_to_id), "Похожий на заголовок абзац оформлен не стилем заголовка."

    list_kind = ""
    if num_key:
        list_kind = "list-numbered"
    elif LIST_DASH_RE.match(text):
        list_kind = "list-dash"
    elif LIST_BULLET_RE.match(text):
        list_kind = "list-bullet"
    elif LIST_NUMBER_RE.match(text):
        list_kind = "list-numbered"
    elif LIST_LETTER_RE.match(text):
        list_kind = "list-lettered"
    if list_kind:
        if _is_list_style(style_id, style_name):
            return list_kind, "", ""
        return list_kind, _list_target(list_kind, id_to_name, name_to_id, profile), "Похожий на список абзац оформлен не списочным стилем."

    if _is_list_style(style_id, style_name):
        return "list-numbered" if "1)" in style_name or "number" in style_name.lower() else "list-bullet", "", ""

    return "", "", ""


def _toc_marker_count(files: dict[str, bytes]) -> int:
    total = 0
    for xml in files.values():
        total += xml.count(b"TOC")
    return total


def _image_count(files: dict[str, bytes]) -> int:
    return sum(1 for name in files if name.lower().startswith("word/media/") and not name.endswith("/"))


def _nearby_caption_counts(body_children: list[etree._Element]) -> tuple[int, int]:
    tables_without_caption = 0
    images_without_caption = 0
    para_texts: dict[int, str] = {}
    para_indexes: dict[int, int] = {}
    para_idx = 0
    for child in body_children:
        if child.tag == f"{W}p":
            para_idx += 1
            para_indexes[id(child)] = para_idx
            para_texts[para_idx] = _p_text(child)
        elif child.tag == f"{W}tbl":
            nearby = [para_texts.get(para_idx - 1, ""), para_texts.get(para_idx, "")]
            if not any(TABLE_CAPTION_RE.match(text or "") for text in nearby):
                tables_without_caption += 1
    for child in body_children:
        if child.tag != f"{W}p":
            continue
        if not child.xpath(".//w:drawing | .//w:pict", namespaces=NS):
            continue
        idx = para_indexes.get(id(child), 0)
        nearby = [para_texts.get(idx - 1, ""), para_texts.get(idx, ""), para_texts.get(idx + 1, "")]
        if not any(FIGURE_CAPTION_RE.match(text or "") for text in nearby):
            images_without_caption += 1
    return tables_without_caption, images_without_caption


def build_style_profile(reference_docx: Path | None) -> StyleProfile:
    if not reference_docx:
        return StyleProfile(heading_by_level={}, style_names={})
    if not reference_docx.exists():
        raise RuntimeError(f"Reference DOCX was not found: {reference_docx}")
    files = read_zip_map(reference_docx)
    tree = _etree_from_bytes(files["word/document.xml"])
    id_to_name, name_to_id = _style_maps(files)
    heading_styles: Counter[int] = Counter()
    table_styles: Counter[str] = Counter()
    figure_styles: Counter[str] = Counter()
    toc_styles: Counter[str] = Counter()
    appendix_styles: Counter[str] = Counter()
    list_bullet_styles: Counter[str] = Counter()
    list_number_styles: Counter[str] = Counter()
    list_paragraph_styles: Counter[str] = Counter()
    for p in _paragraphs(tree):
        text = _p_text(p)
        style_id = _p_style_id(p)
        if not style_id:
            continue
        style_name = id_to_name.get(style_id, style_id)
        if TABLE_CAPTION_RE.match(text):
            table_styles[style_id] += 1
        elif FIGURE_CAPTION_RE.match(text):
            figure_styles[style_id] += 1
        elif TOC_TITLE_RE.match(text):
            toc_styles[style_id] += 1
        elif APPENDIX_RE.match(text):
            appendix_styles[style_id] += 1
        elif _p_num_key(p) or LIST_NUMBER_RE.match(text):
            if _is_list_style(style_id, style_name):
                list_number_styles[style_id] += 1
        elif LIST_DASH_RE.match(text) or LIST_BULLET_RE.match(text) or LIST_LETTER_RE.match(text):
            if _is_list_style(style_id, style_name):
                list_bullet_styles[style_id] += 1
        elif _is_list_style(style_id, style_name):
            list_paragraph_styles[style_id] += 1
        else:
            m = NUMBERED_HEADING_RE.match(text)
            if m:
                level = min(m.group(1).count(".") + 1, 6)
                if _is_heading_style(style_id, style_name):
                    heading_styles[(level, style_id)] += 1
    heading_by_level: dict[int, str] = {}
    for (level, style_id), _count in heading_styles.most_common():
        heading_by_level.setdefault(level, style_id)
    return StyleProfile(
        table_caption=table_styles.most_common(1)[0][0] if table_styles else "",
        figure_caption=figure_styles.most_common(1)[0][0] if figure_styles else "",
        toc_title=toc_styles.most_common(1)[0][0] if toc_styles else "",
        appendix=appendix_styles.most_common(1)[0][0] if appendix_styles else _appendix_target(id_to_name, name_to_id),
        list_bullet=list_bullet_styles.most_common(1)[0][0] if list_bullet_styles else _list_bullet_target(id_to_name, name_to_id),
        list_number=list_number_styles.most_common(1)[0][0] if list_number_styles else _list_number_target(id_to_name, name_to_id),
        list_paragraph=list_paragraph_styles.most_common(1)[0][0] if list_paragraph_styles else _pick_style(id_to_name, name_to_id, ("List Paragraph", "List", "Список")),
        heading_by_level=heading_by_level,
        style_names=id_to_name,
    )


def analyze_docx(docx_path: Path, profile: StyleProfile | None = None) -> tuple[FileReport, dict[str, bytes], etree._ElementTree]:
    files = read_zip_map(docx_path)
    tree = _etree_from_bytes(files["word/document.xml"])
    id_to_name, name_to_id = _style_maps(files)
    paragraphs = _paragraphs(tree)
    body_children = _body_children(tree)
    issues: list[ParagraphInfo] = []
    style_counts: Counter[str] = Counter()
    headings = 0
    appendices = 0
    lists = 0
    caption_tables = 0
    caption_figures = 0
    for idx, p in enumerate(paragraphs, start=1):
        text = _p_text(p)
        style_id = _p_style_id(p)
        num_key = _p_num_key(p)
        style_name = id_to_name.get(style_id, style_id or "(none)")
        style_counts[style_name or "(none)"] += 1
        kind, target_style, issue = _classify_paragraph(text, style_id, style_name, id_to_name, name_to_id, profile, num_key)
        if kind.startswith("heading"):
            headings += 1
        elif kind == "appendix":
            appendices += 1
        elif kind.startswith("list-"):
            lists += 1
        elif kind == "table-caption":
            caption_tables += 1
        elif kind == "figure-caption":
            caption_figures += 1
        if issue:
            issues.append(ParagraphInfo(idx, text, style_id, style_name, kind, target_style, issue))

    tables_without_caption, images_without_caption = _nearby_caption_counts(body_children)
    if tables_without_caption:
        issues.append(ParagraphInfo(0, f"Таблиц без названия рядом: {tables_without_caption}", "", "", "table-caption-proximity", "", "У части таблиц рядом нет подписи вида 'Таблица N'."))
    if images_without_caption:
        issues.append(ParagraphInfo(0, f"Рисунков без названия рядом: {images_without_caption}", "", "", "figure-caption-proximity", "", "У части рисунков рядом нет подписи вида 'Рисунок N'."))

    report = FileReport(
        path=docx_path,
        paragraphs=len(paragraphs),
        tables=len(tree.getroot().xpath(".//w:tbl", namespaces=NS)),
        images=_image_count(files),
        sections=len(tree.getroot().xpath(".//w:sectPr", namespaces=NS)),
        toc_markers=_toc_marker_count(files),
        appendices=appendices,
        headings=headings,
        lists=lists,
        caption_tables=caption_tables,
        caption_figures=caption_figures,
        issues=issues,
        style_counts=style_counts,
    )
    return report, files, tree


def fix_docx(docx_path: Path, out_path: Path, profile: StyleProfile | None = None) -> FileReport:
    report, files, tree = analyze_docx(docx_path, profile)
    id_to_name, name_to_id = _style_maps(files)
    fixes = 0
    for p in _paragraphs(tree):
        text = _p_text(p)
        style_id = _p_style_id(p)
        num_key = _p_num_key(p)
        style_name = id_to_name.get(style_id, style_id or "(none)")
        _kind, target_style, issue = _classify_paragraph(text, style_id, style_name, id_to_name, name_to_id, profile, num_key)
        if issue and target_style and target_style != style_id:
            _set_p_style_id(p, target_style)
            fixes += 1
    fixed_files = dict(files)
    fixed_files["word/document.xml"] = _etree_to_bytes(tree)
    safe_mkdir(out_path.parent)
    write_zip_map(out_path, fixed_files)
    fixed_report, _, _ = analyze_docx(out_path, profile)
    fixed_report.fixes = fixes
    fixed_report.output = out_path
    return fixed_report


def _display_target_style(style_id: str, profile: StyleProfile | None) -> str:
    if not style_id:
        return "-"
    return profile.display(style_id) if profile else style_id


def _write_docx_report(out_path: Path, reports: list[FileReport], profile: StyleProfile | None = None) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Отчёт стилевого процессинга DOCX", level=1)
    summary = doc.add_table(rows=1, cols=13)
    summary.style = "Table Grid"
    headers = ["Файл", "Абзацы", "Таблицы", "Рисунки", "Секции", "Оглавление", "Приложения", "Заголовки", "Списки", "Названия таблиц", "Названия рисунков", "Проблемы", "Исправления"]
    for idx, header in enumerate(headers):
        summary.rows[0].cells[idx].text = header
    for report in reports:
        cells = summary.add_row().cells
        values = [report.path.name, report.paragraphs, report.tables, report.images, report.sections, report.toc_markers, report.appendices, report.headings, report.lists, report.caption_tables, report.caption_figures, len(report.issues), report.fixes]
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
    for report in reports:
        doc.add_heading(report.path.name, level=2)
        if report.output:
            doc.add_paragraph(f"Выходной файл: {report.output}")
        if report.issues:
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            for idx, header in enumerate(["Абзац", "Тип", "Текущий стиль", "Целевой стиль", "Проблема", "Текст"]):
                table.rows[0].cells[idx].text = header
            for issue in report.issues[:80]:
                cells = table.add_row().cells
                values = [issue.index or "-", _kind_label(issue.kind), issue.style_name, _display_target_style(issue.target_style, profile), issue.issue, truncate(issue.text, 160)]
                for idx, value in enumerate(values):
                    cells[idx].text = str(value)
        else:
            doc.add_paragraph("Стилевые проблемы не найдены.")
    doc.save(str(out_path))


def _profile_lines(profile: StyleProfile, reference_docx: Path | None) -> list[str]:
    if not reference_docx:
        return []
    lines = ["## Профиль стилей по эталону\n"]
    lines.append(f"- Эталон: `{reference_docx}`")
    for label, style_id in [("Название таблицы", profile.table_caption), ("Название рисунка", profile.figure_caption), ("Заголовок оглавления", profile.toc_title), ("Приложение", profile.appendix), ("Маркированный список", profile.list_bullet), ("Нумерованный список", profile.list_number)]:
        if style_id:
            lines.append(f"- {label}: `{profile.display(style_id)}`")
    for level, style_id in sorted((profile.heading_by_level or {}).items()):
        lines.append(f"- Заголовок {level}: `{profile.display(style_id)}`")
    lines.append("")
    return lines


def write_report_md(out_path: Path, reports: list[FileReport], profile: StyleProfile | None = None, reference_docx: Path | None = None) -> None:
    safe_mkdir(out_path.parent)
    lines = ["# Отчёт стилевого процессинга DOCX\n"]
    lines.append("| Файл | Абзацы | Таблицы | Рисунки | Секции | Метки оглавления | Приложения | Заголовки | Списки | Названия таблиц | Названия рисунков | Проблемы | Исправления |")
    lines.append("|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|")
    for report in reports:
        lines.append(f"| `{md_escape(report.path.name)}` | {report.paragraphs} | {report.tables} | {report.images} | {report.sections} | {report.toc_markers} | {report.appendices} | {report.headings} | {report.lists} | {report.caption_tables} | {report.caption_figures} | {len(report.issues)} | {report.fixes} |")
    lines.append("")
    if profile and reference_docx:
        lines.extend(_profile_lines(profile, reference_docx))
    for report in reports:
        lines.append(f"## {md_escape(report.path.name)}\n")
        if report.output:
            lines.append(f"- Выходной файл: `{report.output}`")
        top_styles = report.style_counts.most_common(12)
        if top_styles:
            lines.append("- Частые стили абзацев: " + ", ".join(f"`{md_escape(name)}` ({count})" for name, count in top_styles))
        lines.append("")
        if not report.issues:
            lines.append("_Стилевые проблемы не найдены._\n")
            continue
        lines.append("| Абзац | Тип | Текущий стиль | Целевой стиль | Проблема | Текст |")
        lines.append("|---:|---|---|---|---|---|")
        for issue in report.issues[:200]:
            target_style = _display_target_style(issue.target_style, profile)
            lines.append(f"| {issue.index or '-'} | {md_escape(_kind_label(issue.kind))} | `{md_escape(issue.style_name)}` | `{md_escape(target_style)}` | {md_escape(issue.issue)} | {md_escape(truncate(issue.text, 180))} |")
        if len(report.issues) > 200:
            lines.append(f"\n_... показаны первые 200 проблем из {len(report.issues)}._")
        lines.append("")
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    _write_docx_report(out_path.with_suffix(".docx"), reports, profile)


def _report_path_for_json(report: FileReport, input_dir: Path) -> str:
    try:
        return report.path.relative_to(input_dir).as_posix()
    except ValueError:
        return report.path.name


def build_json_payload(input_dir: Path, reports: list[FileReport], fix_enabled: bool) -> dict[str, object]:
    files_payload = []
    for report in reports:
        status = "FAIL" if report.issues else "PASS"
        files_payload.append(
            {
                "path": _report_path_for_json(report, input_dir),
                "status": status,
                "metrics": {
                    "paragraphs": report.paragraphs,
                    "tables": report.tables,
                    "images": report.images,
                    "sections": report.sections,
                    "toc_markers": report.toc_markers,
                    "appendices": report.appendices,
                    "headings": report.headings,
                    "lists": report.lists,
                    "caption_tables": report.caption_tables,
                    "caption_figures": report.caption_figures,
                    "issues": len(report.issues),
                    "fixes": report.fixes,
                },
                "output": str(report.output) if report.output else "",
                "issues": [
                    {
                        "paragraph": issue.index,
                        "kind": issue.kind,
                        "current_style": issue.style_name,
                        "target_style": issue.target_style,
                        "issue": issue.issue,
                        "text": issue.text,
                    }
                    for issue in report.issues
                ],
            }
        )
    fail_files = sum(1 for item in files_payload if item["status"] == "FAIL")
    return {
        "tool": "docx_style_processor",
        "version": 1,
        "input_dir": str(input_dir),
        "files": files_payload,
        "summary": {
            "total_files": len(reports),
            "pass_files": len(reports) - fail_files,
            "fail_files": fail_files,
            "fix_enabled": bool(fix_enabled),
            "paragraphs": sum(report.paragraphs for report in reports),
            "tables": sum(report.tables for report in reports),
            "images": sum(report.images for report in reports),
            "sections": sum(report.sections for report in reports),
            "toc_markers": sum(report.toc_markers for report in reports),
            "appendices": sum(report.appendices for report in reports),
            "headings": sum(report.headings for report in reports),
            "lists": sum(report.lists for report in reports),
            "caption_tables": sum(report.caption_tables for report in reports),
            "caption_figures": sum(report.caption_figures for report in reports),
            "issues": sum(len(report.issues) for report in reports),
            "fixes": sum(report.fixes for report in reports),
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Scan and conservatively fix DOCX structural styles.")
    parser.add_argument("--input", default="input", help="Input folder")
    parser.add_argument("--outdir", default="output/style_processed", help="Output folder for fixed DOCX")
    parser.add_argument("--out", default="report/docx_style_processing.md", help="Markdown report path")
    parser.add_argument("--fix", action="store_true", help="Assign existing heading/list/caption/TOC/appendix styles where safe")
    parser.add_argument("--reference-docx", default="", help="Optional reference DOCX used to choose target styles")
    parser.add_argument("--json-out", default="", help="Optional JSON report path")
    args = parser.parse_args()

    input_dir = Path(args.input).resolve()
    out_dir = Path(args.outdir).resolve()
    report_path = Path(args.out).resolve()
    docx_files = find_docx_files(input_dir)
    reference_docx = Path(args.reference_docx).resolve() if args.reference_docx else None
    if reference_docx:
        docx_files = [path for path in docx_files if path.resolve() != reference_docx]
    if not docx_files:
        safe_mkdir(report_path.parent)
        report_path.write_text("# Отчёт стилевого процессинга DOCX\n\nФайлы DOCX не найдены.\n", encoding="utf-8")
        if args.json_out:
            json_out = Path(args.json_out).resolve()
            write_json_file(json_out, build_json_payload(input_dir, [], args.fix))
            print(f"[OK] JSON report: {json_out}")
        print(f"[WARN] No DOCX files found: {input_dir}")
        return 0

    reports: list[FileReport] = []
    profile = build_style_profile(reference_docx)
    for docx_path in docx_files:
        if args.fix:
            out_path = mirrored_output_path(docx_path, input_dir, out_dir)
            reports.append(fix_docx(docx_path, out_path, profile))
        else:
            report, _, _ = analyze_docx(docx_path, profile)
            reports.append(report)

    write_report_md(report_path, reports, profile, reference_docx)
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(json_out, build_json_payload(input_dir, reports, args.fix))
        print(f"[OK] JSON report: {json_out}")
    print(f"[OK] Проверено файлов: {len(reports)}")
    if args.fix:
        print(f"[OK] Output folder: {out_dir}")
        print(f"[OK] Style assignments: {sum(report.fixes for report in reports)}")
    print(f"[OK] Report: {report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
