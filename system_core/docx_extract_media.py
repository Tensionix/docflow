#!/usr/bin/env python3
r"""
Office Media Extractor

Recursively extracts media files from DOCX/PPTX files under input.
Also writes DOCX/MD index reports in report.

Usage:
  python docx_extract_media.py --input input --outdir output/media --out report/office_media_index.md
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import hashlib
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

from _office_common import safe_mkdir, rel_posix, md_escape, write_json_file


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def write_docx_index(out_path: Path, rows: list[dict[str, Any]]) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Индекс медиафайлов Office", level=1)
    if not rows:
        doc.add_paragraph("Медиафайлы не найдены.")
        doc.save(str(out_path))
        return
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, title in enumerate(("Документ", "Формат", "Внутренний путь", "Размер", "SHA-256", "Выходной файл", "Статус")):
        hdr[idx].text = title
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row["document"])
        cells[1].text = str(row["format"])
        cells[2].text = str(row["internal"])
        cells[3].text = str(row["size"])
        cells[4].text = str(row["sha256"])[:16]
        cells[5].text = str(row["output"])
        cells[6].text = str(row["status"])
    doc.save(str(out_path))

def find_office_media_files(root: Path) -> list[Path]:
    if root.is_file():
        eligible = not root.name.startswith("~$") and root.suffix.lower() in {".docx", ".pptx"}
        return [root] if eligible else []
    return sorted(
        p for p in root.rglob("*")
        if p.is_file() and not p.name.startswith("~$") and p.suffix.lower() in {".docx", ".pptx"}
    )

def media_prefix_for(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "word/media/"
    if suffix == ".pptx":
        return "ppt/media/"
    return ""


def build_json_payload(input_dir: Path, out_dir: Path, office_files: list[Path], rows: list[dict[str, Any]]) -> dict[str, object]:
    rows_by_doc: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        rows_by_doc.setdefault(str(row["document"]), []).append(row)
    files_payload = []
    for office_path in office_files:
        rel_doc = rel_posix(office_path, input_dir)
        doc_rows = rows_by_doc.get(rel_doc, [])
        error_count = sum(1 for row in doc_rows if str(row["status"]).startswith("ERROR"))
        ok_rows = [row for row in doc_rows if row["status"] == "OK"]
        files_payload.append(
            {
                "path": rel_doc,
                "status": "FAIL" if error_count else "PASS",
                "metrics": {
                    "media_extracted": len(ok_rows),
                    "bytes_extracted": sum(int(row["size"]) for row in ok_rows),
                    "errors": error_count,
                },
            }
        )
    media = [
        {
            "document": row["document"],
            "format": row["format"],
            "internal": row["internal"],
            "size": row["size"],
            "sha256": row["sha256"],
            "output": row["output"],
            "status": row["status"],
        }
        for row in rows
    ]
    return {
        "tool": "docx_extract_media",
        "version": 1,
        "input_dir": str(input_dir),
        "output_dir": str(out_dir),
        "files": files_payload,
        "media": media,
        "summary": {
            "total_files": len(office_files),
            "pass_files": sum(1 for item in files_payload if item["status"] == "PASS"),
            "fail_files": sum(1 for item in files_payload if item["status"] == "FAIL"),
            "media_extracted": sum(1 for row in rows if row["status"] == "OK"),
            "bytes_extracted": sum(int(row["size"]) for row in rows if row["status"] == "OK"),
            "errors": sum(1 for row in rows if str(row["status"]).startswith("ERROR")),
        },
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract DOCX/PPTX embedded media recursively.")
    ap.add_argument("--input", default="input", help="Input folder with DOCX/PPTX files (recursive)")
    ap.add_argument("--outdir", default="output/media", help="Output folder for extracted media")
    ap.add_argument("--out", default="report/office_media_index.md", help="Output Markdown index path")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    args = ap.parse_args()

    in_dir = Path(args.input).resolve()
    out_dir = Path(args.outdir).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_dir)
    safe_mkdir(out_path.parent)

    rows: list[dict[str, Any]] = []
    office_files = find_office_media_files(in_dir)
    input_base = in_dir.parent if in_dir.is_file() else in_dir

    for office_path in office_files:
        rel_doc = rel_posix(office_path, in_dir)
        doc_out_dir = out_dir / office_path.relative_to(input_base).with_suffix("")
        media_prefix = media_prefix_for(office_path)
        fmt = office_path.suffix.lower().lstrip(".").upper()
        try:
            with zipfile.ZipFile(office_path) as zf:
                media_names = sorted(
                    name for name in zf.namelist()
                    if name.lower().startswith(media_prefix) and not name.endswith("/")
                )
                for idx, name in enumerate(media_names, start=1):
                    data = zf.read(name)
                    ext = Path(name).suffix or ".bin"
                    out_file = doc_out_dir / f"{idx:03d}_{Path(name).stem}{ext}"
                    safe_mkdir(out_file.parent)
                    out_file.write_bytes(data)
                    sha = hashlib.sha256(data).hexdigest()
                    rows.append(
                        {
                            "document": rel_doc,
                            "format": fmt,
                            "internal": name,
                            "size": len(data),
                            "sha256": sha,
                            "output": rel_posix(out_file, out_dir),
                            "status": "OK",
                        }
                    )
        except Exception as exc:
            rows.append(
                {
                    "document": rel_doc,
                    "format": fmt,
                    "internal": "",
                    "size": 0,
                    "sha256": "",
                    "output": "",
                    "status": f"ERROR: {exc}",
                }
            )

    lines: list[str] = []
    lines.append("# Индекс медиафайлов Office\n")
    lines.append(f"- Входная папка: `{in_dir}`")
    lines.append(f"- Файлов DOCX/PPTX проверено: **{len(office_files)}**")
    lines.append(f"- Медиафайлов извлечено: **{sum(1 for r in rows if r['status'] == 'OK')}**")
    lines.append(f"- Выходная папка: `{out_dir}`\n")
    if not rows:
        lines.append("Медиафайлы не найдены.\n")
    else:
        lines.append("| Документ | Формат | Внутренний путь | Размер | SHA-256 | Выходной файл | Статус |")
        lines.append("|---|---|---|---:|---|---|---|")
        for row in rows:
            lines.append(
                f"| `{md_escape(row['document'])}` | {md_escape(row['format'])} | `{md_escape(row['internal'])}` | {row['size']} | "
                f"`{str(row['sha256'])[:16]}` | `{md_escape(row['output'])}` | {md_escape(row['status'])} |"
            )
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    docx_out = out_path.with_suffix(".docx")
    write_docx_index(docx_out, rows)
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(json_out, build_json_payload(in_dir, out_dir, office_files, rows))
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Extracted media files: {sum(1 for r in rows if r['status'] == 'OK')}")
    print(f"[OK] Output folder: {out_dir}")
    print(f"[OK] Wrote DOCX report: {docx_out}")
    print(f"[OK] Wrote report: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
