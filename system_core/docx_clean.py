#!/usr/bin/env python3
"""
Create cleaned DOCX copies without changing visible document content.

The cleaner removes Word session rsid markup and orphaned comment sidecar
parts that often make Microsoft Word Compare less stable on large documents.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import json
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document

from _office_common import safe_mkdir
from docx_word_compare import _clean_docx_copy


def _docx_files(input_dir: Path) -> list[Path]:
    if input_dir.is_file():
        return [input_dir] if input_dir.suffix.lower() == ".docx" and not input_dir.name.startswith("~$") else []
    return sorted(
        path
        for path in input_dir.rglob("*.docx")
        if path.is_file() and not path.name.startswith("~$")
    )


def _write_reports(report_path: Path, json_path: Path, rows: list[dict[str, object]], outdir: Path) -> None:
    safe_mkdir(report_path.parent)
    safe_mkdir(json_path.parent)
    ok_count = sum(1 for row in rows if row.get("status") == "OK")
    fail_count = len(rows) - ok_count
    lines = [
        "# Очистка DOCX",
        "",
        f"- Выходная папка: `{outdir}`",
        f"- Обработано успешно: {ok_count}",
        f"- Ошибок: {fail_count}",
        "",
        "## Что очищается",
        "",
        "- `w:rsid*`-метки сеансов редактирования Word.",
        "- Осиротевшие comment-парты, если в DOCX нет основного `comments.xml`.",
        "- Связи и content-type записи к удалённым осиротевшим частям.",
        "",
        "Видимое содержимое документа не перестраивается.",
        "",
        "## Файлы",
        "",
        "| Статус | Файл | rsid | Осиротевшие comment-парты | Выход |",
        "|---|---|---:|---:|---|",
    ]
    for row in rows:
        source = str(row.get("source", "")).replace("|", "\\|")
        output = str(row.get("output", row.get("error", ""))).replace("|", "\\|")
        lines.append(
            "| {status} | `{source}` | {rsid} | {orphaned} | `{output}` |".format(
                status=row.get("status", "FAILED"),
                source=source,
                rsid=row.get("removed_rsid_markup", 0),
                orphaned=row.get("removed_orphan_comment_part_count", 0),
                output=output,
            )
        )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    json_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Clean DOCX copies from input into output.")
    parser.add_argument("--input", default="input", help="Input folder")
    parser.add_argument("--outdir", default="output/cleaned_docx", help="Output folder")
    parser.add_argument("--report", default="report/docx_clean.md", help="Markdown report path")
    parser.add_argument("--json-out", default="report/docx_clean.json", help="JSON report path")
    args = parser.parse_args()

    input_dir = Path(args.input).resolve()
    outdir = Path(args.outdir).resolve()
    report_path = Path(args.report).resolve()
    json_path = Path(args.json_out).resolve()
    safe_mkdir(outdir)

    files = _docx_files(input_dir)
    if not files:
        raise FileNotFoundError(f"В input нет DOCX файлов: {input_dir}")

    rows: list[dict[str, object]] = []
    input_base = input_dir.parent if input_dir.is_file() else input_dir
    for source in files:
        relative = source.relative_to(input_base)
        output = outdir / relative
        try:
            stats = _clean_docx_copy(source, output)
            doc = Document(str(output))
            rows.append(
                {
                    "status": "OK",
                    "source": str(source),
                    "output": str(output),
                    "removed_rsid_markup": stats["removed_rsid_markup"],
                    "removed_orphan_comment_parts": stats["removed_orphan_comment_parts"],
                    "removed_orphan_comment_part_count": len(stats["removed_orphan_comment_parts"]),
                    "removed_relationships": stats["removed_relationships"],
                    "removed_content_types": stats["removed_content_types"],
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                }
            )
            print(f"[OK] {source} -> {output}")
        except Exception as exc:
            rows.append(
                {
                    "status": "FAILED",
                    "source": str(source),
                    "output": str(output),
                    "error": str(exc),
                }
            )
            print(f"[FAILED] {source}: {exc}")

    _write_reports(report_path, json_path, rows, outdir)
    print(f"[OK] Report: {report_path}")
    print(f"[OK] JSON: {json_path}")
    return 1 if any(row.get("status") != "OK" for row in rows) else 0


if __name__ == "__main__":
    raise SystemExit(main())
