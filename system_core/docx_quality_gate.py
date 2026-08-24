#!/usr/bin/env python3
r"""
DOCX Quality Gate (scan-only)

Checks a folder of .docx files for:
- Non-black text color (explicit run color not 000000/auto)
- Highlight (w:highlight)
- Shading / background fill (w:shd)
- Comments (word/comments*.xml)
- Tracked changes (w:ins / w:del / w:moveFrom / w:moveTo)
- Strikethrough formatting (w:strike / w:dstrike)

Outputs a Markdown report.

Usage examples:
  python docx_quality_gate.py --input input --out report/docx_quality_gate_report.md
  python docx_quality_gate.py --input input\case_01 --out report\case_01_quality.md
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from pathlib import Path

from _office_common import safe_mkdir, md_escape, find_docx_files, rel_posix, write_json_file
from docx_xml_tools import (
    read_zip_map, list_xml_parts, list_style_parts,
    find_comments, find_tracked_changes, NS, _etree_from_bytes,
)

from lxml import etree
from docx import Document


QUALITY_REPORT_TITLE_RU = "Отчёт проверки качества DOCX"
QUALITY_HARD_FAIL_TITLE_RU = "Строгая проверка качества DOCX"
QUALITY_HEADERS_RU = [
    "Файл",
    "Статус",
    "Цветной текст",
    "Цветные стили",
    "Подсветка текста",
    "Подсветка в стилях",
    "Заливка текста",
    "Заливка в стилях",
    "Комментарии",
    "Зачёркивание",
    "Зачёркивание в стилях",
    "Вставки Word",
    "Удаления Word",
    "Перемещения из",
    "Перемещения в",
]
QUALITY_SUMMARY_LABELS_RU = [
    ("Цветной текст (прямой цвет символов)", "nonblack"),
    ("Цветные стили", "style_nonblack"),
    ("Подсветка текста", "highlight"),
    ("Подсветка в стилях", "style_highlight"),
    ("Заливка текста", "shading"),
    ("Заливка в стилях", "style_shading"),
    ("Файлы с комментариями", "comments"),
    ("Зачёркивание текста", "strike"),
    ("Зачёркивание в стилях", "style_strike"),
    ("Вставки Word (режим правок)", "ins"),
    ("Удаления Word (режим правок)", "del"),
    ("Перемещения из (режим правок)", "moveFrom"),
    ("Перемещения в (режим правок)", "moveTo"),
]
QUALITY_METRIC_KEYS = [
    "nonblack",
    "style_nonblack",
    "highlight",
    "style_highlight",
    "shading",
    "style_shading",
    "comments",
    "strike",
    "style_strike",
    "ins",
    "del",
    "moveFrom",
    "moveTo",
]


def count_nonblack_runs(files: dict[str, bytes]) -> int:
    # Count explicit colors that are not black/auto
    count = 0
    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for el in root.xpath(".//w:rPr/w:color", namespaces=NS):
            val = el.get(f"{{{NS['w']}}}val") or el.get("val")
            if not val:
                continue
            v = val.strip().lower()
            if v not in ("000000", "auto"):
                count += 1
    return count


def count_nonblack_styles(files: dict[str, bytes]) -> int:
    """Count explicit non-black colors defined via styles (styles.xml / stylesWithEffects.xml)."""
    count = 0
    for part in list_style_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for el in root.xpath(".//w:rPr/w:color", namespaces=NS):
            val = el.get(f"{{{NS['w']}}}val") or el.get("val")
            theme_color = el.get(f"{{{NS['w']}}}themeColor")
            if theme_color and (val is None or (val or "").strip().lower() != "000000"):
                count += 1
                continue
            if not val:
                continue
            v = val.strip().lower()
            if v not in ("000000", "auto"):
                count += 1
    return count


def count_tag(files: dict[str, bytes], xpath: str) -> int:
    count = 0
    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        count += len(root.xpath(xpath, namespaces=NS))
    return count


def count_tag_in_styles(files: dict[str, bytes], xpath: str) -> int:
    count = 0
    for part in list_style_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        count += len(root.xpath(xpath, namespaces=NS))
    return count


def count_highlight_parts(files: dict[str, bytes], parts: list[str]) -> int:
    """Count highlight elements that are not explicitly 'none'."""
    count = 0
    for part in parts:
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for el in root.xpath(".//w:highlight", namespaces=NS):
            val = el.get(f"{{{NS['w']}}}val") or el.get("val")
            if val is None:
                count += 1
                continue
            if val.strip().lower() != "none":
                count += 1
    return count


def count_shading_parts(files: dict[str, bytes], parts: list[str]) -> int:
    """Count shading elements with a non-white/non-auto fill."""
    count = 0
    for part in parts:
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for el in root.xpath(".//w:shd", namespaces=NS):
            fill = el.get(f"{{{NS['w']}}}fill") or el.get("fill")
            if fill is None:
                count += 1
                continue
            f = fill.strip().lower()
            if f not in ("auto", "ffffff"):
                count += 1
    return count


def count_highlights(files: dict[str, bytes]) -> int:
    return count_highlight_parts(files, list_xml_parts(files))


def count_highlights_in_styles(files: dict[str, bytes]) -> int:
    return count_highlight_parts(files, list_style_parts(files))


def count_shading(files: dict[str, bytes]) -> int:
    return count_shading_parts(files, list_xml_parts(files))


def count_shading_in_styles(files: dict[str, bytes]) -> int:
    return count_shading_parts(files, list_style_parts(files))


def count_strikethrough(files: dict[str, bytes]) -> int:
    count = 0
    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for rpr in root.xpath(".//w:rPr", namespaces=NS):
            from docx_xml_tools import _has_enabled_strike
            if _has_enabled_strike(rpr):
                count += 1
    return count


def count_strikethrough_in_styles(files: dict[str, bytes]) -> int:
    count = 0
    for part in list_style_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for rpr in root.xpath(".//w:rPr", namespaces=NS):
            from docx_xml_tools import _has_enabled_strike
            if _has_enabled_strike(rpr):
                count += 1
    return count


def write_quality_docx_report(out_path: Path, title: str, in_dir: Path, rows: list[tuple], totals: dict[str, int] | None = None) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Папка input: {in_dir}")
    doc.add_paragraph(f"Файлов проверено: {len(rows)}")

    if totals:
        doc.add_heading("Сводка", level=2)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Показатель"
        t.rows[0].cells[1].text = "Всего"
        for label, key in QUALITY_SUMMARY_LABELS_RU:
            cells = t.add_row().cells
            cells[0].text = label
            cells[1].text = str(totals.get(key, 0))

    doc.add_heading("Результаты по файлам", level=2)
    t = doc.add_table(rows=1, cols=15)
    t.style = "Table Grid"
    for idx, header in enumerate(QUALITY_HEADERS_RU):
        t.rows[0].cells[idx].text = header
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.save(str(out_path))


def quality_json_payload(tool: str, in_dir: Path, rows: list[tuple], totals: dict[str, int] | None = None) -> dict[str, object]:
    metric_totals = {key: 0 for key in QUALITY_METRIC_KEYS}
    files_payload = []
    for row in rows:
        metrics = dict(zip(QUALITY_METRIC_KEYS, row[2:]))
        for key, value in metrics.items():
            metric_totals[key] += int(value)
        files_payload.append(
            {
                "path": row[0],
                "status": row[1],
                "metrics": metrics,
            }
        )
    if totals:
        metric_totals.update({key: int(totals.get(key, 0)) for key in QUALITY_METRIC_KEYS})
    fail_files = sum(1 for row in rows if row[1] == "FAIL")
    summary = {
        "total_files": len(rows),
        "pass_files": len(rows) - fail_files,
        "fail_files": fail_files,
        **metric_totals,
    }
    return {
        "tool": tool,
        "version": 1,
        "input_dir": str(in_dir),
        "files": files_payload,
        "summary": summary,
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Scan DOCX files for quality issues (colors/highlights/comments/changes).")
    ap.add_argument("--input", default="input", help="Input folder with .docx files (default: input, recursive)")
    ap.add_argument("--out", default="report/docx_quality_gate_report.md", help="Output Markdown report path")
    ap.add_argument("--docx-out", default="", help="Optional DOCX report path")
    ap.add_argument("--no-docx-report", action="store_true", help="Do not write a DOCX report")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    args = ap.parse_args()

    in_dir = Path(args.input).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)

    if not in_dir.exists():
        print(f"[ERROR] Input folder does not exist: {in_dir}")
        return 2

    files = find_docx_files(in_dir)
    if not files:
        print(f"[WARN] No .docx files found in: {in_dir}")
        out_path.write_text(f"# {QUALITY_REPORT_TITLE_RU}\n\nФайлы DOCX не найдены.\n", encoding="utf-8")
        if not args.no_docx_report:
            docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
            write_quality_docx_report(docx_out, QUALITY_REPORT_TITLE_RU, in_dir, [], {})
            print(f"[OK] Wrote DOCX report: {docx_out}")
        if args.json_out:
            json_out = Path(args.json_out).resolve()
            write_json_file(json_out, quality_json_payload("docx_quality_gate", in_dir, [], {}))
            print(f"[OK] Wrote JSON report: {json_out}")
        print(f"[OK] Wrote report: {out_path}")
        return 0

    rows = []
    totals = {
        "nonblack": 0,
        "style_nonblack": 0,
        "highlight": 0,
        "style_highlight": 0,
        "shading": 0,
        "style_shading": 0,
        "comments": 0,
        "strike": 0,
        "style_strike": 0,
        "ins": 0,
        "del": 0,
        "moveFrom": 0,
        "moveTo": 0,
    }

    for p in files:
        z = read_zip_map(p)

        nonblack = count_nonblack_runs(z)
        style_nonblack = count_nonblack_styles(z)
        highlight = count_highlights(z)
        style_highlight = count_highlights_in_styles(z)
        shading = count_shading(z)
        style_shading = count_shading_in_styles(z)
        comments_parts = find_comments(z)
        comments = 1 if comments_parts else 0
        strike = count_strikethrough(z)
        style_strike = count_strikethrough_in_styles(z)

        tc = find_tracked_changes(z)
        ins = tc.get("<w:ins", 0)
        dele = tc.get("<w:del", 0)
        mvf = tc.get("<w:moveFrom", 0)
        mvt = tc.get("<w:moveTo", 0)

        totals["nonblack"] += nonblack
        totals["style_nonblack"] += style_nonblack
        totals["highlight"] += highlight
        totals["style_highlight"] += style_highlight
        totals["shading"] += shading
        totals["style_shading"] += style_shading
        totals["comments"] += comments
        totals["strike"] += strike
        totals["style_strike"] += style_strike
        totals["ins"] += ins
        totals["del"] += dele
        totals["moveFrom"] += mvf
        totals["moveTo"] += mvt

        status = "PASS"
        if any([nonblack, style_nonblack, highlight, style_highlight, shading, style_shading, comments, strike, style_strike, ins, dele, mvf, mvt]):
            status = "FAIL"

        rows.append((rel_posix(p, in_dir), status, nonblack, style_nonblack, highlight, style_highlight, shading, style_shading, comments, strike, style_strike, ins, dele, mvf, mvt))

    # Write report
    lines = []
    lines.append(f"# {QUALITY_REPORT_TITLE_RU}\n")
    lines.append(f"- Папка input: `{in_dir}`")
    lines.append(f"- Файлов проверено: **{len(files)}**\n")

    lines.append("## Сводка\n")
    lines.append("| Показатель | Всего |")
    lines.append("|---|---:|")
    for label, key in QUALITY_SUMMARY_LABELS_RU:
        lines.append(f"| {label} | {totals[key]} |")
    lines.append("")

    lines.append("## Результаты по файлам\n")
    lines.append("| " + " | ".join(QUALITY_HEADERS_RU) + " |")
    lines.append("|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|")

    for r in rows:
        (
            fn, st,
            nb, snb,
            hi, shi,
            sh, ssh,
            co,
            strike, style_strike,
            ins, dele, mvf, mvt,
        ) = r
        lines.append(
            f"| `{md_escape(fn)}` | **{st}** | {nb} | {snb} | {hi} | {shi} | {sh} | {ssh} | {co} | {strike} | {style_strike} | {ins} | {dele} | {mvf} | {mvt} |"
        )

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    if not args.no_docx_report:
        docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
        write_quality_docx_report(docx_out, QUALITY_REPORT_TITLE_RU, in_dir, rows, totals)
        print(f"[OK] Wrote DOCX report: {docx_out}")
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(json_out, quality_json_payload("docx_quality_gate", in_dir, rows, totals))
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Wrote report: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
