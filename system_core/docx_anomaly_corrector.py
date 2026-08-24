#!/usr/bin/env python3
"""
DOCX anomaly corrector.

Applies a narrow set of safe formatting fixes to copied DOCX files:
- collapse runs of duplicated empty body paragraphs to one paragraph;
- remove exact table row heights;
- remove table noWrap markers so text can wrap;
- optionally normalize table borders and cell margins.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from _office_common import find_docx_files, md_escape, mirrored_output_path, safe_mkdir, write_json_file
from docx_xml_tools import NS, _etree_from_bytes, _etree_to_bytes, read_zip_map, write_zip_map


W_NS = NS["w"]
W = f"{{{W_NS}}}"
TWIPS_PER_CM = 1440 / 2.54
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


@dataclass
class FileCorrection:
    file: str
    output: str
    status: str = "OK"
    removed_empty_paragraphs: int = 0
    fixed_exact_row_heights: int = 0
    removed_nowrap: int = 0
    normalized_table_borders: int = 0
    normalized_cell_margins: int = 0
    changed: bool = False
    error: str = ""


def w_tag(name: str) -> str:
    return f"{W}{name}"


def w_attr(name: str) -> str:
    return f"{W}{name}"


def cm_to_twips(value: float) -> int:
    return int(round(value * TWIPS_PER_CM))


def _w_val(element: etree._Element | None, name: str = "val") -> str:
    if element is None:
        return ""
    return element.get(w_attr(name)) or element.get(name) or ""


def _body(root: etree._Element) -> etree._Element | None:
    return root.find("w:body", namespaces=NS)


def _visible_text(element: etree._Element) -> str:
    return "".join(node.text or "" for node in element.findall(".//w:t", namespaces=NS))


def _has_nonempty_text(element: etree._Element) -> bool:
    return bool(_visible_text(element).replace("\u00a0", " ").strip())


def _is_removable_blank_paragraph(paragraph: etree._Element) -> bool:
    if paragraph.tag != w_tag("p"):
        return False
    if _has_nonempty_text(paragraph):
        return False
    blocked_tags = {
        "sectPr",
        "br",
        "lastRenderedPageBreak",
        "drawing",
        "pict",
        "bookmarkStart",
        "bookmarkEnd",
        "fldChar",
        "instrText",
        "hyperlink",
        "commentRangeStart",
        "commentRangeEnd",
    }
    for node in paragraph.iter():
        local = etree.QName(node).localname
        if local in blocked_tags:
            return False
    return True


def collapse_duplicated_empty_paragraphs(root: etree._Element) -> int:
    body = _body(root)
    if body is None:
        return 0
    removed = 0
    blank_run: list[etree._Element] = []
    for child in list(body):
        if child.tag == w_tag("p") and _is_removable_blank_paragraph(child):
            blank_run.append(child)
            continue
        removed += _trim_blank_run(body, blank_run)
        blank_run = []
    removed += _trim_blank_run(body, blank_run)
    return removed


def _trim_blank_run(body: etree._Element, blank_run: list[etree._Element]) -> int:
    if len(blank_run) < 2:
        return 0
    removed = 0
    for paragraph in blank_run[1:]:
        body.remove(paragraph)
        removed += 1
    return removed


def fix_exact_row_heights(root: etree._Element) -> int:
    fixed = 0
    for height in root.findall(".//w:trPr/w:trHeight", namespaces=NS):
        rule = (_w_val(height, "hRule") or "").lower()
        if rule != "exact":
            continue
        # Removing only the exact-height rule keeps any stored height value as
        # a soft hint instead of a hard crop.
        for key in (w_attr("hRule"), "hRule"):
            if key in height.attrib:
                del height.attrib[key]
        fixed += 1
    return fixed


def remove_nowrap(root: etree._Element) -> int:
    removed = 0
    for nowrap in list(root.findall(".//w:tcPr/w:noWrap", namespaces=NS)):
        parent = nowrap.getparent()
        if parent is not None:
            parent.remove(nowrap)
            removed += 1
    return removed


def _get_or_add_child(parent: etree._Element, name: str, index: int | None = None) -> etree._Element:
    child = parent.find(f"w:{name}", namespaces=NS)
    if child is not None:
        return child
    child = etree.Element(w_tag(name))
    if index is None:
        parent.append(child)
    else:
        parent.insert(index, child)
    return child


def _get_or_add_tbl_pr(table: etree._Element) -> etree._Element:
    tbl_pr = table.find("w:tblPr", namespaces=NS)
    if tbl_pr is not None:
        return tbl_pr
    tbl_pr = etree.Element(w_tag("tblPr"))
    table.insert(0, tbl_pr)
    return tbl_pr


def normalize_table_borders(root: etree._Element, border_size: int, border_color: str) -> int:
    changed = 0
    for table in root.findall(".//w:tbl", namespaces=NS):
        tbl_pr = _get_or_add_tbl_pr(table)
        for existing in tbl_pr.findall("w:tblBorders", namespaces=NS):
            tbl_pr.remove(existing)
        borders = etree.Element(w_tag("tblBorders"))
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = etree.SubElement(borders, w_tag(side))
            element.set(w_attr("val"), "single")
            element.set(w_attr("sz"), str(border_size))
            element.set(w_attr("space"), "0")
            element.set(w_attr("color"), border_color)
        tbl_pr.append(borders)
        changed += 1
    return changed


def normalize_cell_margins(root: etree._Element, margin_cm: float) -> int:
    margin_twips = cm_to_twips(margin_cm)
    changed = 0
    for table in root.findall(".//w:tbl", namespaces=NS):
        tbl_pr = _get_or_add_tbl_pr(table)
        for existing in tbl_pr.findall("w:tblCellMar", namespaces=NS):
            tbl_pr.remove(existing)
        margin = etree.Element(w_tag("tblCellMar"))
        for side in ("top", "start", "bottom", "end", "left", "right"):
            element = etree.SubElement(margin, w_tag(side))
            element.set(w_attr("w"), str(margin_twips))
            element.set(w_attr("type"), "dxa")
        tbl_pr.append(margin)
        changed += 1
    return changed


def process_docx(source: Path, target: Path, rel_name: str, args: argparse.Namespace) -> FileCorrection:
    row = FileCorrection(file=rel_name, output="")
    try:
        files = read_zip_map(source)
        if "word/document.xml" not in files:
            raise RuntimeError("word/document.xml не найден")
        tree = _etree_from_bytes(files["word/document.xml"])
        root = tree.getroot()

        if args.fix_empty_paragraphs:
            row.removed_empty_paragraphs = collapse_duplicated_empty_paragraphs(root)
        if args.fix_row_heights:
            row.fixed_exact_row_heights = fix_exact_row_heights(root)
        if args.fix_nowrap:
            row.removed_nowrap = remove_nowrap(root)
        if args.normalize_table_borders:
            row.normalized_table_borders = normalize_table_borders(root, max(1, int(args.border_size)), args.border_color)
        if args.normalize_cell_margins:
            row.normalized_cell_margins = normalize_cell_margins(root, max(0.0, float(args.cell_margin_cm)))

        row.changed = any(
            (
                row.removed_empty_paragraphs,
                row.fixed_exact_row_heights,
                row.removed_nowrap,
                row.normalized_table_borders,
                row.normalized_cell_margins,
            )
        )
        files["word/document.xml"] = _etree_to_bytes(tree)
        safe_mkdir(target.parent)
        write_zip_map(target, files)
        row.output = str(target)
    except Exception as exc:
        row.status = "ERROR"
        row.error = str(exc)
    return row


def _resolve_docx_files(input_path: Path) -> tuple[list[Path], Path]:
    resolved = input_path.resolve()
    if resolved.is_file():
        if resolved.suffix.lower() != ".docx" or resolved.name.startswith("~$"):
            return [], resolved.parent
        return [resolved], resolved.parent
    if resolved.is_dir():
        return find_docx_files(resolved), resolved
    return [], resolved


def _target_path(source: Path, input_root: Path, outdir: Path) -> Path:
    if input_root.is_dir():
        return mirrored_output_path(source, input_root, outdir)
    return outdir / source.name


def _rel_name(source: Path, input_root: Path) -> str:
    try:
        return source.relative_to(input_root).as_posix()
    except ValueError:
        return source.name


def _summary(rows: list[FileCorrection]) -> dict[str, Any]:
    return {
        "files": len(rows),
        "ok": sum(1 for row in rows if row.status == "OK"),
        "errors": sum(1 for row in rows if row.status != "OK"),
        "changed_files": sum(1 for row in rows if row.changed),
        "removed_empty_paragraphs": sum(row.removed_empty_paragraphs for row in rows),
        "fixed_exact_row_heights": sum(row.fixed_exact_row_heights for row in rows),
        "removed_nowrap": sum(row.removed_nowrap for row in rows),
        "normalized_table_borders": sum(row.normalized_table_borders for row in rows),
        "normalized_cell_margins": sum(row.normalized_cell_margins for row in rows),
    }


def _json_payload(input_path: Path, outdir: Path, rows: list[FileCorrection], options: dict[str, Any]) -> dict[str, Any]:
    return {
        "tool": "docx_anomaly_corrector",
        "version": 1,
        "input": str(input_path),
        "outdir": str(outdir),
        "options": options,
        "summary": _summary(rows),
        "files": [asdict(row) for row in rows],
    }


def _write_markdown_report(report_path: Path, input_path: Path, outdir: Path, rows: list[FileCorrection], options: dict[str, Any]) -> None:
    safe_mkdir(report_path.parent)
    summary = _summary(rows)
    lines: list[str] = []
    lines.append("# Корректировка аномалий DOCX\n")
    lines.append(f"- Вход: `{input_path}`")
    lines.append(f"- Выходная папка: `{outdir}`")
    lines.append(f"- Файлов обработано: **{summary['files']}**")
    lines.append(f"- Файлов с изменениями: **{summary['changed_files']}**")
    lines.append("")
    lines.append("## Включённые безопасные правки\n")
    for key, label in (
        ("fix_empty_paragraphs", "лишние пустые абзацы"),
        ("fix_row_heights", "точная высота строк таблиц"),
        ("fix_nowrap", "запрет переноса текста в ячейках"),
        ("normalize_table_borders", "унификация границ таблиц"),
        ("normalize_cell_margins", "унификация полей ячеек"),
    ):
        lines.append(f"- {label}: **{'да' if options.get(key) else 'нет'}**")
    lines.append("")
    lines.append("## Сводка\n")
    lines.append("| Метрика | Значение |")
    lines.append("|---|---:|")
    for key in (
        "removed_empty_paragraphs",
        "fixed_exact_row_heights",
        "removed_nowrap",
        "normalized_table_borders",
        "normalized_cell_margins",
    ):
        lines.append(f"| {key} | {summary[key]} |")
    lines.append("")
    lines.append("## Файлы\n")
    lines.append("| Файл | Статус | Изменён | Пустые абзацы | Высота строк | Запрет переноса | Границы | Поля ячеек | Выход |")
    lines.append("|---|---|---|---:|---:|---:|---:|---:|---|")
    for row in rows:
        lines.append(
            f"| `{md_escape(row.file)}` | {md_escape(row.status)} | {'да' if row.changed else 'нет'} | "
            f"{row.removed_empty_paragraphs} | {row.fixed_exact_row_heights} | {row.removed_nowrap} | "
            f"{row.normalized_table_borders} | {row.normalized_cell_margins} | `{md_escape(row.output)}` |"
        )
        if row.error:
            lines.append(f"| `{md_escape(row.file)}` | ERROR detail |  |  |  |  |  |  | {md_escape(row.error)} |")
    report_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _write_docx_report(report_path: Path, input_path: Path, outdir: Path, rows: list[FileCorrection], options: dict[str, Any]) -> None:
    safe_mkdir(report_path.parent)
    summary = _summary(rows)
    doc = Document()
    doc.add_heading("Корректировка аномалий DOCX", level=1)
    doc.add_paragraph(f"Вход: {input_path}")
    doc.add_paragraph(f"Выходная папка: {outdir}")
    doc.add_paragraph(f"Файлов обработано: {summary['files']}")
    doc.add_paragraph(f"Файлов с изменениями: {summary['changed_files']}")
    doc.add_heading("Файлы", level=2)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["Файл", "Статус", "Изменён", "Пустые абзацы", "Высота строк", "Запрет переноса", "Границы", "Поля ячеек", "Выход"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        values = [
            row.file,
            row.status,
            "да" if row.changed else "нет",
            row.removed_empty_paragraphs,
            row.fixed_exact_row_heights,
            row.removed_nowrap,
            row.normalized_table_borders,
            row.normalized_cell_margins,
            row.output,
        ]
        for index, value in enumerate(values):
            cells[index].text = str(value)
    doc.save(str(report_path))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Apply safe DOCX anomaly corrections to copied files.")
    parser.add_argument("--input", default="input", help="DOCX file or folder with DOCX files.")
    parser.add_argument("--outdir", default="output/docx_anomaly_fixed", help="Output folder.")
    parser.add_argument("--report", default="report/docx_anomaly_corrections.md", help="Markdown report path.")
    parser.add_argument("--docx-report", default="", help="Optional DOCX report path.")
    parser.add_argument("--json-out", default="", help="Optional JSON report path.")
    parser.add_argument("--fix-empty-paragraphs", action="store_true", default=True, help="Collapse duplicated empty paragraphs.")
    parser.add_argument("--no-fix-empty-paragraphs", action="store_false", dest="fix_empty_paragraphs", help="Do not collapse duplicated empty paragraphs.")
    parser.add_argument("--fix-row-heights", action="store_true", default=True, help="Remove exact table row-height rules.")
    parser.add_argument("--no-fix-row-heights", action="store_false", dest="fix_row_heights", help="Do not touch exact table row-height rules.")
    parser.add_argument("--fix-nowrap", action="store_true", default=True, help="Remove table-cell noWrap markers.")
    parser.add_argument("--no-fix-nowrap", action="store_false", dest="fix_nowrap", help="Do not touch table-cell noWrap markers.")
    parser.add_argument("--normalize-table-borders", action="store_true", help="Normalize table borders.")
    parser.add_argument("--normalize-cell-margins", action="store_true", help="Normalize table cell margins.")
    parser.add_argument("--border-size", type=int, default=4, help="OOXML border size in eighths of a point; 4 = 0.5 pt.")
    parser.add_argument("--border-color", default="000000", help="Border color hex RGB.")
    parser.add_argument("--cell-margin-cm", type=float, default=0.2, help="Cell margin in centimeters.")
    args = parser.parse_args(argv)

    input_path = Path(args.input).resolve()
    outdir = Path(args.outdir).resolve()
    report_path = Path(args.report).resolve()
    docx_report = Path(args.docx_report).resolve() if args.docx_report else report_path.with_suffix(".docx")
    json_report = Path(args.json_out).resolve() if args.json_out else report_path.with_suffix(".json")
    options = {
        "fix_empty_paragraphs": bool(args.fix_empty_paragraphs),
        "fix_row_heights": bool(args.fix_row_heights),
        "fix_nowrap": bool(args.fix_nowrap),
        "normalize_table_borders": bool(args.normalize_table_borders),
        "normalize_cell_margins": bool(args.normalize_cell_margins),
        "border_size": int(args.border_size),
        "border_color": str(args.border_color),
        "cell_margin_cm": float(args.cell_margin_cm),
    }

    if not input_path.exists():
        safe_mkdir(report_path.parent)
        report_path.write_text(f"# Корректировка аномалий DOCX\n\nВход не найден: `{input_path}`\n", encoding="utf-8")
        write_json_file(json_report, _json_payload(input_path, outdir, [], options))
        print(f"[ERROR] Input does not exist: {input_path}")
        return 2

    docx_files, input_root = _resolve_docx_files(input_path)
    safe_mkdir(outdir)
    rows: list[FileCorrection] = []
    for source in docx_files:
        target = _target_path(source, input_root, outdir)
        rel_name = _rel_name(source, input_root if input_root.is_dir() else source.parent)
        rows.append(process_docx(source, target, rel_name, args))

    _write_markdown_report(report_path, input_path, outdir, rows, options)
    _write_docx_report(docx_report, input_path, outdir, rows, options)
    write_json_file(json_report, _json_payload(input_path, outdir, rows, options))

    summary = _summary(rows)
    print(f"[OK] Processed: {summary['files']} file(s)")
    print(f"[OK] Changed files: {summary['changed_files']}")
    print(f"[OK] Output folder: {outdir}")
    print(f"[OK] Wrote report: {report_path}")
    print(f"[OK] Wrote DOCX report: {docx_report}")
    print(f"[OK] Wrote JSON report: {json_report}")
    return 0 if summary["errors"] == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
