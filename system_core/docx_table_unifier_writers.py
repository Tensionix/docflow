from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from docx_table_unifier_model import OriginCell, TableGroup, normalize_text

KNOWN_WIDTHS_CM: dict[int, list[float]] = {
    5: [1.0, 5.1, 5.7, 4.0, 12.4],
    9: [1.1, 3.6, 3.6, 4.2, 2.3, 3.4, 2.1, 5.4, 1.6],
}

KNOWN_CENTER_COLUMNS: dict[int, set[int]] = {
    5: {0, 3},
    9: {0, 4, 6, 8},
}

HEADER_FILL_TOP = "D9EAF7"
HEADER_FILL_SUB = "EAF3FA"
SECTION_FILL = "E6E6E6"
SUBSECTION_FILL = "F2F2F2"
BODY_FONT_PT = 8.0
SECTION_FONT_PT = 8.5
DEFAULT_TOTAL_WIDTH_CM = 27.7
PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
}
LAYOUT_PROFILES = {
    "standard": {
        "left_margin_cm": 1.0,
        "right_margin_cm": 1.0,
        "top_margin_cm": 1.0,
        "bottom_margin_cm": 1.0,
        "total_width_cm": 27.7,
        "merge_section_rows": False,
    },
    "merged-sections": {
        "left_margin_cm": 2.0,
        "right_margin_cm": 1.0,
        "top_margin_cm": 1.5,
        "bottom_margin_cm": 1.5,
        "total_width_cm": 26.7,
        "merge_section_rows": True,
    },
}


def _margin_cm(value_mm: float | None, fallback_cm: float) -> float:
    if value_mm is None:
        return fallback_cm
    return max(0.0, float(value_mm) / 10.0)


def _page_dimensions_cm(page_size: str, orientation: str) -> tuple[float, float]:
    width, height = PAGE_SIZES_CM.get(page_size.upper(), PAGE_SIZES_CM["A4"])
    if orientation == "landscape":
        return max(width, height), min(width, height)
    return min(width, height), max(width, height)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is not None:
        return
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_table_cell_margins(table, top=40, start=65, bottom=40, end=65) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m_type, m_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m_type}"))
        if node is None:
            node = OxmlElement(f"w:{m_type}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(m_value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_spacing(paragraph, before=0, after=0, line=1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _set_font(run, size: float = BODY_FONT_PT, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def _set_cell_text(cell, text: str, *, font_size: float = BODY_FONT_PT, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(paragraph)
    run = paragraph.add_run(text)
    _set_font(run, size=font_size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _non_empty_indexes(row: list[str]) -> list[int]:
    return [idx for idx, value in enumerate(row) if normalize_text(value)]


def _looks_numeric(value: str) -> bool:
    value = normalize_text(value)
    return bool(value) and re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?", value) is not None


def _header_row_count(group: TableGroup) -> int:
    if not group.tables:
        return 0
    first = group.tables[0]
    if first.header_row_count > 0:
        return first.header_row_count
    return min(2, len(first.first_rows))


def _preheader_row_count(group: TableGroup, preheader_mode: str) -> int:
    if preheader_mode != "separate" or not group.tables:
        return 0
    first = group.tables[0]
    return min(first.preheader_row_count, _header_row_count(group))


def _repeat_header_bounds(group: TableGroup, preheader_mode: str) -> tuple[int, int]:
    header_rows = _header_row_count(group)
    preheader_rows = _preheader_row_count(group, preheader_mode)
    return preheader_rows, max(0, header_rows - preheader_rows)


def _header_origins(group: TableGroup, repeat_start: int, repeat_count: int) -> list[OriginCell]:
    if not group.tables or repeat_count <= 0:
        return []
    repeat_end = repeat_start + repeat_count
    first = group.tables[0]
    origins = [
        origin
        for origin in first.origins
        if repeat_start <= origin.row < repeat_end and (origin.rowspan > 1 or origin.colspan > 1)
    ]
    return sorted(origins, key=lambda item: (item.rowspan * item.colspan, item.row, item.col), reverse=True)


def _apply_header_merges(table, group: TableGroup, repeat_start: int, repeat_count: int, widths_cm: list[float]) -> None:
    repeat_end = repeat_start + repeat_count
    for origin in _header_origins(group, repeat_start, repeat_count):
        start_row = max(0, origin.row)
        start_col = max(0, origin.col)
        end_row = min(repeat_end - 1, origin.row + origin.rowspan - 1)
        end_col = min(group.column_count - 1, origin.col + origin.colspan - 1)
        if end_row < start_row or end_col < start_col:
            continue
        try:
            merged = table.rows[start_row].cells[start_col].merge(table.rows[end_row].cells[end_col])
        except Exception:
            continue
        _set_cell_shading(merged, HEADER_FILL_TOP if start_row == 0 else HEADER_FILL_SUB)
        _set_cell_text(merged, origin.text, font_size=BODY_FONT_PT, bold=True, center=True)
        try:
            if start_col < len(widths_cm):
                merged.width = Cm(sum(widths_cm[start_col : end_col + 1]))
        except Exception:
            pass


def _section_kind(row: list[str], body_start_row: int, row_index: int) -> str | None:
    if row_index < body_start_row:
        return None
    non_empty = _non_empty_indexes(row)
    if len(non_empty) != 1 or non_empty[0] != 0:
        return None
    text = normalize_text(row[0])
    if not text:
        return None
    if row_index == body_start_row:
        return "section"
    return "subsection"


def _known_center_columns(group: TableGroup, preheader_mode: str = "include") -> set[int]:
    if group.column_count in KNOWN_CENTER_COLUMNS:
        return KNOWN_CENTER_COLUMNS[group.column_count]

    scores: list[float] = [0.0] * group.column_count
    counted_rows = 0
    body_start_row = _header_row_count(group)
    for row_index, row in enumerate(group.rows):
        if row_index < body_start_row or _section_kind(row, body_start_row, row_index) is not None:
            continue
        counted_rows += 1
        for col, value in enumerate(row):
            text = normalize_text(value)
            if not text:
                continue
            if _looks_numeric(text):
                scores[col] += 1.0
            elif len(text) <= 8:
                scores[col] += 0.45
    centers: set[int] = {0} if group.column_count else set()
    if counted_rows > 0:
        for idx, score in enumerate(scores):
            if score / counted_rows >= 0.60:
                centers.add(idx)
    return centers


def _scale_widths(widths: list[float], target_total_width_cm: float) -> list[float]:
    total = sum(widths) or 1.0
    factor = target_total_width_cm / total
    scaled = [round(width * factor, 4) for width in widths]
    if scaled:
        scaled[-1] = round(scaled[-1] + (target_total_width_cm - sum(scaled)), 4)
    return scaled


def _adaptive_widths_cm(group: TableGroup, total_width_cm: float) -> list[float]:
    if group.column_count in KNOWN_WIDTHS_CM:
        return _scale_widths(KNOWN_WIDTHS_CM[group.column_count], total_width_cm)

    weights: list[float] = [1.0] * group.column_count
    header_rows = _header_row_count(group)
    preview_rows = group.rows[: min(len(group.rows), 24)]
    for row_index, row in enumerate(preview_rows):
        row_weight = 1.35 if row_index < header_rows else 1.0
        for idx in range(group.column_count):
            text = normalize_text(row[idx] if idx < len(row) else "")
            if not text:
                continue
            compact_len = min(len(text.replace("\n", " ")), 60)
            addition = 0.55 + compact_len / 18.0
            if _looks_numeric(text):
                addition *= 0.75
            weights[idx] += addition * row_weight

    total_weight = sum(weights) or 1.0
    min_width = 1.0 if group.column_count <= 8 else 0.75
    widths = [max(min_width, total_width_cm * (weight / total_weight)) for weight in weights]
    total = sum(widths)
    if total > total_width_cm:
        shrinkable = [max(0.0, width - min_width) for width in widths]
        shrink_total = sum(shrinkable)
        overflow = total - total_width_cm
        if shrink_total > 0:
            widths = [width - overflow * (room / shrink_total) for width, room in zip(widths, shrinkable)]
    total = sum(widths)
    if total < total_width_cm and widths:
        widths[-1] += total_width_cm - total
    return widths


def _apply_cell_widths(row, widths_cm: list[float]) -> None:
    for idx, width in enumerate(widths_cm):
        row.cells[idx].width = Cm(width)
        row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _first_non_empty_text(row_values: list[str]) -> str:
    for value in row_values:
        text = normalize_text(value)
        if text:
            return text
    return ""


def write_docx(
    groups: list[TableGroup],
    out_path: Path,
    source_name: str,
    mode: str,
    *,
    layout: str = "standard",
    preheader_mode: str = "include",
    page_size: str = "A4",
    page_orientation: str = "landscape",
    margin_top_mm: float | None = None,
    margin_right_mm: float | None = None,
    margin_bottom_mm: float | None = None,
    margin_left_mm: float | None = None,
) -> None:
    profile = LAYOUT_PROFILES.get(layout, LAYOUT_PROFILES["standard"])
    document = Document()
    section = document.sections[0]
    page_width_cm, page_height_cm = _page_dimensions_cm(page_size, page_orientation)
    section.orientation = WD_ORIENT.LANDSCAPE if page_orientation == "landscape" else WD_ORIENT.PORTRAIT
    section.page_width = Cm(page_width_cm)
    section.page_height = Cm(page_height_cm)
    left_margin_cm = _margin_cm(margin_left_mm, float(profile["left_margin_cm"]))
    right_margin_cm = _margin_cm(margin_right_mm, float(profile["right_margin_cm"]))
    top_margin_cm = _margin_cm(margin_top_mm, float(profile["top_margin_cm"]))
    bottom_margin_cm = _margin_cm(margin_bottom_mm, float(profile["bottom_margin_cm"]))
    section.left_margin = Cm(left_margin_cm)
    section.right_margin = Cm(right_margin_cm)
    section.top_margin = Cm(top_margin_cm)
    section.bottom_margin = Cm(bottom_margin_cm)

    merge_section_rows = bool(profile["merge_section_rows"])
    total_width_cm = max(6.0, page_width_cm - left_margin_cm - right_margin_cm)

    for group_index, group in enumerate(groups):
        if group_index > 0:
            document.add_paragraph()
        table = document.add_table(rows=0, cols=group.column_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_cell_margins(table)

        widths_cm = _adaptive_widths_cm(group, total_width_cm)
        body_start_row = _header_row_count(group)
        repeat_start, repeat_count = _repeat_header_bounds(group, preheader_mode)
        repeat_end = repeat_start + repeat_count
        center_columns = _known_center_columns(group, preheader_mode)

        for row_index, row_values in enumerate(group.rows):
            row = table.add_row()
            _apply_cell_widths(row, widths_cm)
            row_kind = _section_kind(row_values, body_start_row, row_index)
            is_preheader = preheader_mode == "separate" and row_index < repeat_start
            is_header = repeat_start <= row_index < repeat_end
            fill = None
            if is_preheader:
                fill = SECTION_FILL
            elif is_header:
                fill = HEADER_FILL_TOP if row_index == 0 else HEADER_FILL_SUB
                _set_repeat_table_header(row)
            elif row_kind == "section":
                fill = SECTION_FILL
            elif row_kind == "subsection":
                fill = SUBSECTION_FILL

            if row_kind is not None and merge_section_rows and group.column_count > 1:
                merged = row.cells[0]
                for col_idx in range(1, group.column_count):
                    merged = merged.merge(row.cells[col_idx])
                if fill is not None:
                    _set_cell_shading(merged, fill)
                _set_cell_text(
                    merged,
                    _first_non_empty_text(row_values),
                    font_size=SECTION_FONT_PT,
                    bold=True,
                    center=False,
                )
                continue

            for col_idx in range(group.column_count):
                value = normalize_text(row_values[col_idx] if col_idx < len(row_values) else "")
                cell = row.cells[col_idx]
                if fill is not None:
                    _set_cell_shading(cell, fill)
                if is_header:
                    _set_cell_text(cell, value, font_size=BODY_FONT_PT, bold=True, center=True)
                elif is_preheader:
                    _set_cell_text(cell, value, font_size=SECTION_FONT_PT, bold=True, center=False)
                elif row_kind is not None:
                    _set_cell_text(
                        cell,
                        value,
                        font_size=SECTION_FONT_PT,
                        bold=(col_idx == 0 and bool(value)),
                        center=(col_idx != 0),
                    )
                else:
                    _set_cell_text(cell, value, font_size=BODY_FONT_PT, center=(col_idx in center_columns))

        if repeat_start > 0:
            _apply_header_merges(table, group, 0, repeat_start, widths_cm)
        _apply_header_merges(table, group, repeat_start, repeat_count, widths_cm)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def _escape_md(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", "<br>")


def write_markdown(groups: list[TableGroup], out_path: Path, source_name: str, mode: str) -> None:
    lines: list[str] = []
    lines.append(f"# Audion DOCX Table Unifier — {source_name}")
    lines.append("")
    lines.append(f"Mode: `{mode}`")
    lines.append("")

    for group in groups:
        lines.append(f"## Group {group.group_index:02d}")
        lines.append("")
        lines.append(f"- Columns: {group.column_count}")
        lines.append(f"- Source tables: {', '.join(str(i) for i in group.source_table_indexes())}")
        lines.append(f"- Group start reason: {group.break_reason}")
        lines.append("")

        header = "| " + " | ".join([f"C{idx+1}" for idx in range(group.column_count)]) + " |"
        separator = "| " + " | ".join(["---"] * group.column_count) + " |"
        lines.append(header)
        lines.append(separator)
        for row in group.rows:
            lines.append("| " + " | ".join(_escape_md(normalize_text(cell)) for cell in row) + " |")
        lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def write_json(groups: list[TableGroup], out_path: Path, source_name: str, mode: str, *, layout: str = "standard", preheader_mode: str = "include", page_size: str = "A4", page_orientation: str = "landscape") -> None:
    payload = {
        "source_name": source_name,
        "mode": mode,
        "layout": layout,
        "preheader_mode": preheader_mode,
        "page_size": page_size,
        "page_orientation": page_orientation,
        "groups": [],
    }
    for group in groups:
        header_rows = _header_row_count(group)
        preheader_rows = _preheader_row_count(group, preheader_mode)
        repeat_start, repeat_count = _repeat_header_bounds(group, preheader_mode)
        payload["groups"].append(
            {
                "group_index": group.group_index,
                "column_count": group.column_count,
                "break_reason": group.break_reason,
                "source_tables": [
                    {
                        "table_index": table.table_index,
                        "width": table.width,
                        "height": table.height,
                        "merged_cells": table.merged_cells,
                        "continuation_slots": table.continuation_slots,
                        "header_preview": table.header_preview,
                        "detected_preheader_rows": table.preheader_row_count,
                        "dropped_header_rows": group.dropped_header_rows.get(table.table_index, 0),
                    }
                    for table in group.tables
                ],
                "row_count": len(group.rows),
                "header_rows": header_rows,
                "preheader_rows": preheader_rows,
                "repeat_header_start": repeat_start,
                "repeat_header_rows": repeat_count,
                "section_rows_detected": sum(
                    1 for idx, row in enumerate(group.rows) if _section_kind(row, header_rows, idx) is not None
                ),
            }
        )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
