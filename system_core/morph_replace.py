from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any
import argparse
import json
import re

from docx import Document
from openpyxl import load_workbook


RU_WORD_RE = re.compile(r"[А-Яа-яЁё]+")
CASE_GRAMMEMES = {"nomn", "gent", "datv", "accs", "ablt", "loct", "voct", "gen2", "loc2"}
NUMBER_GRAMMEMES = {"sing", "plur"}
GENDER_GRAMMEMES = {"masc", "femn", "neut"}
ANIMACY_GRAMMEMES = {"anim", "inan"}
INVALID_REPORT_NAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
NON_SENTENCE_ABBREVIATIONS = {
    "абз",
    "вл",
    "г",
    "гг",
    "д",
    "др",
    "им",
    "кв",
    "корп",
    "куб",
    "млн",
    "млрд",
    "п",
    "пп",
    "подп",
    "пос",
    "пр",
    "рис",
    "с",
    "см",
    "ст",
    "стр",
    "табл",
    "т",
    "тыс",
    "ул",
    "ч",
}


@dataclass(frozen=True)
class Token:
    text: str
    start: int
    end: int
    lemma: str


@dataclass(frozen=True)
class TextMatch:
    start: int
    end: int
    found: str
    tokens: tuple[Token, ...]


@dataclass(frozen=True)
class ReplacementSpan:
    start: int
    end: int
    found: str
    replacement: str
    after_start: int
    after_end: int


@dataclass(frozen=True)
class TextReplacement:
    before: str
    after: str
    spans: tuple[ReplacementSpan, ...]


class Lemmatizer:
    def __init__(self) -> None:
        try:
            import pymorphy3
        except ImportError as exc:
            raise RuntimeError("pymorphy3 не установлен. Установите зависимость для morph-replace.") from exc
        self._morph = pymorphy3.MorphAnalyzer()
        self._cache: dict[str, str] = {}
        self._parse_cache: dict[str, Any] = {}

    def lemma(self, word: str) -> str:
        key = word.lower().replace("ё", "е")
        if key not in self._cache:
            self._cache[key] = self._morph.parse(word)[0].normal_form.replace("ё", "е")
        return self._cache[key]

    def parse(self, word: str) -> Any:
        key = word.lower().replace("ё", "е")
        if key not in self._parse_cache:
            self._parse_cache[key] = self._morph.parse(word)[0]
        return self._parse_cache[key]

    def grammar(self, word: str) -> set[str]:
        tag = self.parse(word).tag
        return {
            grammeme
            for group in (CASE_GRAMMEMES, NUMBER_GRAMMEMES, GENDER_GRAMMEMES, ANIMACY_GRAMMEMES)
            for grammeme in group
            if grammeme in tag
        }

    def inflect_word(self, word: str, grammemes: set[str]) -> str:
        parse = self.parse(word)
        candidates = [
            grammemes,
            grammemes - ANIMACY_GRAMMEMES,
            grammemes - GENDER_GRAMMEMES,
            grammemes - ANIMACY_GRAMMEMES - GENDER_GRAMMEMES,
            grammemes & (CASE_GRAMMEMES | NUMBER_GRAMMEMES),
            grammemes & CASE_GRAMMEMES,
        ]
        seen: set[frozenset[str]] = set()
        for candidate in candidates:
            key = frozenset(candidate)
            if not candidate or key in seen:
                continue
            seen.add(key)
            inflected = parse.inflect(candidate)
            if inflected is not None:
                return restore_word_case(word, inflected.word)
        return word


def tokens(text: str, lemmatizer: Lemmatizer) -> list[Token]:
    return [
        Token(match.group(0), match.start(), match.end(), lemmatizer.lemma(match.group(0)))
        for match in RU_WORD_RE.finditer(text)
    ]


def phrase_lemmas(text: str, lemmatizer: Lemmatizer) -> list[str]:
    result = [token.lemma for token in tokens(text, lemmatizer)]
    if not result:
        raise RuntimeError("--find должен содержать хотя бы одно русское слово.")
    return result


def find_matches(text: str, find_lemmas: list[str], lemmatizer: Lemmatizer) -> list[TextMatch]:
    text_tokens = tokens(text, lemmatizer)
    size = len(find_lemmas)
    matches: list[TextMatch] = []
    if len(text_tokens) < size:
        return matches
    for index in range(0, len(text_tokens) - size + 1):
        window = text_tokens[index : index + size]
        if [token.lemma for token in window] == find_lemmas:
            start = window[0].start
            end = window[-1].end
            matches.append(TextMatch(start=start, end=end, found=text[start:end], tokens=tuple(window)))
    return matches


def restore_word_case(template: str, word: str) -> str:
    if template.isupper():
        return word.upper()
    if template[:1].isupper():
        return word[:1].upper() + word[1:]
    return word


def capitalize_like_found(replacement_word: str, inflected_word: str, match: TextMatch, word_index: int) -> str:
    if word_index == 0 and match.tokens and match.tokens[0].text[:1].isupper() and replacement_word[:1].islower():
        return inflected_word[:1].upper() + inflected_word[1:]
    return inflected_word


def inflect_replacement_phrase(
    replacement: str,
    match: TextMatch,
    lemmatizer: Lemmatizer,
    *,
    capitalize_first_like_match: bool = True,
) -> str:
    if not match.tokens:
        return replacement
    target_grammar = lemmatizer.grammar(match.tokens[-1].text)
    if not target_grammar:
        return replacement

    parts: list[str] = []
    cursor = 0
    word_index = 0
    for word_match in RU_WORD_RE.finditer(replacement):
        parts.append(replacement[cursor : word_match.start()])
        word = word_match.group(0)
        inflected = lemmatizer.inflect_word(word, target_grammar)
        if capitalize_first_like_match:
            parts.append(capitalize_like_found(word, inflected, match, word_index))
        else:
            parts.append(inflected)
        cursor = word_match.end()
        word_index += 1
    parts.append(replacement[cursor:])
    return "".join(parts)


def replacement_for_match(
    match: TextMatch,
    replacement: str,
    lemmatizer: Lemmatizer,
    inflect_replacement: bool,
    *,
    capitalize_first_like_match: bool = True,
) -> str:
    if not inflect_replacement:
        return replacement
    return inflect_replacement_phrase(
        replacement,
        match,
        lemmatizer,
        capitalize_first_like_match=capitalize_first_like_match,
    )


def text_equals_at(text: str, start: int, end: int, expected: str) -> bool:
    if start < 0 or end > len(text):
        return False
    return text[start:end].casefold().replace("ё", "е") == expected.casefold().replace("ё", "е")


def already_applied(text: str, match: TextMatch, applied_phrase: str, mode: str, append_separator: str) -> bool:
    if text_equals_at(text, match.end - len(applied_phrase), match.end, applied_phrase):
        return True
    if mode == "append":
        expected_right = append_separator + applied_phrase
        return text_equals_at(text, match.end, match.end + len(expected_right), expected_right)
    return False


def replace_text(
    text: str,
    find_lemmas: list[str],
    replacement: str,
    lemmatizer: Lemmatizer,
    inflect_replacement: bool,
    mode: str = "replace",
    append_separator: str = ", ",
) -> TextReplacement | None:
    matches = find_matches(text, find_lemmas, lemmatizer)
    if not matches:
        return None

    parts: list[str] = []
    cursor = 0
    after_position = 0
    spans: list[ReplacementSpan] = []
    for match in matches:
        added_or_replacement = replacement_for_match(
            match,
            replacement,
            lemmatizer,
            inflect_replacement,
            capitalize_first_like_match=mode != "append",
        )
        if already_applied(text, match, added_or_replacement, mode, append_separator):
            continue
        unchanged = text[cursor : match.start]
        parts.append(unchanged)
        after_position += len(unchanged)
        applied_replacement = (
            match.found + append_separator + added_or_replacement
            if mode == "append"
            else added_or_replacement
        )
        parts.append(applied_replacement)
        spans.append(
            ReplacementSpan(
                start=match.start,
                end=match.end,
                found=match.found,
                replacement=applied_replacement,
                after_start=after_position,
                after_end=after_position + len(applied_replacement),
            )
        )
        after_position += len(applied_replacement)
        cursor = match.end
    parts.append(text[cursor:])
    if not spans:
        return None
    return TextReplacement(before=text, after="".join(parts), spans=tuple(spans))


def snippet(text: str, start: int, end: int, radius: int = 70) -> str:
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    prefix = "..." if left else ""
    suffix = "..." if right < len(text) else ""
    return f"{prefix}{text[left:right]}{suffix}"


def one_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def nearest_non_space(text: str, index: int, step: int) -> int:
    while 0 <= index < len(text) and text[index].isspace():
        index += step
    return index


def word_before_dot(text: str, dot_index: int) -> str:
    match = re.search(r"[A-Za-zА-Яа-яЁё]+$", text[:dot_index])
    return match.group(0) if match else ""


def is_sentence_dot(text: str, dot_index: int) -> bool:
    previous_index = nearest_non_space(text, dot_index - 1, -1)
    next_index = nearest_non_space(text, dot_index + 1, 1)
    previous_char = text[previous_index] if previous_index >= 0 else ""
    next_char = text[next_index] if next_index < len(text) else ""
    if previous_char.isdigit() and next_char.isdigit():
        return False
    previous_word = word_before_dot(text, dot_index)
    previous_word_lower = previous_word.lower().replace("ё", "е")
    if previous_word_lower in NON_SENTENCE_ABBREVIATIONS:
        return False
    if len(previous_word) == 1 and previous_word.isupper():
        return False
    if next_char and next_char.islower():
        return False
    return True


def previous_sentence_dot(text: str, start: int) -> int:
    index = text.rfind(".", 0, start)
    while index >= 0:
        if is_sentence_dot(text, index):
            return index
        index = text.rfind(".", 0, index)
    return -1


def next_sentence_dot(text: str, end: int) -> int:
    index = text.find(".", end)
    while index >= 0:
        if is_sentence_dot(text, index):
            return index
        index = text.find(".", index + 1)
    return -1


def sentence_bounds(text: str, start: int, end: int) -> tuple[int, int]:
    previous_dot = previous_sentence_dot(text, start)
    sentence_start = previous_dot + 1 if previous_dot >= 0 else 0
    while sentence_start < len(text) and text[sentence_start].isspace():
        sentence_start += 1
    next_dot = next_sentence_dot(text, max(end, start))
    sentence_end = next_dot + 1 if next_dot >= 0 else len(text)
    return sentence_start, sentence_end


def sentence_fragment(text: str, start: int, end: int) -> dict[str, Any]:
    sentence_start, sentence_end = sentence_bounds(text, start, end)
    return {
        "start": sentence_start,
        "end": sentence_end,
        "text": one_line(text[sentence_start:sentence_end]),
    }


def format_location(location: dict[str, Any]) -> str:
    if location.get("kind") == "paragraph":
        return f"абзац {location.get('paragraph')}"
    if location.get("kind") == "table_cell":
        return (
            f"таблица {location.get('table')}, строка {location.get('row')}, "
            f"ячейка {location.get('cell')}, абзац {location.get('paragraph')}"
        )
    if location.get("kind") == "cell":
        return f"лист {location.get('sheet')}, ячейка {location.get('cell')}"
    return str(location.get("kind", "document"))


def location_record(
    location: dict[str, Any],
    before: str,
    after: str,
    span: ReplacementSpan,
    hit_index: int,
) -> dict[str, Any]:
    before_sentence = sentence_fragment(before, span.start, span.end)
    after_sentence = sentence_fragment(after, span.after_start, span.after_end)
    return {
        **location,
        "hit": hit_index,
        "found": span.found,
        "replacement": span.replacement,
        "before": snippet(before, span.start, span.end),
        "after": snippet(after, span.after_start, span.after_end),
        "sentence_before": before_sentence["text"],
        "sentence_after": after_sentence["text"],
        "sentence_before_range": {"start": before_sentence["start"], "end": before_sentence["end"]},
        "sentence_after_range": {"start": after_sentence["start"], "end": after_sentence["end"]},
    }


def render_markdown_report(payload: dict[str, Any]) -> str:
    mode_label = "Добавить рядом" if payload.get("mode") == "append" else "Заменить"
    lines = [
        "# Морфологический поиск и замена",
        "",
        f"- Режим: {mode_label}",
        f"- Найти: {payload['find']}",
        f"- Замена / добавка: {payload['replace']}",
        f"- Файлов: {payload['files']}",
        f"- Файлов с заменами: {payload['changed_files']}",
        f"- Замен: {payload['replacements']}",
        f"- Dry-run: {'да' if payload['dry_run'] else 'нет'}",
        "",
        "Каждое срабатывание ниже показывает полное предложение с уже смоделированной заменой: от предыдущей точки до следующей точки, с защитой типичных дат и сокращений.",
        "",
    ]
    for result in payload["results"]:
        file_name = Path(str(result["file"])).name
        lines.extend([f"## {file_name}", "", f"- Тип: {result.get('type', '')}", f"- Замен: {result['replacements']}", ""])
        if int(result["replacements"]) == 0:
            lines.extend(["Срабатываний не найдено.", ""])
            continue
        for location in result.get("locations", []):
            lines.extend(
                [
                    f"### {location['hit']}. {location['found']} -> {location['replacement']}",
                    "",
                    f"- Где: {format_location(location)}",
                    "",
                    "Предложение после замены:",
                    "",
                    "```text",
                    str(location.get("sentence_after", "")),
                    "```",
                    "",
                ]
            )
    return "\n".join(lines).rstrip() + "\n"


def safe_report_stem(path: str) -> str:
    stem = Path(path).stem.strip() or "document"
    safe = INVALID_REPORT_NAME_RE.sub("_", stem).strip(" .")
    return safe or "document"


def report_paths_for_result(report_dir: Path, result: dict[str, Any], used: set[str]) -> tuple[Path, Path]:
    base = safe_report_stem(str(result["file"]))
    index = 1
    while True:
        suffix = "" if index == 1 else f"_{index}"
        stem = f"{base}{suffix}"
        key = stem.casefold()
        if key not in used:
            used.add(key)
            return report_dir / f"{stem}.json", report_dir / f"{stem}.md"
        index += 1


def single_result_payload(payload: dict[str, Any], result: dict[str, Any]) -> dict[str, Any]:
    replacements = int(result["replacements"])
    return {
        key: value
        for key, value in payload.items()
        if key not in {"results", "per_document_report_dir", "per_document_reports"}
    } | {
        "files": 1,
        "changed_files": 1 if replacements > 0 else 0,
        "replacements": replacements,
        "results": [result],
    }


def write_per_document_reports(payload: dict[str, Any], report_dir: Path) -> list[dict[str, Any]]:
    report_dir.mkdir(parents=True, exist_ok=True)
    used: set[str] = set()
    reports: list[dict[str, Any]] = []
    for result in payload["results"]:
        json_path, md_path = report_paths_for_result(report_dir, result, used)
        document_payload = single_result_payload(payload, result)
        json_path.write_text(json.dumps(document_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        md_path.write_text(render_markdown_report(document_payload), encoding="utf-8")
        reports.append(
            {
                "file": str(result["file"]),
                "replacements": int(result["replacements"]),
                "json": str(json_path),
                "markdown": str(md_path),
            }
        )
    return reports


def default_per_document_report_dir(report_path: Path) -> Path:
    return report_path.with_suffix("") if report_path.suffix else report_path.with_name(f"{report_path.name}_files")


def is_supported(path: Path, extensions: set[str]) -> bool:
    return path.is_file() and path.suffix.lower() in extensions and not path.name.startswith("~$")


def collect_sources(input_path: Path, recursive: bool, extensions: set[str]) -> list[Path]:
    if input_path.is_file():
        return [input_path] if is_supported(input_path, extensions) else []
    if not input_path.is_dir():
        raise FileNotFoundError(f"Input path was not found: {input_path}")
    pattern = "**/*" if recursive else "*"
    return sorted(path for path in input_path.glob(pattern) if is_supported(path, extensions))


def unique_path(path: Path, overwrite: bool) -> Path:
    if overwrite or not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(2, 1000):
        candidate = path.with_name(f"{stem}_{index}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Could not allocate unique output path for {path.name}")


def output_path_for(source: Path, input_root: Path, output: Path, multiple: bool, overwrite: bool, suffix_tag: str) -> Path:
    if not multiple and output.suffix.lower() in {".docx", ".xlsx"}:
        return unique_path(output, overwrite)
    relative_parent = source.parent.relative_to(input_root) if input_root.is_dir() else Path()
    target_dir = output / relative_parent
    return unique_path(target_dir / f"{source.stem}_{suffix_tag}{source.suffix}", overwrite)


def paragraph_text_and_map(paragraph: Any) -> tuple[str, list[tuple[int, int]]]:
    chunks: list[str] = []
    char_map: list[tuple[int, int]] = []
    for run_index, run in enumerate(paragraph.runs):
        text = run.text or ""
        chunks.append(text)
        char_map.extend((run_index, char_index) for char_index in range(len(text)))
    return "".join(chunks), char_map


def apply_docx_replacements(paragraph: Any, spans: tuple[ReplacementSpan, ...]) -> bool:
    _text, char_map = paragraph_text_and_map(paragraph)
    if not char_map:
        return False
    for span in sorted(spans, key=lambda item: item.start, reverse=True):
        if span.start >= len(char_map) or span.end - 1 >= len(char_map):
            return False
        start_run_index, start_char_index = char_map[span.start]
        end_run_index, end_char_index = char_map[span.end - 1]
        start_run = paragraph.runs[start_run_index]
        end_run = paragraph.runs[end_run_index]
        if start_run_index == end_run_index:
            text = start_run.text or ""
            start_run.text = text[:start_char_index] + span.replacement + text[end_char_index + 1 :]
            continue
        start_text = start_run.text or ""
        end_text = end_run.text or ""
        start_run.text = start_text[:start_char_index] + span.replacement
        for run_index in range(start_run_index + 1, end_run_index):
            paragraph.runs[run_index].text = ""
        end_run.text = end_text[end_char_index + 1 :]
    return True


def paragraph_xml_key(paragraph: Any) -> str:
    try:
        return str(paragraph._p.getroottree().getpath(paragraph._p))
    except Exception:
        return str(id(paragraph._p))


def append_paragraph_target(
    targets: list[tuple[Any, dict[str, Any]]],
    paragraph: Any,
    location: dict[str, Any],
    seen: set[str],
) -> None:
    key = paragraph_xml_key(paragraph)
    if key in seen:
        return
    seen.add(key)
    targets.append((paragraph, location))


def iter_table_paragraphs(table: Any, table_path: str, seen: set[str]) -> list[tuple[Any, dict[str, Any]]]:
    targets: list[tuple[Any, dict[str, Any]]] = []
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                append_paragraph_target(
                    targets,
                    paragraph,
                    {
                        "kind": "table_cell",
                        "table": table_path,
                        "row": row_index,
                        "cell": cell_index,
                        "paragraph": paragraph_index,
                    },
                    seen,
                )
            for nested_index, nested_table in enumerate(cell.tables, start=1):
                targets.extend(iter_table_paragraphs(nested_table, f"{table_path}.{row_index}.{cell_index}.{nested_index}", seen))
    return targets


def iter_docx_paragraphs(document: Any) -> list[tuple[Any, dict[str, Any]]]:
    targets: list[tuple[Any, dict[str, Any]]] = []
    seen: set[str] = set()
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        append_paragraph_target(targets, paragraph, {"kind": "paragraph", "paragraph": paragraph_index}, seen)
    for table_index, table in enumerate(document.tables, start=1):
        targets.extend(iter_table_paragraphs(table, str(table_index), seen))
    return targets


def process_docx(
    source: Path,
    target: Path,
    find_lemmas: list[str],
    replacement: str,
    lemmatizer: Lemmatizer,
    dry_run: bool,
    inflect_replacement: bool,
    mode: str,
    append_separator: str,
) -> dict[str, Any]:
    document = Document(str(source))
    locations: list[dict[str, Any]] = []
    replacements = 0
    for paragraph, location in iter_docx_paragraphs(document):
        before, _char_map = paragraph_text_and_map(paragraph)
        if not before:
            continue
        result = replace_text(
            before,
            find_lemmas,
            replacement,
            lemmatizer,
            inflect_replacement,
            mode,
            append_separator,
        )
        if result is None:
            continue
        replacements += len(result.spans)
        if not dry_run:
            applied = apply_docx_replacements(paragraph, result.spans)
            if not applied:
                paragraph.text = result.after
        for span in result.spans:
            record = location_record(location, before, result.after, span, len(locations) + 1)
            record["warning"] = "Replacement text uses the formatting of the first matched run; formatting inside replaced fragments may be simplified."
            locations.append(record)
    if replacements and not dry_run:
        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(target))
    return {
        "file": str(source),
        "output": str(target) if replacements and not dry_run else "",
        "type": "docx",
        "replacements": replacements,
        "locations": locations,
    }


def process_xlsx(
    source: Path,
    target: Path,
    find_lemmas: list[str],
    replacement: str,
    lemmatizer: Lemmatizer,
    dry_run: bool,
    inflect_replacement: bool,
    mode: str,
    append_separator: str,
) -> dict[str, Any]:
    workbook = load_workbook(str(source))
    locations: list[dict[str, Any]] = []
    replacements = 0
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if cell.data_type == "f" or not isinstance(cell.value, str):
                    continue
                before = cell.value
                result = replace_text(
                    before,
                    find_lemmas,
                    replacement,
                    lemmatizer,
                    inflect_replacement,
                    mode,
                    append_separator,
                )
                if result is None:
                    continue
                replacements += len(result.spans)
                if not dry_run:
                    cell.value = result.after
                for span in result.spans:
                    locations.append(
                        location_record(
                            {"kind": "cell", "sheet": worksheet.title, "cell": cell.coordinate},
                            before,
                            result.after,
                            span,
                            len(locations) + 1,
                        )
                    )
    if replacements and not dry_run:
        target.parent.mkdir(parents=True, exist_ok=True)
        workbook.save(str(target))
    return {
        "file": str(source),
        "output": str(target) if replacements and not dry_run else "",
        "type": "xlsx",
        "replacements": replacements,
        "locations": locations,
    }


def parse_extensions(raw: str) -> set[str]:
    extensions = {item.strip().lower() for item in raw.split(",") if item.strip()}
    normalized = {item if item.startswith(".") else f".{item}" for item in extensions}
    allowed = {".docx", ".xlsx"}
    invalid = normalized - allowed
    if invalid:
        raise RuntimeError(f"Unsupported extensions: {', '.join(sorted(invalid))}. Allowed: .docx,.xlsx")
    return normalized or allowed


def build_parser(add_help: bool = True) -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Morphological Russian text replacement in DOCX/XLSX files.",
        add_help=add_help,
    )
    parser.add_argument("--input", required=True, help="DOCX/XLSX file or folder.")
    parser.add_argument("--output", default="output", help="Output file or folder.")
    parser.add_argument("--find", required=True, help="Russian phrase to find by lemmas.")
    parser.add_argument("--replace", default="", help="Replacement text or appended phrase.")
    parser.add_argument("--append", default="", help="Phrase to append next to each found phrase. Implies --mode append.")
    parser.add_argument(
        "--mode",
        choices=("replace", "append"),
        default="replace",
        help="replace = replace found phrase; append = keep found phrase and add the replacement phrase after it.",
    )
    parser.add_argument("--append-separator", default=", ", help="Separator used by --mode append.")
    parser.add_argument(
        "--inflect-replacement",
        action="store_true",
        help="Try to inflect replacement words to the grammar of the found phrase.",
    )
    parser.add_argument("--recursive", action="store_true", help="Search input folder recursively.")
    parser.add_argument("--extensions", default=".docx,.xlsx", help="Comma-separated extensions: .docx,.xlsx.")
    parser.add_argument("--dry-run", action="store_true", help="Write only the report, do not save files.")
    parser.add_argument("--report", default="report/morph_replace.json", help="JSON report path.")
    parser.add_argument("--md-report", default="", help="Optional Markdown report path with full replacement sentences.")
    parser.add_argument(
        "--per-document-report-dir",
        default="",
        help="Optional folder for per-document JSON/Markdown reports. Defaults to the JSON report path without extension.",
    )
    parser.add_argument("--overwrite", action="store_true", help="Allow overwriting explicit output paths.")
    return parser


def run(args: argparse.Namespace) -> int:
    input_path = Path(args.input).resolve()
    output = Path(args.output).resolve()
    report_path = Path(args.report).resolve()
    md_report_path = Path(args.md_report).resolve() if str(args.md_report or "").strip() else report_path.with_suffix(".md")
    per_document_report_dir = (
        Path(args.per_document_report_dir).resolve()
        if str(args.per_document_report_dir or "").strip()
        else default_per_document_report_dir(report_path)
    )
    extensions = parse_extensions(str(args.extensions))
    lemmatizer = Lemmatizer()
    find_lemmas = phrase_lemmas(str(args.find), lemmatizer)
    mode = "append" if str(args.append or "").strip() else str(args.mode)
    replacement = str(args.append or args.replace).strip()
    if not replacement:
        raise RuntimeError("--replace обязателен для mode=replace; --replace или --append обязателен для mode=append.")
    append_separator = str(args.append_separator)
    suffix_tag = "morph_appended" if mode == "append" else "morph_replaced"
    sources = collect_sources(input_path, bool(args.recursive), extensions)
    input_root = input_path if input_path.is_dir() else input_path.parent
    multiple = len(sources) != 1 or input_path.is_dir()
    results: list[dict[str, Any]] = []
    for source in sources:
        target = output_path_for(source, input_root, output, multiple, bool(args.overwrite), suffix_tag)
        if source.suffix.lower() == ".docx":
            results.append(
                process_docx(
                    source,
                    target,
                    find_lemmas,
                    replacement,
                    lemmatizer,
                    bool(args.dry_run),
                    bool(args.inflect_replacement),
                    mode,
                    append_separator,
                )
            )
        elif source.suffix.lower() == ".xlsx":
            results.append(
                process_xlsx(
                    source,
                    target,
                    find_lemmas,
                    replacement,
                    lemmatizer,
                    bool(args.dry_run),
                    bool(args.inflect_replacement),
                    mode,
                    append_separator,
                )
            )
    payload = {
        "command": "morph-replace",
        "dry_run": bool(args.dry_run),
        "input": str(input_path),
        "output": str(output),
        "find": str(args.find),
        "find_lemmas": find_lemmas,
        "replace": replacement,
        "mode": mode,
        "append_separator": append_separator if mode == "append" else "",
        "replacement_mode": "inflected" if args.inflect_replacement else "exact",
        "extensions": sorted(extensions),
        "files": len(sources),
        "changed_files": sum(1 for result in results if int(result["replacements"]) > 0),
        "replacements": sum(int(result["replacements"]) for result in results),
        "warnings": [
            "DOCX replacement text uses the formatting of the first matched run; formatting inside replaced fragments may be simplified.",
        ],
        "results": results,
    }
    payload["per_document_report_dir"] = str(per_document_report_dir)
    payload["per_document_reports"] = write_per_document_reports(payload, per_document_report_dir)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    md_report_path.parent.mkdir(parents=True, exist_ok=True)
    md_report_path.write_text(render_markdown_report(payload), encoding="utf-8")
    print(f"Files: {payload['files']}; changed files: {payload['changed_files']}; replacements: {payload['replacements']}")
    print(f"Report: {report_path}")
    print(f"Markdown report: {md_report_path}")
    if args.dry_run:
        print("Dry-run: DOCX/XLSX files were not written.")
    return 0


def main(argv: list[str] | None = None) -> int:
    return run(build_parser().parse_args(argv))


if __name__ == "__main__":
    raise SystemExit(main())
