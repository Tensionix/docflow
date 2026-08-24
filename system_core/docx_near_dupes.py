#!/usr/bin/env python3
"""
DOCX Similarity Map

- Extracts body text, tables, heading sections, and zone codes
- Combines text shingles, heading vocabulary, zone-code overlap, and section alignment
- Searches DOCX files recursively under the input folder
- Outputs Markdown/DOCX/JSON reports with recommendations

Usage:
  python docx_near_dupes.py --input input --threshold 0.30 --out report/docx_similarity_map.md --diff
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import difflib
import hashlib
import itertools
import json
import re
import zlib
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import List, Set, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from _office_common import safe_mkdir, norm_key
from docx_pair_diff import (
    Section,
    alignment_confidence_counts,
    alignment_map_rows,
    canonical_zone_codes,
    extract_content_blocks,
    match_sections_by_similarity,
    section_match_counts,
    section_matches_similarity,
    split_sections,
    title_without_zone_codes,
    tokenize_for_match,
)

WORD_RE = re.compile(r"[0-9A-Za-zА-Яа-яЁё_]+")
REPORT_HEADER_FILL = "D9EAF7"
PAIR_TABLE_WIDTHS_IN = (0.62, 0.62, 0.76, 0.70, 0.78, 1.70, 4.92)
PAIR_RECOMMENDATION_FILL = "EEF5FB"


@dataclass(frozen=True)
class DocProfile:
    doc_id: str
    path: Path
    text: str
    sketch: Set[int]
    sections: list[Section]
    heading_words: Set[str]
    zone_codes: Set[str]
    blocks_total: int
    section_count: int
    table_count: int


@dataclass(frozen=True)
class PairCandidate:
    score: float
    content_similarity: float
    heading_similarity: float
    zone_similarity: float
    structural_similarity: float
    file_a: str
    file_b: str
    confidence: dict[str, int]
    section_status: dict[str, int]
    matched_sections: int
    rows_total: int
    recommendation: str


def rel_docx_id(path: Path, root: Path) -> str:
    base = root.parent if root.is_file() else root
    return path.relative_to(base).as_posix()

def safe_diff_filename(a: str, b: str) -> str:
    name = f"{a}__VS__{b}.diff"
    name = re.sub(r'[<>:"/\\\\|?*]+', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    if len(name) > 180:
        name = f"{name[:170]}__truncated.diff"
    return name


def format_counts(values: dict[str, int] | Counter[str]) -> str:
    return ", ".join(f"{key}: {value}" for key, value in sorted(values.items())) or "нет"


def set_doc_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_table_column_widths(table, widths_in: tuple[float, ...]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths_in):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def write_report_cell(cell, value: object, *, bold: bool = False, fill: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("" if value is None else str(value))
    run.bold = bold
    run.font.size = Pt(8)


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: List[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def section_text(section: Section) -> str:
    return "\n".join(block.normalized for block in section.blocks if block.normalized)


def profile_text_from_sections(sections: list[Section]) -> str:
    return "\n".join(section_text(section) for section in sections)


def heading_words_from_sections(sections: list[Section]) -> Set[str]:
    words: Set[str] = set()
    for section in sections:
        if section.title == "Без раздела":
            continue
        words.update(tokenize_for_match(title_without_zone_codes(section.title)))
    return words


def zone_codes_from_sections(sections: list[Section]) -> Set[str]:
    codes: Set[str] = set()
    for section in sections:
        codes.update(canonical_zone_codes(section.title))
    return codes


def analyze_docx_profile(path: Path, doc_id: str, *, k_words: int, bottom_k: int) -> DocProfile:
    blocks = extract_content_blocks(path)
    sections = split_sections(blocks, depth=1)
    text = norm_key(profile_text_from_sections(sections))
    words = tokenize_words(text)
    return DocProfile(
        doc_id=doc_id,
        path=path,
        text=text,
        sketch=shingle_sketch(words, k_words=k_words, bottom_k=bottom_k),
        sections=sections,
        heading_words=heading_words_from_sections(sections),
        zone_codes=zone_codes_from_sections(sections),
        blocks_total=len(blocks),
        section_count=len(sections),
        table_count=sum(1 for block in blocks if block.kind == "table"),
    )

def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()

def tokenize_words(text: str) -> List[str]:
    return WORD_RE.findall(text)

def shingle_sketch(words: List[str], *, k_words: int, bottom_k: int) -> Set[int]:
    if len(words) < k_words:
        h = zlib.crc32(" ".join(words).encode("utf-8")) & 0xFFFFFFFF
        return {h}
    hashes = []
    for i in range(len(words) - k_words + 1):
        sh = " ".join(words[i:i+k_words])
        h = zlib.crc32(sh.encode("utf-8")) & 0xFFFFFFFF
        hashes.append(h)
    hashes.sort()
    return set(hashes[: min(bottom_k, len(hashes))])

def jaccard(a: Set[int], b: Set[int]) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def nonempty_jaccard(a: Set[str], b: Set[str]) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def recommendation_for_pair(
    *,
    content_similarity: float,
    structural_similarity: float,
    confidence: dict[str, int],
    rows_total: int,
) -> str:
    if content_similarity >= 0.995:
        return "Почти нормализованный дубль: проверить как дубль или версию без технических отличий."
    reliable = confidence.get("уверенно", 0)
    check = confidence.get("проверить", 0)
    doubtful = confidence.get("сомнительно", 0) + confidence.get("нет пары", 0)
    reliable_ratio = reliable / rows_total if rows_total else 0.0
    if reliable_ratio >= 0.80 and doubtful <= 2:
        return "Связанные версии: подходит для сравнения DocFlow или Word COM по кускам."
    if reliable_ratio >= 0.50 or structural_similarity >= 0.55:
        return "Похожий кандидат: перед сравнением просмотреть карту сопоставления."
    if check or doubtful:
        return "Слабый кандидат: возможна связь, но нужны ручная проверка и карта разделов."
    return "Кандидат по общему сходству: проверить вручную."


def base_pair_scores(a: DocProfile, b: DocProfile) -> tuple[float, float, float]:
    return (
        jaccard(a.sketch, b.sketch),
        nonempty_jaccard(a.heading_words, b.heading_words),
        nonempty_jaccard(a.zone_codes, b.zone_codes),
    )


def build_pair_candidate(a: DocProfile, b: DocProfile, *, base_scores: tuple[float, float, float] | None = None) -> PairCandidate:
    if base_scores is None:
        base_scores = base_pair_scores(a, b)
    content_similarity, heading_similarity, zone_similarity = base_scores
    matches = match_sections_by_similarity(a.sections, b.sections, min_similarity=0.30)
    rows = alignment_map_rows(matches)
    confidence = dict(alignment_confidence_counts(rows))
    status = dict(section_match_counts(matches))
    structural_similarity = section_matches_similarity(matches)
    matched_sections = sum(1 for match in matches if match.a is not None and match.b is not None)
    rows_total = len(rows)
    score = max(content_similarity, heading_similarity, zone_similarity, structural_similarity)
    return PairCandidate(
        score=score,
        content_similarity=content_similarity,
        heading_similarity=heading_similarity,
        zone_similarity=zone_similarity,
        structural_similarity=structural_similarity,
        file_a=a.doc_id,
        file_b=b.doc_id,
        confidence=confidence,
        section_status=status,
        matched_sections=matched_sections,
        rows_total=rows_total,
        recommendation=recommendation_for_pair(
            content_similarity=content_similarity,
            structural_similarity=structural_similarity,
            confidence=confidence,
            rows_total=rows_total,
        ),
    )


def unified_diff(a_text: str, b_text: str, a_name: str, b_name: str, context: int = 3) -> str:
    a_lines = a_text.splitlines()
    b_lines = b_text.splitlines()
    diff = difflib.unified_diff(a_lines, b_lines, fromfile=a_name, tofile=b_name, n=context, lineterm="")
    return "\n".join(diff)

def add_docx_diff_block(doc: Document, diff_text: str, max_lines: int) -> None:
    lines = diff_text.splitlines()
    if not lines:
        doc.add_paragraph("Различия нормализованного текста не найдены.")
        return

    if len(lines) > max_lines:
        lines = lines[:max_lines] + [f"... (показана часть diff, всего строк: {len(diff_text.splitlines())})"]

    p = doc.add_paragraph()
    for line in lines:
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8)

def write_docx_report(
    out_path: Path,
    *,
    input_dir: Path,
    files_scanned: int,
    threshold: float,
    results: List[PairCandidate],
    extracted: dict[str, str],
    include_diffs: bool,
    diff_context: int,
    max_diff_lines: int,
    exact_duplicate_groups: List[Tuple[str, List[str]]],
    unreadable_files: List[Tuple[str, str]],
) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    set_doc_landscape(doc)
    doc.add_heading("Карта сходства DOCX", level=1)
    doc.add_paragraph(f"Входная папка: {input_dir}")
    doc.add_paragraph(f"Файлов проверено: {files_scanned}")
    doc.add_paragraph(f"Порог кандидатности: {threshold:.4f}")
    doc.add_paragraph("Итоговый балл = максимум из сходства текста, заголовков, кодов зон и структурного сопоставления.")

    if exact_duplicate_groups:
        doc.add_heading("Точные дубли по SHA-256", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "SHA-256"
        hdr[1].text = "Файлы"
        hdr[2].text = "Количество"
        for sha, paths in exact_duplicate_groups:
            cells = table.add_row().cells
            cells[0].text = sha[:16]
            cells[1].text = "\n".join(paths)
            cells[2].text = str(len(paths))

    if unreadable_files:
        doc.add_heading("Нечитаемые файлы DOCX", level=2)
        for name, err in unreadable_files:
            doc.add_paragraph(f"{name}: {err}", style="List Bullet")

    if not results:
        doc.add_paragraph("Пар выше порога сходства не найдено.")
        doc.save(str(out_path))
        return

    doc.add_heading("Кандидаты и рекомендации", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Итог", "Текст", "Заг.", "Зоны", "Структ.", "Уверенность", "Файлы"]
    for index, header in enumerate(headers):
        align = WD_ALIGN_PARAGRAPH.LEFT if index in {5, 6} else WD_ALIGN_PARAGRAPH.CENTER
        write_report_cell(hdr[index], header, bold=True, fill=REPORT_HEADER_FILL, align=align)
    set_table_column_widths(table, PAIR_TABLE_WIDTHS_IN)
    for item in results:
        cells = table.add_row().cells
        values = [
            f"{item.score:.4f}",
            f"{item.content_similarity:.4f}",
            f"{item.heading_similarity:.4f}",
            f"{item.zone_similarity:.4f}",
            f"{item.structural_similarity:.4f}",
            format_counts(item.confidence),
            f"{item.file_a}\n{item.file_b}",
        ]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if index in {5, 6} else WD_ALIGN_PARAGRAPH.CENTER
            write_report_cell(cells[index], value, align=align)
        recommendation_cells = table.add_row().cells
        merged = recommendation_cells[0].merge(recommendation_cells[-1])
        write_report_cell(
            merged,
            f"Рекомендация: {item.recommendation}",
            bold=True,
            fill=PAIR_RECOMMENDATION_FILL,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
    set_table_column_widths(table, PAIR_TABLE_WIDTHS_IN)

    if include_diffs:
        doc.add_heading("Фрагменты diff", level=2)
        for idx, item in enumerate(results, start=1):
            doc.add_heading(f"{idx}. {item.file_a} vs {item.file_b} ({item.score:.4f})", level=3)
            diff_text = unified_diff(extracted[item.file_a], extracted[item.file_b], item.file_a, item.file_b, context=diff_context)
            add_docx_diff_block(doc, diff_text, max_diff_lines)

    doc.save(str(out_path))

def write_json_report(
    out_path: Path,
    *,
    input_dir: Path,
    files_scanned: int,
    threshold: float,
    exact_duplicate_groups: List[Tuple[str, List[str]]],
    unreadable_files: List[Tuple[str, str]],
    results: List[PairCandidate],
) -> None:
    payload = {
        "input": str(input_dir),
        "files_scanned": files_scanned,
        "threshold": threshold,
        "exact_duplicate_groups": [
            {"sha256": sha, "files": paths} for sha, paths in exact_duplicate_groups
        ],
        "unreadable_files": [
            {"file": name, "error": err} for name, err in unreadable_files
        ],
        "pairs_above_threshold": [
            {
                "score": item.score,
                "content_similarity": item.content_similarity,
                "heading_similarity": item.heading_similarity,
                "zone_similarity": item.zone_similarity,
                "structural_similarity": item.structural_similarity,
                "file_a": item.file_a,
                "file_b": item.file_b,
                "confidence": item.confidence,
                "section_status": item.section_status,
                "matched_sections": item.matched_sections,
                "rows_total": item.rows_total,
                "recommendation": item.recommendation,
            }
            for item in results
        ],
    }
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

def main() -> int:
    ap = argparse.ArgumentParser(description="Build a DOCX similarity map recursively in a folder.")
    ap.add_argument("--input", default="input", help="Input folder with .docx files (default: input, recursive)")
    ap.add_argument("--threshold", type=float, default=0.30, help="Candidate score threshold (default: 0.30)")
    ap.add_argument("--k-words", type=int, default=5, help="Shingle size in words (default: 5)")
    ap.add_argument("--bottom-k", type=int, default=6000, help="Sketch size (default: 6000)")
    ap.add_argument("--out", default="report/docx_similarity_map.md", help="Output Markdown report path")
    ap.add_argument("--diff", action="store_true", help="Write unified diffs for matched pairs into output/diffs")
    ap.add_argument("--docx-report", action="store_true", default=True, help="Write a DOCX report next to the Markdown report (default: on)")
    ap.add_argument("--no-docx-report", action="store_false", dest="docx_report", help="Do not write a DOCX report")
    ap.add_argument("--docx-out", default="", help="Optional DOCX report path")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    ap.add_argument("--diff-context", type=int, default=3, help="Diff context lines (default: 3)")
    ap.add_argument("--max-diff-lines", type=int, default=4000, help="Truncate diffs longer than this (default: 4000)")
    args = ap.parse_args()

    in_dir = Path(args.input).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)

    if in_dir.is_file():
        eligible = in_dir.suffix.lower() == ".docx" and not in_dir.name.startswith("~$")
        docx_files = [in_dir] if eligible else []
    else:
        docx_files = sorted(
            p for p in in_dir.rglob("*.docx")
            if p.is_file() and not p.name.startswith("~$")
        )

    hash_groups: dict[str, list[str]] = defaultdict(list)
    for p in docx_files:
        hash_groups[file_sha256(p)].append(rel_docx_id(p, in_dir))
    exact_duplicate_groups = sorted(
        (sha, sorted(paths)) for sha, paths in hash_groups.items() if len(paths) > 1
    )

    if len(docx_files) < 2:
        lines = ["# Карта сходства DOCX\n", "Недостаточно файлов DOCX для сравнения.\n"]
        if exact_duplicate_groups:
            lines.append("## Точные дубли по SHA-256\n")
            for sha, paths in exact_duplicate_groups:
                lines.append(f"- `{sha}`")
                for p in paths:
                    lines.append(f"  - `{p}`")
        out_path.write_text("\n".join(lines), encoding="utf-8")
        json_out = Path(args.json_out).resolve() if args.json_out else out_path.with_suffix(".json")
        write_json_report(
            json_out,
            input_dir=in_dir,
            files_scanned=len(docx_files),
            threshold=args.threshold,
            exact_duplicate_groups=exact_duplicate_groups,
            unreadable_files=[],
            results=[],
        )
        if args.docx_report:
            docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
            write_docx_report(
                docx_out,
                input_dir=in_dir,
                files_scanned=len(docx_files),
                threshold=args.threshold,
                results=[],
                extracted={},
                include_diffs=False,
                diff_context=args.diff_context,
                max_diff_lines=args.max_diff_lines,
                exact_duplicate_groups=exact_duplicate_groups,
                unreadable_files=[],
            )
            print(f"[OK] Wrote DOCX report: {docx_out}")
        print(f"[OK] Wrote: {out_path}")
        print(f"[OK] Wrote JSON report: {json_out}")
        return 0

    extracted: dict[str, str] = {}
    profiles: dict[str, DocProfile] = {}
    unreadable_files: List[Tuple[str, str]] = []
    for p in docx_files:
        doc_id = rel_docx_id(p, in_dir)
        try:
            profile = analyze_docx_profile(p, doc_id, k_words=args.k_words, bottom_k=args.bottom_k)
        except Exception as exc:
            unreadable_files.append((doc_id, str(exc)))
            continue
        profiles[doc_id] = profile
        extracted[doc_id] = profile.text

    results: List[PairCandidate] = []
    names = sorted(profiles)
    for a, b in itertools.combinations(names, 2):
        base_scores = base_pair_scores(profiles[a], profiles[b])
        if max(base_scores) < args.threshold:
            continue
        candidate = build_pair_candidate(profiles[a], profiles[b], base_scores=base_scores)
        if candidate.score >= args.threshold:
            results.append(candidate)
    results.sort(reverse=True, key=lambda item: item.score)

    lines = []
    lines.append("# Карта сходства DOCX\n")
    lines.append(f"- Входная папка: `{in_dir}`")
    lines.append(f"- Файлов проверено: **{len(docx_files)}**")
    lines.append(f"- Порог кандидатности: **{args.threshold:.4f}**")
    lines.append("- Итоговый балл = максимум из сходства текста, заголовков, кодов зон и структурного сопоставления.\n")

    lines.append("## Точные дубли по SHA-256\n")
    if not exact_duplicate_groups:
        lines.append("Побайтово одинаковые дубли DOCX не найдены.\n")
    else:
        for sha, paths in exact_duplicate_groups:
            lines.append(f"- SHA-256 `{sha}`")
            for p in paths:
                lines.append(f"  - `{p}`")
            lines.append("")

    lines.append("## Нечитаемые файлы DOCX\n")
    if not unreadable_files:
        lines.append("Нечитаемые файлы DOCX не найдены.\n")
    else:
        for name, err in unreadable_files:
            lines.append(f"- `{name}`: {err}")
        lines.append("")

    if not results:
        lines.append("Пар выше порога сходства не найдено.\n")
        out_path.write_text("\n".join(lines), encoding="utf-8")
        json_out = Path(args.json_out).resolve() if args.json_out else out_path.with_suffix(".json")
        write_json_report(
            json_out,
            input_dir=in_dir,
            files_scanned=len(docx_files),
            threshold=args.threshold,
            exact_duplicate_groups=exact_duplicate_groups,
            unreadable_files=unreadable_files,
            results=[],
        )
        if args.docx_report:
            docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
            write_docx_report(
                docx_out,
                input_dir=in_dir,
                files_scanned=len(docx_files),
                threshold=args.threshold,
                results=[],
                extracted=extracted,
                include_diffs=args.diff,
                diff_context=args.diff_context,
                max_diff_lines=args.max_diff_lines,
                exact_duplicate_groups=exact_duplicate_groups,
                unreadable_files=unreadable_files,
            )
            print(f"[OK] Wrote DOCX report: {docx_out}")
        print(f"[OK] Wrote: {out_path}")
        print(f"[OK] Wrote JSON report: {json_out}")
        return 0

    lines.append("## Кандидаты и рекомендации\n")
    lines.append("| Итог | Текст | Заголовки | Коды зон | Структура | Уверенность | Файл A | Файл B | Рекомендация |")
    lines.append("|---:|---:|---:|---:|---:|---|---|---|---|")
    for item in results:
        lines.append(
            f"| {item.score:.4f} | {item.content_similarity:.4f} | {item.heading_similarity:.4f} | "
            f"{item.zone_similarity:.4f} | {item.structural_similarity:.4f} | {format_counts(item.confidence)} | "
            f"`{item.file_a}` | `{item.file_b}` | {item.recommendation} |"
        )
    lines.append("")

    if args.diff:
        diffs_dir = out_path.parent / "diffs"
        safe_mkdir(diffs_dir)
        lines.append("## Файлы diff\n")
        for item in results:
            diff_text = unified_diff(extracted[item.file_a], extracted[item.file_b], item.file_a, item.file_b, context=args.diff_context)
            dl = diff_text.splitlines()
            if len(dl) > args.max_diff_lines:
                diff_text = "\n".join(dl[:args.max_diff_lines]) + f"\n... (показана часть diff, всего строк: {len(dl)})"
            diff_name = safe_diff_filename(item.file_a, item.file_b)
            diff_path = diffs_dir / diff_name
            diff_path.write_text(diff_text, encoding="utf-8")
            lines.append(f"- `{diff_path.relative_to(out_path.parent)}` (score {item.score:.4f})")
        lines.append("")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    json_out = Path(args.json_out).resolve() if args.json_out else out_path.with_suffix(".json")
    write_json_report(
        json_out,
        input_dir=in_dir,
        files_scanned=len(docx_files),
        threshold=args.threshold,
        exact_duplicate_groups=exact_duplicate_groups,
        unreadable_files=unreadable_files,
        results=results,
    )
    if args.docx_report:
        docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
        write_docx_report(
            docx_out,
            input_dir=in_dir,
            files_scanned=len(docx_files),
            threshold=args.threshold,
            results=results,
            extracted=extracted,
            include_diffs=args.diff,
            diff_context=args.diff_context,
            max_diff_lines=args.max_diff_lines,
            exact_duplicate_groups=exact_duplicate_groups,
            unreadable_files=unreadable_files,
        )
        print(f"[OK] Wrote DOCX report: {docx_out}")
    print(f"[OK] Wrote: {out_path}")
    print(f"[OK] Wrote JSON report: {json_out}")
    if args.diff:
        print(f"[OK] Diffs: {out_path.parent / 'diffs'}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
