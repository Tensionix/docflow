#!/usr/bin/env python3
"""
DOCX comparison through Microsoft Word COM.

This wraps Word's native CompareDocuments command and writes a third DOCX
containing Word revisions. It is intentionally separate from the deterministic
section-aware diff because Word Compare can fail or hang on very large pairs.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import csv
from copy import deepcopy
from dataclasses import dataclass
import json
import os
from pathlib import Path
import shutil
import subprocess
import sys
import time
from zipfile import ZIP_DEFLATED, ZipFile

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt
from lxml import etree

from _office_common import safe_mkdir
from docx_pair_diff import (
    Section,
    SectionMatch,
    extract_content_blocks,
    match_confidence,
    match_sections_by_similarity,
    section_identity_reason,
    split_sections,
)


SCRIPT_RUNNER = Path(__file__).resolve().parent / "word_com" / "Compare-DocxWord.ps1"
MERGE_SCRIPT = Path(__file__).resolve().parent / "docx_merge.py"
PAIR_DIFF_SCRIPT = Path(__file__).resolve().parent / "docx_pair_diff.py"
WORD_UNAVAILABLE_MESSAGE = (
    "[ERROR] Microsoft Word COM недоступен из текущего сеанса. Этот инструмент\n"
    "требует установленного Microsoft Word и доступного COM-запуска Word.Application\n"
    "на этой машине. Если Word не может обработать пару документов, используйте\n"
    "section-aware сравнение DocFlow вместо Word COM."
)
DEFAULT_TIMEOUT_SEC = 600
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_BODY = f"{{{W_NS}}}body"
W_P = f"{{{W_NS}}}p"
W_TBL = f"{{{W_NS}}}tbl"
W_SECT_PR = f"{{{W_NS}}}sectPr"
SMART_INITIAL_CHUNKS = 8
MAX_PROBLEM_BISECT_DEPTH = 8
WORD_COMPARE_ATTEMPTS = 3
WORD_COMPARE_RETRY_DELAY_SEC = 2


@dataclass(frozen=True)
class ChunkPlan:
    ordinal: int
    title: str
    matches: tuple[SectionMatch, ...]


@dataclass(frozen=True)
class ChunkResult:
    ordinal: int
    title: str
    status: str
    returncode: int
    old_chunk: Path
    new_chunk: Path
    result_docx: Path
    report: Path
    fallback: bool = False
    attempts: int = 1
    label: str = ""


@dataclass(frozen=True)
class ProblemChunk:
    slot_index: int
    ordinal: int
    title: str
    old_chunk: Path
    new_chunk: Path
    returncode: int
    attempts: int
    old_ranges: list[tuple[int, int]]
    new_ranges: list[tuple[int, int]]


@dataclass(frozen=True)
class ChunkPiece:
    title: str
    old_ranges: list[tuple[int, int]]
    new_ranges: list[tuple[int, int]]
    source_ordinal: int
    subordinal: int
    subcount: int


def _pwsh_candidates(root: Path) -> list[str]:
    return [
        str(root / "system_core" / "powershell" / "pwsh.exe"),
        "pwsh",
        "powershell",
    ]


def _resolve_pwsh(root: Path) -> str:
    for candidate in _pwsh_candidates(root):
        if Path(candidate).exists():
            return candidate
        found = shutil.which(candidate)
        if found:
            return found
    raise RuntimeError("PowerShell was not found. Install portable PowerShell or add pwsh/powershell to PATH.")


def _check_word_available(pwsh: str, root: Path) -> int:
    script = f"""
$word = $null
try {{
    $word = New-Object -ComObject Word.Application
    exit 0
}} catch {{
    Write-Output @'
{WORD_UNAVAILABLE_MESSAGE}
'@
    Write-Output $_.Exception.Message
    exit 2
}} finally {{
    if ($word -ne $null) {{
        try {{ $word.Quit() | Out-Null }} catch {{}}
        try {{ [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null }} catch {{}}
    }}
}}
"""
    process = subprocess.run(
        [pwsh, "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        cwd=str(root),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if process.stdout:
        print(process.stdout.rstrip())
    return process.returncode


def _validate_docx(path: Path, label: str) -> Path:
    resolved = path.resolve()
    if not resolved.exists():
        raise FileNotFoundError(f"{label}: файл не найден: {resolved}")
    if resolved.suffix.lower() != ".docx":
        raise RuntimeError(f"{label}: нужен файл .docx: {resolved}")
    if resolved.name.startswith("~$"):
        raise RuntimeError(f"{label}: временный Office lock-файл не подходит: {resolved}")
    return resolved


def _winword_pids() -> set[int]:
    if os.name != "nt":
        return set()
    ps_exe = shutil.which("powershell") or shutil.which("pwsh")
    if ps_exe:
        try:
            completed = subprocess.run(
                [
                    ps_exe,
                    "-NoProfile",
                    "-Command",
                    "[Console]::OutputEncoding=[System.Text.UTF8Encoding]::new($false); "
                    "Get-Process WINWORD -ErrorAction SilentlyContinue | ForEach-Object { $_.Id }",
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.DEVNULL,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=10,
            )
            return {int(line.strip()) for line in completed.stdout.splitlines() if line.strip().isdigit()}
        except Exception:
            pass
    try:
        completed = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq WINWORD.EXE", "/FO", "CSV", "/NH"],
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=10,
        )
    except Exception:
        return set()
    pids: set[int] = set()
    for row in csv.reader(completed.stdout.splitlines()):
        if len(row) < 2:
            continue
        try:
            pids.add(int(row[1]))
        except ValueError:
            continue
    return pids


def _kill_processes(pids: set[int]) -> None:
    if os.name != "nt":
        return
    for pid in sorted(pids):
        subprocess.run(
            ["taskkill", "/PID", str(pid), "/F", "/T"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=10,
        )


def _kill_word_before_compare() -> None:
    pids = _winword_pids()
    if not pids:
        return
    print(f"[INFO] Перед Word Compare остановлены процессы WINWORD: {', '.join(str(pid) for pid in sorted(pids))}")
    _kill_processes(pids)
    deadline = time.monotonic() + 10
    while time.monotonic() < deadline:
        remaining = _winword_pids() & pids
        if not remaining:
            return
        time.sleep(0.5)


def _write_failure_report(
    report_path: Path,
    original: Path,
    revised: Path,
    out_path: Path,
    message: str,
) -> None:
    safe_mkdir(report_path.parent)
    report_path.write_text(
        "\n".join(
            [
                "# Отчёт сравнения DOCX через Microsoft Word COM",
                "",
                "- Статус: **FAILED**",
                f"- Выходной файл: `{out_path}`",
                f"- Старый документ: `{original}`",
                f"- Новый документ: `{revised}`",
                "- Метод: Microsoft Word COM CompareDocuments.",
                "",
                "## Ошибка",
                "",
                "```text",
                message,
                "```",
                "",
                "Если Word COM не может обработать эту пару документов, используйте сравнение DocFlow по структуре разделов.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def _remove_rsid_markup(xml_bytes: bytes) -> tuple[bytes, int]:
    root = etree.fromstring(xml_bytes)
    removed = 0
    for element in root.iter():
        for attr_name in list(element.attrib):
            qname = etree.QName(attr_name)
            if qname.namespace == W_NS and qname.localname.startswith("rsid"):
                del element.attrib[attr_name]
                removed += 1
    for rsids in root.findall(f".//{{{W_NS}}}rsids"):
        parent = rsids.getparent()
        if parent is not None:
            parent.remove(rsids)
            removed += 1
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=False), removed


def _remove_comment_content_types(xml_bytes: bytes, removed_parts: set[str]) -> tuple[bytes, int]:
    if not removed_parts:
        return xml_bytes, 0
    removed_parts_lower = {part.lower() for part in removed_parts}
    root = etree.fromstring(xml_bytes)
    removed = 0
    for override in list(root):
        part_name = override.get("PartName", "").lstrip("/").lower()
        if part_name in removed_parts_lower:
            root.remove(override)
            removed += 1
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=False), removed


def _remove_comment_relationships(xml_bytes: bytes, removed_parts: set[str]) -> tuple[bytes, int]:
    if not removed_parts:
        return xml_bytes, 0
    removed_names = {Path(part).name.lower() for part in removed_parts}
    root = etree.fromstring(xml_bytes)
    removed = 0
    for relationship in list(root):
        target = relationship.get("Target", "").replace("\\", "/")
        rel_type = relationship.get("Type", "").lower()
        target_name = Path(target).name.lower()
        if target_name in removed_names or any(name.removesuffix(".xml") in rel_type for name in removed_names):
            root.remove(relationship)
            removed += 1
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=False), removed


def _clean_docx_copy(src: Path, dst: Path) -> dict[str, object]:
    """Make a safer DOCX copy for Word Compare without changing visible content."""
    safe_mkdir(dst.parent)
    with ZipFile(src, "r") as zin:
        names = set(zin.namelist())
        known_orphan_paths = {
            "word/commentsExtended.xml",
            "word/commentsExtensible.xml",
            "word/commentsIds.xml",
            "word/people.xml",
        }
        orphan_comment_parts: set[str] = set()
        orphan_comment_targets: set[str] = set()
        if "word/comments.xml" not in names:
            orphan_names = {
                "commentsextended.xml",
                "commentsextensible.xml",
                "commentsids.xml",
                "people.xml",
            }
            orphan_comment_parts = {
                name
                for name in names
                if name.startswith("word/") and Path(name).name.lower() in orphan_names
            }
            orphan_comment_targets = set(orphan_comment_parts) | known_orphan_paths

        removed_rsid = 0
        removed_rels = 0
        removed_content_types = 0
        with ZipFile(dst, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                name = info.filename
                if name in orphan_comment_parts:
                    continue
                data = zin.read(name)
                if name in {"word/document.xml", "word/settings.xml"}:
                    data, count = _remove_rsid_markup(data)
                    removed_rsid += count
                elif name == "[Content_Types].xml":
                    data, count = _remove_comment_content_types(data, orphan_comment_targets)
                    removed_content_types += count
                elif name.endswith(".rels"):
                    data, count = _remove_comment_relationships(data, orphan_comment_targets)
                    removed_rels += count
                zout.writestr(info, data)
    return {
        "path": str(dst),
        "removed_orphan_comment_parts": sorted(orphan_comment_parts),
        "removed_rsid_markup": removed_rsid,
        "removed_relationships": removed_rels,
        "removed_content_types": removed_content_types,
    }


def _body_ranges(sections: tuple[Section, ...]) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for section in sections:
        indices = [block.body_index for block in section.blocks]
        if indices:
            ranges.append((min(indices), max(indices)))
    ranges.sort()
    merged: list[tuple[int, int]] = []
    for start, end in ranges:
        if not merged or start > merged[-1][1] + 1:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
    return merged


def _range_item_count(ranges: list[tuple[int, int]]) -> int:
    return sum(max(0, end - start + 1) for start, end in ranges)


def _document_body_tags(src: Path) -> list[str]:
    with ZipFile(src, "r") as zin:
        xml_bytes = zin.read("word/document.xml")
    root = etree.fromstring(xml_bytes)
    body = root.find(W_BODY)
    if body is None:
        raise RuntimeError("DOCX body not found in word/document.xml")
    return [child.tag for child in list(body)]


def _indices_from_ranges(ranges: list[tuple[int, int]]) -> list[int]:
    indices: list[int] = []
    for start, end in ranges:
        indices.extend(range(start, end + 1))
    return sorted(set(indices))


def _ranges_from_indices(indices: list[int]) -> list[tuple[int, int]]:
    if not indices:
        return []
    sorted_indices = sorted(indices)
    ranges: list[tuple[int, int]] = []
    start = previous = sorted_indices[0]
    for index in sorted_indices[1:]:
        if index == previous + 1:
            previous = index
            continue
        ranges.append((start, previous))
        start = previous = index
    ranges.append((start, previous))
    return ranges


def _split_ranges_safely(
    ranges: list[tuple[int, int]],
    body_tags: list[str],
    pieces: int,
) -> list[list[tuple[int, int]]]:
    indices = _indices_from_ranges(ranges)
    if not indices:
        return [[] for _ in range(max(1, pieces))]
    pieces = max(1, min(pieces, len(indices)))
    groups: list[list[int]] = []
    for piece_index in range(pieces):
        start_pos = round(piece_index * len(indices) / pieces)
        end_pos = round((piece_index + 1) * len(indices) / pieces)
        group = indices[start_pos:end_pos]
        if group:
            groups.append(group)

    # Keep a paragraph immediately before a table with that table. This avoids
    # chunks that start with a naked table after its caption/preheader.
    for index in range(1, len(groups)):
        start = groups[index][0]
        previous = start - 1
        if (
            0 <= start < len(body_tags)
            and body_tags[start] == W_TBL
            and previous >= 0
            and body_tags[previous] == W_P
            and previous in groups[index - 1]
            and len(groups[index - 1]) > 1
        ):
            groups[index - 1].remove(previous)
            groups[index].insert(0, previous)

    return [_ranges_from_indices(group) for group in groups if group]


def _split_sections_into_ranges(src: Path, sections: tuple[Section, ...], pieces: int) -> list[list[tuple[int, int]]]:
    ranges = _body_ranges(sections)
    if pieces <= 1:
        return [ranges]
    return _split_ranges_safely(ranges, _document_body_tags(src), pieces)


def _full_docx_body_ranges(src: Path) -> list[tuple[int, int]]:
    tags = _document_body_tags(src)
    return _ranges_from_indices([index for index, tag in enumerate(tags) if tag != W_SECT_PR])


def _split_ranges_pair(
    *,
    old_docx: Path,
    new_docx: Path,
    old_ranges: list[tuple[int, int]],
    new_ranges: list[tuple[int, int]],
    pieces: int,
) -> list[tuple[list[tuple[int, int]], list[tuple[int, int]]]]:
    old_parts = (
        _split_ranges_safely(old_ranges, _document_body_tags(old_docx), pieces)
        if old_ranges
        else _blank_range_list(pieces)
    )
    new_parts = (
        _split_ranges_safely(new_ranges, _document_body_tags(new_docx), pieces)
        if new_ranges
        else _blank_range_list(pieces)
    )
    count = max(len(old_parts), len(new_parts))
    while len(old_parts) < count:
        old_parts.append([])
    while len(new_parts) < count:
        new_parts.append([])
    return list(zip(old_parts, new_parts))


def _contains_index(ranges: list[tuple[int, int]], index: int) -> bool:
    return any(start <= index <= end for start, end in ranges)


def _replace_document_body(xml_bytes: bytes, ranges: list[tuple[int, int]]) -> bytes:
    root = etree.fromstring(xml_bytes)
    body = root.find(W_BODY)
    if body is None:
        raise RuntimeError("DOCX body not found in word/document.xml")

    original_children = list(body)
    sect_pr = next((deepcopy(child) for child in reversed(original_children) if child.tag == W_SECT_PR), None)
    selected: list[etree._Element] = []
    if ranges:
        selected = [
            deepcopy(child)
            for index, child in enumerate(original_children)
            if child.tag != W_SECT_PR and _contains_index(ranges, index)
        ]
    if not selected:
        selected = [etree.Element(W_P)]
    if sect_pr is None:
        sect_pr = etree.Element(W_SECT_PR)

    for child in original_children:
        body.remove(child)
    for child in selected:
        body.append(child)
    body.append(sect_pr)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=False)


def _write_docx_ranges(src: Path, dst: Path, ranges: list[tuple[int, int]]) -> None:
    safe_mkdir(dst.parent)
    with ZipFile(src, "r") as zin, ZipFile(dst, "w", ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "word/document.xml":
                data = _replace_document_body(data, ranges)
            zout.writestr(info, data)


def _write_docx_chunk(src: Path, dst: Path, sections: tuple[Section, ...]) -> None:
    _write_docx_ranges(src, dst, _body_ranges(sections))


def _section_match_sort_key(match: SectionMatch, anchors: list[tuple[int, int]]) -> tuple[int, int, int]:
    if match.b is not None:
        return (match.b.ordinal * 2, 0 if match.a else 1, match.a.ordinal if match.a else 0)
    if match.a is None:
        return (10**9, 0, 0)
    later = [b_ord for a_ord, b_ord in anchors if a_ord > match.a.ordinal]
    if later:
        return (min(later) * 2 - 1, 0, match.a.ordinal)
    earlier = [b_ord for a_ord, b_ord in anchors if a_ord < match.a.ordinal]
    if earlier:
        return (max(earlier) * 2 + 1, 0, match.a.ordinal)
    return (match.a.ordinal * 2, 0, match.a.ordinal)


def _ordered_matches(matches: list[SectionMatch]) -> list[SectionMatch]:
    anchors = [(match.a.ordinal, match.b.ordinal) for match in matches if match.a is not None and match.b is not None]
    return sorted(matches, key=lambda match: _section_match_sort_key(match, anchors))


def _match_weight(match: SectionMatch) -> int:
    a_words = match.a.word_count if match.a is not None else 0
    b_words = match.b.word_count if match.b is not None else 0
    return max(a_words, b_words, 1)


def _chunk_title(matches: tuple[SectionMatch, ...], ordinal: int) -> str:
    titles: list[str] = []
    for match in matches:
        section = match.b or match.a
        if section is None:
            continue
        title = section.title.strip() or "Без раздела"
        if title not in titles:
            titles.append(title)
    if not titles:
        return f"Кусок {ordinal}"
    if len(titles) == 1:
        return titles[0]
    return f"{titles[0]} ... +{len(titles) - 1}"


def _build_chunk_plan(matches: list[SectionMatch], chunks: int | None) -> list[ChunkPlan]:
    ordered = _ordered_matches(matches)
    if not ordered:
        return []
    if chunks is None or chunks <= 0:
        return [
            ChunkPlan(ordinal=index, title=_chunk_title((match,), index), matches=(match,))
            for index, match in enumerate(ordered, start=1)
        ]

    chunk_count = max(1, min(chunks, len(ordered)))
    total_weight = sum(_match_weight(match) for match in ordered)
    target_weight = max(1, total_weight // chunk_count)
    groups: list[list[SectionMatch]] = []
    current: list[SectionMatch] = []
    current_weight = 0

    for index, match in enumerate(ordered):
        remaining_matches = len(ordered) - index
        remaining_groups = chunk_count - len(groups) - 1
        if current and (current_weight >= target_weight or remaining_matches <= remaining_groups):
            groups.append(current)
            current = []
            current_weight = 0
        current.append(match)
        current_weight += _match_weight(match)

    if current:
        groups.append(current)

    while len(groups) > chunk_count:
        tail = groups.pop()
        groups[-1].extend(tail)

    return [
        ChunkPlan(ordinal=index, title=_chunk_title(tuple(group), index), matches=tuple(group))
        for index, group in enumerate(groups, start=1)
    ]


def _parse_chunk_count(value: str) -> int | None:
    text = (value or "").strip()
    if not text:
        return None
    try:
        parsed = int(float(text.replace(",", ".")))
    except ValueError as exc:
        raise RuntimeError("Количество кусков должно быть целым числом или пустым полем.") from exc
    if parsed <= 0:
        raise RuntimeError("Количество кусков должно быть больше нуля или пустым полем.")
    return parsed


def _blank_range_list(count: int) -> list[list[tuple[int, int]]]:
    return [[] for _ in range(max(1, count))]


def _expand_plan_into_pieces(old_clean: Path, new_clean: Path, plan: ChunkPlan) -> list[ChunkPiece]:
    old_sections = tuple(match.a for match in plan.matches if match.a is not None)
    new_sections = tuple(match.b for match in plan.matches if match.b is not None)
    old_ranges = _body_ranges(old_sections)
    new_ranges = _body_ranges(new_sections)
    return [
        ChunkPiece(
            title=plan.title,
            old_ranges=old_ranges,
            new_ranges=new_ranges,
            source_ordinal=plan.ordinal,
            subordinal=1,
            subcount=1,
        )
    ]


def _run_word_compare(
    *,
    pwsh: str,
    root: Path,
    original: Path,
    revised: Path,
    out_path: Path,
    report_path: Path,
    author: str,
    timeout_sec: int,
    close_word_before_compare: bool,
) -> tuple[int, str]:
    if close_word_before_compare:
        _kill_word_before_compare()
    command = [
        pwsh,
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(SCRIPT_RUNNER),
        "-Original",
        str(original),
        "-Revised",
        str(revised),
        "-Out",
        str(out_path),
        "-Report",
        str(report_path),
        "-Author",
        str(author),
    ]
    before_pids = _winword_pids()
    try:
        process = subprocess.run(
            command,
            cwd=str(root),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=max(30, int(timeout_sec or DEFAULT_TIMEOUT_SEC)),
        )
    except subprocess.TimeoutExpired as exc:
        new_pids = _winword_pids() - before_pids
        if new_pids:
            _kill_processes(new_pids)
        timeout_message = (
            f"Word COM CompareDocuments не завершился за {max(30, int(timeout_sec or DEFAULT_TIMEOUT_SEC))} секунд.\n"
            "Новые процессы WINWORD, созданные операцией, были остановлены."
        )
        _write_failure_report(report_path, original, revised, out_path, timeout_message)
        output = str(exc.stdout or "").rstrip()
        return 124, "\n".join(part for part in [output, f"[ERROR] {timeout_message}"] if part)
    if process.returncode != 0 and not report_path.exists():
        _write_failure_report(report_path, original, revised, out_path, process.stdout or "Word COM compare failed.")
    return process.returncode, process.stdout or ""


def _run_word_compare_with_retries(
    *,
    pwsh: str,
    root: Path,
    original: Path,
    revised: Path,
    out_path: Path,
    report_path: Path,
    author: str,
    timeout_sec: int,
    close_word_before_compare: bool,
    attempts: int = WORD_COMPARE_ATTEMPTS,
) -> tuple[int, str, int]:
    outputs: list[str] = []
    last_rc = 1
    attempts = max(1, int(attempts or 1))
    for attempt in range(1, attempts + 1):
        if attempt > 1:
            print(
                f"[INFO] Повтор Word Compare: попытка {attempt}/{attempts} "
                f"для {original.name} -> {revised.name}"
            )
            try:
                out_path.unlink(missing_ok=True)
            except OSError:
                pass
        rc, output = _run_word_compare(
            pwsh=pwsh,
            root=root,
            original=original,
            revised=revised,
            out_path=out_path,
            report_path=report_path,
            author=author,
            timeout_sec=timeout_sec,
            close_word_before_compare=close_word_before_compare,
        )
        last_rc = rc
        if output:
            outputs.append(output.rstrip())
        if rc == 0 and out_path.exists():
            return rc, "\n".join(part for part in outputs if part), attempt
        if attempt < attempts:
            time.sleep(WORD_COMPARE_RETRY_DELAY_SEC)
    return last_rc, "\n".join(part for part in outputs if part), attempts


def _run_docflow_redline_fallback(
    *,
    root: Path,
    original: Path,
    revised: Path,
    out_path: Path,
    report_path: Path,
) -> tuple[int, str]:
    stats_path = report_path.with_suffix(".stats.json")
    command = [
        sys.executable,
        str(PAIR_DIFF_SCRIPT),
        "--a",
        str(original),
        "--b",
        str(revised),
        "--out",
        str(report_path),
        "--stats-out",
        str(stats_path),
        "--redline-out",
        str(out_path),
        "--max-diff-lines",
        "300",
        "--max-changes",
        "80",
    ]
    process = subprocess.run(
        command,
        cwd=str(root),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return process.returncode, process.stdout or ""


def _match_confidence(match: SectionMatch) -> str:
    return match_confidence(match)


def _match_confidence_reason(match: SectionMatch) -> str:
    return section_identity_reason(match)


def _alignment_plan_rows(plans: list[ChunkPlan]) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for plan in plans:
        for match in plan.matches:
            rows.append(
                {
                    "chunk": plan.ordinal,
                    "chunk_title": plan.title,
                    "status": match.status,
                    "confidence": _match_confidence(match),
                    "confidence_reason": _match_confidence_reason(match),
                    "similarity": match.similarity if match.a is not None and match.b is not None else None,
                    "old_ordinal": match.a.ordinal if match.a is not None else None,
                    "old_title": match.a.title if match.a is not None else "",
                    "old_words": match.a.word_count if match.a is not None else 0,
                    "new_ordinal": match.b.ordinal if match.b is not None else None,
                    "new_title": match.b.title if match.b is not None else "",
                    "new_words": match.b.word_count if match.b is not None else 0,
                }
            )
    return rows


def _write_cell(cell, text: object, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(8)


def _write_alignment_plan_reports(
    *,
    docx_path: Path,
    md_path: Path,
    json_path: Path,
    original: Path,
    revised: Path,
    plans: list[ChunkPlan],
) -> list[dict[str, object]]:
    rows = _alignment_plan_rows(plans)
    safe_mkdir(docx_path.parent)

    json_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

    md_lines = [
        "# План сопоставления разделов",
        "",
        f"- Старый DOCX: `{original}`",
        f"- Новый DOCX: `{revised}`",
        f"- Строк сопоставления: {len(rows)}",
        "",
        "| Кусок | Статус | Уверенность | Основание | Сходство | Старый раздел | Новый раздел |",
        "|---:|---|---|---|---:|---|---|",
    ]
    for row in rows:
        similarity = "" if row["similarity"] is None else f"{float(row['similarity']):.4f}"
        old_title = f"A{row['old_ordinal']}: {row['old_title']}".replace("|", "\\|") if row["old_ordinal"] else ""
        new_title = f"B{row['new_ordinal']}: {row['new_title']}".replace("|", "\\|") if row["new_ordinal"] else ""
        reason = str(row["confidence_reason"]).replace("|", "\\|")
        md_lines.append(
            f"| {row['chunk']} | {row['status']} | {row['confidence']} | {reason} | {similarity} | {old_title} | {new_title} |"
        )
    md_path.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    doc.add_heading("План сопоставления разделов для Word COM", level=1)
    doc.add_paragraph(f"Старый DOCX: {original}")
    doc.add_paragraph(f"Новый DOCX: {revised}")
    doc.add_paragraph(
        "Таблица показывает, какие разделы старого документа были сопоставлены с разделами нового "
        "перед разрезкой и отправкой в Microsoft Word COM."
    )

    headers = ["Кусок", "Статус", "Уверенность", "Основание", "Сходство", "Старый #", "Старый раздел", "Новый #", "Новый раздел", "Слова A/B"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        _write_cell(table.rows[0].cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        similarity = "" if row["similarity"] is None else f"{float(row['similarity']):.4f}"
        values = [
            row["chunk"],
            row["status"],
            row["confidence"],
            row["confidence_reason"],
            similarity,
            "" if row["old_ordinal"] is None else row["old_ordinal"],
            row["old_title"],
            "" if row["new_ordinal"] is None else row["new_ordinal"],
            row["new_title"],
            f"{row['old_words']}/{row['new_words']}",
        ]
        for index, value in enumerate(values):
            _write_cell(cells[index], value)

    doc.save(str(docx_path))
    return rows


def _write_incomplete_chunk_docx(
    *,
    out_path: Path,
    report_path: Path,
    title: str,
    label: str,
    original: Path,
    revised: Path,
    reason: str,
) -> None:
    safe_mkdir(out_path.parent)
    safe_mkdir(report_path.parent)
    doc = Document()
    heading = doc.add_paragraph()
    run = heading.add_run("НЕПОЛНОЕ СРАВНЕНИЕ WORD COM")
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph(title)
    doc.add_paragraph(
        "Эта доля документа не была сравнена Word COM: дальнейшее дихотомическое "
        "разрезание уже невозможно или достигнут предел безопасного дробления."
    )
    doc.add_paragraph(f"Метка доли: {label}")
    doc.add_paragraph(f"Причина: {reason}")
    doc.add_paragraph(f"Старый DOCX-кусок: {original}")
    doc.add_paragraph(f"Новый DOCX-кусок: {revised}")
    doc.save(str(out_path))
    report_path.write_text(
        "\n".join(
            [
                "# Неполное сравнение Word COM",
                "",
                f"- Статус: **PARTIAL**",
                f"- Доля: `{label}`",
                f"- Раздел: {title}",
                f"- Причина: {reason}",
                f"- Старый DOCX-кусок: `{original}`",
                f"- Новый DOCX-кусок: `{revised}`",
                f"- В итоговый DOCX вставлен жёлтый предупреждающий блок: `{out_path}`",
                "",
            ]
        ),
        encoding="utf-8",
    )


def _merge_chunk_results(root: Path, result_dir: Path, files: list[Path], out_path: Path, report_path: Path) -> tuple[int, str]:
    if len(files) == 1:
        shutil.copy2(files[0], out_path)
        report_path.write_text(
            "# Склейка DOCX-кусков\n\nОдин кусок, склейка не потребовалась.\n",
            encoding="utf-8",
        )
        return 0, ""
    command = [
        sys.executable,
        str(MERGE_SCRIPT),
        "--input",
        str(result_dir),
        "--files",
        ";".join(path.name for path in files),
        "--out",
        str(out_path),
        "--report",
        str(report_path),
    ]
    process = subprocess.run(
        command,
        cwd=str(root),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return process.returncode, process.stdout or ""


def _compare_problem_chunk(
    *,
    root: Path,
    pwsh: str,
    author: str,
    problem: ProblemChunk,
    old_chunk_dir: Path,
    new_chunk_dir: Path,
    result_dir: Path,
    chunk_report_dir: Path,
    timeout_sec: int,
    close_word_before_compare: bool,
) -> list[ChunkResult]:
    problem_timeout = max(300, int(timeout_sec or DEFAULT_TIMEOUT_SEC))
    first_parts = _split_ranges_pair(
        old_docx=problem.old_chunk,
        new_docx=problem.new_chunk,
        old_ranges=_full_docx_body_ranges(problem.old_chunk),
        new_ranges=_full_docx_body_ranges(problem.new_chunk),
        pieces=2,
    )
    print(f"[INFO] Дихотомия проблемного куска {problem.ordinal}: стартуем с {len(first_parts)} долей")

    def compare_ranges(
        *,
        old_ranges: list[tuple[int, int]],
        new_ranges: list[tuple[int, int]],
        label: str,
        depth: int,
    ) -> list[ChunkResult]:
        file_label = label.replace(".", "_")
        old_chunk = old_chunk_dir / f"{file_label}_old.docx"
        new_chunk = new_chunk_dir / f"{file_label}_new.docx"
        result_docx = result_dir / f"{file_label}_word_compare.docx"
        chunk_report = chunk_report_dir / f"{file_label}_word_compare.md"
        title = f"{problem.title} — доля {label}"

        _write_docx_ranges(problem.old_chunk, old_chunk, old_ranges)
        _write_docx_ranges(problem.new_chunk, new_chunk, new_ranges)
        rc, output, attempts = _run_word_compare_with_retries(
            pwsh=pwsh,
            root=root,
            original=old_chunk,
            revised=new_chunk,
            out_path=result_docx,
            report_path=chunk_report,
            author=author,
            timeout_sec=problem_timeout,
            close_word_before_compare=close_word_before_compare,
            attempts=1,
        )
        if output:
            print(output.rstrip())

        if rc != 0 or not result_docx.exists():
            if depth < MAX_PROBLEM_BISECT_DEPTH:
                halves = _split_ranges_pair(
                    old_docx=problem.old_chunk,
                    new_docx=problem.new_chunk,
                    old_ranges=old_ranges,
                    new_ranges=new_ranges,
                    pieces=2,
                )
                if len(halves) > 1:
                    print(f"[WARN] Доля {label} не прошла Word COM; делим пополам.")
                    nested: list[ChunkResult] = []
                    for index, (old_half, new_half) in enumerate(halves, start=1):
                        nested.extend(
                            compare_ranges(
                                old_ranges=old_half,
                                new_ranges=new_half,
                                label=f"{label}.{index}",
                                depth=depth + 1,
                            )
                        )
                    return nested

            incomplete_report = chunk_report_dir / f"{file_label}_incomplete.md"
            incomplete_docx = result_dir / f"{file_label}_incomplete.docx"
            reason = (
                f"Word COM завершился с кодом {rc}; глубина дихотомии {depth}/"
                f"{MAX_PROBLEM_BISECT_DEPTH}, дальше долю не делим."
            )
            _write_incomplete_chunk_docx(
                out_path=incomplete_docx,
                report_path=incomplete_report,
                title=title,
                label=label,
                original=old_chunk,
                revised=new_chunk,
                reason=reason,
            )
            print(f"[WARN] Доля {label} вставлена как неполное сравнение: {incomplete_docx}")
            return [
                ChunkResult(
                    problem.ordinal,
                    title,
                    "incomplete",
                    rc,
                    old_chunk,
                    new_chunk,
                    incomplete_docx,
                    incomplete_report,
                    False,
                    attempts,
                    label,
                )
            ]

        return [
            ChunkResult(
                problem.ordinal,
                title,
                "word",
                rc,
                old_chunk,
                new_chunk,
                result_docx,
                chunk_report,
                False,
                attempts,
                label,
            )
        ]

    results: list[ChunkResult] = []
    for index, (old_part, new_part) in enumerate(first_parts, start=1):
        results.extend(compare_ranges(old_ranges=old_part, new_ranges=new_part, label=f"{problem.ordinal}.{index}", depth=1))
    return results


def _write_chunked_report(
    report_path: Path,
    stats_path: Path,
    original: Path,
    revised: Path,
    out_path: Path,
    clean_stats: dict[str, object],
    plans: list[ChunkPlan],
    results: list[ChunkResult],
    merge_report: Path,
    alignment_docx: Path,
    alignment_md: Path,
    alignment_json: Path,
    chunk_count: int | None,
    split_mode: str,
    old_sections: int,
    new_sections: int,
) -> None:
    status = "OK" if all(result.status == "word" for result in results) else "PARTIAL"
    if any(result.status == "failed" for result in results):
        status = "FAILED"
    split_text = f"умная разрезка: стартовые {SMART_INITIAL_CHUNKS} кусков, затем дихотомия проблемных долей" if split_mode == "smart" else (
        "по верхнему порядку разделов" if chunk_count is None else f"примерно на {chunk_count} кусков"
    )
    lines = [
        "# Отчёт сравнения DOCX через Microsoft Word COM",
        "",
        f"- Статус: **{status}**",
        "- Метод: Word COM CompareDocuments по кускам.",
        f"- Разрезание: {split_text}",
        f"- Выходной файл: `{out_path}`",
        f"- Старый документ: `{original}`",
        f"- Новый документ: `{revised}`",
        f"- Разделов в старом документе: {old_sections}",
        f"- Разделов в новом документе: {new_sections}",
        f"- Кусков сравнения: {len(results)}",
        f"- Статистика: `{stats_path}`",
        f"- План сопоставления DOCX: `{alignment_docx}`",
        f"- План сопоставления Markdown: `{alignment_md}`",
        f"- План сопоставления JSON: `{alignment_json}`",
        f"- Отчёт склейки: `{merge_report}`",
        "",
        "## Очистка перед резкой",
        "",
        f"- Старый DOCX: удалено rsid-меток: {clean_stats['old']['removed_rsid_markup']}; "
        f"осиротевших comment-партов: {len(clean_stats['old']['removed_orphan_comment_parts'])}.",
        f"- Новый DOCX: удалено rsid-меток: {clean_stats['new']['removed_rsid_markup']}; "
        f"осиротевших comment-партов: {len(clean_stats['new']['removed_orphan_comment_parts'])}.",
        "",
        "## Куски",
        "",
        "| # | Статус | Раздел | Отчёт |",
        "|---:|---|---|---|",
    ]
    for result in results[:250]:
        method = {
            "word": "Word COM",
            "fallback": "DocFlow fallback",
            "incomplete": "PARTIAL: жёлтый предупреждающий блок",
            "failed": "Ошибка",
        }.get(result.status, result.status)
        title = result.title.replace("|", "\\|")
        label = result.label or str(result.ordinal)
        if result.attempts > 1:
            method = f"{method}, попыток {result.attempts}"
        lines.append(f"| {label} | {method} | {title} | `{result.report}` |")
    if len(results) > 250:
        lines.append(f"| ... | ... | Показаны первые 250 из {len(results)} кусков | ... |")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _compare_chunked(
    *,
    root: Path,
    pwsh: str,
    original: Path,
    revised: Path,
    out_path: Path,
    report_path: Path,
    author: str,
    timeout_sec: int,
    split_mode: str,
    chunks_text: str,
    close_word_before_compare: bool,
) -> int:
    split_mode = split_mode if split_mode in {"smart", "chunks"} else "smart"
    chunk_count = SMART_INITIAL_CHUNKS if split_mode == "smart" else _parse_chunk_count(chunks_text)
    work_dir = report_path.parent / "word_compare_chunked"
    clean_dir = work_dir / "cleaned"
    old_chunk_dir = work_dir / "old_chunks"
    new_chunk_dir = work_dir / "new_chunks"
    problem_old_chunk_dir = work_dir / "problem_old_chunks"
    problem_new_chunk_dir = work_dir / "problem_new_chunks"
    result_dir = work_dir / "result_chunks"
    chunk_report_dir = work_dir / "chunk_reports"
    for folder in (clean_dir, old_chunk_dir, new_chunk_dir, problem_old_chunk_dir, problem_new_chunk_dir, result_dir, chunk_report_dir):
        safe_mkdir(folder)

    old_clean = clean_dir / "old_cleaned.docx"
    new_clean = clean_dir / "new_cleaned.docx"
    clean_stats = {
        "old": _clean_docx_copy(original, old_clean),
        "new": _clean_docx_copy(revised, new_clean),
    }

    old_sections = split_sections(extract_content_blocks(old_clean), depth=1)
    new_sections = split_sections(extract_content_blocks(new_clean), depth=1)
    matches = match_sections_by_similarity(old_sections, new_sections, min_similarity=0.30)
    plans = _build_chunk_plan(matches, chunk_count)
    if not plans:
        raise RuntimeError("Не удалось построить план разрезания DOCX: разделы не найдены.")
    alignment_docx = report_path.with_name("word_compare_alignment_plan.docx")
    alignment_md = report_path.with_name("word_compare_alignment_plan.md")
    alignment_json = report_path.with_name("word_compare_alignment_plan.json")
    alignment_rows = _write_alignment_plan_reports(
        docx_path=alignment_docx,
        md_path=alignment_md,
        json_path=alignment_json,
        original=original,
        revised=revised,
        plans=plans,
    )

    result_slots: list[list[ChunkResult]] = []
    problems: list[ProblemChunk] = []
    chunk_timeout = max(60, min(max(30, int(timeout_sec or DEFAULT_TIMEOUT_SEC)), 240))
    piece_counter = 0
    for plan in plans:
        for piece in _expand_plan_into_pieces(old_clean, new_clean, plan):
            piece_counter += 1
            slot_index = len(result_slots)
            result_slots.append([])
            old_chunk = old_chunk_dir / f"{piece_counter:03d}_old.docx"
            new_chunk = new_chunk_dir / f"{piece_counter:03d}_new.docx"
            result_docx = result_dir / f"{piece_counter:03d}_word_compare.docx"
            chunk_report = chunk_report_dir / f"{piece_counter:03d}_word_compare.md"

            _write_docx_ranges(old_clean, old_chunk, piece.old_ranges)
            _write_docx_ranges(new_clean, new_chunk, piece.new_ranges)
            rc, output, attempts = _run_word_compare_with_retries(
                pwsh=pwsh,
                root=root,
                original=old_chunk,
                revised=new_chunk,
                out_path=result_docx,
                report_path=chunk_report,
                author=author,
                timeout_sec=chunk_timeout,
                close_word_before_compare=close_word_before_compare,
                attempts=1,
            )
            if output:
                print(output.rstrip())
            if rc != 0 or not result_docx.exists():
                problems.append(
                    ProblemChunk(
                        slot_index=slot_index,
                        ordinal=piece_counter,
                        title=piece.title,
                        old_chunk=old_chunk,
                        new_chunk=new_chunk,
                        returncode=rc,
                        attempts=attempts,
                        old_ranges=piece.old_ranges,
                        new_ranges=piece.new_ranges,
                    )
                )
                print(f"[WARN] Кусок {piece_counter} отложен на дополнительное дробление после основного массива.")
                continue

            result_slots[slot_index] = [
                ChunkResult(
                    piece_counter,
                    piece.title,
                    "word",
                    rc,
                    old_chunk,
                    new_chunk,
                    result_docx,
                    chunk_report,
                    False,
                    attempts,
                    str(piece_counter),
                )
            ]

    if problems:
        print(f"[INFO] Основной массив завершён. Проблемных кусков: {len(problems)}. Запускаем дополнительное дробление.")
    for problem in problems:
        result_slots[problem.slot_index] = _compare_problem_chunk(
            root=root,
            pwsh=pwsh,
            author=author,
            problem=problem,
            old_chunk_dir=problem_old_chunk_dir,
            new_chunk_dir=problem_new_chunk_dir,
            result_dir=result_dir,
            chunk_report_dir=chunk_report_dir,
            timeout_sec=chunk_timeout,
            close_word_before_compare=close_word_before_compare,
        )

    results = [result for slot in result_slots for result in slot]
    result_files = [result.result_docx for result in results]
    incomplete_count = sum(1 for result in results if result.status == "incomplete")

    if any(result.status == "failed" for result in results):
        failed = ", ".join(str(result.ordinal) for result in results if result.status == "failed")
        raise RuntimeError(f"Не удалось сравнить куски DOCX: {failed}")
    if incomplete_count:
        print(f"[WARN] Неполное завершение: {incomplete_count} долей вставлены жёлтыми предупреждающими блоками.")

    merge_report = work_dir / "merge_chunks.md"
    merge_rc, merge_output = _merge_chunk_results(root, result_dir, result_files, out_path, merge_report)
    if merge_output:
        print(merge_output.rstrip())
    if merge_rc != 0:
        raise RuntimeError(f"Склейка результатов сравнения завершилась с кодом {merge_rc}")

    doc = Document(str(out_path))
    stats_path = report_path.with_suffix(".stats.json")
    stats = {
        "status": "OK" if all(result.status == "word" for result in results) else "PARTIAL",
        "method": "chunked_word_com_compare",
        "split_mode": split_mode,
        "requested_chunks": chunk_count,
        "section_chunks": len(plans),
        "initial_chunks": piece_counter,
        "problem_chunks": len(problems),
        "incomplete_chunks": incomplete_count,
        "created_chunks": len(results),
        "alignment_rows": len(alignment_rows),
        "alignment_docx": str(alignment_docx),
        "alignment_md": str(alignment_md),
        "alignment_json": str(alignment_json),
        "old_sections": len(old_sections),
        "new_sections": len(new_sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "cleaned": clean_stats,
        "chunks": [
            {
                "ordinal": result.ordinal,
                "label": result.label or str(result.ordinal),
                "title": result.title,
                "status": result.status,
                "incomplete": result.status == "incomplete",
                "fallback": result.fallback,
                "attempts": result.attempts,
                "old_chunk": str(result.old_chunk),
                "new_chunk": str(result.new_chunk),
                "result_docx": str(result.result_docx),
                "report": str(result.report),
            }
            for result in results
        ],
    }
    stats_path.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_chunked_report(
        report_path,
        stats_path,
        original,
        revised,
        out_path,
        clean_stats,
        plans,
        results,
        merge_report,
        alignment_docx,
        alignment_md,
        alignment_json,
        chunk_count,
        split_mode,
        len(old_sections),
        len(new_sections),
    )
    print(
        f"[OK] Chunked Word compare: chunks={len(results)}, "
        f"section_chunks={len(plans)}, paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}"
    )
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Compare two DOCX files using Microsoft Word COM CompareDocuments.")
    parser.add_argument("--a", required=True, help="Original / old DOCX")
    parser.add_argument("--b", required=True, help="Revised / new DOCX")
    parser.add_argument("--out", default="output/word_compare/word_compare_result.docx", help="Output DOCX with Word revisions")
    parser.add_argument("--report", default="report/docx_word_compare.md", help="Output Markdown report")
    parser.add_argument("--author", default="Audion DocFlow", help="Revision author shown by Word")
    parser.add_argument("--timeout-sec", type=int, default=DEFAULT_TIMEOUT_SEC, help="Word COM timeout in seconds")
    parser.add_argument("--chunked", action="store_true", help="Clean, split, compare chunks, then merge results")
    parser.add_argument("--split-mode", choices=("smart", "chunks"), default="smart", help="Chunked split mode")
    parser.add_argument(
        "--auto-smart-on-problem",
        action="store_true",
        help="If direct Word COM compare fails, retry through smart chunked compare.",
    )
    parser.add_argument(
        "--close-word-before-compare",
        action="store_true",
        help="Force-close WINWORD before each Word COM compare call.",
    )
    parser.add_argument("--chunks", default="", help="Optional target chunk count. Empty means one chunk per matched section.")
    args = parser.parse_args()

    root = Path(__file__).resolve().parents[1]
    original = _validate_docx(Path(args.a), "DOCX A")
    revised = _validate_docx(Path(args.b), "DOCX B")
    if original == revised:
        raise RuntimeError("DOCX A и DOCX B должны быть разными файлами.")

    out_path = Path(args.out).resolve()
    report_path = Path(args.report).resolve()
    safe_mkdir(out_path.parent)
    safe_mkdir(report_path.parent)

    if not SCRIPT_RUNNER.exists():
        raise RuntimeError(f"Word COM compare script was not found: {SCRIPT_RUNNER}")

    pwsh = _resolve_pwsh(root)
    word_rc = _check_word_available(pwsh, root)
    if word_rc != 0:
        return 2 if word_rc == 2 else word_rc

    if args.chunked:
        return _compare_chunked(
            root=root,
            pwsh=pwsh,
            original=original,
            revised=revised,
            out_path=out_path,
            report_path=report_path,
            author=args.author,
            timeout_sec=args.timeout_sec,
            split_mode=args.split_mode,
            chunks_text=args.chunks,
            close_word_before_compare=args.close_word_before_compare,
        )

    rc, output, _attempts = _run_word_compare_with_retries(
        pwsh=pwsh,
        root=root,
        original=original,
        revised=revised,
        out_path=out_path,
        report_path=report_path,
        author=args.author,
        timeout_sec=args.timeout_sec,
        close_word_before_compare=args.close_word_before_compare,
        attempts=1 if args.auto_smart_on_problem else WORD_COMPARE_ATTEMPTS,
    )
    if output:
        print(output.rstrip())
    if rc != 0 or not out_path.exists():
        if args.auto_smart_on_problem:
            direct_report = report_path.with_name(f"{report_path.stem}.direct_word_compare_failed.md")
            if report_path.exists():
                try:
                    shutil.copy2(report_path, direct_report)
                except OSError:
                    pass
            print("[WARN] Прямое сравнение Word COM не завершилось успешно. Запускаем умную разрезку.")
            return _compare_chunked(
                root=root,
                pwsh=pwsh,
                original=original,
                revised=revised,
                out_path=out_path,
                report_path=report_path,
                author=args.author,
                timeout_sec=args.timeout_sec,
                split_mode="smart",
                chunks_text="",
                close_word_before_compare=args.close_word_before_compare,
            )
        if not report_path.exists():
            _write_failure_report(report_path, original, revised, out_path, output or "Word COM compare failed.")
        return rc

    doc = Document(str(out_path))
    print(f"[OK] Validated with python-docx: paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
