#!/usr/bin/env python3
"""
DOCX Text Hygiene (Fix)

Edits DOCX XML text nodes (w:t) in-place and writes a new DOCX.
Fixes are applied within each text node to avoid breaking formatting.

Fix rules:
- collapse multiple spaces
- remove spaces before punctuation
- add missing space after punctuation (, ; : ! ?) with basic heuristics
- optional: add missing space after dot (.) (more false positives, disabled by default)
- optional: remove strikethrough text (explicit flag, disabled by default)

Usage:
  python docx_text_hygiene_fix.py --input input --outdir output/hygiene_fixed
  python docx_text_hygiene_fix.py --file x.docx --out output/x_fixed.docx
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import re
from pathlib import Path

from docx import Document
from _office_common import safe_mkdir, find_docx_files, mirrored_output_path, rel_posix, write_json_file
from docx_xml_tools import (
    read_zip_map, write_zip_map, list_xml_parts, NS,
    _etree_from_bytes, _etree_to_bytes,
    find_strikethrough_style_ids, remove_strikethrough_text,
)
try:
    from docx_text_hygiene_scan import (
        find_text_issues as _find_text_issues,
        is_missing_after_dot_candidate as _is_missing_after_dot_candidate,
    )
except ImportError:  # pragma: no cover - package import fallback
    from system_core.docx_text_hygiene_scan import (
        find_text_issues as _find_text_issues,
        is_missing_after_dot_candidate as _is_missing_after_dot_candidate,
    )

PAT_DOUBLE_SPACE = re.compile(r" {2,}")
PAT_SPACE_BEFORE_PUNCT = re.compile(r"\s+([,.;:!?])")
PAT_MISSING_AFTER_PUNCT = re.compile(r"([,;:!?])(?=[^\s\]\)\}\>\"'”’])")
PAT_MISSING_AFTER_DOT = re.compile(r"(?<!\d)(\.)(?=[A-Za-zА-Яа-яЁё])")  # heuristic

TEMP_OFFICE_PATTERNS = ("~$*.docx", "~$*.xlsx", "~$*.pptx", "~$*.doc", "~$*.xls", "~$*.ppt")


def fix_missing_after_dot(text: str) -> str:
    return PAT_MISSING_AFTER_DOT.sub(lambda match: ". " if _is_missing_after_dot_candidate(text, match) else ".", text)


def fix_text(s: str, fix_dot: bool) -> str:
    if not s:
        return s

    # Normalize NBSP to space
    s = s.replace("\u00A0", " ")

    # 1) collapse multiple spaces
    s = PAT_DOUBLE_SPACE.sub(" ", s)

    # 2) remove spaces before punctuation
    s = PAT_SPACE_BEFORE_PUNCT.sub(r"\1", s)

    # 3) add missing space after punctuation (but avoid cases like "...,)" etc.)
    s = PAT_MISSING_AFTER_PUNCT.sub(r"\1 ", s)

    # 4) optional: after dot
    if fix_dot:
        s = fix_missing_after_dot(s)

    return s


def find_text_issues(part: str, text_node_index: int, text: str, fix_dot: bool) -> list[dict[str, object]]:
    return _find_text_issues(part, text_node_index, text, fix_dot)


def process_one(in_path: Path, out_path: Path, fix_dot: bool, remove_strikethrough: bool = False) -> tuple[int, int, list[dict[str, object]]]:
    files = read_zip_map(in_path)
    strike_style_ids = find_strikethrough_style_ids(files) if remove_strikethrough else set()
    changed_nodes = 0
    changed_chars = 0
    findings: list[dict[str, object]] = []

    for part in list_xml_parts(files):
        if part not in files:
            continue
        data = files[part]
        if remove_strikethrough:
            data = remove_strikethrough_text(
                data,
                mode="delete",
                strike_style_ids=strike_style_ids,
            )
        tree = _etree_from_bytes(data)
        root = tree.getroot()
        modified = data != files[part]

        for text_node_index, t in enumerate(root.xpath(".//w:t", namespaces=NS)):
            if not t.text:
                continue
            if t.text.strip() == "" and t.get("{http://www.w3.org/XML/1998/namespace}space") == "preserve":
                continue
            planned_findings = find_text_issues(part, text_node_index, t.text, fix_dot)
            new = fix_text(t.text, fix_dot)
            if new != t.text:
                changed_nodes += 1
                changed_chars += abs(len(new) - len(t.text))
                for finding in planned_findings:
                    finding["action"] = "FIX"
                    finding["autofix"] = "applied"
                findings.extend(planned_findings)
                t.text = new
                modified = True

        if modified:
            files[part] = _etree_to_bytes(tree)

    write_zip_map(out_path, files)
    return changed_nodes, changed_chars, findings

def list_temp_office_files(root: Path) -> list[tuple[Path, str]]:
    root = root.resolve()
    found: list[tuple[Path, str]] = []
    for pattern in TEMP_OFFICE_PATTERNS:
        for path in root.rglob(pattern):
            if not path.is_file():
                continue
            resolved = path.resolve()
            if not str(resolved).lower().startswith(str(root).lower()):
                continue
            rel = resolved.relative_to(root).as_posix()
            found.append((resolved, rel))
    return sorted(found, key=lambda item: item[1])

def purge_temp_office_files(root: Path) -> list[str]:
    removed: list[str] = []
    for path, rel in list_temp_office_files(root):
        path.unlink()
        removed.append(rel)
    return removed

def write_change_report(report_path: Path, rows: list[dict[str, object]], removed_temp: list[str]) -> None:
    safe_mkdir(report_path.parent)
    lines: list[str] = []
    lines.append("# Изменения текстовой гигиены DOCX\n")
    lines.append(f"- Временных файлов Office удалено: **{len(removed_temp)}**\n")
    if removed_temp:
        lines.append("## Удалённые временные файлы Office\n")
        for item in removed_temp:
            lines.append(f"- `{item}`")
        lines.append("")
    lines.append("## Обработанные файлы DOCX\n")
    lines.append("| Файл | Выходной файл | Изменённые текстовые узлы | Примерное изменение символов | Статус |")
    lines.append("|---|---|---:|---:|---|")
    for row in rows:
        lines.append(
            f"| `{row['file']}` | `{row['output']}` | {row['changed_nodes']} | {row['changed_chars']} | {row['status']} |"
        )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    doc = Document()
    doc.add_heading("Изменения текстовой гигиены DOCX", level=1)
    doc.add_paragraph(f"Временных файлов Office удалено: {len(removed_temp)}")
    if removed_temp:
        doc.add_heading("Удалённые временные файлы Office", level=2)
        for item in removed_temp:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Обработанные файлы DOCX", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Файл"
    hdr[1].text = "Выходной файл"
    hdr[2].text = "Изменённые текстовые узлы"
    hdr[3].text = "Примерное изменение символов"
    hdr[4].text = "Статус"
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row["file"])
        cells[1].text = str(row["output"])
        cells[2].text = str(row["changed_nodes"])
        cells[3].text = str(row["changed_chars"])
        cells[4].text = str(row["status"])
    doc.save(str(report_path.with_suffix(".docx")))


def hygiene_fix_json_payload(input_path: Path, outdir: Path | None, rows: list[dict[str, object]], removed_temp: list[str], options: dict[str, object]) -> dict[str, object]:
    changed_nodes = sum(int(row.get("changed_nodes", 0) or 0) for row in rows)
    changed_chars = sum(int(row.get("changed_chars", 0) or 0) for row in rows)
    errors = sum(1 for row in rows if str(row.get("status", "")) != "OK")
    return {
        "tool": "docx_text_hygiene_fix",
        "version": 1,
        "input": str(input_path),
        "outdir": str(outdir) if outdir is not None else "",
        "options": options,
        "summary": {
            "files": len(rows),
            "ok": len(rows) - errors,
            "errors": errors,
            "changed_nodes": changed_nodes,
            "changed_chars": changed_chars,
            "removed_temp_office_files": len(removed_temp),
        },
        "removed_temp_office_files": removed_temp,
        "files": rows,
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Fix basic text hygiene issues in DOCX files.")
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--input", help="Input folder (process all .docx recursively)")
    g.add_argument("--file", help="Single DOCX file")
    ap.add_argument("--outdir", default="output/hygiene_fixed", help="Output folder for batch mode")
    ap.add_argument("--out", help="Output DOCX path for single-file mode")
    ap.add_argument("--report", default="report/docx_text_hygiene_changes.md", help="Output changes report path")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    ap.add_argument("--fix-dot", action="store_true", help="Also fix missing space after '.' (more false positives)")
    ap.add_argument("--remove-strikethrough", action="store_true", help="Remove visible strikethrough text. Disabled by default because it changes document content.")
    ap.add_argument("--dry-run", action="store_true", help="Show planned removals/outputs without deleting files or writing output")
    args = ap.parse_args()

    if args.input:
        in_dir = Path(args.input).resolve()
        out_dir = Path(args.outdir).resolve()
        if args.dry_run:
            temp_files = list_temp_office_files(in_dir)
            docx_files = find_docx_files(in_dir)
            for _path, rel in temp_files:
                print(f"[DRY-RUN] would remove: {rel}")
            for p in docx_files:
                out_path = mirrored_output_path(p, in_dir, out_dir)
                print(f"[DRY-RUN] would write: {out_path}")
            print(f"[OK] Dry-run temporary Office files: {len(temp_files)}")
            print(f"[OK] Dry-run DOCX files: {len(docx_files)}")
            return 0
        safe_mkdir(out_dir)

        removed_temp = purge_temp_office_files(in_dir)
        docx_files = find_docx_files(in_dir)
        if not docx_files:
            print(f"[WARN] No .docx files found in: {in_dir}")
            report_path = Path(args.report).resolve()
            write_change_report(report_path, [], removed_temp)
            if args.json_out:
                json_out = Path(args.json_out).resolve()
                write_json_file(
                    json_out,
                    hygiene_fix_json_payload(
                        in_dir,
                        out_dir,
                        [],
                        removed_temp,
                        {"fix_dot": bool(args.fix_dot), "remove_strikethrough": bool(args.remove_strikethrough)},
                    ),
                )
                print(f"[OK] Wrote JSON report: {json_out}")
            print(f"[OK] Removed temporary Office files: {len(removed_temp)}")
            print(f"[OK] Wrote changes report: {report_path}")
            return 0

        total_nodes = 0
        rows = []
        for p in docx_files:
            out_path = mirrored_output_path(p, in_dir, out_dir)
            safe_mkdir(out_path.parent)
            try:
                nodes, chars, findings = process_one(p, out_path, args.fix_dot, args.remove_strikethrough)
                output_rel = out_path.relative_to(out_dir).as_posix()
                status = "OK"
            except Exception as exc:
                nodes, chars, findings = 0, 0, []
                output_rel = ""
                status = f"ERROR: {exc}"
            total_nodes += nodes
            rows.append(
                {
                    "file": rel_posix(p, in_dir),
                    "output": output_rel,
                    "changed_nodes": nodes,
                    "changed_chars": chars,
                    "findings": findings,
                    "status": status,
                }
            )
        report_path = Path(args.report).resolve()
        write_change_report(report_path, rows, removed_temp)
        if args.json_out:
            json_out = Path(args.json_out).resolve()
            write_json_file(
                json_out,
                hygiene_fix_json_payload(
                    in_dir,
                    out_dir,
                    rows,
                    removed_temp,
                    {"fix_dot": bool(args.fix_dot), "remove_strikethrough": bool(args.remove_strikethrough)},
                ),
            )
            print(f"[OK] Wrote JSON report: {json_out}")

        print(f"[OK] Processed: {len(docx_files)} file(s)")
        print(f"[OK] Removed temporary Office files: {len(removed_temp)}")
        print(f"[OK] Output folder: {out_dir}")
        print("[OK] Output mirrors the input folder structure.")
        print(f"[OK] Changed text nodes: {total_nodes}")
        print(f"[OK] Wrote changes report: {report_path}")
        return 0

    in_file = Path(args.file).resolve()
    if not in_file.exists():
        print(f"[ERROR] File not found: {in_file}")
        return 2
    if not args.out:
        print("[ERROR] --out is required when using --file")
        return 2
    out_path = Path(args.out).resolve()
    if args.dry_run:
        print(f"[DRY-RUN] would write: {out_path}")
        return 0
    safe_mkdir(out_path.parent)
    nodes, chars, findings = process_one(in_file, out_path, args.fix_dot, args.remove_strikethrough)
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(
            json_out,
            hygiene_fix_json_payload(
                in_file,
                None,
                [{"file": in_file.name, "output": str(out_path), "changed_nodes": nodes, "changed_chars": chars, "findings": findings, "status": "OK"}],
                [],
                {"fix_dot": bool(args.fix_dot), "remove_strikethrough": bool(args.remove_strikethrough)},
            ),
        )
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Wrote: {out_path}")
    print(f"[OK] Changed text nodes: {nodes}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
