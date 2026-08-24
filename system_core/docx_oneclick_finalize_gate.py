#!/usr/bin/env python3
r"""
DOCX One-Click Finalize + Gate

Pipeline for a folder of DOCX files:
1) (optional) strip comments
2) (optional) accept tracked changes (simple)
3) force black text + remove highlight/shading
4) text hygiene fix (spaces/punctuation; optional dot rule)
5) quality gate scan BEFORE and AFTER; summary PASS/FAIL

This tool is intentionally deterministic (no AI).
It creates new DOCX files and never edits originals.

Usage:
  python docx_oneclick_finalize_gate.py --input input --outdir output/final --summary report/summary_pass_fail.md
  python docx_oneclick_finalize_gate.py --input input\case_01 --outdir output\case_01_final

Options:
  --keep-comments         Do not strip comments (default: strip)
  --keep-changes          Do not accept tracked changes (default: accept)
  --fix-dot               Also fix missing space after '.' (more false positives)
  --strikethrough-preserve-layout
                         Replace strikethrough text with spaces instead of deleting it
  --keep-work             Keep intermediate files in output/_work
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import shutil
import subprocess
import sys
from pathlib import Path

from _office_common import safe_mkdir, find_docx_files, mirrored_output_path, rel_posix
from docx import Document

# Import gate helpers to compute per-file PASS/FAIL summary
from docx_quality_gate import (
    count_nonblack_runs, count_nonblack_styles,
    count_highlights, count_highlights_in_styles,
    count_shading, count_shading_in_styles,
    count_strikethrough, count_strikethrough_in_styles,
    QUALITY_HEADERS_RU,
)
from docx_xml_tools import read_zip_map, find_comments, find_tracked_changes


def scan_gate(docx_paths: list[Path], root: Path | None = None) -> list[dict]:
    rows = []
    for p in docx_paths:
        z = read_zip_map(p)
        nonblack = count_nonblack_runs(z)
        style_nonblack = count_nonblack_styles(z)
        highlight = count_highlights(z)
        style_highlight = count_highlights_in_styles(z)
        shading = count_shading(z)
        style_shading = count_shading_in_styles(z)
        comments_parts = find_comments(z)
        comments = 1 if comments_parts else 0
        strike = count_strikethrough(z)
        style_strike = count_strikethrough_in_styles(z)
        tc = find_tracked_changes(z)
        ins = tc.get("<w:ins", 0)
        dele = tc.get("<w:del", 0)
        mvf = tc.get("<w:moveFrom", 0)
        mvt = tc.get("<w:moveTo", 0)
        status = "PASS"
        if any([nonblack, style_nonblack, highlight, style_highlight, shading, style_shading, comments, strike, style_strike, ins, dele, mvf, mvt]):
            status = "FAIL"
        rows.append({
            "file": rel_posix(p, root) if root is not None else p.name,
            "status": status,
            "nonblack": nonblack,
            "style_nonblack": style_nonblack,
            "highlight": highlight,
            "style_highlight": style_highlight,
            "shading": shading,
            "style_shading": style_shading,
            "comments": comments,
            "strike": strike,
            "style_strike": style_strike,
            "ins": ins,
            "del": dele,
            "moveFrom": mvf,
            "moveTo": mvt,
        })
    return rows


def write_gate_report_md(out_path: Path, title: str, folder: Path, rows: list[dict]) -> None:
    safe_mkdir(out_path.parent)
    lines = []
    lines.append(f"# {title}\n")
    lines.append(f"- Папка: `{folder}`")
    lines.append(f"- Файлов проверено: **{len(rows)}**\n")
    lines.append("| " + " | ".join(QUALITY_HEADERS_RU) + " |")
    lines.append("|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|")
    for r in rows:
        lines.append(
            f"| `{r['file']}` | {r['status']} | {r['nonblack']} | {r['style_nonblack']} | {r['highlight']} | {r['style_highlight']} | "
            f"{r['shading']} | {r['style_shading']} | {r['comments']} | {r['strike']} | {r['style_strike']} | "
            f"{r['ins']} | {r['del']} | {r['moveFrom']} | {r['moveTo']} |"
        )
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_gate_report_docx(out_path: Path, title: str, folder: Path, rows: list[dict]) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Папка: {folder}")
    doc.add_paragraph(f"Файлов проверено: {len(rows)}")
    table = doc.add_table(rows=1, cols=15)
    table.style = "Table Grid"
    for idx, header in enumerate(QUALITY_HEADERS_RU):
        table.rows[0].cells[idx].text = header
    for r in rows:
        values = [r["file"], r["status"], r["nonblack"], r["style_nonblack"], r["highlight"], r["style_highlight"], r["shading"], r["style_shading"], r["comments"], r["strike"], r["style_strike"], r["ins"], r["del"], r["moveFrom"], r["moveTo"]]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
    doc.save(str(out_path))


def write_summary_docx(out_path: Path, lines: list[str]) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    for line in lines:
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].replace("**", "").replace("`", ""))
        elif line.startswith("|") or line.startswith("_"):
            continue
        else:
            doc.add_paragraph(line.replace("**", "").replace("`", ""))
    doc.save(str(out_path))


def run_step(script: Path, args: list[str]) -> int:
    runner = (
        "import runpy, sys; "
        "from pathlib import Path; "
        "target = Path(sys.argv[1]).resolve(); "
        "sys.path.insert(0, str(target.parent)); "
        "sys.argv = [str(target), *sys.argv[2:]]; "
        "runpy.run_path(str(target), run_name='__main__')"
    )
    cmd = [sys.executable, "-c", runner, str(script), *args]
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, encoding="utf-8", errors="replace")
    # Re-emit output so the launcher shows it
    if p.stdout:
        print(p.stdout.rstrip())
    return p.returncode


def main() -> int:
    ap = argparse.ArgumentParser(description="One-click finalize + gate for a folder of DOCX files.")
    ap.add_argument("--input", default="input", help="Input folder with .docx files (recursive)")
    ap.add_argument("--outdir", default="output/final", help="Output folder for final DOCX files")
    ap.add_argument("--summary", default="report/summary_pass_fail.md", help="Output summary Markdown path")

    ap.add_argument("--keep-comments", action="store_true", help="Do NOT strip comments (default strips comments)")
    ap.add_argument("--keep-changes", action="store_true", help="Do NOT accept tracked changes (default accepts changes)")
    ap.add_argument("--fix-dot", action="store_true", help="Also fix missing space after '.'")
    ap.add_argument(
        "--strikethrough-delete",
        action="store_true",
        help="Deprecated: strikethrough text is deleted by default",
    )
    ap.add_argument(
        "--strikethrough-preserve-layout",
        action="store_true",
        help="Replace strikethrough text with same-length spaces instead of deleting it",
    )
    ap.add_argument("--keep-work", action="store_true", help="Keep intermediate files in output/_work")
    args = ap.parse_args()

    in_dir = Path(args.input).resolve()
    out_dir = Path(args.outdir).resolve()
    summary_path = Path(args.summary).resolve()

    if not in_dir.exists():
        print(f"[ERROR] Input folder does not exist: {in_dir}")
        return 2

    files = find_docx_files(in_dir)
    if not files:
        print(f"[WARN] No .docx files found in: {in_dir}")
        safe_mkdir(summary_path.parent)
        summary_path.write_text("# Финализация + проверка DOCX\n\nФайлы DOCX не найдены.\n", encoding="utf-8")
        return 0

    base = Path(__file__).resolve().parent
    s_strip = base / "docx_strip_comments.py"
    s_accept = base / "docx_accept_changes_simple.py"
    s_black = base / "docx_finalize_black_clean.py"
    s_hfix = base / "docx_text_hygiene_fix.py"
    s_gate = base / "docx_quality_gate.py"

    work_dir = out_dir.parent / "_work"
    safe_mkdir(out_dir)
    safe_mkdir(work_dir)

    # Gate BEFORE
    before_rows = scan_gate(files, in_dir)
    gate_before_path = summary_path.parent / "gate_before.md"
    write_gate_report_md(gate_before_path, "Проверка DOCX - ДО", in_dir, before_rows)
    write_gate_report_docx(gate_before_path.with_suffix(".docx"), "Проверка DOCX - ДО", in_dir, before_rows)

    # Pipeline directories
    current_dir = in_dir
    current_files = files

    def step_dir(name: str) -> Path:
        d = work_dir / name
        safe_mkdir(d)
        return d

    # Step 1: strip comments
    if not args.keep_comments:
        d = step_dir("01_no_comments")
        rc = run_step(s_strip, ["--input", str(current_dir), "--outdir", str(d)])
        if rc != 0:
            print("[ERROR] Strip comments step failed.")
            return rc
        current_dir = d
        current_files = find_docx_files(current_dir)

    # Step 2: accept changes
    if not args.keep_changes:
        d = step_dir("02_accepted_changes")
        rc = run_step(s_accept, ["--input", str(current_dir), "--outdir", str(d)])
        if rc != 0:
            print("[ERROR] Accept changes step failed.")
            return rc
        current_dir = d
        current_files = find_docx_files(current_dir)

    # Step 3: finalize black + clean
    d = step_dir("03_final_black")
    args_black = ["--input", str(current_dir), "--outdir", str(d)]
    if args.strikethrough_preserve_layout:
        args_black.extend(["--strikethrough-mode", "preserve-layout"])
    rc = run_step(s_black, args_black)
    if rc != 0:
        print("[ERROR] Finalize black step failed.")
        return rc
    current_dir = d
    current_files = find_docx_files(current_dir)

    # Step 4: text hygiene fix
    d = step_dir("04_hygiene_fixed")
    args_h = ["--input", str(current_dir), "--outdir", str(d)]
    if args.fix_dot:
        args_h.append("--fix-dot")
    rc = run_step(s_hfix, args_h)
    if rc != 0:
        print("[ERROR] Text hygiene fix step failed.")
        return rc
    current_dir = d
    current_files = find_docx_files(current_dir)

    # Copy finals
    for p in current_files:
        final_path = mirrored_output_path(p, current_dir, out_dir)
        safe_mkdir(final_path.parent)
        shutil.copy2(p, final_path)

    # Gate AFTER (scan final folder)
    after_rows = scan_gate(find_docx_files(out_dir), out_dir)
    gate_after_path = summary_path.parent / "gate_after.md"
    write_gate_report_md(gate_after_path, "Проверка DOCX - ПОСЛЕ", out_dir, after_rows)
    write_gate_report_docx(gate_after_path.with_suffix(".docx"), "Проверка DOCX - ПОСЛЕ", out_dir, after_rows)

    # Summary
    safe_mkdir(summary_path.parent)
    fail_before = [r for r in before_rows if r["status"] == "FAIL"]
    fail_after = [r for r in after_rows if r["status"] == "FAIL"]

    lines = []
    lines.append("# Финализация + проверка DOCX\n")
    lines.append(f"- Входная папка: `{in_dir}`")
    lines.append(f"- Итоговая папка: `{out_dir}`")
    lines.append(f"- Файлов обработано: **{len(files)}**\n")

    lines.append("## Параметры\n")
    lines.append(f"- Удалять комментарии: **{'НЕТ' if args.keep_comments else 'ДА'}**")
    lines.append(f"- Принять правки Word: **{'НЕТ' if args.keep_changes else 'ДА'}**")
    lines.append(f"- Зачёркнутый текст: **{'заменять пробелами' if args.strikethrough_preserve_layout else 'удалять'}**")
    lines.append(f"- Исправлять пропуск после точки: **{'ДА' if args.fix_dot else 'НЕТ'}**\n")

    lines.append("## Отчёты проверки\n")
    lines.append(f"- ДО: `{gate_before_path}`")
    lines.append(f"- ПОСЛЕ: `{gate_after_path}`\n")

    lines.append("## Результат\n")
    if not fail_after:
        lines.append("- **PASS**: после финализации проверка не нашла проблем качества.\n")
    else:
        lines.append(f"- **FAIL**: после финализации остались проблемы качества: файлов {len(fail_after)}.\n")

    if fail_after:
        lines.append("### Файлы с проблемами после финализации\n")
        fail_headers = QUALITY_HEADERS_RU[:1] + QUALITY_HEADERS_RU[2:]
        lines.append("| " + " | ".join(fail_headers) + " |")
        lines.append("|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|")
        for r in fail_after:
            lines.append(
                f"| `{r['file']}` | {r['nonblack']} | {r['style_nonblack']} | {r['highlight']} | {r['style_highlight']} | "
                f"{r['shading']} | {r['style_shading']} | {r['comments']} | {r['strike']} | {r['style_strike']} | "
                f"{r['ins']} | {r['del']} | {r['moveFrom']} | {r['moveTo']} |"
            )
        lines.append("")

    summary_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    write_summary_docx(summary_path.with_suffix(".docx"), lines)
    print(f"[OK] Final folder: {out_dir}")
    print(f"[OK] Summary: {summary_path}")

    if not args.keep_work:
        # keep only summary + gate reports, remove work
        try:
            shutil.rmtree(work_dir)
            print("[OK] Removed intermediate folder: output/_work")
        except Exception:
            pass

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
