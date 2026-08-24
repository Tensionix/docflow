#!/usr/bin/env python3
"""
DOCX Finalize: Black + Clean

Creates a new DOCX where:
- strikethrough text is removed
- explicit non-black run colors are forced to black (000000)
- highlights are removed
- shading/background fills (w:shd) are removed

This script is intended for final compliance: "black text on white background".

Usage:
  python docx_finalize_black_clean.py --input input --outdir output/final_black
  python docx_finalize_black_clean.py --file x.docx --out output/x_final.docx
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from pathlib import Path

from docx import Document
from _office_common import safe_mkdir, find_docx_files, mirrored_output_path
from docx_xml_tools import (
    read_zip_map, write_zip_map, list_xml_parts,
    list_style_parts,
    strip_shading_and_highlight, force_black_color,
    find_strikethrough_style_ids, remove_strikethrough_text,
    strip_strikethrough_marks,
)

def _count_token(files: dict[str, bytes], token: bytes) -> int:
    return sum(data.count(token) for data in files.values())

def process_one(docx_path: Path, out_path: Path, strikethrough_mode: str) -> dict[str, int]:
    files = read_zip_map(docx_path)
    strike_style_ids = find_strikethrough_style_ids(files)
    before_len = sum(len(data) for data in files.values())
    before_highlight = _count_token(files, b"<w:highlight")
    before_shading = _count_token(files, b"<w:shd")
    before_color = _count_token(files, b"<w:color")
    before_strike = _count_token(files, b"<w:strike") + _count_token(files, b"<w:dstrike")

    # Clean document content parts
    for part in list_xml_parts(files):
        data = files[part]
        data = remove_strikethrough_text(
            data,
            mode=strikethrough_mode,
            strike_style_ids=strike_style_ids,
        )
        data = strip_shading_and_highlight(data)
        data = force_black_color(data)
        files[part] = data

    # Clean style-level formatting sources (styles may enforce color/highlight/shading)
    for part in list_style_parts(files):
        data = files[part]
        data = strip_strikethrough_marks(data)
        data = strip_shading_and_highlight(data)
        data = force_black_color(data)
        files[part] = data
    write_zip_map(out_path, files)
    after_len = sum(len(data) for data in files.values())
    after_highlight = _count_token(files, b"<w:highlight")
    after_shading = _count_token(files, b"<w:shd")
    after_color = _count_token(files, b"<w:color")
    after_strike = _count_token(files, b"<w:strike") + _count_token(files, b"<w:dstrike")
    return {
        "highlight_removed": max(0, before_highlight - after_highlight),
        "shading_removed": max(0, before_shading - after_shading),
        "color_nodes_seen": before_color,
        "color_nodes_remaining": after_color,
        "strike_marks_removed": max(0, before_strike - after_strike),
        "package_size_delta": after_len - before_len,
    }

def write_change_report(report_path: Path, rows: list[dict[str, object]]) -> None:
    safe_mkdir(report_path.parent)
    headers = ("Файл", "Выходной файл", "Подсветка удалена", "Заливка удалена", "Узлов цвета найдено", "Зачёркиваний удалено", "Изменение размера пакета")
    lines = ["# Изменения финализации DOCX\n", "| " + " | ".join(headers) + " |", "|---|---|---:|---:|---:|---:|---:|"]
    for row in rows:
        lines.append(
            f"| `{row['file']}` | `{row['output']}` | {row['highlight_removed']} | {row['shading_removed']} | "
            f"{row['color_nodes_seen']} | {row['strike_marks_removed']} | {row['package_size_delta']} |"
        )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    doc = Document()
    doc.add_heading("Изменения финализации DOCX", level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title
    for row in rows:
        cells = table.add_row().cells
        values = (
            row["file"], row["output"], row["highlight_removed"], row["shading_removed"],
            row["color_nodes_seen"], row["strike_marks_removed"], row["package_size_delta"],
        )
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
    doc.save(str(report_path.with_suffix(".docx")))

def main() -> int:
    ap = argparse.ArgumentParser(description="Force black text and remove highlights/shading in DOCX.")
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--input", help="Input folder (process all .docx recursively)")
    g.add_argument("--file", help="Single DOCX file")
    ap.add_argument("--outdir", default="output/final_black", help="Output folder for batch mode")
    ap.add_argument("--out", help="Output DOCX path for single-file mode")
    ap.add_argument("--report", default="report/docx_finalize_changes.md", help="Output changes report path")
    ap.add_argument(
        "--strikethrough-mode",
        choices=("preserve-layout", "delete"),
        default="delete",
        help="How to remove strikethrough text: delete text nodes, or replace with spaces to preserve layout",
    )
    args = ap.parse_args()

    if args.input:
        in_dir = Path(args.input).resolve()
        out_dir = Path(args.outdir).resolve()
        safe_mkdir(out_dir)

        docx_files = find_docx_files(in_dir)
        if not docx_files:
            print(f"[WARN] No .docx files found in: {in_dir}")
            return 0

        rows = []
        for p in docx_files:
            out_path = mirrored_output_path(p, in_dir, out_dir)
            safe_mkdir(out_path.parent)
            stats = process_one(p, out_path, args.strikethrough_mode)
            rows.append(
                {
                    "file": p.relative_to(in_dir).as_posix(),
                    "output": out_path.relative_to(out_dir).as_posix(),
                    **stats,
                }
            )
        report_path = Path(args.report).resolve()
        write_change_report(report_path, rows)

        print(f"[OK] Processed: {len(docx_files)} file(s)")
        print(f"[OK] Output folder: {out_dir}")
        print("[OK] Output mirrors the input folder structure.")
        print(f"[OK] Wrote changes report: {report_path}")
        return 0

    in_file = Path(args.file).resolve()
    if not in_file.exists():
        print(f"[ERROR] File not found: {in_file}")
        return 2
    if not args.out:
        print("[ERROR] --out is required when using --file")
        return 2
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)
    process_one(in_file, out_path, args.strikethrough_mode)
    print(f"[OK] Wrote: {out_path}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
