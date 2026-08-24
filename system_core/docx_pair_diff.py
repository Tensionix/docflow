#!/usr/bin/env python3
r"""
DOCX Pair Diff

Compares two explicit DOCX files by user-visible content only:
- body text paragraphs;
- body tables;
- structural heading path for every changed block.

By default it matches top-level sections by content similarity before diffing,
so renumbered or moved sections do not create a positional cascade.

The report intentionally ignores table-of-contents blocks, media, and
low-level DOCX package noise.

Usage:
  python docx_pair_diff.py --a input/a.docx --b input/folder/b.docx --out report/docx_pair_diff.md
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from copy import deepcopy
import difflib
import hashlib
import json
import re
import zlib
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from _office_common import safe_mkdir


SPACE_RE = re.compile(r"\s+")
HEADING_STYLE_RE = re.compile(r"(?:heading|заголовок)\s*(\d+)", re.IGNORECASE)
HEADING_ID_RE = re.compile(r"heading(\d+)", re.IGNORECASE)
TOC_STYLE_RE = re.compile(r"^(toc|оглавление|содержание)(?:\s*\d+)?$", re.IGNORECASE)
TOC_TITLE_RE = re.compile(r"^(оглавление|содержание|table of contents)\s*[:.]?$", re.IGNORECASE)
TOC_ENTRY_RE = re.compile(r".{4,}(?:\t|\.{2,}|\s{2,})\s*(?:\d+|[ivxlcdm]+)$", re.IGNORECASE)
MATCH_WORD_RE = re.compile(r"[0-9A-Za-zА-Яа-яЁё_]+")
REDLINE_TOKEN_RE = re.compile(r"\s+|[0-9A-Za-zА-Яа-яЁё_]+|[^\s0-9A-Za-zА-Яа-яЁё_]+")
ZONE_CODE_PAREN_RE = re.compile(r"\(([A-ZА-ЯЁ]{1,6})\s*[-–—.]?\s*(\d+(?:\.\d+)?[A-ZА-ЯЁ]?)\)", re.IGNORECASE)
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


@dataclass(frozen=True)
class ContentBlock:
    kind: str
    ordinal: int
    kind_ordinal: int
    path: tuple[str, ...]
    text: str
    rows: tuple[tuple[str, ...], ...] = ()
    anchor: str = ""
    body_index: int = -1

    @property
    def path_text(self) -> str:
        return " / ".join(self.path) if self.path else "Без раздела"

    @property
    def kind_ru(self) -> str:
        return {
            "heading": "Заголовок",
            "paragraph": "Текст",
            "table": "Таблица",
        }.get(self.kind, self.kind)

    @property
    def label(self) -> str:
        suffix = f"{self.kind_ru} {self.kind_ordinal}"
        if self.kind == "table" and self.anchor:
            return f"{suffix}: {self.anchor}"
        return suffix

    @property
    def normalized(self) -> str:
        if self.kind == "table":
            row_text = "\n".join(" | ".join(row) for row in self.rows)
            return normalize_text(row_text)
        return normalize_text(self.text)

    @property
    def sequence_key(self) -> str:
        digest = hashlib.sha256(self.normalized.encode("utf-8")).hexdigest()
        return f"{self.kind}:{digest}"


@dataclass(frozen=True)
class Change:
    tag: str
    a: ContentBlock | None
    b: ContentBlock | None
    location: str
    summary: str
    details: str


@dataclass(frozen=True)
class Section:
    ordinal: int
    key: tuple[str, ...]
    title: str
    blocks: tuple[ContentBlock, ...]
    match_text: str
    words: tuple[str, ...]
    shingles: frozenset[int]

    @property
    def word_count(self) -> int:
        return len(self.words)


@dataclass(frozen=True)
class SectionMatch:
    a: Section | None
    b: Section | None
    similarity: float
    status: str


def normalize_text(value: str) -> str:
    return SPACE_RE.sub(" ", (value or "").replace("\xa0", " ")).strip()


def truncate(value: str, limit: int = 1000) -> str:
    text = normalize_text(value)
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def iter_block_items(doc: DocxDocument) -> Iterator[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style
    candidates = []
    if style is not None:
        candidates.extend([str(getattr(style, "name", "") or ""), str(getattr(style, "style_id", "") or "")])
    for candidate in candidates:
        match = HEADING_STYLE_RE.search(candidate) or HEADING_ID_RE.search(candidate.replace(" ", ""))
        if match:
            try:
                level = int(match.group(1))
            except ValueError:
                continue
            return max(1, min(9, level))
    return None


def paragraph_style_candidates(paragraph: Paragraph) -> list[str]:
    style = paragraph.style
    if style is None:
        return []
    return [
        str(getattr(style, "name", "") or "").strip(),
        str(getattr(style, "style_id", "") or "").strip(),
    ]


def paragraph_has_toc_field(paragraph: Paragraph) -> bool:
    try:
        xml = paragraph._p.xml
    except Exception:
        return False
    return "TOC " in xml or ('w:fldCharType="begin"' in xml and "instrText" in xml and "TOC" in xml)


def is_toc_paragraph(paragraph: Paragraph, text: str, inside_toc: bool) -> bool:
    candidates = paragraph_style_candidates(paragraph)
    if any(TOC_STYLE_RE.match(candidate) for candidate in candidates if candidate):
        return True
    if paragraph_has_toc_field(paragraph):
        return True
    if TOC_TITLE_RE.match(text):
        return True
    if inside_toc and TOC_ENTRY_RE.match(text):
        return True
    return False


def table_rows(table: Table) -> tuple[tuple[str, ...], ...]:
    rows: list[tuple[str, ...]] = []
    for row in table.rows:
        values = tuple(normalize_text(cell.text) for cell in row.cells)
        if any(values):
            rows.append(values)
    return tuple(rows)


def table_shape(rows: tuple[tuple[str, ...], ...]) -> str:
    if not rows:
        return "0x0"
    return f"{len(rows)}x{max(len(row) for row in rows)}"


def table_to_text(rows: tuple[tuple[str, ...], ...], limit_rows: int = 8) -> str:
    if not rows:
        return "[пустая таблица]"
    lines = []
    for index, row in enumerate(rows[:limit_rows], start=1):
        lines.append(f"{index}: " + " | ".join(row))
    if len(rows) > limit_rows:
        lines.append(f"... строк всего: {len(rows)}")
    return "\n".join(lines)


def normalize_for_match(value: str) -> str:
    return normalize_text(value).lower().replace("ё", "е")


def canonical_zone_codes(value: str) -> tuple[str, ...]:
    codes: list[str] = []
    for prefix, number in ZONE_CODE_PAREN_RE.findall(value.upper().replace("Ё", "Е")):
        code = f"{prefix}{number}".replace(" ", "").replace("-", "").replace("–", "").replace("—", "")
        if code and code not in codes:
            codes.append(code)
    return tuple(codes)


def title_without_zone_codes(value: str) -> str:
    return normalize_for_match(ZONE_CODE_PAREN_RE.sub("", value))


def meaningful_title(value: str) -> str:
    title = normalize_for_match(value)
    return "" if title == normalize_for_match("Без раздела") else title


def tokenize_for_match(value: str) -> tuple[str, ...]:
    return tuple(MATCH_WORD_RE.findall(normalize_for_match(value)))


def title_word_bag(value: str) -> Counter[str]:
    return Counter(tokenize_for_match(title_without_zone_codes(value)))


def same_title_word_bag(a: str, b: str) -> bool:
    if not meaningful_title(a) or not meaningful_title(b):
        return False
    a_words = title_word_bag(a)
    b_words = title_word_bag(b)
    return bool(a_words and b_words and a_words == b_words)


def title_similarity(a: str, b: str) -> float:
    a_title = normalize_for_match(a)
    b_title = normalize_for_match(b)
    if a_title and b_title and a_title == b_title:
        return 1.0
    if same_title_word_bag(a, b):
        return 1.0
    return difflib.SequenceMatcher(None, a_title, b_title).ratio()


def shingle_fingerprint(words: tuple[str, ...], k_words: int = 5) -> frozenset[int]:
    if not words:
        return frozenset()
    if len(words) < k_words:
        return frozenset(zlib.crc32(word.encode("utf-8")) & 0xFFFFFFFF for word in words)
    hashes = (
        zlib.crc32(" ".join(words[index : index + k_words]).encode("utf-8")) & 0xFFFFFFFF
        for index in range(len(words) - k_words + 1)
    )
    return frozenset(hashes)


def jaccard(a: frozenset[int], b: frozenset[int]) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def section_title(key: tuple[str, ...]) -> str:
    return " / ".join(key) if key else "Без раздела"


def section_match_text(blocks: tuple[ContentBlock, ...]) -> str:
    content_parts: list[str] = []
    heading_parts: list[str] = []
    for block in blocks:
        if block.kind == "heading":
            heading_parts.append(block.normalized)
        else:
            content_parts.append(block.normalized)
    return normalize_for_match("\n".join(content_parts or heading_parts))


def build_section(ordinal: int, key: tuple[str, ...], blocks: list[ContentBlock]) -> Section:
    block_tuple = tuple(blocks)
    match_text = section_match_text(block_tuple)
    words = tokenize_for_match(match_text)
    return Section(
        ordinal=ordinal,
        key=key,
        title=section_title(key),
        blocks=block_tuple,
        match_text=match_text,
        words=words,
        shingles=shingle_fingerprint(words),
    )


def split_sections(blocks: list[ContentBlock], depth: int = 1) -> list[Section]:
    depth = max(1, min(9, depth))
    sections: list[Section] = []
    current_key: tuple[str, ...] | None = None
    current_blocks: list[ContentBlock] = []

    for block in blocks:
        key = tuple(block.path[:depth]) if block.path else ()
        if current_key is not None and key != current_key:
            sections.append(build_section(len(sections) + 1, current_key, current_blocks))
            current_blocks = []
        current_key = key
        current_blocks.append(block)

    if current_key is not None and current_blocks:
        sections.append(build_section(len(sections) + 1, current_key, current_blocks))
    return sections


def section_similarity(a: Section, b: Section) -> float:
    content_score = jaccard(a.shingles, b.shingles)
    title_score = title_similarity(a.title, b.title)
    if not a.shingles or not b.shingles:
        return title_score
    return max(content_score, content_score * 0.70 + title_score * 0.30)


def section_status(a: Section, b: Section) -> str:
    moved = a.ordinal != b.ordinal
    renamed = normalize_for_match(a.title) != normalize_for_match(b.title)
    if moved and renamed:
        return "перемещён/переименован"
    if moved:
        return "перемещён"
    if renamed:
        return "переименован"
    return "сопоставлен"


def match_sections_by_similarity(
    a_sections: list[Section],
    b_sections: list[Section],
    *,
    min_similarity: float,
) -> list[SectionMatch]:
    candidates: list[tuple[float, Section, Section]] = []
    for a in a_sections:
        for b in b_sections:
            score = section_similarity(a, b)
            if score >= min_similarity:
                candidates.append((score, a, b))

    candidates.sort(key=lambda item: (-item[0], item[1].ordinal, item[2].ordinal))
    used_a: set[int] = set()
    used_b: set[int] = set()
    matched: list[SectionMatch] = []
    for score, a, b in candidates:
        if a.ordinal in used_a or b.ordinal in used_b:
            continue
        used_a.add(a.ordinal)
        used_b.add(b.ordinal)
        matched.append(SectionMatch(a=a, b=b, similarity=score, status=section_status(a, b)))

    deleted = [
        SectionMatch(a=section, b=None, similarity=0.0, status="удалён")
        for section in a_sections
        if section.ordinal not in used_a
    ]
    inserted = [
        SectionMatch(a=None, b=section, similarity=0.0, status="добавлен")
        for section in b_sections
        if section.ordinal not in used_b
    ]

    return (
        sorted(matched, key=lambda item: item.a.ordinal if item.a else 10**9)
        + sorted(deleted, key=lambda item: item.a.ordinal if item.a else 10**9)
        + sorted(inserted, key=lambda item: item.b.ordinal if item.b else 10**9)
    )


def extract_content_blocks(docx_path: Path) -> list[ContentBlock]:
    doc = Document(str(docx_path))
    heading_stack: list[str] = []
    blocks: list[ContentBlock] = []
    counts: Counter[str] = Counter()
    last_paragraph = ""
    inside_toc = False

    for body_index, child in enumerate(doc.element.body.iterchildren()):
        if isinstance(child, CT_P):
            item = Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            item = Table(child, doc)
        else:
            continue

        if isinstance(item, Paragraph):
            text = normalize_text(item.text)
            if not text:
                continue

            level = heading_level(item)
            if is_toc_paragraph(item, text, inside_toc):
                if TOC_TITLE_RE.match(text) or paragraph_has_toc_field(item):
                    inside_toc = True
                continue
            if inside_toc and level is not None:
                inside_toc = False
            elif inside_toc:
                continue

            if level is not None:
                heading_stack = heading_stack[: level - 1]
                heading_stack.append(text)
                counts["heading"] += 1
                blocks.append(
                    ContentBlock(
                        kind="heading",
                        ordinal=len(blocks) + 1,
                        kind_ordinal=counts["heading"],
                        path=tuple(heading_stack),
                        text=text,
                        body_index=body_index,
                    )
                )
            else:
                counts["paragraph"] += 1
                blocks.append(
                    ContentBlock(
                        kind="paragraph",
                        ordinal=len(blocks) + 1,
                        kind_ordinal=counts["paragraph"],
                        path=tuple(heading_stack),
                        text=text,
                        body_index=body_index,
                    )
                )
                last_paragraph = text
            continue

        if inside_toc:
            continue

        rows = table_rows(item)
        counts["table"] += 1
        anchor = ""
        if last_paragraph:
            anchor = truncate(last_paragraph, 120)
        blocks.append(
            ContentBlock(
                kind="table",
                ordinal=len(blocks) + 1,
                kind_ordinal=counts["table"],
                path=tuple(heading_stack),
            text=table_to_text(rows, limit_rows=999_999),
            rows=rows,
            anchor=anchor,
            body_index=body_index,
        )
    )

    return blocks


def block_diff_lines(blocks: list[ContentBlock]) -> list[str]:
    lines: list[str] = []
    current_path = ""
    for block in blocks:
        if block.path_text != current_path:
            current_path = block.path_text
            lines.append(f"@@ {current_path}")
        if block.kind == "table":
            lines.append(f"▦ {block.label} ({table_shape(block.rows)})")
            for index, row in enumerate(block.rows, start=1):
                lines.append(f"  R{index}: " + " | ".join(row))
        elif block.kind == "heading":
            lines.append(f"# {block.text}")
        else:
            lines.append(f"¶ {block.text}")
    return lines


def unified_diff(a_lines: list[str], b_lines: list[str], a_name: str, b_name: str, context: int) -> list[str]:
    return list(difflib.unified_diff(a_lines, b_lines, fromfile=a_name, tofile=b_name, n=context, lineterm=""))


def block_similarity(a: ContentBlock, b: ContentBlock) -> float:
    a_text = a.normalized[:20_000]
    b_text = b.normalized[:20_000]
    if not a_text and not b_text:
        return 1.0
    return difflib.SequenceMatcher(None, a_text, b_text).ratio()


def table_change_details(a: ContentBlock, b: ContentBlock) -> str:
    a_rows = [" | ".join(row) for row in a.rows]
    b_rows = [" | ".join(row) for row in b.rows]
    matcher = difflib.SequenceMatcher(None, a_rows, b_rows)
    counts: Counter[str] = Counter()
    examples: list[str] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag == "delete":
            counts["deleted"] += i2 - i1
            for row in a_rows[i1:i2]:
                if len(examples) < 8:
                    examples.append(f"- удалена строка: {truncate(row, 220)}")
        elif tag == "insert":
            counts["inserted"] += j2 - j1
            for row in b_rows[j1:j2]:
                if len(examples) < 8:
                    examples.append(f"- добавлена строка: {truncate(row, 220)}")
        else:
            counts["changed"] += max(i2 - i1, j2 - j1)
            for old, new in zip(a_rows[i1:i2], b_rows[j1:j2]):
                if len(examples) < 8:
                    examples.append(f"- было: {truncate(old, 160)}\n  стало: {truncate(new, 160)}")

    summary = (
        f"Размер A: {table_shape(a.rows)}; размер B: {table_shape(b.rows)}; "
        f"изменено строк: {counts['changed']}; добавлено: {counts['inserted']}; удалено: {counts['deleted']}."
    )
    if examples:
        return summary + "\n" + "\n".join(examples)
    return summary


def make_change(tag: str, a: ContentBlock | None, b: ContentBlock | None) -> Change:
    if a is not None and b is not None:
        if a.path_text == b.path_text:
            location = a.path_text
        else:
            location = f"A: {a.path_text} / B: {b.path_text}"
        if a.kind == "table" and b.kind == "table":
            summary = f"Изменена таблица: {a.label} → {b.label}; сходство {block_similarity(a, b):.3f}"
            details = table_change_details(a, b)
        elif a.kind == b.kind:
            summary = f"Изменён блок: {a.label} → {b.label}; сходство {block_similarity(a, b):.3f}"
            details = f"Было: {truncate(a.text)}\nСтало: {truncate(b.text)}"
        else:
            summary = f"Тип блока изменился: {a.label} → {b.label}"
            details = f"Было: {truncate(a.text)}\nСтало: {truncate(b.text)}"
        return Change("changed", a, b, location, summary, details)

    if a is not None:
        details = table_to_text(a.rows) if a.kind == "table" else truncate(a.text)
        return Change("deleted", a, None, a.path_text, f"Удалён блок: {a.label}", details)

    assert b is not None
    details = table_to_text(b.rows) if b.kind == "table" else truncate(b.text)
    return Change("inserted", None, b, b.path_text, f"Добавлен блок: {b.label}", details)


def compare_blocks(a_blocks: list[ContentBlock], b_blocks: list[ContentBlock]) -> tuple[float, Counter[str], list[Change]]:
    a_keys = [block.sequence_key for block in a_blocks]
    b_keys = [block.sequence_key for block in b_blocks]
    matcher = difflib.SequenceMatcher(None, a_keys, b_keys)
    counts: Counter[str] = Counter()
    changes: list[Change] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        counts[tag] += max(i2 - i1, j2 - j1)
        if tag == "equal":
            continue
        if tag == "delete":
            for block in a_blocks[i1:i2]:
                changes.append(make_change(tag, block, None))
            continue
        if tag == "insert":
            for block in b_blocks[j1:j2]:
                changes.append(make_change(tag, None, block))
            continue

        a_chunk = a_blocks[i1:i2]
        b_chunk = b_blocks[j1:j2]
        paired = min(len(a_chunk), len(b_chunk))
        for index in range(paired):
            changes.append(make_change(tag, a_chunk[index], b_chunk[index]))
        for block in a_chunk[paired:]:
            changes.append(make_change("delete", block, None))
        for block in b_chunk[paired:]:
            changes.append(make_change("insert", None, block))

    return matcher.ratio(), counts, changes


def section_preview(section: Section, limit: int = 1800) -> str:
    lines: list[str] = []
    for block in section.blocks:
        if block.kind == "table":
            lines.append(f"[Таблица {table_shape(block.rows)}]\n{table_to_text(block.rows, limit_rows=6)}")
        elif block.kind == "heading":
            lines.append(f"# {block.text}")
        else:
            lines.append(block.text)
        if len("\n".join(lines)) >= limit:
            break
    return truncate("\n".join(lines), limit)


def make_section_change(tag: str, section: Section) -> Change:
    first = section.blocks[0]
    if tag == "deleted":
        return Change("deleted", first, None, section.title, f"Удалён раздел: A{section.ordinal}. {section.title}", section_preview(section))
    return Change("inserted", None, first, section.title, f"Добавлен раздел: B{section.ordinal}. {section.title}", section_preview(section))


def merge_counter(target: Counter[str], source: Counter[str]) -> None:
    for key, value in source.items():
        target[key] += value


def section_matches_similarity(matches: list[SectionMatch]) -> float:
    total_weight = 0
    weighted_score = 0.0
    for match in matches:
        if match.a is not None and match.b is not None:
            weight = max(match.a.word_count, match.b.word_count, 1)
            weighted_score += match.similarity * weight
            total_weight += weight
        elif match.a is not None:
            total_weight += max(match.a.word_count, 1)
        elif match.b is not None:
            total_weight += max(match.b.word_count, 1)
    if total_weight == 0:
        return 1.0
    return weighted_score / total_weight


def section_diff_lines(matches: list[SectionMatch], context: int) -> list[str]:
    lines: list[str] = []
    for match in matches:
        if match.a is not None and match.b is not None:
            a_name = f"A{match.a.ordinal}: {match.a.title}"
            b_name = f"B{match.b.ordinal}: {match.b.title}"
            lines.append(f"@@ {match.status}: {a_name} -> {b_name}; сходство {match.similarity:.4f}")
            diff = unified_diff(block_diff_lines(list(match.a.blocks)), block_diff_lines(list(match.b.blocks)), a_name, b_name, context)
            lines.extend(diff or ["  Различия текста и таблиц внутри сопоставленного раздела не найдены."])
            continue
        if match.a is not None:
            lines.append(f"@@ удалён раздел: A{match.a.ordinal}: {match.a.title}")
            lines.extend("- " + line for line in block_diff_lines(list(match.a.blocks)))
            continue
        if match.b is not None:
            lines.append(f"@@ добавлен раздел: B{match.b.ordinal}: {match.b.title}")
            lines.extend("+ " + line for line in block_diff_lines(list(match.b.blocks)))
    return lines


def compare_blocks_by_sections(
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
    *,
    section_depth: int,
    section_threshold: float,
    diff_context: int,
) -> tuple[float, Counter[str], list[Change], list[SectionMatch], list[str]]:
    a_sections = split_sections(a_blocks, depth=section_depth)
    b_sections = split_sections(b_blocks, depth=section_depth)
    matches = match_sections_by_similarity(a_sections, b_sections, min_similarity=section_threshold)

    opcode_counts: Counter[str] = Counter()
    changes: list[Change] = []
    for match in matches:
        if match.a is not None and match.b is not None:
            _, counts, section_changes = compare_blocks(list(match.a.blocks), list(match.b.blocks))
            merge_counter(opcode_counts, counts)
            changes.extend(section_changes)
            continue
        if match.a is not None:
            opcode_counts["delete"] += len(match.a.blocks)
            changes.append(make_section_change("deleted", match.a))
            continue
        if match.b is not None:
            opcode_counts["insert"] += len(match.b.blocks)
            changes.append(make_section_change("inserted", match.b))

    return section_matches_similarity(matches), opcode_counts, changes, matches, section_diff_lines(matches, diff_context)


def block_counts(blocks: list[ContentBlock]) -> Counter[str]:
    return Counter(block.kind for block in blocks)


def format_counter(counter: Counter[str]) -> str:
    return ", ".join(f"{key}: {value}" for key, value in sorted(counter.items())) or "нет"


def section_match_counts(matches: list[SectionMatch]) -> Counter[str]:
    return Counter(match.status for match in matches)


def section_match_rows(matches: list[SectionMatch]) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for match in matches:
        a_title = f"A{match.a.ordinal}: {match.a.title}" if match.a is not None else ""
        b_title = f"B{match.b.ordinal}: {match.b.title}" if match.b is not None else ""
        similarity = f"{match.similarity:.4f}" if match.a is not None and match.b is not None else ""
        rows.append((match.status, similarity, a_title, b_title))
    return rows


def section_identity_reason(match: SectionMatch) -> str:
    if match.a is None or match.b is None:
        return "раздел без пары"

    a_title = meaningful_title(match.a.title)
    b_title = meaningful_title(match.b.title)
    a_codes = set(canonical_zone_codes(match.a.title))
    b_codes = set(canonical_zone_codes(match.b.title))
    shared_codes = sorted(a_codes & b_codes)
    exact_title = bool(a_title and b_title and a_title == b_title)
    a_base_title = title_without_zone_codes(match.a.title)
    b_base_title = title_without_zone_codes(match.b.title)
    same_base_title = bool(a_base_title and b_base_title and a_base_title == b_base_title)
    same_word_bag = same_title_word_bag(match.a.title, match.b.title)

    if exact_title and shared_codes:
        return "совпал заголовок и код зоны: " + ", ".join(shared_codes)
    if same_base_title and shared_codes:
        return "совпала основа заголовка и код зоны: " + ", ".join(shared_codes)
    if same_word_bag and shared_codes:
        return "совпал код зоны и набор слов заголовка: " + ", ".join(shared_codes)
    if exact_title:
        return "совпал заголовок"
    if same_word_bag:
        return "совпал набор слов заголовка"
    if shared_codes and difflib.SequenceMatcher(None, a_base_title, b_base_title).ratio() >= 0.75:
        return "совпал код зоны и близкий заголовок: " + ", ".join(shared_codes)
    if match.similarity >= 0.80:
        return "высокая похожесть содержимого"
    if match.similarity >= 0.55:
        return "средняя похожесть содержимого"
    return "низкая похожесть содержимого"


def match_confidence(match: SectionMatch) -> str:
    if match.a is None or match.b is None:
        return "нет пары"
    reason = section_identity_reason(match)
    if reason.startswith(("совпал заголовок", "совпала основа заголовка", "совпал код зоны", "совпал набор слов")):
        return "уверенно"
    if match.similarity >= 0.80:
        return "уверенно"
    if match.similarity >= 0.55:
        return "проверить"
    return "сомнительно"


def alignment_map_rows(matches: list[SectionMatch]) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for index, match in enumerate(matches, start=1):
        rows.append(
            {
                "row": index,
                "status": match.status,
                "confidence": match_confidence(match),
                "confidence_reason": section_identity_reason(match),
                "similarity": match.similarity if match.a is not None and match.b is not None else None,
                "old_ordinal": match.a.ordinal if match.a is not None else None,
                "old_title": match.a.title if match.a is not None else "",
                "old_zone_codes": list(canonical_zone_codes(match.a.title)) if match.a is not None else [],
                "old_words": match.a.word_count if match.a is not None else 0,
                "new_ordinal": match.b.ordinal if match.b is not None else None,
                "new_title": match.b.title if match.b is not None else "",
                "new_zone_codes": list(canonical_zone_codes(match.b.title)) if match.b is not None else [],
                "new_words": match.b.word_count if match.b is not None else 0,
            }
        )
    return rows


def alignment_confidence_counts(rows: list[dict[str, object]]) -> Counter[str]:
    return Counter(str(row["confidence"]) for row in rows)


ALIGNMENT_MAP_COLUMN_WIDTHS_IN = (0.27, 0.88, 0.89, 1.08, 0.69, 0.89, 2.07, 0.69, 1.83, 0.96)
ALIGNMENT_TABLE_HEADER_FILL = "D9EAF7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_table_column_widths(table, widths_in: tuple[float, ...]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths_in):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def add_labeled_paragraph(doc: Document, label: str, value: object) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(str(value))


def write_table_cell(
    cell,
    text: object,
    *,
    bold: bool = False,
    size: int = 8,
    fill: str | None = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)


def write_alignment_map_reports(
    *,
    docx_path: Path,
    md_path: Path,
    json_path: Path,
    a_path: Path,
    b_path: Path,
    compare_mode_title: str,
    section_depth: int,
    section_threshold: float,
    section_matches: list[SectionMatch],
) -> list[dict[str, object]]:
    rows = alignment_map_rows(section_matches)
    safe_mkdir(docx_path.parent)
    safe_mkdir(md_path.parent)
    safe_mkdir(json_path.parent)

    payload = {
        "a": str(a_path),
        "b": str(b_path),
        "compare_mode": compare_mode_title,
        "section_depth": section_depth,
        "section_threshold": section_threshold,
        "rows_total": len(rows),
        "confidence": dict(alignment_confidence_counts(rows)),
        "rows": rows,
    }
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    md_lines = [
        "# Карта сопоставления DOCX",
        "",
        f"- Старый DOCX: `{a_path}`",
        f"- Новый DOCX: `{b_path}`",
        f"- Режим сравнения: **{compare_mode_title}**",
        f"- Глубина разделов: {section_depth}",
        f"- Порог похожести: {section_threshold:.2f}",
        f"- Строк сопоставления: {len(rows)}",
        f"- Уверенность: {format_counter(alignment_confidence_counts(rows))}",
        "",
    ]
    if not rows:
        md_lines.append("_В этом режиме разделы не сопоставлялись; сравнение выполнялось по порядку блоков документа._")
    else:
        md_lines.extend(
            [
                "| # | Статус | Уверенность | Основание | Сходство | Старый раздел | Новый раздел | Слова A/B |",
                "|---:|---|---|---|---:|---|---|---:|",
            ]
        )
        for row in rows:
            similarity = "" if row["similarity"] is None else f"{float(row['similarity']):.4f}"
            old_title = f"A{row['old_ordinal']}: {row['old_title']}".replace("|", "\\|") if row["old_ordinal"] else ""
            new_title = f"B{row['new_ordinal']}: {row['new_title']}".replace("|", "\\|") if row["new_ordinal"] else ""
            reason = str(row["confidence_reason"]).replace("|", "\\|")
            md_lines.append(
                f"| {row['row']} | {row['status']} | {row['confidence']} | {reason} | {similarity} | "
                f"{old_title} | {new_title} | {row['old_words']}/{row['new_words']} |"
            )
    md_path.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    doc.add_heading("Карта сопоставления DOCX", level=1)
    intro = doc.add_paragraph(
        "Таблица показывает, какие разделы DocFlow реально сопоставил перед построением отчёта "
        "и DOCX с правками. Низкая уверенность означает, что пару стоит проверить вручную."
    )
    intro.paragraph_format.space_after = Pt(8)

    doc.add_heading("Документы", level=2)
    add_labeled_paragraph(doc, "Старый DOCX", a_path)
    add_labeled_paragraph(doc, "Новый DOCX", b_path)

    doc.add_heading("Параметры сравнения", level=2)
    for key, value in (
        ("Режим сравнения", compare_mode_title),
        ("Глубина разделов", section_depth),
        ("Порог похожести", f"{section_threshold:.2f}"),
        ("Строк сопоставления", len(rows)),
        ("Уверенность", format_counter(alignment_confidence_counts(rows))),
    ):
        add_labeled_paragraph(doc, key, value)

    if not rows:
        doc.add_paragraph("В этом режиме разделы не сопоставлялись; сравнение выполнялось по порядку блоков документа.")
    else:
        doc.add_heading("Таблица сопоставления", level=2)
        headers = ["#", "Статус", "Уверенность", "Основание", "Сходство", "Старый #", "Старый раздел", "Новый #", "Новый раздел", "Слова A/B"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        set_row_repeat_header(table.rows[0])
        for index, header in enumerate(headers):
            align = WD_ALIGN_PARAGRAPH.LEFT if header in {"Основание", "Старый раздел", "Новый раздел"} else WD_ALIGN_PARAGRAPH.CENTER
            write_table_cell(table.rows[0].cells[index], header, bold=True, fill=ALIGNMENT_TABLE_HEADER_FILL, align=align)
        set_table_column_widths(table, ALIGNMENT_MAP_COLUMN_WIDTHS_IN)

        for row in rows:
            cells = table.add_row().cells
            similarity = "" if row["similarity"] is None else f"{float(row['similarity']):.4f}"
            values = [
                row["row"],
                row["status"],
                row["confidence"],
                row["confidence_reason"],
                similarity,
                "" if row["old_ordinal"] is None else row["old_ordinal"],
                row["old_title"],
                "" if row["new_ordinal"] is None else row["new_ordinal"],
                row["new_title"],
                f"{row['old_words']}/{row['new_words']}",
            ]
            for index, value in enumerate(values):
                align = WD_ALIGN_PARAGRAPH.LEFT if index in {3, 6, 8} else WD_ALIGN_PARAGRAPH.CENTER
                write_table_cell(cells[index], value, align=align)
        set_table_column_widths(table, ALIGNMENT_MAP_COLUMN_WIDTHS_IN)

    doc.save(str(docx_path))
    return rows


def text_stats(blocks: list[ContentBlock]) -> dict[str, int]:
    text = "\n".join(block.text for block in blocks if block.kind in {"heading", "paragraph"})
    return {
        "characters": len(text),
        "words": len(tokenize_for_match(text)),
        "headings": sum(1 for block in blocks if block.kind == "heading"),
        "paragraphs": sum(1 for block in blocks if block.kind == "paragraph"),
    }


def table_stats(blocks: list[ContentBlock]) -> dict[str, int]:
    tables = [block for block in blocks if block.kind == "table"]
    rows = sum(len(block.rows) for block in tables)
    cells = sum(len(row) for block in tables for row in block.rows)
    max_cols = max((len(row) for block in tables for row in block.rows), default=0)
    return {
        "tables": len(tables),
        "rows": rows,
        "cells": cells,
        "max_columns": max_cols,
    }


def change_stats(changes: list[Change]) -> dict[str, object]:
    by_tag = Counter(change.tag for change in changes)
    by_kind: Counter[str] = Counter()
    by_location: Counter[str] = Counter()
    for change in changes:
        kind = change.b.kind if change.b is not None else change.a.kind if change.a is not None else "unknown"
        by_kind[kind] += 1
        by_location[change.location] += 1
    return {
        "total": len(changes),
        "by_tag": dict(by_tag),
        "by_block_kind": dict(by_kind),
        "top_locations": [
            {"location": location, "changes": count}
            for location, count in by_location.most_common(20)
        ],
    }


def stats_payload(
    *,
    a_path: Path,
    b_path: Path,
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
    similarity: float,
    opcode_counts: Counter[str],
    changes: list[Change],
    diff_lines: list[str],
    compare_mode: str,
    compare_mode_title: str,
    section_matches: list[SectionMatch],
    outputs: dict[str, str],
) -> dict[str, object]:
    return {
        "a": str(a_path),
        "b": str(b_path),
        "scope": "body_text_and_tables_without_toc",
        "compare_mode": compare_mode,
        "compare_mode_title": compare_mode_title,
        "similarity": similarity,
        "similarity_percent": round(similarity * 100, 4),
        "blocks": {
            "a_total": len(a_blocks),
            "b_total": len(b_blocks),
            "a_by_kind": dict(block_counts(a_blocks)),
            "b_by_kind": dict(block_counts(b_blocks)),
        },
        "text": {
            "a": text_stats(a_blocks),
            "b": text_stats(b_blocks),
        },
        "tables": {
            "a": table_stats(a_blocks),
            "b": table_stats(b_blocks),
        },
        "operations": dict(opcode_counts),
        "changes": change_stats(changes),
        "alignment_map": {
            "rows": len(section_matches),
            "confidence": dict(alignment_confidence_counts(alignment_map_rows(section_matches))),
        },
        "section_matching": {
            "total": len(section_matches),
            "by_status": dict(section_match_counts(section_matches)),
            "matches": [
                {
                    "status": match.status,
                    "similarity": match.similarity,
                    "similarity_percent": round(match.similarity * 100, 4),
                    "a_ordinal": match.a.ordinal if match.a else None,
                    "a_title": match.a.title if match.a else None,
                    "b_ordinal": match.b.ordinal if match.b else None,
                    "b_title": match.b.title if match.b else None,
                    "a_words": match.a.word_count if match.a else 0,
                    "b_words": match.b.word_count if match.b else 0,
                }
                for match in section_matches
            ],
        },
        "diff_lines": len(diff_lines),
        "outputs": outputs,
    }


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


class RevisionWriter:
    def __init__(self, *, author: str = "Audion DocFlow") -> None:
        self.author = author
        self.date = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        self.revision_id = 1

    def _next_id(self) -> str:
        value = str(self.revision_id)
        self.revision_id += 1
        return value

    def append_revision_text(self, paragraph, text: str, kind: str) -> None:
        if not text:
            return
        tag = "w:ins" if kind == "insert" else "w:del"
        text_tag = "w:t" if kind == "insert" else "w:delText"
        container = OxmlElement(tag)
        container.set(qn("w:id"), self._next_id())
        container.set(qn("w:author"), self.author)
        container.set(qn("w:date"), self.date)

        parts = text.split("\n")
        for index, part in enumerate(parts):
            if index:
                br_run = OxmlElement("w:r")
                br_run.append(OxmlElement("w:br"))
                container.append(br_run)
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0070C0" if kind == "insert" else "C00000")
            rpr.append(color)
            if kind == "insert":
                underline = OxmlElement("w:u")
                underline.set(qn("w:val"), "single")
                rpr.append(underline)
            else:
                rpr.append(OxmlElement("w:strike"))
            run.append(rpr)
            text_node = OxmlElement(text_tag)
            if part.startswith(" ") or part.endswith(" "):
                text_node.set(XML_SPACE, "preserve")
            text_node.text = part
            run.append(text_node)
            container.append(run)
        paragraph._p.append(container)

    def append_segments(self, paragraph, segments: list[tuple[str, str]]) -> None:
        for tag, text in segments:
            if not text:
                continue
            if tag == "equal":
                paragraph.add_run(text)
            elif tag == "delete":
                self.append_revision_text(paragraph, text, "delete")
            elif tag == "insert":
                self.append_revision_text(paragraph, text, "insert")


class XmlRevisionWriter:
    def __init__(self, *, author: str = "Audion DocFlow") -> None:
        self.author = author
        self.date = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        self.revision_id = 1

    def next_id(self) -> str:
        value = str(self.revision_id)
        self.revision_id += 1
        return value

    def text_run(self, text: str, rpr_template=None, *, text_tag: str = "w:t", revision_kind: str | None = None):
        run = OxmlElement("w:r")
        if rpr_template is not None:
            run.append(deepcopy(rpr_template))
        elif revision_kind is not None:
            rpr = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0070C0" if revision_kind == "insert" else "C00000")
            rpr.append(color)
            if revision_kind == "insert":
                underline = OxmlElement("w:u")
                underline.set(qn("w:val"), "single")
                rpr.append(underline)
            else:
                rpr.append(OxmlElement("w:strike"))
            run.append(rpr)
        text_node = OxmlElement(text_tag)
        if text.startswith(" ") or text.endswith(" "):
            text_node.set(XML_SPACE, "preserve")
        text_node.text = text
        run.append(text_node)
        return run

    def revision_element(self, text: str, kind: str, rpr_template=None):
        container = OxmlElement("w:ins" if kind == "insert" else "w:del")
        container.set(qn("w:id"), self.next_id())
        container.set(qn("w:author"), self.author)
        container.set(qn("w:date"), self.date)
        text_tag = "w:t" if kind == "insert" else "w:delText"
        parts = text.split("\n")
        for index, part in enumerate(parts):
            if index:
                br_run = OxmlElement("w:r")
                br_run.append(OxmlElement("w:br"))
                container.append(br_run)
            container.append(self.text_run(part, rpr_template, text_tag=text_tag, revision_kind=kind))
        return container


def w_child(element, tag: str):
    return element.find(qn(tag))


def first_rpr_template(paragraph):
    found = paragraph.find(".//" + qn("w:rPr"))
    return deepcopy(found) if found is not None else None


def paragraph_text_xml(paragraph) -> str:
    parts: list[str] = []
    for node in paragraph.iter():
        if node.tag in {qn("w:t"), qn("w:delText")} and node.text:
            parts.append(node.text)
        elif node.tag == qn("w:tab"):
            parts.append("\t")
        elif node.tag == qn("w:br"):
            parts.append("\n")
    return normalize_text("".join(parts))


def reset_paragraph_to_ppr(paragraph) -> None:
    ppr = w_child(paragraph, "w:pPr")
    for child in list(paragraph):
        if ppr is not None and child is ppr:
            continue
        paragraph.remove(child)


def append_plain_text(paragraph, text: str, writer: XmlRevisionWriter, rpr_template=None) -> None:
    if text:
        paragraph.append(writer.text_run(text, rpr_template))


def append_revision_text_xml(paragraph, text: str, kind: str, writer: XmlRevisionWriter, rpr_template=None) -> None:
    if text:
        paragraph.append(writer.revision_element(text, kind, rpr_template))


def patch_paragraph_with_segments(paragraph, segments: list[tuple[str, str]], writer: XmlRevisionWriter) -> None:
    rpr_template = first_rpr_template(paragraph)
    reset_paragraph_to_ppr(paragraph)
    for tag, text in segments:
        if not text:
            continue
        if tag == "equal":
            append_plain_text(paragraph, text, writer, rpr_template)
        elif tag == "delete":
            append_revision_text_xml(paragraph, text, "delete", writer, rpr_template)
        elif tag == "insert":
            append_revision_text_xml(paragraph, text, "insert", writer, rpr_template)


def mark_paragraph_revision(paragraph, kind: str, writer: XmlRevisionWriter) -> None:
    text = paragraph_text_xml(paragraph)
    rpr_template = first_rpr_template(paragraph)
    reset_paragraph_to_ppr(paragraph)
    append_revision_text_xml(paragraph, text, kind, writer, rpr_template)


def mark_text_in_element(element, kind: str, writer: XmlRevisionWriter) -> None:
    for paragraph in element.xpath(".//w:p", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
        mark_paragraph_revision(paragraph, kind, writer)


def clone_as_revision(element, kind: str, writer: XmlRevisionWriter):
    cloned = deepcopy(element)
    mark_text_in_element(cloned, kind, writer)
    return cloned


def table_row_texts(table_element) -> list[str]:
    rows: list[str] = []
    for row in table_element.xpath("./w:tr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
        values = []
        for cell in row.xpath("./w:tc", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            values.append(paragraph_text_xml(cell))
        rows.append(" | ".join(values))
    return rows


def row_cells(row_element) -> list:
    return row_element.xpath("./w:tc", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})


def patch_row_cells(old_row, new_row, writer: XmlRevisionWriter) -> None:
    old_cells = row_cells(old_row)
    new_cells = row_cells(new_row)
    paired = min(len(old_cells), len(new_cells))
    for index in range(paired):
        old_text = paragraph_text_xml(old_cells[index])
        new_text = paragraph_text_xml(new_cells[index])
        if normalize_text(old_text) == normalize_text(new_text):
            continue
        paragraphs = new_cells[index].xpath(".//w:p", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        if not paragraphs:
            paragraph = OxmlElement("w:p")
            new_cells[index].append(paragraph)
        else:
            paragraph = paragraphs[0]
            for extra in paragraphs[1:]:
                parent = extra.getparent()
                if parent is not None:
                    parent.remove(extra)
        patch_paragraph_with_segments(paragraph, redline_text_segments(old_text, new_text), writer)
    for cell in new_cells[paired:]:
        mark_text_in_element(cell, "insert", writer)


def patch_table_element(old_table, new_table, writer: XmlRevisionWriter) -> None:
    old_rows = old_table.xpath("./w:tr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    new_rows = new_table.xpath("./w:tr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    old_texts = table_row_texts(old_table)
    new_texts = table_row_texts(new_table)
    matcher = difflib.SequenceMatcher(None, old_texts, new_texts)
    for tag, i1, i2, j1, j2 in reversed(matcher.get_opcodes()):
        if tag == "equal":
            continue
        if tag == "insert":
            for row in new_rows[j1:j2]:
                mark_text_in_element(row, "insert", writer)
            continue
        if tag == "delete":
            insert_at = j1 if j1 < len(new_rows) else len(new_table)
            for row in old_rows[i1:i2]:
                cloned = clone_as_revision(row, "delete", writer)
                new_table.insert(insert_at, cloned)
                insert_at += 1
            continue
        paired = min(i2 - i1, j2 - j1)
        for offset in range(paired):
            patch_row_cells(old_rows[i1 + offset], new_rows[j1 + offset], writer)
        insert_at = j1 + paired
        for row in old_rows[i1 + paired : i2]:
            cloned = clone_as_revision(row, "delete", writer)
            new_table.insert(insert_at, cloned)
            insert_at += 1
        for row in new_rows[j1 + paired : j2]:
            mark_text_in_element(row, "insert", writer)


def body_children_by_index(root) -> dict[int, object]:
    body = root.find(qn("w:body"))
    return {index: child for index, child in enumerate(list(body))} if body is not None else {}


def insert_body_element(body, element, before_element=None) -> None:
    if before_element is not None and before_element.getparent() is body:
        body.insert(body.index(before_element), element)
        return
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.insert(body.index(sect_pr), element)
    else:
        body.append(element)


def block_element(block: ContentBlock, elements: dict[int, object]):
    return elements.get(block.body_index)


def mark_block_inserted(block: ContentBlock, elements: dict[int, object], writer: XmlRevisionWriter) -> None:
    element = block_element(block, elements)
    if element is not None:
        mark_text_in_element(element, "insert", writer)


def insert_deleted_block(block: ContentBlock, a_elements: dict[int, object], body, writer: XmlRevisionWriter, before_element=None) -> None:
    element = block_element(block, a_elements)
    if element is None:
        return
    insert_body_element(body, clone_as_revision(element, "delete", writer), before_element)


def patch_changed_block(a: ContentBlock, b: ContentBlock, a_elements: dict[int, object], b_elements: dict[int, object], body, writer: XmlRevisionWriter) -> None:
    old_element = block_element(a, a_elements)
    new_element = block_element(b, b_elements)
    if old_element is None or new_element is None:
        return
    if a.kind == b.kind and a.kind in {"heading", "paragraph"}:
        patch_paragraph_with_segments(new_element, redline_text_segments(a.text, b.text), writer)
        return
    if a.kind == "table" and b.kind == "table":
        patch_table_element(old_element, new_element, writer)
        return
    insert_deleted_block(a, a_elements, body, writer, before_element=new_element)
    mark_text_in_element(new_element, "insert", writer)


def before_element_for_b_index(b_blocks: list[ContentBlock], b_elements: dict[int, object], index: int):
    if 0 <= index < len(b_blocks):
        return block_element(b_blocks[index], b_elements)
    return None


def patch_block_sequences_in_body(
    *,
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
    a_elements: dict[int, object],
    b_elements: dict[int, object],
    body,
    writer: XmlRevisionWriter,
) -> None:
    a_keys = [block.sequence_key for block in a_blocks]
    b_keys = [block.sequence_key for block in b_blocks]
    matcher = difflib.SequenceMatcher(None, a_keys, b_keys)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag == "insert":
            for block in b_blocks[j1:j2]:
                mark_block_inserted(block, b_elements, writer)
            continue
        if tag == "delete":
            before = before_element_for_b_index(b_blocks, b_elements, j1)
            for block in a_blocks[i1:i2]:
                insert_deleted_block(block, a_elements, body, writer, before_element=before)
            continue

        a_chunk = a_blocks[i1:i2]
        b_chunk = b_blocks[j1:j2]
        paired = min(len(a_chunk), len(b_chunk))
        for index in range(paired):
            patch_changed_block(a_chunk[index], b_chunk[index], a_elements, b_elements, body, writer)
        before = before_element_for_b_index(b_blocks, b_elements, j1 + paired)
        for block in a_chunk[paired:]:
            insert_deleted_block(block, a_elements, body, writer, before_element=before)
        for block in b_chunk[paired:]:
            mark_block_inserted(block, b_elements, writer)


def enable_track_revisions_xml(settings_xml: bytes | None) -> bytes:
    if settings_xml is None:
        root = OxmlElement("w:settings")
    else:
        root = etree.fromstring(settings_xml)
    if root.find(qn("w:trackRevisions")) is None:
        root.append(OxmlElement("w:trackRevisions"))
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def write_docx_package_with_document(source_docx: Path, out_path: Path, document_xml: bytes, settings_xml: bytes | None) -> None:
    safe_mkdir(out_path.parent)
    with ZipFile(source_docx, "r") as zin, ZipFile(out_path, "w", ZIP_DEFLATED) as zout:
        written_settings = False
        for info in zin.infolist():
            if info.filename == "word/document.xml":
                zout.writestr(info, document_xml)
            elif info.filename == "word/settings.xml":
                zout.writestr(info, settings_xml or zin.read(info.filename))
                written_settings = True
            else:
                zout.writestr(info, zin.read(info.filename))
        if not written_settings and settings_xml is not None:
            zout.writestr("word/settings.xml", settings_xml)


def redline_tokens(value: str) -> list[str]:
    return REDLINE_TOKEN_RE.findall(value or "")


def redline_text_segments(old: str, new: str) -> list[tuple[str, str]]:
    old_tokens = redline_tokens(old)
    new_tokens = redline_tokens(new)
    matcher = difflib.SequenceMatcher(None, old_tokens, new_tokens)
    segments: list[tuple[str, str]] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            segments.append(("equal", "".join(old_tokens[i1:i2])))
        elif tag == "delete":
            segments.append(("delete", "".join(old_tokens[i1:i2])))
        elif tag == "insert":
            segments.append(("insert", "".join(new_tokens[j1:j2])))
        else:
            segments.append(("delete", "".join(old_tokens[i1:i2])))
            segments.append(("insert", "".join(new_tokens[j1:j2])))
    return segments


def redline_heading_level(block: ContentBlock) -> int:
    if block.kind != "heading":
        return 2
    return max(1, min(9, len(block.path) or 1))


def add_review_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)


def write_text_block(doc: Document, writer: RevisionWriter, block: ContentBlock, revision: str | None = None) -> None:
    if block.kind == "heading":
        paragraph = doc.add_heading("", level=redline_heading_level(block))
    else:
        paragraph = doc.add_paragraph()
    if revision:
        writer.append_revision_text(paragraph, block.text, revision)
    else:
        paragraph.add_run(block.text)


def write_table_as_docx(doc: Document, writer: RevisionWriter, block: ContentBlock, revision: str | None = None) -> None:
    add_review_note(doc, f"{block.label} ({table_shape(block.rows)})")
    if not block.rows:
        paragraph = doc.add_paragraph()
        if revision:
            writer.append_revision_text(paragraph, "[пустая таблица]", revision)
        else:
            paragraph.add_run("[пустая таблица]")
        return
    max_cols = max(len(row) for row in block.rows)
    table = doc.add_table(rows=len(block.rows), cols=max_cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(block.rows):
        cells = table.rows[row_index].cells
        for col_index in range(max_cols):
            value = row[col_index] if col_index < len(row) else ""
            paragraph = cells[col_index].paragraphs[0]
            if revision:
                writer.append_revision_text(paragraph, value, revision)
            else:
                paragraph.add_run(value)


def write_block(doc: Document, writer: RevisionWriter, block: ContentBlock, revision: str | None = None) -> None:
    if block.kind == "table":
        write_table_as_docx(doc, writer, block, revision)
    else:
        write_text_block(doc, writer, block, revision)


def write_changed_table(doc: Document, writer: RevisionWriter, a: ContentBlock, b: ContentBlock) -> None:
    add_review_note(doc, f"Изменена таблица: {a.label} -> {b.label}; {table_shape(a.rows)} -> {table_shape(b.rows)}")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Тип"
    hdr[1].text = "Было"
    hdr[2].text = "Стало"
    a_rows = [" | ".join(row) for row in a.rows]
    b_rows = [" | ".join(row) for row in b.rows]
    matcher = difflib.SequenceMatcher(None, a_rows, b_rows)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag == "delete":
            for old in a_rows[i1:i2]:
                cells = table.add_row().cells
                cells[0].text = "Удалено"
                writer.append_revision_text(cells[1].paragraphs[0], old, "delete")
        elif tag == "insert":
            for new in b_rows[j1:j2]:
                cells = table.add_row().cells
                cells[0].text = "Добавлено"
                writer.append_revision_text(cells[2].paragraphs[0], new, "insert")
        else:
            paired = max(i2 - i1, j2 - j1)
            for offset in range(paired):
                old = a_rows[i1 + offset] if i1 + offset < i2 else ""
                new = b_rows[j1 + offset] if j1 + offset < j2 else ""
                cells = table.add_row().cells
                cells[0].text = "Изменено"
                if old:
                    writer.append_revision_text(cells[1].paragraphs[0], old, "delete")
                if new:
                    writer.append_revision_text(cells[2].paragraphs[0], new, "insert")


def write_changed_block(doc: Document, writer: RevisionWriter, a: ContentBlock, b: ContentBlock) -> None:
    if a.kind == "table" and b.kind == "table":
        write_changed_table(doc, writer, a, b)
        return
    if a.kind == b.kind and a.kind in {"heading", "paragraph"}:
        paragraph = doc.add_heading("", level=redline_heading_level(b)) if b.kind == "heading" else doc.add_paragraph()
        writer.append_segments(paragraph, redline_text_segments(a.text, b.text))
        return
    write_block(doc, writer, a, "delete")
    write_block(doc, writer, b, "insert")


def write_redline_block_sequence(doc: Document, writer: RevisionWriter, a_blocks: list[ContentBlock], b_blocks: list[ContentBlock]) -> None:
    a_keys = [block.sequence_key for block in a_blocks]
    b_keys = [block.sequence_key for block in b_blocks]
    matcher = difflib.SequenceMatcher(None, a_keys, b_keys)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for block in b_blocks[j1:j2]:
                write_block(doc, writer, block)
            continue
        if tag == "delete":
            for block in a_blocks[i1:i2]:
                write_block(doc, writer, block, "delete")
            continue
        if tag == "insert":
            for block in b_blocks[j1:j2]:
                write_block(doc, writer, block, "insert")
            continue

        a_chunk = a_blocks[i1:i2]
        b_chunk = b_blocks[j1:j2]
        paired = min(len(a_chunk), len(b_chunk))
        for index in range(paired):
            write_changed_block(doc, writer, a_chunk[index], b_chunk[index])
        for block in a_chunk[paired:]:
            write_block(doc, writer, block, "delete")
        for block in b_chunk[paired:]:
            write_block(doc, writer, block, "insert")


def write_redline_docx(
    out_path: Path,
    *,
    a_path: Path,
    b_path: Path,
    similarity: float,
    compare_mode_title: str,
    section_matches: list[SectionMatch],
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
) -> None:
    # Preserve the revised/new document package as the carrier, then patch only
    # the matched body paragraphs/tables with Word-compatible revision markup.
    # Statistics and explanatory notes stay outside this DOCX in report files.
    with ZipFile(a_path, "r") as a_zip, ZipFile(b_path, "r") as b_zip:
        a_root = etree.fromstring(a_zip.read("word/document.xml"))
        b_root = etree.fromstring(b_zip.read("word/document.xml"))
        b_settings = b_zip.read("word/settings.xml") if "word/settings.xml" in b_zip.namelist() else None

    b_body = b_root.find(qn("w:body"))
    if b_body is None:
        raise RuntimeError("Не удалось найти word/document.xml body в новом DOCX.")

    a_elements = body_children_by_index(a_root)
    b_elements = body_children_by_index(b_root)
    writer = XmlRevisionWriter()

    if section_matches:
        matched = [match for match in section_matches if match.a is not None and match.b is not None]
        for match in sorted(matched, key=lambda item: item.b.ordinal if item.b else 10**9):
            patch_block_sequences_in_body(
                a_blocks=list(match.a.blocks),
                b_blocks=list(match.b.blocks),
                a_elements=a_elements,
                b_elements=b_elements,
                body=b_body,
                writer=writer,
            )

        for match in sorted((item for item in section_matches if item.b is not None and item.a is None), key=lambda item: item.b.ordinal):
            for block in match.b.blocks:
                mark_block_inserted(block, b_elements, writer)

        matched_by_a = sorted(matched, key=lambda item: item.a.ordinal if item.a else 10**9)
        for match in sorted((item for item in section_matches if item.a is not None and item.b is None), key=lambda item: item.a.ordinal):
            before = None
            for candidate in matched_by_a:
                if candidate.a is not None and candidate.b is not None and candidate.a.ordinal > match.a.ordinal:
                    if candidate.b.blocks:
                        before = block_element(candidate.b.blocks[0], b_elements)
                    break
            for block in match.a.blocks:
                insert_deleted_block(block, a_elements, b_body, writer, before_element=before)
    else:
        patch_block_sequences_in_body(
            a_blocks=a_blocks,
            b_blocks=b_blocks,
            a_elements=a_elements,
            b_elements=b_elements,
            body=b_body,
            writer=writer,
        )

    document_xml = etree.tostring(b_root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    settings_xml = enable_track_revisions_xml(b_settings)
    write_docx_package_with_document(b_path, out_path, document_xml, settings_xml)


def markdown_report(
    *,
    a_path: Path,
    b_path: Path,
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
    similarity: float,
    opcode_counts: Counter[str],
    changes: list[Change],
    diff_lines: list[str],
    max_changes: int,
    max_diff_lines: int,
    compare_mode_title: str,
    section_matches: list[SectionMatch],
    alignment_outputs: dict[str, str] | None = None,
) -> str:
    lines: list[str] = []
    lines.append("# Отчёт сравнения двух DOCX\n")
    lines.append("Сравниваются только текстовые абзацы и таблицы основного документа. Служебная навигационная часть, media и внутренние части DOCX-пакета не анализируются.\n")
    lines.append(f"- A: `{a_path}`")
    lines.append(f"- B: `{b_path}`")
    lines.append(f"- Режим сравнения: **{compare_mode_title}**")
    lines.append(f"- Сходство: **{similarity:.4f}**")
    lines.append(f"- Блоков A: **{len(a_blocks)}** ({format_counter(block_counts(a_blocks))})")
    lines.append(f"- Блоков B: **{len(b_blocks)}** ({format_counter(block_counts(b_blocks))})")
    lines.append(f"- Сводка операций: **{format_counter(opcode_counts)}**")
    lines.append(f"- Изменений в отчёте: **{min(len(changes), max_changes)} из {len(changes)}**\n")
    if alignment_outputs:
        lines.append(f"- Карта сопоставления DOCX: `{alignment_outputs['docx']}`")
        lines.append(f"- Карта сопоставления Markdown: `{alignment_outputs['markdown']}`")
        lines.append(f"- Карта сопоставления JSON: `{alignment_outputs['json']}`\n")

    if section_matches:
        lines.append("## Сопоставление разделов\n")
        lines.append(f"- Разделов/блоков сопоставления: **{len(section_matches)}**")
        lines.append(f"- Сводка: **{format_counter(section_match_counts(section_matches))}**\n")
        lines.append("| Статус | Сходство | Раздел A | Раздел B |")
        lines.append("|---|---:|---|---|")
        for status, score, a_title, b_title in section_match_rows(section_matches):
            lines.append(f"| {status} | {score} | `{a_title}` | `{b_title}` |")
        lines.append("")

    lines.append("## Изменения по структурным местам\n")
    current_location = None
    for index, change in enumerate(changes[:max_changes], start=1):
        if change.location != current_location:
            current_location = change.location
            lines.append(f"### {current_location}\n")
        lines.append(f"#### {index}. {change.summary}\n")
        lines.append("```text")
        lines.append(change.details)
        lines.append("```\n")
    if len(changes) > max_changes:
        lines.append(f"_Показана часть изменений: {max_changes} из {len(changes)}._\n")

    lines.append("## Raw diff текста и таблиц\n")
    shown = diff_lines
    if len(shown) > max_diff_lines:
        shown = shown[:max_diff_lines] + [f"... показана часть diff, всего строк: {len(diff_lines)}"]
    if not shown:
        lines.append("_Различия текста и таблиц не найдены._\n")
    else:
        lines.append("```diff")
        lines.extend(shown)
        lines.append("```\n")
    return "\n".join(lines) + "\n"


def add_mono_paragraph(doc: Document, text: str, font_size: int = 8) -> None:
    p = doc.add_paragraph()
    for line in text.splitlines() or [""]:
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(font_size)


def write_docx_report(
    out_path: Path,
    *,
    a_path: Path,
    b_path: Path,
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
    similarity: float,
    opcode_counts: Counter[str],
    changes: list[Change],
    diff_lines: list[str],
    max_changes: int,
    max_diff_lines: int,
    compare_mode_title: str,
    section_matches: list[SectionMatch],
) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Отчёт сравнения двух DOCX", level=1)
    doc.add_paragraph("Сравниваются только текстовые абзацы и таблицы основного документа. Служебная навигационная часть, media и внутренние части DOCX-пакета не анализируются.")

    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "Параметр"
    summary.rows[0].cells[1].text = "Значение"
    for key, value in (
        ("A", str(a_path)),
        ("B", str(b_path)),
        ("Режим сравнения", compare_mode_title),
        ("Сходство", f"{similarity:.4f}"),
        ("Блоков A", f"{len(a_blocks)} ({format_counter(block_counts(a_blocks))})"),
        ("Блоков B", f"{len(b_blocks)} ({format_counter(block_counts(b_blocks))})"),
        ("Сводка операций", format_counter(opcode_counts)),
        ("Изменений в отчёте", f"{min(len(changes), max_changes)} из {len(changes)}"),
    ):
        cells = summary.add_row().cells
        cells[0].text = key
        cells[1].text = value

    if section_matches:
        doc.add_heading("Сопоставление разделов", level=2)
        doc.add_paragraph(f"Сводка: {format_counter(section_match_counts(section_matches))}.")
        section_table = doc.add_table(rows=1, cols=4)
        section_table.style = "Table Grid"
        hdr = section_table.rows[0].cells
        hdr[0].text = "Статус"
        hdr[1].text = "Сходство"
        hdr[2].text = "Раздел A"
        hdr[3].text = "Раздел B"
        for status, score, a_title, b_title in section_match_rows(section_matches):
            cells = section_table.add_row().cells
            cells[0].text = status
            cells[1].text = score
            cells[2].text = a_title
            cells[3].text = b_title

    doc.add_heading("Изменения по структурным местам", level=2)
    current_location = None
    for index, change in enumerate(changes[:max_changes], start=1):
        if change.location != current_location:
            current_location = change.location
            doc.add_heading(current_location, level=3)
        doc.add_paragraph(f"{index}. {change.summary}")
        add_mono_paragraph(doc, change.details, font_size=8)

    if len(changes) > max_changes:
        doc.add_paragraph(f"Показана часть изменений: {max_changes} из {len(changes)}.")

    doc.add_heading("Raw diff текста и таблиц", level=2)
    shown = diff_lines
    if len(shown) > max_diff_lines:
        shown = shown[:max_diff_lines] + [f"... показана часть diff, всего строк: {len(diff_lines)}"]
    if not shown:
        doc.add_paragraph("Различия текста и таблиц не найдены.")
    else:
        add_mono_paragraph(doc, "\n".join(shown), font_size=7)

    doc.save(str(out_path))


def json_payload(
    *,
    a_path: Path,
    b_path: Path,
    a_blocks: list[ContentBlock],
    b_blocks: list[ContentBlock],
    similarity: float,
    opcode_counts: Counter[str],
    changes: list[Change],
    diff_lines: list[str],
    max_changes: int,
    compare_mode: str,
    section_matches: list[SectionMatch],
) -> dict[str, object]:
    return {
        "a": str(a_path),
        "b": str(b_path),
        "scope": "body_text_and_tables",
        "compare_mode": compare_mode,
        "similarity": similarity,
        "block_similarity": similarity,
        "block_counts_a": dict(block_counts(a_blocks)),
        "block_counts_b": dict(block_counts(b_blocks)),
        "opcode_counts": dict(opcode_counts),
        "changes_total": len(changes),
        "changes_reported": min(len(changes), max_changes),
        "diff_lines": len(diff_lines),
        "alignment_map": {
            "rows": len(section_matches),
            "confidence": dict(alignment_confidence_counts(alignment_map_rows(section_matches))),
            "items": alignment_map_rows(section_matches),
        },
        "section_matches": [
            {
                "status": match.status,
                "similarity": match.similarity,
                "a_ordinal": match.a.ordinal if match.a else None,
                "a_title": match.a.title if match.a else None,
                "b_ordinal": match.b.ordinal if match.b else None,
                "b_title": match.b.title if match.b else None,
            }
            for match in section_matches
        ],
        "changes": [
            {
                "tag": change.tag,
                "location": change.location,
                "summary": change.summary,
                "details": change.details,
                "a_block": change.a.label if change.a else None,
                "b_block": change.b.label if change.b else None,
            }
            for change in changes[:max_changes]
        ],
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Create a text/table structural diff report for two DOCX files.")
    ap.add_argument("--a", required=True, help="DOCX file A")
    ap.add_argument("--b", required=True, help="DOCX file B")
    ap.add_argument("--out", default="report/docx_pair_diff.md", help="Output Markdown report path")
    ap.add_argument("--docx-out", default="", help="Optional DOCX report path")
    ap.add_argument("--diff-out", default="", help="Optional raw text/table diff path")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    ap.add_argument("--stats-out", default="", help="Optional compact statistics JSON path")
    ap.add_argument("--redline-out", default="", help="Optional DOCX with Word-compatible tracked changes")
    ap.add_argument("--alignment-map-out", default="", help="Optional DOCX alignment map path")
    ap.add_argument("--alignment-map-md", default="", help="Optional Markdown alignment map path")
    ap.add_argument("--alignment-map-json", default="", help="Optional JSON alignment map path")
    ap.add_argument("--diff-context", type=int, default=2, help="Diff context lines (default: 2)")
    ap.add_argument("--max-diff-lines", type=int, default=3000, help="Max raw diff lines in reports (default: 3000)")
    ap.add_argument("--max-changes", type=int, default=250, help="Max structural changes listed in reports (default: 250)")
    ap.add_argument("--compare-mode", choices=("sections", "sequence"), default="sections", help="Comparison mode: sections or sequence (default: sections)")
    ap.add_argument("--section-depth", type=int, default=1, help="Heading depth for section matching (default: 1)")
    ap.add_argument("--section-threshold", type=float, default=0.30, help="Minimum section similarity for matching (default: 0.30)")
    args = ap.parse_args()

    a_path = Path(args.a).resolve()
    b_path = Path(args.b).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)

    if not a_path.exists() or not b_path.exists():
        print("[ERROR] One of the DOCX files does not exist.")
        print(f"A: {a_path}")
        print(f"B: {b_path}")
        return 2
    if a_path.suffix.lower() != ".docx" or b_path.suffix.lower() != ".docx":
        print("[ERROR] Both files must be .docx")
        return 2

    a_blocks = extract_content_blocks(a_path)
    b_blocks = extract_content_blocks(b_path)
    section_matches: list[SectionMatch] = []
    section_threshold = max(0.0, min(1.0, args.section_threshold))
    if args.compare_mode == "sections":
        similarity, opcode_counts, changes, section_matches, diff_lines = compare_blocks_by_sections(
            a_blocks,
            b_blocks,
            section_depth=args.section_depth,
            section_threshold=section_threshold,
            diff_context=args.diff_context,
        )
        compare_mode_title = "Сопоставление разделов по похожести"
    else:
        similarity, opcode_counts, changes = compare_blocks(a_blocks, b_blocks)
        diff_lines = unified_diff(
            block_diff_lines(a_blocks),
            block_diff_lines(b_blocks),
            str(a_path),
            str(b_path),
            args.diff_context,
        )
        compare_mode_title = "По порядку документа"

    diff_path = Path(args.diff_out).resolve() if args.diff_out else out_path.with_name(out_path.stem + ".diff")
    docx_path = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
    json_path = Path(args.json_out).resolve() if args.json_out else out_path.with_suffix(".json")
    stats_path = Path(args.stats_out).resolve() if args.stats_out else out_path.with_name(out_path.stem + ".stats.json")
    redline_path = Path(args.redline_out).resolve() if args.redline_out else None
    alignment_docx_path = Path(args.alignment_map_out).resolve() if args.alignment_map_out else out_path.with_name("docx_pair_alignment_map.docx")
    alignment_md_path = Path(args.alignment_map_md).resolve() if args.alignment_map_md else out_path.with_name("docx_pair_alignment_map.md")
    alignment_json_path = Path(args.alignment_map_json).resolve() if args.alignment_map_json else out_path.with_name("docx_pair_alignment_map.json")
    alignment_outputs = {
        "docx": str(alignment_docx_path),
        "markdown": str(alignment_md_path),
        "json": str(alignment_json_path),
    }
    outputs = {
        "markdown_report": str(out_path),
        "docx_report": str(docx_path),
        "raw_diff": str(diff_path),
        "json_report": str(json_path),
        "stats": str(stats_path),
        "alignment_map_docx": str(alignment_docx_path),
        "alignment_map_markdown": str(alignment_md_path),
        "alignment_map_json": str(alignment_json_path),
    }
    if redline_path is not None:
        outputs["redline_docx"] = str(redline_path)

    for path in (diff_path, json_path, stats_path, alignment_docx_path, alignment_md_path, alignment_json_path):
        safe_mkdir(path.parent)

    write_alignment_map_reports(
        docx_path=alignment_docx_path,
        md_path=alignment_md_path,
        json_path=alignment_json_path,
        a_path=a_path,
        b_path=b_path,
        compare_mode_title=compare_mode_title,
        section_depth=args.section_depth,
        section_threshold=section_threshold,
        section_matches=section_matches,
    )

    markdown = markdown_report(
        a_path=a_path,
        b_path=b_path,
        a_blocks=a_blocks,
        b_blocks=b_blocks,
        similarity=similarity,
        opcode_counts=opcode_counts,
        changes=changes,
        diff_lines=diff_lines,
        max_changes=max(0, args.max_changes),
        max_diff_lines=max(0, args.max_diff_lines),
        compare_mode_title=compare_mode_title,
        section_matches=section_matches,
        alignment_outputs=alignment_outputs,
    )
    out_path.write_text(markdown, encoding="utf-8")

    diff_path.write_text("\n".join(diff_lines) + "\n", encoding="utf-8")

    write_docx_report(
        docx_path,
        a_path=a_path,
        b_path=b_path,
        a_blocks=a_blocks,
        b_blocks=b_blocks,
        similarity=similarity,
        opcode_counts=opcode_counts,
        changes=changes,
        diff_lines=diff_lines,
        max_changes=max(0, args.max_changes),
        max_diff_lines=max(0, args.max_diff_lines),
        compare_mode_title=compare_mode_title,
        section_matches=section_matches,
    )

    json_path.write_text(
        json.dumps(
            json_payload(
                a_path=a_path,
                b_path=b_path,
                a_blocks=a_blocks,
                b_blocks=b_blocks,
                similarity=similarity,
                opcode_counts=opcode_counts,
                changes=changes,
                diff_lines=diff_lines,
                max_changes=max(0, args.max_changes),
                compare_mode=args.compare_mode,
                section_matches=section_matches,
            ),
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    stats_path.write_text(
        json.dumps(
            stats_payload(
                a_path=a_path,
                b_path=b_path,
                a_blocks=a_blocks,
                b_blocks=b_blocks,
                similarity=similarity,
                opcode_counts=opcode_counts,
                changes=changes,
                diff_lines=diff_lines,
                compare_mode=args.compare_mode,
                compare_mode_title=compare_mode_title,
                section_matches=section_matches,
                outputs=outputs,
            ),
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    if redline_path is not None:
        write_redline_docx(
            redline_path,
            a_path=a_path,
            b_path=b_path,
            similarity=similarity,
            compare_mode_title=compare_mode_title,
            section_matches=section_matches,
            a_blocks=a_blocks,
            b_blocks=b_blocks,
        )

    print(f"[OK] Wrote DOCX report: {docx_path}")
    print(f"[OK] Wrote Markdown report: {out_path}")
    print(f"[OK] Wrote raw text/table diff: {diff_path}")
    print(f"[OK] Wrote JSON report: {json_path}")
    print(f"[OK] Wrote stats JSON: {stats_path}")
    print(f"[OK] Wrote alignment map DOCX: {alignment_docx_path}")
    print(f"[OK] Wrote alignment map Markdown: {alignment_md_path}")
    print(f"[OK] Wrote alignment map JSON: {alignment_json_path}")
    if redline_path is not None:
        print(f"[OK] Wrote redline DOCX: {redline_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
