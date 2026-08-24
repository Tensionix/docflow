from __future__ import annotations

from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn

_SERVICE_PREFIX_RE = re.compile(r"^\s*<\*>\s*")
_ZERO_WIDTH_RE = re.compile(r"[\u200b\u200c\u200d\ufeff]")
_SOFT_HYPHEN_RE = re.compile(r"\u00ad")


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = text.replace("\r", "\n")
    text = text.replace("\t", " ")
    text = _ZERO_WIDTH_RE.sub("", text)
    text = _SOFT_HYPHEN_RE.sub("", text)
    text = _SERVICE_PREFIX_RE.sub("", text)
    raw_lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    lines = [line for line in raw_lines if line]
    if len(lines) > 1:
        collapsed: list[str] = []
        for line in lines:
            if not collapsed or collapsed[-1] != line:
                collapsed.append(line)
        lines = collapsed
        if len(set(lines)) == 1:
            lines = [lines[0]]
    return "\n".join(lines).strip()


def normalize_for_match(text: str) -> str:
    text = normalize_text(text).casefold()
    text = text.replace("ё", "е")
    text = re.sub(r"\s+", " ", text)
    return text


def row_similarity(a: list[str], b: list[str]) -> float:
    sa = " | ".join(normalize_for_match(x) for x in a if normalize_for_match(x))
    sb = " | ".join(normalize_for_match(x) for x in b if normalize_for_match(x))
    if not sa and not sb:
        return 1.0
    if not sa or not sb:
        return 0.0
    return SequenceMatcher(None, sa, sb).ratio()


@dataclass
class OriginCell:
    origin_id: int
    row: int
    col: int
    text: str
    rowspan: int = 1
    colspan: int = 1

    @property
    def is_merged(self) -> bool:
        return self.rowspan > 1 or self.colspan > 1


@dataclass
class TableModel:
    table_index: int
    width: int
    height: int
    origins: list[OriginCell]
    origin_grid: list[list[int]]
    display_grid: list[list[str]]
    header_preview: list[list[str]]
    header_row_count: int
    preheader_row_count: int
    merged_cells: int
    continuation_slots: int

    @property
    def first_rows(self) -> list[list[str]]:
        return self.display_grid[:2]

    @property
    def header_rows(self) -> list[list[str]]:
        return self.display_grid[: self.header_row_count]

    def preview_text(self) -> str:
        parts: list[str] = []
        source = self.header_preview if self.header_preview else self.first_rows
        for row in source:
            text = " | ".join(cell for cell in row if cell)
            if text:
                parts.append(text)
        return " || ".join(parts)


@dataclass
class TableGroup:
    group_index: int
    column_count: int
    mode: str
    tables: list[TableModel] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    break_reason: str = "first group"
    dropped_header_rows: dict[int, int] = field(default_factory=dict)
    anchor_preview: list[list[str]] = field(default_factory=list)

    def source_table_indexes(self) -> list[int]:
        return [table.table_index for table in self.tables]


def _cell_text_from_tc(tc) -> str:
    texts = [node.text for node in tc.iter(qn("w:t")) if node.text]
    return normalize_text("\n".join(texts))


def _row_preview(row: list[str]) -> list[str]:
    return [normalize_text(cell) for cell in row]


def _looks_numeric(value: str) -> bool:
    value = normalize_for_match(value)
    return bool(value) and re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?", value) is not None


def _looks_record_number(value: str) -> bool:
    value = normalize_for_match(value)
    value = value.rstrip(".")
    return bool(value) and re.fullmatch(r"\d+(?:\.\d+){1,}", value) is not None


def detect_header_rows(display_grid: list[list[str]], max_rows: int = 4) -> int:
    header_count = 0
    for row in display_grid[:max_rows]:
        non_empty = [normalize_text(cell) for cell in row if normalize_text(cell)]
        if not non_empty:
            break
        numeric_cells = sum(1 for cell in non_empty if _looks_numeric(cell))
        first = non_empty[0]
        if _looks_numeric(first) or _looks_record_number(first):
            break
        if numeric_cells / max(len(non_empty), 1) >= 0.5:
            break
        header_count += 1
    return header_count


def detect_preheader_rows(display_grid: list[list[str]], header_row_count: int) -> int:
    if header_row_count < 2:
        return 0
    count = 0
    for row in display_grid[: header_row_count - 1]:
        non_empty_indexes = [idx for idx, cell in enumerate(row) if normalize_text(cell)]
        if len(non_empty_indexes) != 1:
            break
        text = normalize_for_match(row[non_empty_indexes[0]])
        if text in {"№", "n"} or _looks_numeric(text) or _looks_record_number(text):
            break
        count += 1
    return count


def table_header_rows(table: TableModel, preheader_mode: str = "include") -> list[list[str]]:
    if preheader_mode == "separate" and table.preheader_row_count > 0:
        return table.display_grid[table.preheader_row_count : table.header_row_count]
    return table.header_rows


def build_table_model(table, table_index: int) -> TableModel:
    tbl = table._tbl
    grid_cols = getattr(getattr(tbl, "tblGrid", None), "gridCol_lst", [])
    width = len(grid_cols)

    origin_grid: list[list[int]] = []
    origins: list[OriginCell] = []
    active_vmerge: dict[int, int] = {}

    def new_origin(row: int, col: int, text: str) -> int:
        origin_id = len(origins)
        origins.append(OriginCell(origin_id=origin_id, row=row, col=col, text=text))
        return origin_id

    for row_idx, tr in enumerate(tbl.tr_lst):
        if width == 0:
            width = max(
                width,
                sum(
                    int(tc.tcPr.gridSpan.val)
                    if tc.tcPr is not None and tc.tcPr.gridSpan is not None
                    else 1
                    for tc in tr.tc_lst
                ),
            )
        row_slots = [-1] * width
        col_idx = 0
        for tc in tr.tc_lst:
            while col_idx < width and row_slots[col_idx] != -1:
                col_idx += 1
            tc_pr = tc.tcPr
            grid_span = int(tc_pr.gridSpan.val) if tc_pr is not None and tc_pr.gridSpan is not None else 1
            vmerge = None
            if tc_pr is not None and tc_pr.vMerge is not None:
                vmerge = tc_pr.vMerge.val or "continue"

            if vmerge == "continue":
                for offset in range(grid_span):
                    col = col_idx + offset
                    origin_id = active_vmerge.get(col)
                    if origin_id is None:
                        origin_id = new_origin(row_idx, col, "")
                    row_slots[col] = origin_id
                    active_vmerge[col] = origin_id
            else:
                origin_id = new_origin(row_idx, col_idx, _cell_text_from_tc(tc))
                for offset in range(grid_span):
                    col = col_idx + offset
                    row_slots[col] = origin_id
                    if vmerge == "restart":
                        active_vmerge[col] = origin_id
                    else:
                        active_vmerge.pop(col, None)
            col_idx += grid_span

        for col in range(width):
            if row_slots[col] == -1 and col in active_vmerge:
                row_slots[col] = active_vmerge[col]
            elif row_slots[col] != -1 and col not in active_vmerge:
                active_vmerge.pop(col, None)
        origin_grid.append(row_slots)

    if not origin_grid:
        return TableModel(
            table_index=table_index,
            width=0,
            height=0,
            origins=[],
            origin_grid=[],
            display_grid=[],
            header_preview=[],
            header_row_count=0,
            preheader_row_count=0,
            merged_cells=0,
            continuation_slots=0,
        )

    coverage: dict[int, list[tuple[int, int]]] = {origin.origin_id: [] for origin in origins}
    for row_idx, row in enumerate(origin_grid):
        for col_idx, origin_id in enumerate(row):
            if origin_id >= 0:
                coverage[origin_id].append((row_idx, col_idx))

    for origin in origins:
        coords = coverage[origin.origin_id]
        if coords:
            rows = [r for r, _ in coords]
            cols = [c for _, c in coords]
            origin.row = min(rows)
            origin.col = min(cols)
            origin.rowspan = max(rows) - origin.row + 1
            origin.colspan = max(cols) - origin.col + 1

    display_grid: list[list[str]] = []
    continuation_slots = 0
    for row_idx, row in enumerate(origin_grid):
        display_row: list[str] = []
        for col_idx, origin_id in enumerate(row):
            if origin_id < 0:
                display_row.append("")
                continue
            origin = origins[origin_id]
            if origin.row == row_idx and origin.col == col_idx:
                display_row.append(origin.text)
            else:
                display_row.append("")
                continuation_slots += 1
        display_grid.append(display_row)

    merged_cells = sum(1 for origin in origins if origin.is_merged)
    header_row_count = detect_header_rows(display_grid)
    preheader_row_count = detect_preheader_rows(display_grid, header_row_count)
    header_preview = [_row_preview(row) for row in display_grid[:header_row_count]]

    return TableModel(
        table_index=table_index,
        width=width,
        height=len(display_grid),
        origins=origins,
        origin_grid=origin_grid,
        display_grid=display_grid,
        header_preview=header_preview,
        header_row_count=header_row_count,
        preheader_row_count=preheader_row_count,
        merged_cells=merged_cells,
        continuation_slots=continuation_slots,
    )


def parse_docx_tables(docx_path: Path) -> list[TableModel]:
    document = Document(str(docx_path))
    models: list[TableModel] = []
    for idx, table in enumerate(document.tables, start=1):
        model = build_table_model(table, idx)
        if model.width == 0 or model.height == 0:
            continue
        models.append(model)
    return models


def compatibility_reason(anchor: TableModel, candidate: TableModel, mode: str, threshold: float, preheader_mode: str = "include") -> tuple[bool, str, float]:
    if anchor.width != candidate.width:
        return False, f"column count changed: {anchor.width} -> {candidate.width}", 0.0
    if mode == "width-only":
        return True, f"column count matched: {anchor.width}", 1.0

    anchor_rows = table_header_rows(anchor, preheader_mode) or anchor.first_rows[:1]
    candidate_rows = table_header_rows(candidate, preheader_mode) or candidate.first_rows[:1]
    pairs = min(len(anchor_rows), len(candidate_rows), 4)
    if pairs == 0:
        return True, "no header preview available; width matched", 1.0

    scores = [row_similarity(anchor_rows[i], candidate_rows[i]) for i in range(pairs)]
    score = sum(scores) / len(scores)
    if score >= threshold:
        return True, f"width matched and header similarity={score:.2f}", score
    return False, f"header similarity too low: {score:.2f} < {threshold:.2f}", score


def redundant_leading_rows(anchor: TableModel, candidate: TableModel, threshold: float = 0.93) -> int:
    skip = 0
    anchor_headers = anchor.header_rows or anchor.first_rows[:1]
    candidate_headers = candidate.header_rows or candidate.first_rows[:1]
    for idx in range(min(len(anchor_headers), len(candidate_headers), 4)):
        if row_similarity(anchor_headers[idx], candidate_headers[idx]) >= threshold:
            skip += 1
        else:
            break
    return skip


def group_tables(tables: list[TableModel], mode: str = "safe", preheader_mode: str = "include") -> list[TableGroup]:
    if mode not in {"safe", "balanced", "width-only"}:
        raise ValueError(f"Unsupported mode: {mode}")
    threshold = {"safe": 0.82, "balanced": 0.65, "width-only": 0.0}[mode]

    groups: list[TableGroup] = []
    current: TableGroup | None = None
    anchor: TableModel | None = None

    for table in tables:
        if current is None:
            current = TableGroup(
                group_index=1,
                column_count=table.width,
                mode=mode,
                tables=[table],
                rows=[row[:] for row in table.display_grid],
                break_reason="first group",
                anchor_preview=table.first_rows,
            )
            current.dropped_header_rows[table.table_index] = 0
            anchor = table
            continue

        assert anchor is not None
        compatible, reason, _score = compatibility_reason(anchor, table, mode, threshold, preheader_mode=preheader_mode)
        if compatible:
            skip = redundant_leading_rows(anchor, table)
            current.tables.append(table)
            current.rows.extend(row[:] for row in table.display_grid[skip:])
            current.dropped_header_rows[table.table_index] = skip
        else:
            groups.append(current)
            current = TableGroup(
                group_index=len(groups) + 1,
                column_count=table.width,
                mode=mode,
                tables=[table],
                rows=[row[:] for row in table.display_grid],
                break_reason=reason,
                anchor_preview=table.first_rows,
            )
            current.dropped_header_rows[table.table_index] = 0
            anchor = table

    if current is not None:
        groups.append(current)
    return groups
