from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any
import argparse
import json
import re

from docx import Document


WORD_RE = re.compile(r"[A-Za-zА-Яа-яЁё]+")
COMMA_RE = re.compile(r",")
ROMAN_NUMERAL_RE = re.compile(r"[IVXLCDM]+")
SPACE_CHARS = {" ", "\t", "\u00A0"}
INVALID_REPORT_NAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')


@dataclass(frozen=True)
class ParagraphTarget:
    paragraph: Any
    location: dict[str, Any]


def positions_to_lower(text: str, keep_words: set[str]) -> list[int]:
    keep_words_lower = {word.lower() for word in keep_words}
    positions: list[int] = []
    for comma in COMMA_RE.finditer(text):
        pos = comma.end()
        while pos < len(text):
            while pos < len(text) and text[pos] in SPACE_CHARS:
                pos += 1
            word_match = WORD_RE.match(text, pos)
            if not word_match:
                break
            word = word_match.group(0)
            if (
                word in keep_words
                or word.lower() in keep_words_lower
                or (len(word) >= 2 and word.isupper())
                or ROMAN_NUMERAL_RE.fullmatch(word)
            ):
                pos = word_match.end()
                continue
            if word[0].isupper():
                positions.append(word_match.start())
            break
    return positions


def parse_keep_words(raw: str) -> set[str]:
    if not raw.strip():
        return set()
    candidate = Path(raw.strip())
    if candidate.exists():
        text = candidate.read_text(encoding="utf-8")
        if candidate.suffix.lower() == ".json":
            payload = json.loads(text)
            if isinstance(payload, dict):
                values = payload.get("keep_words", payload.get("words", []))
            else:
                values = payload
            return {str(item).strip() for item in values if str(item).strip()} if isinstance(values, list) else set()
        if candidate.suffix.lower() in {".yaml", ".yml"}:
            import yaml

            payload = yaml.safe_load(text)
            if isinstance(payload, dict):
                values = payload.get("keep_words", payload.get("words", []))
            else:
                values = payload
            return {str(item).strip() for item in values if str(item).strip()} if isinstance(values, list) else set()
        return {
            item.strip()
            for line in text.splitlines()
            for item in line.split(",")
            if item.strip() and not item.strip().startswith("#")
        }
    return {item.strip() for item in raw.split(",") if item.strip()}


def paragraph_text_and_map(paragraph: Any) -> tuple[str, list[tuple[int, int]]]:
    chunks: list[str] = []
    char_map: list[tuple[int, int]] = []
    for run_index, run in enumerate(paragraph.runs):
        text = run.text or ""
        chunks.append(text)
        char_map.extend((run_index, char_index) for char_index in range(len(text)))
    return "".join(chunks), char_map


def snippet(text: str, positions: list[int], radius: int = 70) -> str:
    if not positions:
        return text[: radius * 2]
    start = max(0, min(positions) - radius)
    end = min(len(text), max(positions) + radius)
    prefix = "..." if start else ""
    suffix = "..." if end < len(text) else ""
    return f"{prefix}{text[start:end]}{suffix}"


def one_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def after_fragment(text: str, position: int, context_chars: int) -> str:
    end = min(len(text), position + max(context_chars, 1))
    suffix = "..." if end < len(text) else ""
    return one_line(text[position:end]) + suffix


def hit_records(
    location: dict[str, Any],
    before: str,
    after: str,
    positions: list[int],
    start_index: int,
    context_chars: int,
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for offset, position in enumerate(positions, start=0):
        before_word_match = WORD_RE.match(before, position)
        after_word_match = WORD_RE.match(after, position)
        before_word = before_word_match.group(0) if before_word_match else ""
        after_word = after_word_match.group(0) if after_word_match else ""
        comma_position = before.rfind(",", 0, position)
        following_start = after_word_match.end() if after_word_match else position
        following_end = min(len(after), following_start + max(context_chars, 1))
        records.append(
            {
                **location,
                "hit": start_index + offset,
                "position": position,
                "comma_position": comma_position,
                "word_before": before_word,
                "word_after": after_word,
                "before": snippet(before, [position]),
                "after": snippet(after, [position]),
                "after_fragment": after_fragment(after, position, context_chars),
                "following_text": one_line(after[following_start:following_end]) + ("..." if following_end < len(after) else ""),
            }
        )
    return records


def apply_lowercase_positions(paragraph: Any, positions: list[int], dry_run: bool) -> str:
    _text, char_map = paragraph_text_and_map(paragraph)
    run_chars = [list(run.text or "") for run in paragraph.runs]
    for position in positions:
        if position >= len(char_map):
            continue
        run_index, char_index = char_map[position]
        run_chars[run_index][char_index] = run_chars[run_index][char_index].lower()
    after = "".join("".join(chars) for chars in run_chars)
    if not dry_run:
        for run, chars in zip(paragraph.runs, run_chars):
            run.text = "".join(chars)
    return after


def paragraph_xml_key(paragraph: Any) -> str:
    try:
        return str(paragraph._p.getroottree().getpath(paragraph._p))
    except Exception:
        return str(id(paragraph._p))


def append_target(targets: list[ParagraphTarget], paragraph: Any, location: dict[str, Any], seen: set[str]) -> None:
    key = paragraph_xml_key(paragraph)
    if key in seen:
        return
    seen.add(key)
    targets.append(ParagraphTarget(paragraph=paragraph, location=location))


def iter_table_paragraphs(table: Any, table_path: str, seen: set[str]) -> list[ParagraphTarget]:
    targets: list[ParagraphTarget] = []
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                append_target(
                    targets,
                    paragraph,
                    {
                        "kind": "table_cell",
                        "table": table_path,
                        "row": row_index,
                        "cell": cell_index,
                        "paragraph": paragraph_index,
                    },
                    seen,
                )
            for nested_index, nested_table in enumerate(cell.tables, start=1):
                targets.extend(iter_table_paragraphs(nested_table, f"{table_path}.{row_index}.{cell_index}.{nested_index}", seen))
    return targets


def iter_document_targets(document: Any, scope: str = "table-cells") -> list[ParagraphTarget]:
    targets: list[ParagraphTarget] = []
    seen: set[str] = set()
    if scope == "all":
        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            append_target(targets, paragraph, {"kind": "paragraph", "paragraph": paragraph_index}, seen)
    for table_index, table in enumerate(document.tables, start=1):
        targets.extend(iter_table_paragraphs(table, str(table_index), seen))
    return targets


def is_docx(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() == ".docx" and not path.name.startswith("~$")


def collect_docx(input_path: Path, recursive: bool) -> list[Path]:
    if input_path.is_file():
        return [input_path] if is_docx(input_path) else []
    if not input_path.is_dir():
        raise FileNotFoundError(f"Input path was not found: {input_path}")
    pattern = "**/*" if recursive else "*"
    return sorted(path for path in input_path.glob(pattern) if is_docx(path))


def unique_path(path: Path, overwrite: bool) -> Path:
    if overwrite or not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(2, 1000):
        candidate = path.with_name(f"{stem}_{index}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Could not allocate unique output path for {path.name}")


def output_path_for(source: Path, input_root: Path, output: Path, multiple: bool, overwrite: bool) -> Path:
    if not multiple and output.suffix.lower() == ".docx":
        return unique_path(output, overwrite)
    relative_parent = source.parent.relative_to(input_root) if input_root.is_dir() else Path()
    target_dir = output / relative_parent
    return unique_path(target_dir / f"{source.stem}_comma_lowercase.docx", overwrite)


def process_docx(source: Path, target: Path, keep_words: set[str], dry_run: bool, context_chars: int, scope: str) -> dict[str, Any]:
    document = Document(str(source))
    changes: list[dict[str, Any]] = []
    hits: list[dict[str, Any]] = []
    for item in iter_document_targets(document, scope):
        before, _char_map = paragraph_text_and_map(item.paragraph)
        if not before:
            continue
        positions = positions_to_lower(before, keep_words)
        if not positions:
            continue
        after = apply_lowercase_positions(item.paragraph, positions, dry_run)
        paragraph_hits = hit_records(item.location, before, after, positions, len(hits) + 1, context_chars)
        hits.extend(paragraph_hits)
        changes.append(
            {
                **item.location,
                "changes": len(positions),
                "before": snippet(before, positions),
                "after": snippet(after, positions),
                "hits": paragraph_hits,
            }
        )
    if changes and not dry_run:
        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(target))
    return {
        "file": str(source),
        "output": str(target) if changes and not dry_run else "",
        "scope": scope,
        "changes": sum(int(change["changes"]) for change in changes),
        "locations": changes,
        "hits": hits,
    }


def format_location(hit: dict[str, Any]) -> str:
    if hit.get("kind") == "paragraph":
        return f"абзац {hit.get('paragraph')}"
    if hit.get("kind") == "table_cell":
        return (
            f"таблица {hit.get('table')}, строка {hit.get('row')}, "
            f"ячейка {hit.get('cell')}, абзац {hit.get('paragraph')}"
        )
    return str(hit.get("kind", "document"))


def render_markdown_report(payload: dict[str, Any]) -> str:
    lines = [
        "# Регистр после запятых",
        "",
        f"- Файлов: {payload['files']}",
        f"- Файлов с потенциальными правками: {payload['changed_files']}",
        f"- Срабатываний: {payload['changes']}",
        f"- Dry-run: {'да' if payload['dry_run'] else 'нет'}",
        f"- Область проверки: {'только ячейки таблиц' if payload['scope'] == 'table-cells' else 'весь документ'}",
        f"- Контекст после срабатывания: {payload['context_chars']} символов",
        "",
        "Каждое срабатывание ниже показывает слово до/после и фрагмент текста после смоделированного понижения регистра.",
        "",
    ]
    for result in payload["results"]:
        file_name = Path(str(result["file"])).name
        lines.extend([f"## {file_name}", "", f"- Срабатываний: {result['changes']}", ""])
        if int(result["changes"]) == 0:
            lines.extend(["Срабатываний не найдено.", ""])
            continue
        for hit in result.get("hits", []):
            lines.extend(
                [
                    f"### {hit['hit']}. {hit['word_before']} -> {hit['word_after']}",
                    "",
                    f"- Где: {format_location(hit)}",
                    f"- Позиция в абзаце: {hit['position']}",
                    "",
                    "После правки:",
                    "",
                    "```text",
                    str(hit["after_fragment"]),
                    "```",
                    "",
                ]
            )
    return "\n".join(lines).rstrip() + "\n"


def safe_report_stem(path: str) -> str:
    stem = Path(path).stem.strip() or "document"
    safe = INVALID_REPORT_NAME_RE.sub("_", stem).strip(" .")
    return safe or "document"


def report_paths_for_result(report_dir: Path, result: dict[str, Any], used: set[str]) -> tuple[Path, Path]:
    base = safe_report_stem(str(result["file"]))
    index = 1
    while True:
        suffix = "" if index == 1 else f"_{index}"
        stem = f"{base}{suffix}"
        key = stem.casefold()
        if key not in used:
            used.add(key)
            return report_dir / f"{stem}.json", report_dir / f"{stem}.md"
        index += 1


def single_result_payload(payload: dict[str, Any], result: dict[str, Any]) -> dict[str, Any]:
    changes = int(result["changes"])
    return {
        key: value
        for key, value in payload.items()
        if key not in {"results", "per_document_report_dir", "per_document_reports"}
    } | {
        "files": 1,
        "changed_files": 1 if changes > 0 else 0,
        "changes": changes,
        "results": [result],
    }


def write_per_document_reports(payload: dict[str, Any], report_dir: Path) -> list[dict[str, Any]]:
    report_dir.mkdir(parents=True, exist_ok=True)
    used: set[str] = set()
    reports: list[dict[str, Any]] = []
    for result in payload["results"]:
        json_path, md_path = report_paths_for_result(report_dir, result, used)
        document_payload = single_result_payload(payload, result)
        json_path.write_text(json.dumps(document_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        md_path.write_text(render_markdown_report(document_payload), encoding="utf-8")
        reports.append(
            {
                "file": str(result["file"]),
                "changes": int(result["changes"]),
                "json": str(json_path),
                "markdown": str(md_path),
            }
        )
    return reports


def default_per_document_report_dir(report_path: Path) -> Path:
    return report_path.with_suffix("") if report_path.suffix else report_path.with_name(f"{report_path.name}_files")


def build_parser(add_help: bool = True) -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Lowercase first word letters after commas inside DOCX table cells.",
        add_help=add_help,
    )
    parser.add_argument("--input", required=True, help="DOCX file or folder.")
    parser.add_argument("--output", default="output", help="Output DOCX file or folder.")
    parser.add_argument("--recursive", action="store_true", help="Search input folder recursively.")
    parser.add_argument("--dry-run", action="store_true", help="Write only the report, do not save DOCX files.")
    parser.add_argument("--report", default="report/comma_lowercase.json", help="JSON report path.")
    parser.add_argument("--md-report", default="", help="Optional Markdown report path.")
    parser.add_argument(
        "--per-document-report-dir",
        default="",
        help="Optional folder for per-document JSON/Markdown reports. Defaults to the JSON report path without extension.",
    )
    parser.add_argument("--context-chars", type=int, default=240, help="Characters to include after each hit in reports.")
    parser.add_argument(
        "--scope",
        choices=("table-cells", "all"),
        default="table-cells",
        help="Where to search: table-cells = only inside DOCX table cells; all = body paragraphs plus table cells.",
    )
    parser.add_argument("--keep-words", default="", help="CSV allowlist or path to txt/json/yaml allowlist.")
    parser.add_argument("--overwrite", action="store_true", help="Allow overwriting explicit output paths.")
    return parser


def run(args: argparse.Namespace) -> int:
    input_path = Path(args.input).resolve()
    output = Path(args.output).resolve()
    report_path = Path(args.report).resolve()
    md_report_path = Path(args.md_report).resolve() if str(args.md_report or "").strip() else report_path.with_suffix(".md")
    per_document_report_dir = (
        Path(args.per_document_report_dir).resolve()
        if str(args.per_document_report_dir or "").strip()
        else default_per_document_report_dir(report_path)
    )
    keep_words = parse_keep_words(str(args.keep_words or ""))
    context_chars = max(int(args.context_chars), 1)
    scope = str(args.scope)
    sources = collect_docx(input_path, bool(args.recursive))
    input_root = input_path if input_path.is_dir() else input_path.parent
    multiple = len(sources) != 1 or input_path.is_dir()
    results = [
        process_docx(
            source,
            output_path_for(source, input_root, output, multiple, bool(args.overwrite)),
            keep_words,
            bool(args.dry_run),
            context_chars,
            scope,
        )
        for source in sources
    ]
    payload = {
        "command": "comma-lowercase-docx",
        "dry_run": bool(args.dry_run),
        "input": str(input_path),
        "output": str(output),
        "scope": scope,
        "files": len(sources),
        "changed_files": sum(1 for result in results if int(result["changes"]) > 0),
        "changes": sum(int(result["changes"]) for result in results),
        "context_chars": context_chars,
        "keep_words": sorted(keep_words),
        "results": results,
    }
    payload["per_document_report_dir"] = str(per_document_report_dir)
    payload["per_document_reports"] = write_per_document_reports(payload, per_document_report_dir)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    md_report_path.parent.mkdir(parents=True, exist_ok=True)
    md_report_path.write_text(render_markdown_report(payload), encoding="utf-8")
    print(f"Files: {payload['files']}; changed files: {payload['changed_files']}; changes: {payload['changes']}")
    print(f"Report: {report_path}")
    print(f"Markdown report: {md_report_path}")
    if args.dry_run:
        print("Dry-run: DOCX files were not written.")
    return 0


def main(argv: list[str] | None = None) -> int:
    return run(build_parser().parse_args(argv))


if __name__ == "__main__":
    raise SystemExit(main())
