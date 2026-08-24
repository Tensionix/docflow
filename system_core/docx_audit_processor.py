#!/usr/bin/env python3
"""
DOCX Audit Processor.

Deterministic audit/fix layer for narrow corporate writing rules. The processor
edits only visible DOCX text nodes (w:t) and writes processed copies under
output/audit_processed.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Match, Pattern

from docx import Document
from lxml import etree

from _office_common import find_docx_files, md_escape, mirrored_output_path, safe_mkdir, truncate, write_json_file
from docx_xml_tools import NS, _etree_from_bytes, _etree_to_bytes, list_xml_parts, read_zip_map, write_zip_map


XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
TEXT_WINDOW = 70
MAX_FINDINGS_PER_FILE = 300
RULE_CATALOG = Path(__file__).resolve().parents[1] / "config" / "rules" / "rules.yaml"
ANCHOR_COLOR = "C05600"

NUM_RE = r"(?P<num>(?<![\wА-Яа-яЁё])[-+]?\d+(?:[ \u00A0]?\d{3})*(?:[,.]\d+)?)"
AFTER_UNIT_RE = r"(?=$|[\s\u00A0,.;:!?()\[\]{}<>«»\"'“”‘’/-])"


@dataclass(frozen=True)
class AuditRule:
    code: str
    title: str
    severity: str
    description: str
    pattern: Pattern[str]
    replacement: str | Callable[[Match[str]], str]
    fixable: bool = True

    def replacement_for(self, match: Match[str]) -> str:
        if callable(self.replacement):
            return self.replacement(match)
        return match.expand(self.replacement)


@dataclass
class AuditFinding:
    rule_code: str
    rule_title: str
    severity: str
    action: str
    part: str
    before: str
    after: str
    context: str
    start: int = 0
    end: int = 0
    text_node_index: int = 0
    anchor_id: str = ""


@dataclass
class FileAuditReport:
    path: Path
    output: Path | None = None
    annotated_output: Path | None = None
    status: str = "OK"
    text_nodes: int = 0
    changed_nodes: int = 0
    fixes: int = 0
    anchors: int = 0
    findings: list[AuditFinding] = field(default_factory=list)
    rule_counts: dict[str, int] = field(default_factory=dict)


def _clean_spaces(value: str) -> str:
    return re.sub(r"[ \u00A0]+", " ", value.replace("\u00A0", " ")).strip()


def _same_case_unit(unit: str) -> str:
    compact = _clean_spaces(unit).replace(" ", "")
    lower = compact.lower().replace("ё", "е")
    if lower == "гкал/ч":
        return "Гкал/ч"
    if lower == "гкал":
        return "Гкал"
    if lower == "квт":
        return "кВт"
    if lower == "мвт":
        return "МВт"
    if lower in {"руб", "руб."}:
        return "руб."
    return _clean_spaces(unit)


def _sqm_replacement(match: Match[str]) -> str:
    return f"{_clean_spaces(match.group('num'))} кв. м"


def _cubic_replacement(match: Match[str]) -> str:
    return f"{_clean_spaces(match.group('num'))} куб. м"


def _common_unit_replacement(match: Match[str]) -> str:
    return f"{_clean_spaces(match.group('num'))} {_same_case_unit(match.group('unit'))}"


def _degree_replacement(match: Match[str]) -> str:
    suffix = match.group("scale") or ""
    return f"{_clean_spaces(match.group('num'))}°{'C' if suffix else ''}"


def _date_without_year_suffix(match: Match[str]) -> str:
    return match.group("date")


def _russian_federation_suggestion(_match: Match[str]) -> str:
    return "Российская Федерация / Российской Федерации (проверить падеж)"


def _build_rules() -> list[AuditRule]:
    sqm_units = (
        r"(?:"
        r"кв\.?м\.?|"
        r"кв\.\s+м|"
        r"кв\s+м|"
        r"кв\.?\s*метр(?:а|ов)?|"
        r"квадратн(?:ый|ого|ому|ым|ом|ая|ой|ую|ые|ых|ыми)?\s+метр(?:а|ов)?|"
        r"м\s*(?:2|²)"
        r")"
    )
    cubic_units = (
        r"(?:"
        r"куб\.?м\.?|"
        r"куб\.\s+м|"
        r"куб\s+м|"
        r"куб\.?\s*метр(?:а|ов)?|"
        r"кубическ(?:ий|ого|ому|им|ом|ая|ой|ую|ие|их|ими)?\s+метр(?:а|ов)?|"
        r"м\s*(?:3|³)"
        r")"
    )
    common_units = (
        r"(?P<unit>"
        r"мм|см|км|м|га|т|кг|л|мл|"
        r"руб\.?|"
        r"Гкал/ч|Гкал|гкал/ч|гкал|кВт|квт|МВт|мвт"
        r")"
    )

    return [
        AuditRule(
            code="AUDIT-UNIT-SQM",
            title="Единицы площади: кв. м",
            severity="major",
            description="Нормализует варианты после числа: кв.м, кв м, м2, м², квадратных метров -> кв. м.",
            pattern=re.compile(rf"{NUM_RE}[ \u00A0]*{sqm_units}{AFTER_UNIT_RE}", re.IGNORECASE),
            replacement=_sqm_replacement,
        ),
        AuditRule(
            code="AUDIT-UNIT-CUBIC",
            title="Единицы объёма: куб. м",
            severity="major",
            description="Нормализует варианты после числа: куб.м, куб м, м3, м³, кубических метров -> куб. м.",
            pattern=re.compile(rf"{NUM_RE}[ \u00A0]*{cubic_units}{AFTER_UNIT_RE}", re.IGNORECASE),
            replacement=_cubic_replacement,
        ),
        AuditRule(
            code="AUDIT-UNIT-SPACE",
            title="Пробел между числом и единицей измерения",
            severity="minor",
            description="Добавляет пробел в безопасных случаях: 10м, 5км, 20га, 3Гкал/ч, 100руб.",
            pattern=re.compile(rf"{NUM_RE}{common_units}{AFTER_UNIT_RE}"),
            replacement=_common_unit_replacement,
        ),
        AuditRule(
            code="AUDIT-PERCENT",
            title="Процент без пробела",
            severity="minor",
            description="Нормализует запись процента: 50 % -> 50%.",
            pattern=re.compile(rf"{NUM_RE}[ \u00A0]+%"),
            replacement=lambda match: f"{_clean_spaces(match.group('num'))}%",
        ),
        AuditRule(
            code="AUDIT-DEGREE",
            title="Градусы без пробела",
            severity="minor",
            description="Нормализует запись градусов: 5 °, 5 °C -> 5°, 5°C.",
            pattern=re.compile(rf"{NUM_RE}[ \u00A0]*°[ \u00A0]*(?P<scale>[CcСс])?"),
            replacement=_degree_replacement,
        ),
        AuditRule(
            code="AUDIT-NUMBER-SIGN",
            title="Пробел после знака номера",
            severity="minor",
            description="Нормализует запись номера: №1 -> № 1.",
            pattern=re.compile(r"№[ \u00A0]*(?P<num>\d+)"),
            replacement=lambda match: f"№ {match.group('num')}",
        ),
        AuditRule(
            code="AUDIT-DATE-YEAR-SUFFIX",
            title="Дата без суффикса г.",
            severity="minor",
            description="Убирает лишнее 'г.' после полной даты: 20.12.2012 г. -> 20.12.2012.",
            pattern=re.compile(r"(?P<date>\b\d{1,2}\.\d{1,2}\.\d{4})[ \u00A0]*г\."),
            replacement=_date_without_year_suffix,
        ),
        AuditRule(
            code="AUDIT-CAPTION-TABLE-TYPO",
            title="Опечатки в названии таблицы",
            severity="major",
            description="Исправляет безопасные варианты: Таблрица 1, Таблица. 1 -> Таблица 1.",
            pattern=re.compile(r"\b(?:Таблрица|Таблица\.)[ \u00A0]*(?P<num>\d+)", re.IGNORECASE),
            replacement=lambda match: f"Таблица {match.group('num')}",
        ),
        AuditRule(
            code="AUDIT-RF-SCAN",
            title="Сокращение РФ",
            severity="warning",
            description="Фиксирует сокращение РФ для ручной замены полным названием в нужном падеже.",
            pattern=re.compile(r"\bРФ\b"),
            replacement=_russian_federation_suggestion,
            fixable=False,
        ),
    ]


RULES = _build_rules()


def _context(text: str, start: int, end: int) -> str:
    left = max(0, start - TEXT_WINDOW)
    right = min(len(text), end + TEXT_WINDOW)
    return text[left:right]


def _apply_rules_to_text(text: str, part: str, *, fix: bool) -> tuple[str, list[AuditFinding]]:
    current = text
    findings: list[AuditFinding] = []
    for rule in RULES:
        if rule.fixable:

            def replace(match: Match[str], audit_rule: AuditRule = rule) -> str:
                before = match.group(0)
                after = audit_rule.replacement_for(match)
                if after != before:
                    findings.append(
                        AuditFinding(
                            rule_code=audit_rule.code,
                            rule_title=audit_rule.title,
                            severity=audit_rule.severity,
                            action="FIX" if fix else "PLAN",
                            part=part,
                            before=before,
                            after=after,
                            context=_context(current, match.start(), match.end()),
                            start=match.start(),
                            end=match.end(),
                        )
                    )
                return after if fix else before

            current = rule.pattern.sub(replace, current)
            continue

        for match in rule.pattern.finditer(current):
            before = match.group(0)
            after = rule.replacement_for(match)
            findings.append(
                AuditFinding(
                    rule_code=rule.code,
                    rule_title=rule.title,
                    severity=rule.severity,
                    action="SUGGEST",
                    part=part,
                before=before,
                after=after,
                context=_context(current, match.start(), match.end()),
                start=match.start(),
                end=match.end(),
            )
        )
    return current, findings


def _set_text_preserve_space(node, value: str) -> None:
    node.text = value
    if value.startswith(" ") or value.endswith(" "):
        node.set(XML_SPACE, "preserve")


def _w_tag(local_name: str) -> str:
    return f"{{{NS['w']}}}{local_name}"


def _assign_anchor_ids(findings: list[AuditFinding], start_index: int) -> int:
    index = start_index
    for finding in findings:
        index += 1
        finding.anchor_id = f"A{index:04d} {finding.rule_code}"
    return index


def _copy_anchor_ids(source: list[AuditFinding], target: list[AuditFinding]) -> None:
    for planned, actual in zip(source, target):
        actual.anchor_id = planned.anchor_id


def _insert_anchor_run_after_text_node(text_node, findings: list[AuditFinding]) -> bool:
    anchors = [finding.anchor_id for finding in findings if finding.anchor_id]
    if not anchors:
        return False

    run = text_node.getparent()
    parent = run.getparent() if run is not None else None
    if parent is None:
        return False

    anchor_run = etree.Element(_w_tag("r"))
    run_props = etree.Element(_w_tag("rPr"))
    bold = etree.Element(_w_tag("b"))
    color = etree.Element(_w_tag("color"))
    color.set(_w_tag("val"), ANCHOR_COLOR)
    run_props.append(bold)
    run_props.append(color)
    anchor_text = etree.Element(_w_tag("t"))
    anchor_text.set(XML_SPACE, "preserve")
    anchor_text.text = " " + " ".join(f"⟦{anchor}⟧" for anchor in anchors)
    anchor_run.append(run_props)
    anchor_run.append(anchor_text)
    parent.insert(parent.index(run) + 1, anchor_run)
    return True


def _annotated_output_path(docx_path: Path, input_root: Path, out_dir: Path) -> Path:
    base = mirrored_output_path(docx_path, input_root, out_dir)
    return base.with_name(f"{base.stem}__annotated{base.suffix}")


def process_docx(
    docx_path: Path,
    out_path: Path | None,
    *,
    fix: bool,
    dry_run: bool = False,
    annotated_out_path: Path | None = None,
) -> FileAuditReport:
    files = read_zip_map(docx_path)
    annotated_files = dict(files) if annotated_out_path is not None and not dry_run else None
    report = FileAuditReport(
        path=docx_path,
        output=out_path if fix and not dry_run else None,
        annotated_output=annotated_out_path if annotated_files is not None else None,
    )
    changed_any_part = False
    annotated_any_part = False
    anchor_index = 0

    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        annotated_tree = _etree_from_bytes(annotated_files[part]) if annotated_files is not None else None
        annotated_nodes = annotated_tree.getroot().xpath(".//w:t", namespaces=NS) if annotated_tree is not None else []
        part_modified = False
        part_annotated = False
        for node_index, text_node in enumerate(root.xpath(".//w:t", namespaces=NS)):
            if not text_node.text:
                continue
            if text_node.text.strip() == "" and text_node.get(XML_SPACE) == "preserve":
                continue
            report.text_nodes += 1
            _, planned_findings = _apply_rules_to_text(text_node.text, part, fix=False)
            if planned_findings:
                anchor_index = _assign_anchor_ids(planned_findings, anchor_index)
            new_text, node_findings = _apply_rules_to_text(text_node.text, part, fix=fix)
            _copy_anchor_ids(planned_findings, node_findings)
            for finding in node_findings:
                finding.text_node_index = node_index
                report.rule_counts[finding.rule_code] = report.rule_counts.get(finding.rule_code, 0) + 1
            report.findings.extend(node_findings)
            if annotated_tree is not None and planned_findings and node_index < len(annotated_nodes):
                if _insert_anchor_run_after_text_node(annotated_nodes[node_index], planned_findings):
                    report.anchors += len(planned_findings)
                    part_annotated = True
            if new_text != text_node.text:
                report.changed_nodes += 1
                report.fixes += sum(1 for finding in node_findings if finding.action == "FIX")
                _set_text_preserve_space(text_node, new_text)
                part_modified = True
        if part_modified:
            files[part] = _etree_to_bytes(tree)
            changed_any_part = True
        if part_annotated and annotated_tree is not None and annotated_files is not None:
            annotated_files[part] = _etree_to_bytes(annotated_tree)
            annotated_any_part = True

    if fix and not dry_run and out_path is not None:
        safe_mkdir(out_path.parent)
        if changed_any_part:
            write_zip_map(out_path, files)
        else:
            write_zip_map(out_path, files)
    if annotated_files is not None and annotated_out_path is not None:
        safe_mkdir(annotated_out_path.parent)
        if annotated_any_part:
            write_zip_map(annotated_out_path, annotated_files)
        else:
            write_zip_map(annotated_out_path, annotated_files)
    if dry_run:
        report.status = "DRY-RUN"
    elif not report.findings:
        report.status = "PASS"
    elif fix:
        has_unresolved = any(finding.action != "FIX" for finding in report.findings)
        if has_unresolved and report.fixes:
            report.status = "PARTIAL"
        elif has_unresolved:
            report.status = "FAIL"
        else:
            report.status = "FIXED"
    else:
        report.status = "FAIL"
    return report


def _rel_path(path: Path, root: Path) -> str:
    try:
        return path.relative_to(root).as_posix()
    except ValueError:
        return path.name


def _write_docx_report(out_path: Path, reports: list[FileAuditReport], *, fix: bool, dry_run: bool) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Отчёт audit processor DOCX", level=1)
    doc.add_paragraph(f"Режим исправления: {'ДА' if fix else 'НЕТ'}")
    doc.add_paragraph(f"Dry-run: {'ДА' if dry_run else 'НЕТ'}")
    doc.add_paragraph(f"Карта правил: {RULE_CATALOG}")
    doc.add_paragraph("Ограничение: правила применяются внутри отдельных текстовых узлов DOCX (w:t).")

    doc.add_heading("Сводка", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    for idx, header in enumerate(["Файл", "Проблемы", "Исправления", "Якоря", "Изменённые узлы", "Статус", "Выходной файл", "DOCX с якорями"]):
        table.rows[0].cells[idx].text = header
    for item in reports:
        cells = table.add_row().cells
        cells[0].text = item.path.name
        cells[1].text = str(len(item.findings))
        cells[2].text = str(item.fixes)
        cells[3].text = str(item.anchors)
        cells[4].text = str(item.changed_nodes)
        cells[5].text = item.status
        cells[6].text = str(item.output or "")
        cells[7].text = str(item.annotated_output or "")

    doc.add_heading("Правила", level=2)
    rules_table = doc.add_table(rows=1, cols=4)
    rules_table.style = "Table Grid"
    for idx, header in enumerate(["Код", "Правило", "Severity", "Описание"]):
        rules_table.rows[0].cells[idx].text = header
    for rule in RULES:
        cells = rules_table.add_row().cells
        cells[0].text = rule.code
        cells[1].text = rule.title
        cells[2].text = rule.severity
        cells[3].text = rule.description

    doc.add_heading("Примеры", level=2)
    for item in reports:
        doc.add_heading(item.path.name, level=3)
        if not item.findings:
            doc.add_paragraph("Audit-проблемы не найдены.")
            continue
        details = doc.add_table(rows=1, cols=6)
        details.style = "Table Grid"
        for idx, header in enumerate(["Код", "Действие", "Часть DOCX", "Было", "Стало / рекомендация", "Контекст"]):
            details.rows[0].cells[idx].text = header
        for finding in item.findings[:80]:
            cells = details.add_row().cells
            values = [finding.rule_code, finding.action, finding.part, finding.before, finding.after, truncate(finding.context, 160)]
            for idx, value in enumerate(values):
                cells[idx].text = str(value)
    doc.save(str(out_path))


def write_report_md(out_path: Path, reports: list[FileAuditReport], *, input_root: Path, fix: bool, dry_run: bool) -> None:
    safe_mkdir(out_path.parent)
    lines: list[str] = ["# Отчёт audit processor DOCX\n"]
    lines.append(f"- Режим исправления: **{'ДА' if fix else 'НЕТ'}**")
    lines.append(f"- Dry-run: **{'ДА' if dry_run else 'НЕТ'}**")
    lines.append(f"- Карта правил проекта: `{RULE_CATALOG}`")
    lines.append("- Область правок: только видимый текст в отдельных узлах `w:t`; структура DOCX, стили, таблицы, секции и media не переписываются.")
    lines.append("")
    lines.append("## Сводка\n")
    lines.append("| Файл | Проблемы | Исправления | Якоря | Изменённые текстовые узлы | Статус | Выходной файл | DOCX с якорями |")
    lines.append("|---|---:|---:|---:|---:|---|---|---|")
    for item in reports:
        output = _rel_path(item.output, input_root.parent) if item.output else ""
        annotated_output = _rel_path(item.annotated_output, input_root.parent) if item.annotated_output else ""
        lines.append(
            f"| `{md_escape(_rel_path(item.path, input_root))}` | {len(item.findings)} | {item.fixes} | {item.anchors} | {item.changed_nodes} | {item.status} | `{md_escape(output)}` | `{md_escape(annotated_output)}` |"
        )
    lines.append("")

    totals: dict[str, int] = {}
    for item in reports:
        for code, count in item.rule_counts.items():
            totals[code] = totals.get(code, 0) + count
    lines.append("## Правила и найденные случаи\n")
    lines.append("| Код | Правило | Severity | Найдено | Режим |")
    lines.append("|---|---|---|---:|---|")
    for rule in RULES:
        mode = "FIX" if rule.fixable else "SUGGEST"
        lines.append(f"| `{rule.code}` | {md_escape(rule.title)} | `{rule.severity}` | {totals.get(rule.code, 0)} | `{mode}` |")
    lines.append("")

    for item in reports:
        lines.append(f"## {md_escape(_rel_path(item.path, input_root))}\n")
        if item.output:
            lines.append(f"- Выходной файл: `{md_escape(str(item.output))}`")
        if item.annotated_output:
            lines.append(f"- DOCX с якорями: `{md_escape(str(item.annotated_output))}`")
        if not item.findings:
            lines.append("_Audit-проблемы не найдены._\n")
            continue
        lines.append("| Правило | Действие | Часть DOCX | Было | Стало / рекомендация | Контекст |")
        lines.append("|---|---|---|---|---|---|")
        for finding in item.findings[:MAX_FINDINGS_PER_FILE]:
            lines.append(
                f"| `{finding.rule_code}` | `{finding.action}` | `{md_escape(finding.part)}` | `{md_escape(finding.before)}` | `{md_escape(finding.after)}` | {md_escape(truncate(finding.context, 180))} |"
            )
        if len(item.findings) > MAX_FINDINGS_PER_FILE:
            lines.append(f"\n_... показаны первые {MAX_FINDINGS_PER_FILE} случаев из {len(item.findings)}._")
        lines.append("")
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_json_payload(input_root: Path, reports: list[FileAuditReport], *, fix: bool, dry_run: bool) -> dict[str, object]:
    files = []
    for item in reports:
        files.append(
            {
                "path": _rel_path(item.path, input_root),
                "status": item.status,
                "output": str(item.output) if item.output else "",
                "annotated_output": str(item.annotated_output) if item.annotated_output else "",
                "metrics": {
                    "text_nodes": item.text_nodes,
                    "issues": len(item.findings),
                    "fixes": item.fixes,
                    "anchors": item.anchors,
                    "changed_nodes": item.changed_nodes,
                    "rule_counts": item.rule_counts,
                },
                "findings": [
                    {
                        "rule_code": finding.rule_code,
                        "rule_title": finding.rule_title,
                        "severity": finding.severity,
                        "action": finding.action,
                        "part": finding.part,
                        "before": finding.before,
                        "after": finding.after,
                        "context": finding.context,
                        "start": finding.start,
                        "end": finding.end,
                        "text_node_index": finding.text_node_index,
                        "anchor_id": finding.anchor_id,
                    }
                    for finding in item.findings
                ],
            }
        )

    return {
        "tool": "docx_audit_processor",
        "version": 1,
        "input_dir": str(input_root),
        "rule_catalog": str(RULE_CATALOG),
        "fix_enabled": bool(fix),
        "dry_run": bool(dry_run),
        "anchors_enabled": any(item.annotated_output for item in reports),
        "rules": [
            {
                "code": rule.code,
                "title": rule.title,
                "severity": rule.severity,
                "fixable": rule.fixable,
                "description": rule.description,
            }
            for rule in RULES
        ],
        "files": files,
        "summary": {
            "total_files": len(reports),
            "pass_files": sum(1 for item in reports if item.status in {"PASS", "FIXED"}),
            "fail_files": sum(1 for item in reports if item.status in {"FAIL", "PARTIAL"}),
            "issues": sum(len(item.findings) for item in reports),
            "fixes": sum(item.fixes for item in reports),
            "anchors": sum(item.anchors for item in reports),
            "changed_nodes": sum(item.changed_nodes for item in reports),
        },
    }


def _input_paths(args: argparse.Namespace) -> tuple[Path, list[Path]]:
    if args.file:
        file_path = Path(args.file).resolve()
        return file_path.parent, [file_path]
    input_root = Path(args.input).resolve()
    return input_root, find_docx_files(input_root)


def main() -> int:
    parser = argparse.ArgumentParser(description="Scan and safely fix deterministic DOCX audit rules.")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--input", default="input", help="Input folder with .docx files (recursive)")
    group.add_argument("--file", default="", help="Single DOCX file")
    parser.add_argument("--outdir", default="output/audit_processed", help="Output folder for fixed DOCX files")
    parser.add_argument("--out", default="", help="Output DOCX path for single-file fix mode")
    parser.add_argument("--report", default="report/docx_audit_processor.md", help="Markdown report path")
    parser.add_argument("--docx-out", default="", help="Optional DOCX report path")
    parser.add_argument("--no-docx-report", action="store_true", help="Do not write a DOCX report")
    parser.add_argument("--json-out", default="", help="Optional JSON report path")
    parser.add_argument("--fix", action="store_true", help="Apply safe fixable rules to DOCX copies")
    parser.add_argument("--dry-run", action="store_true", help="Plan fixes without writing DOCX output files")
    parser.add_argument("--annotate", action="store_true", help="Write __annotated DOCX copies with audit anchors")
    parser.add_argument("--annotated-outdir", default="output/audit_annotated", help="Output folder for __annotated DOCX files")
    args = parser.parse_args()

    input_root, docx_files = _input_paths(args)
    report_path = Path(args.report).resolve()
    out_dir = Path(args.outdir).resolve()
    annotated_dir = Path(args.annotated_outdir).resolve()
    if args.file and args.fix and not args.out and not args.dry_run:
        print("[ERROR] --out is required with --file --fix")
        return 2
    if args.file and not docx_files[0].exists():
        print(f"[ERROR] File not found: {docx_files[0]}")
        return 2
    if not args.file and not input_root.exists():
        print(f"[ERROR] Input folder does not exist: {input_root}")
        return 2

    reports: list[FileAuditReport] = []
    if not docx_files:
        safe_mkdir(report_path.parent)
        report_path.write_text("# Отчёт audit processor DOCX\n\nФайлы DOCX не найдены.\n", encoding="utf-8")
        if args.json_out:
            write_json_file(Path(args.json_out).resolve(), build_json_payload(input_root, [], fix=args.fix, dry_run=args.dry_run))
        print(f"[WARN] No DOCX files found: {input_root}")
        print(f"[OK] Report: {report_path}")
        return 0

    for docx_path in docx_files:
        out_path: Path | None = None
        annotated_out_path: Path | None = None
        if args.fix:
            out_path = Path(args.out).resolve() if args.file else mirrored_output_path(docx_path, input_root, out_dir)
        if args.annotate and not args.dry_run:
            annotated_out_path = _annotated_output_path(docx_path, input_root, annotated_dir)
        reports.append(process_docx(docx_path, out_path, fix=args.fix, dry_run=args.dry_run, annotated_out_path=annotated_out_path))

    write_report_md(report_path, reports, input_root=input_root, fix=args.fix, dry_run=args.dry_run)
    if not args.no_docx_report:
        docx_report = Path(args.docx_out).resolve() if args.docx_out else report_path.with_suffix(".docx")
        _write_docx_report(docx_report, reports, fix=args.fix, dry_run=args.dry_run)
        print(f"[OK] DOCX report: {docx_report}")
    if args.json_out:
        json_path = Path(args.json_out).resolve()
        write_json_file(json_path, build_json_payload(input_root, reports, fix=args.fix, dry_run=args.dry_run))
        print(f"[OK] JSON report: {json_path}")

    print(f"[OK] Проверено файлов: {len(reports)}")
    print(f"[OK] Найдено audit-случаев: {sum(len(item.findings) for item in reports)}")
    if args.fix:
        if args.dry_run:
            print("[OK] Dry-run: DOCX output не записывался.")
        else:
            print(f"[OK] Output: {Path(args.out).resolve() if args.file and args.out else out_dir}")
        print(f"[OK] Исправлений: {sum(item.fixes for item in reports)}")
    if args.annotate and not args.dry_run:
        print(f"[OK] DOCX с якорями: {annotated_dir}")
        print(f"[OK] Якорей проставлено: {sum(item.anchors for item in reports)}")
    print(f"[OK] Report: {report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
