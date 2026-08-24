#!/usr/bin/env python3
"""Remove safe invisible nonprinting garbage from copied DOCX documents."""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from _office_common import find_docx_files, md_escape, mirrored_output_path, safe_mkdir, write_json_file
from docx_xml_tools import NS, _etree_from_bytes, _etree_to_bytes, list_xml_parts, read_zip_map, write_zip_map


XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# Conservative allow-list of what this cleaner may remove.
#
# We intentionally do not touch normal spaces, NBSP, tabs, line/page/section
# breaks, visible hyphens/dashes, non-breaking hyphens, zero-width break/control
# characters, or bidi marks: deleting those can move lines, change wrapping, or
# alter text direction. The current safe target is the optional hyphen artifact
# seen in Word as a nonprinting mark inside words.
NONPRINTING_CHARS: dict[str, str] = {
    "\u00AD": "soft_hyphen",
}


@dataclass
class CleanResult:
    source: Path
    output: Path
    status: str = "OK"
    by_kind: dict[str, int] = field(default_factory=dict)
    xml_soft_hyphens: int = 0
    changed_nodes: int = 0
    error: str = ""

    @property
    def text_removed(self) -> int:
        return sum(self.by_kind.values())

    @property
    def total_removed(self) -> int:
        return self.text_removed + self.xml_soft_hyphens


def _set_text_preserve_space(node, value: str) -> None:
    node.text = value
    if value.startswith(" ") or value.endswith(" "):
        node.set(XML_SPACE, "preserve")
    elif XML_SPACE in node.attrib:
        node.attrib.pop(XML_SPACE, None)


def clean_text(text: str, result: CleanResult) -> str:
    if not text:
        return text
    changed = False
    out: list[str] = []
    for char in text:
        kind = NONPRINTING_CHARS.get(char)
        if kind is None:
            out.append(char)
            continue
        result.by_kind[kind] = result.by_kind.get(kind, 0) + 1
        changed = True
    return "".join(out) if changed else text


def remove_nonprinting(source: Path, output: Path) -> CleanResult:
    files = read_zip_map(source)
    result = CleanResult(source=source, output=output)

    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        modified = False

        # Word stores optional hyphen both as U+00AD in w:t and as an empty
        # w:softHyphen element between runs.
        for node in root.xpath(".//w:softHyphen", namespaces=NS):
            parent = node.getparent()
            if parent is None:
                continue
            parent.remove(node)
            result.xml_soft_hyphens += 1
            modified = True

        for text_node in root.xpath(".//w:t", namespaces=NS):
            text = text_node.text or ""
            cleaned = clean_text(text, result)
            if cleaned == text:
                continue
            result.changed_nodes += 1
            _set_text_preserve_space(text_node, cleaned)
            modified = True

        if modified:
            files[part] = _etree_to_bytes(tree)

    safe_mkdir(output.parent)
    write_zip_map(output, files)
    return result


def resolve_inputs(args: argparse.Namespace) -> tuple[Path, list[Path]]:
    if args.file:
        source = Path(args.file).resolve()
        return source.parent, [source] if source.exists() and source.suffix.lower() == ".docx" else []
    input_root = Path(args.input).resolve()
    return input_root, find_docx_files(input_root) if input_root.exists() else []


def _kind_summary(result: CleanResult) -> str:
    parts = [f"{kind}: {count}" for kind, count in sorted(result.by_kind.items()) if count]
    if result.xml_soft_hyphens:
        parts.append(f"w:softHyphen: {result.xml_soft_hyphens}")
    return "; ".join(parts)


def write_report(report_path: Path, input_root: Path, results: list[CleanResult]) -> None:
    safe_mkdir(report_path.parent)
    lines = ["# Удаление непечатаемого мусора DOCX\n"]
    lines.append(f"- Вход: `{md_escape(str(input_root))}`")
    lines.append(f"- Файлов обработано: **{len(results)}**")
    lines.append(f"- Удалено символов/элементов: **{sum(item.total_removed for item in results)}**")
    lines.append("")
    lines.append("| Файл | Статус | Удалено | Изменённые текстовые узлы | Выходной файл | Детали | Ошибка |")
    lines.append("|---|---|---:|---:|---|---|---|")
    for item in results:
        lines.append(
            f"| `{md_escape(str(item.source))}` | `{item.status}` | {item.total_removed} | {item.changed_nodes} | `{md_escape(str(item.output))}` | {md_escape(_kind_summary(item))} | {md_escape(item.error)} |"
        )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    doc = Document()
    doc.add_heading("Удаление непечатаемого мусора DOCX", level=1)
    doc.add_paragraph(f"Вход: {input_root}")
    doc.add_paragraph(f"Файлов обработано: {len(results)}")
    doc.add_paragraph(f"Удалено символов/элементов: {sum(item.total_removed for item in results)}")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for index, header in enumerate(["Файл", "Статус", "Удалено", "Текстовые узлы", "Выходной файл", "Детали"]):
        table.rows[0].cells[index].text = header
    for item in results:
        cells = table.add_row().cells
        cells[0].text = str(item.source)
        cells[1].text = item.status
        cells[2].text = str(item.total_removed)
        cells[3].text = str(item.changed_nodes)
        cells[4].text = str(item.output)
        cells[5].text = _kind_summary(item)
    doc.save(str(report_path.with_suffix(".docx")))


def build_json(input_root: Path, results: list[CleanResult]) -> dict[str, object]:
    all_kinds = sorted({kind for item in results for kind in item.by_kind})
    return {
        "tool": "docx_nonprinting_clean",
        "input_root": str(input_root),
        "summary": {
            "files": len(results),
            "ok": sum(1 for item in results if item.status == "OK"),
            "failed": sum(1 for item in results if item.status != "OK"),
            "removed": sum(item.total_removed for item in results),
            "xml_soft_hyphens": sum(item.xml_soft_hyphens for item in results),
            "by_kind": {kind: sum(item.by_kind.get(kind, 0) for item in results) for kind in all_kinds},
        },
        "files": [
            {
                "source": str(item.source),
                "output": str(item.output),
                "status": item.status,
                "removed": item.total_removed,
                "changed_nodes": item.changed_nodes,
                "by_kind": item.by_kind,
                "xml_soft_hyphens": item.xml_soft_hyphens,
                "error": item.error,
            }
            for item in results
        ],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Remove safe invisible nonprinting garbage from DOCX copies.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--input", help="Input folder with DOCX files")
    group.add_argument("--file", help="Single DOCX file")
    parser.add_argument("--outdir", default="output/nonprinting_cleaned", help="Output folder for batch mode")
    parser.add_argument("--out", default="", help="Output DOCX path for single-file mode")
    parser.add_argument("--report", default="report/docx_nonprinting_clean.md", help="Markdown report path")
    parser.add_argument("--json-out", default="", help="Optional JSON report path")
    args = parser.parse_args()

    input_root, docx_files = resolve_inputs(args)
    report_path = Path(args.report).resolve()
    out_dir = Path(args.outdir).resolve()

    if args.file and not docx_files:
        print(f"[ERROR] DOCX file not found: {Path(args.file).resolve()}")
        return 2
    if args.input and not input_root.exists():
        print(f"[ERROR] Input folder does not exist: {input_root}")
        return 2
    if not docx_files:
        safe_mkdir(report_path.parent)
        report_path.write_text("# Удаление непечатаемого мусора DOCX\n\nФайлы DOCX не найдены.\n", encoding="utf-8")
        print(f"[WARN] No DOCX files found: {input_root}")
        print(f"[OK] Report: {report_path}")
        return 0

    results: list[CleanResult] = []
    for source in docx_files:
        output = Path(args.out).resolve() if args.file and args.out else mirrored_output_path(source, input_root, out_dir)
        try:
            result = remove_nonprinting(source, output)
            print(f"[OK] {source} -> {output} removed={result.total_removed}")
        except Exception as exc:
            result = CleanResult(source=source, output=output, status="FAILED", error=str(exc))
            print(f"[FAILED] {source}: {exc}")
        results.append(result)

    write_report(report_path, input_root, results)
    if args.json_out:
        json_path = Path(args.json_out).resolve()
        write_json_file(json_path, build_json(input_root, results))
        print(f"[OK] JSON: {json_path}")
    print(f"[OK] DOCX report: {report_path.with_suffix('.docx')}")
    print(f"[OK] Report: {report_path}")
    return 1 if any(item.status != "OK" for item in results) else 0


if __name__ == "__main__":
    raise SystemExit(main())
