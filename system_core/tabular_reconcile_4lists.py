#!/usr/bin/env python3
"""
Tabular Reconcile (4 lists)

Given two tabular sources (XLSX/CSV/DOCX table), produce:

1) Rows (by key) present in both AND all compared fields match
2) Rows present in both BUT some compared fields differ (show which fields differ)
3) Rows present only in A
4) Rows present only in B

This is deterministic (no AI). Designed for office workflows.

Examples:
  python tabular_reconcile_4lists.py --a A.xlsx --b B.docx --key "Наименование мероприятия" --fields "Срок реализации" "Местоположение" --out report/reconcile.md

Notes:
- XLSX is read with data_only=True (computed values).
- DOCX: reads table #1 by default (use --docx-table-index to change).
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import csv
from pathlib import Path
from typing import Dict, List, Any, Tuple

from docx import Document
from docx import Document as DocxDocument
from openpyxl import load_workbook

from _office_common import safe_mkdir, norm_space, norm_key, md_escape, DiffField, truncate, write_json_file

AUTO_KEY_PRIORITIES = [
    "№",
    "id",
    "Id",
    "ID",
    "code",
    "Код",
    "key",
    "Ключ",
    "Наименование",
    "Name",
    "Title",
    "Имя",
]

def load_csv(path: Path) -> List[Dict[str, Any]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        r = csv.DictReader(f)
        return [dict(row) for row in r]

def load_xlsx(path: Path, sheet: str | None) -> List[Dict[str, Any]]:
    wb = load_workbook(str(path), data_only=True, read_only=True)
    ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    headers = [norm_space(str(h)) if h is not None else "" for h in rows[0]]
    out: List[Dict[str, Any]] = []
    for row in rows[1:]:
        d = {}
        for i, h in enumerate(headers):
            if not h:
                continue
            v = row[i] if i < len(row) else None
            if isinstance(v, str):
                v = v.replace("\r\n", "\n").replace("\r", "\n")
            d[h] = v
        out.append(d)
    return out

def load_docx_table(path: Path, table_index: int) -> List[Dict[str, Any]]:
    doc = Document(str(path))
    if not doc.tables:
        return []
    idx = max(0, table_index)
    if idx >= len(doc.tables):
        idx = 0
    t = doc.tables[idx]
    if not t.rows:
        return []
    headers = [norm_space(cell.text) for cell in t.rows[0].cells]
    out: List[Dict[str, Any]] = []
    for r in t.rows[1:]:
        d = {}
        for i, h in enumerate(headers):
            if not h:
                continue
            v = r.cells[i].text if i < len(r.cells) else ""
            d[h] = v
        out.append(d)
    return out

def load_any(path: Path, *, xlsx_sheet: str | None, docx_table_index: int) -> Tuple[str, List[Dict[str, Any]]]:
    ext = path.suffix.lower()
    if ext == ".csv":
        return "csv", load_csv(path)
    if ext in (".xlsx", ".xlsm"):
        return "xlsx", load_xlsx(path, xlsx_sheet)
    if ext == ".docx":
        return "docx", load_docx_table(path, docx_table_index)
    raise ValueError(f"Unsupported file type: {ext}")

def make_key(row: Dict[str, Any], key_cols: List[str]) -> str:
    parts = []
    for k in key_cols:
        parts.append(norm_key(str(row.get(k, "") or "")))
    return " | ".join(parts).strip(" |")

def field_value(row: Dict[str, Any], field: str) -> str:
    v = row.get(field, "")
    if v is None:
        return ""
    return norm_space(str(v))

def table_columns(rows: list[dict[str, Any]]) -> list[str]:
    seen: set[str] = set()
    columns: list[str] = []
    for row in rows:
        for col in row.keys():
            name = norm_space(str(col))
            key = norm_key(name)
            if name and key not in seen:
                seen.add(key)
                columns.append(name)
    return columns

def detect_auto_key(a_rows: list[dict[str, Any]], b_rows: list[dict[str, Any]]) -> str:
    a_columns = table_columns(a_rows)
    b_by_norm = {norm_key(col): col for col in table_columns(b_rows)}
    common: list[tuple[str, str]] = []
    for col in a_columns:
        key = norm_key(col)
        if key in b_by_norm:
            common.append((key, col))
    common_by_norm = {key: col for key, col in common}
    for candidate in AUTO_KEY_PRIORITIES:
        found = common_by_norm.get(norm_key(candidate))
        if found:
            return found
    for candidate in AUTO_KEY_PRIORITIES:
        candidate_key = norm_key(candidate)
        for key, col in common:
            if key.startswith(candidate_key) or candidate_key in key:
                return col
    return ""

def write_docx_report(
    out_path: Path,
    *,
    a_path: Path,
    b_path: Path,
    a_kind: str,
    b_kind: str,
    a_rows: list[dict],
    b_rows: list[dict],
    key_cols: list[str],
    fields: list[str],
    exact: list[str],
    mismatched: list[tuple[str, list[DiffField]]],
    only_a: list[str],
    only_b: list[str],
    max_rows: int,
) -> None:
    safe_mkdir(out_path.parent)
    doc = DocxDocument()
    doc.add_heading("Отчёт сверки таблиц (4 списка)", level=1)
    doc.add_paragraph(f"A: {a_path.name} ({a_kind}), строк: {len(a_rows)}")
    doc.add_paragraph(f"B: {b_path.name} ({b_kind}), строк: {len(b_rows)}")
    doc.add_paragraph(f"Ключевые колонки: {', '.join(key_cols)}")
    doc.add_paragraph(f"Сравниваемые поля: {', '.join(fields) if fields else '(нет)'}")

    doc.add_heading("Сводка", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Категория"
    t.rows[0].cells[1].text = "Количество"
    for label, count in [
        ("Есть в обоих, все поля совпадают", len(exact)),
        ("Есть в обоих, часть полей отличается", len(mismatched)),
        ("Только в A", len(only_a)),
        ("Только в B", len(only_b)),
    ]:
        cells = t.add_row().cells
        cells[0].text = label
        cells[1].text = str(count)

    doc.add_heading("1) Есть в обоих, все поля совпадают", level=2)
    for k in exact[:max_rows]:
        doc.add_paragraph(k)

    doc.add_heading("2) Есть в обоих, но поля отличаются", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for idx, header in enumerate(["Ключ", "Поле", "A", "B"]):
        t.rows[0].cells[idx].text = header
    shown = 0
    for k, diffs in mismatched:
        for d in diffs:
            cells = t.add_row().cells
            cells[0].text = k
            cells[1].text = d.field
            cells[2].text = truncate(d.a, 200)
            cells[3].text = truncate(d.b, 200)
            shown += 1
            if shown >= max_rows:
                break
        if shown >= max_rows:
            break

    doc.add_heading("3) Только в A", level=2)
    for k in only_a[:max_rows]:
        doc.add_paragraph(k)

    doc.add_heading("4) Только в B", level=2)
    for k in only_b[:max_rows]:
        doc.add_paragraph(k)

    doc.save(str(out_path))

def main() -> int:
    ap = argparse.ArgumentParser(description="Reconcile two tabular files into 4 lists (match/mismatch/onlyA/onlyB).")
    ap.add_argument("--a", required=True, help="File A (xlsx/csv/docx)")
    ap.add_argument("--b", required=True, help="File B (xlsx/csv/docx)")
    ap.add_argument("--key", nargs="*", default=[], help="Key column(s) used to match rows")
    ap.add_argument("--key-csv", default="", help="Key columns as a comma-separated string (e.g. ColA,ColB)")
    ap.add_argument("--auto-key", action="store_true", help="Detect a key column by header if --key/--key-csv is not provided")
    ap.add_argument("--fields", nargs="*", default=[], help="Fields to compare (default: all non-key fields present)")
    ap.add_argument("--fields-csv", default="", help="Fields to compare as a comma-separated string")
    ap.add_argument("--out", default="report/reconcile_4lists.md", help="Output Markdown report path")
    ap.add_argument("--docx-out", default="", help="Optional DOCX report path")
    ap.add_argument("--no-docx-report", action="store_true", help="Do not write a DOCX report")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")

    ap.add_argument("--xlsx-sheet-a", default=None, help="Sheet name for A (xlsx only). Default: first sheet")
    ap.add_argument("--xlsx-sheet-b", default=None, help="Sheet name for B (xlsx only). Default: first sheet")
    ap.add_argument("--docx-table-index-a", type=int, default=0, help="DOCX table index for A (0-based). Default 0")
    ap.add_argument("--docx-table-index-b", type=int, default=0, help="DOCX table index for B (0-based). Default 0")

    ap.add_argument("--max-rows", type=int, default=5000, help="Max rows per section in report (default: 5000)")
    args = ap.parse_args()

    a_path = Path(args.a).resolve()
    b_path = Path(args.b).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)

    if not a_path.exists() or not b_path.exists():
        print("[ERROR] One of the input files does not exist.")
        print(f"A: {a_path}")
        print(f"B: {b_path}")
        return 2

    a_kind, a_rows = load_any(a_path, xlsx_sheet=args.xlsx_sheet_a, docx_table_index=args.docx_table_index_a)
    b_kind, b_rows = load_any(b_path, xlsx_sheet=args.xlsx_sheet_b, docx_table_index=args.docx_table_index_b)

    key_cols = []
    if args.key_csv.strip():
        key_cols = [norm_space(x) for x in args.key_csv.split(",") if norm_space(x)]
    else:
        key_cols = [norm_space(k) for k in args.key if norm_space(k)]
    auto_key_used = False
    if not key_cols and args.auto_key:
        found_key = detect_auto_key(a_rows, b_rows)
        if found_key:
            key_cols = [found_key]
            auto_key_used = True
            print(f"[INFO] auto-key: {found_key}")
    if not key_cols:
        print("[ERROR] Key columns are required. Use --key or --key-csv.")
        return 2

    # Build field list if not provided: union of columns minus keys
    if args.fields_csv.strip():
        fields = [norm_space(x) for x in args.fields_csv.split(",") if norm_space(x)]
    elif args.fields:
        fields = [norm_space(f) for f in args.fields]
    else:
        cols = set()
        for r in a_rows + b_rows:
            cols.update([norm_space(c) for c in r.keys()])
        fields = sorted([c for c in cols if c and c not in set(key_cols)])

    # Map by key; if duplicates, keep first and warn counts
    a_map: Dict[str, Dict[str, Any]] = {}
    b_map: Dict[str, Dict[str, Any]] = {}
    dup_a = 0
    dup_b = 0

    for r in a_rows:
        k = make_key(r, key_cols)
        if not k:
            continue
        if k in a_map:
            dup_a += 1
            continue
        a_map[k] = r

    for r in b_rows:
        k = make_key(r, key_cols)
        if not k:
            continue
        if k in b_map:
            dup_b += 1
            continue
        b_map[k] = r

    keys_a = set(a_map.keys())
    keys_b = set(b_map.keys())
    common = sorted(keys_a & keys_b)
    only_a = sorted(keys_a - keys_b)
    only_b = sorted(keys_b - keys_a)

    exact = []
    mismatched = []  # (key, diffs)
    for k in common:
        ra = a_map[k]
        rb = b_map[k]
        diffs: List[DiffField] = []
        for f in fields:
            va = field_value(ra, f)
            vb = field_value(rb, f)
            if va != vb:
                diffs.append(DiffField(f, va, vb))
        if diffs:
            mismatched.append((k, diffs))
        else:
            exact.append(k)

    # Report
    lines: List[str] = []
    lines.append("# Отчёт сверки таблиц (4 списка)\n")
    lines.append(f"- A: `{a_path.name}` ({a_kind}), строк: **{len(a_rows)}**")
    lines.append(f"- B: `{b_path.name}` ({b_kind}), строк: **{len(b_rows)}**")
    lines.append(f"- Ключевые колонки: {', '.join([f'`{c}`' for c in key_cols])}")
    lines.append(f"- Сравниваемые поля: {', '.join([f'`{c}`' for c in fields]) if fields else '(нет)'}\n")

    if dup_a or dup_b:
        lines.append("## Предупреждения\n")
        if dup_a:
            lines.append(f"- Дубли ключей в A пропущены: **{dup_a}**")
        if dup_b:
            lines.append(f"- Дубли ключей в B пропущены: **{dup_b}**")
        lines.append("")

    lines.append("## Сводка\n")
    lines.append("| Категория | Количество |")
    lines.append("|---|---:|")
    lines.append(f"| 1) Есть в обоих, все поля совпадают | {len(exact)} |")
    lines.append(f"| 2) Есть в обоих, часть полей отличается | {len(mismatched)} |")
    lines.append(f"| 3) Только в A | {len(only_a)} |")
    lines.append(f"| 4) Только в B | {len(only_b)} |")
    lines.append("")

    # Section 1
    lines.append("## 1) Есть в обоих, все поля совпадают\n")
    if not exact:
        lines.append("_Нет._\n")
    else:
        lines.append("| Ключ |")
        lines.append("|---|")
        for k in exact[:args.max_rows]:
            lines.append(f"| `{md_escape(k)}` |")
        if len(exact) > args.max_rows:
            lines.append(f"\n_... показаны первые {args.max_rows} ключей._\n")
        lines.append("")

    # Section 2
    lines.append("## 2) Есть в обоих, но поля отличаются\n")
    if not mismatched:
        lines.append("_Нет._\n")
    else:
        lines.append("| Ключ | Поле | A | B |")
        lines.append("|---|---|---|---|")
        shown = 0
        for k, diffs in mismatched:
            for d in diffs:
                lines.append(f"| `{md_escape(k)}` | `{md_escape(d.field)}` | {truncate(md_escape(d.a), 200)} | {truncate(md_escape(d.b), 200)} |")
                shown += 1
                if shown >= args.max_rows:
                    break
            if shown >= args.max_rows:
                break
        if shown >= args.max_rows:
            lines.append(f"\n_... показаны первые {args.max_rows} строк отличий._\n")
        lines.append("")

    # Section 3
    lines.append("## 3) Только в A\n")
    if not only_a:
        lines.append("_Нет._\n")
    else:
        lines.append("| Ключ |")
        lines.append("|---|")
        for k in only_a[:args.max_rows]:
            lines.append(f"| `{md_escape(k)}` |")
        if len(only_a) > args.max_rows:
            lines.append(f"\n_... показаны первые {args.max_rows} ключей._\n")
        lines.append("")

    # Section 4
    lines.append("## 4) Только в B\n")
    if not only_b:
        lines.append("_Нет._\n")
    else:
        lines.append("| Ключ |")
        lines.append("|---|")
        for k in only_b[:args.max_rows]:
            lines.append(f"| `{md_escape(k)}` |")
        if len(only_b) > args.max_rows:
            lines.append(f"\n_... показаны первые {args.max_rows} ключей._\n")
        lines.append("")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    if not args.no_docx_report:
        docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
        write_docx_report(
            docx_out,
            a_path=a_path,
            b_path=b_path,
            a_kind=a_kind,
            b_kind=b_kind,
            a_rows=a_rows,
            b_rows=b_rows,
            key_cols=key_cols,
            fields=fields,
            exact=exact,
            mismatched=mismatched,
            only_a=only_a,
            only_b=only_b,
            max_rows=args.max_rows,
        )
        print(f"[OK] Wrote DOCX report: {docx_out}")
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        status = "FAIL" if mismatched or only_a or only_b else "PASS"
        write_json_file(
            json_out,
            {
                "tool": "tabular_reconcile_4lists",
                "version": 1,
                "input_dir": "",
                "inputs": {"a": str(a_path), "b": str(b_path)},
                "files": [
                    {
                        "path": a_path.name,
                        "status": status,
                        "metrics": {"rows": len(a_rows), "duplicate_keys": dup_a},
                    },
                    {
                        "path": b_path.name,
                        "status": status,
                        "metrics": {"rows": len(b_rows), "duplicate_keys": dup_b},
                    },
                ],
                "summary": {
                    "status": status,
                    "key_columns": key_cols,
                    "auto_key": auto_key_used,
                    "fields": fields,
                    "exact": len(exact),
                    "mismatched": len(mismatched),
                    "only_a": len(only_a),
                    "only_b": len(only_b),
                    "duplicate_keys_a": dup_a,
                    "duplicate_keys_b": dup_b,
                },
                "exact": exact[: args.max_rows],
                "mismatched": [
                    {
                        "key": key,
                        "diffs": [
                            {"field": diff.field, "a": diff.a, "b": diff.b}
                            for diff in diffs
                        ],
                    }
                    for key, diffs in mismatched[: args.max_rows]
                ],
                "only_a": only_a[: args.max_rows],
                "only_b": only_b[: args.max_rows],
            },
        )
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Wrote: {out_path}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
