#!/usr/bin/env python3
"""
DOCX merge through Microsoft Word COM.

This delegates merging to Word instead of editing the DOCX ZIP package
directly. The goal is to behave close to a person opening Word and inserting
the next document after the previous one.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from pathlib import Path
import shutil
import subprocess
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document

from _office_common import safe_mkdir


SCRIPT_RUNNER = Path(__file__).resolve().parent / "word_com" / "Merge-DocxWord.ps1"
WORD_UNAVAILABLE_MESSAGE = (
    "[ERROR] Microsoft Word COM недоступен из текущего сеанса. Этот инструмент\n"
    "требует установленного Microsoft Word и доступного COM-запуска Word.Application\n"
    "на этой машине. Альтернативы для склейки в этом тулките нет по дизайну\n"
    "(см. AGENTS.md -> DOCX Processing Decisions)."
)


def _pwsh_candidates(root: Path) -> list[str]:
    return [
        str(root / "system_core" / "powershell" / "pwsh.exe"),
        "pwsh",
        "powershell",
    ]


def _resolve_pwsh(root: Path) -> str:
    for candidate in _pwsh_candidates(root):
        if Path(candidate).exists():
            return candidate
        found = shutil.which(candidate)
        if found:
            return found
    raise RuntimeError("PowerShell was not found. Install portable PowerShell or add pwsh/powershell to PATH.")


def _check_word_available(pwsh: str, root: Path) -> int:
    script = f"""
$word = $null
try {{
    $word = New-Object -ComObject Word.Application
    exit 0
}} catch {{
    Write-Output @'
{WORD_UNAVAILABLE_MESSAGE}
'@
    Write-Output $_.Exception.Message
    exit 2
}} finally {{
    if ($word -ne $null) {{
        try {{ $word.Quit() | Out-Null }} catch {{}}
        try {{ [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null }} catch {{}}
    }}
}}
"""
    process = subprocess.run(
        [pwsh, "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        cwd=str(root),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if process.stdout:
        print(process.stdout.rstrip())
    return process.returncode


def main() -> int:
    parser = argparse.ArgumentParser(description="Merge DOCX files using Microsoft Word COM.")
    parser.add_argument("--input", default="input", help="Input folder")
    parser.add_argument("--files", default="", help="Optional semicolon-separated file list relative to input")
    parser.add_argument("--out", default="output/merged.docx", help="Output DOCX path")
    parser.add_argument("--report", default="report/docx_merge.md", help="Output Markdown report")
    args = parser.parse_args()

    root = Path(__file__).resolve().parents[1]
    input_dir = Path(args.input).resolve()
    out_path = Path(args.out).resolve()
    report_path = Path(args.report).resolve()
    safe_mkdir(out_path.parent)
    safe_mkdir(report_path.parent)

    if not SCRIPT_RUNNER.exists():
        raise RuntimeError(f"Word COM merge script was not found: {SCRIPT_RUNNER}")

    pwsh = _resolve_pwsh(root)
    word_rc = _check_word_available(pwsh, root)
    if word_rc != 0:
        return 2 if word_rc == 2 else word_rc

    command = [
        pwsh,
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(SCRIPT_RUNNER),
        "-InputDir",
        str(input_dir),
        "-Files",
        str(args.files or ""),
        "-Out",
        str(out_path),
        "-Report",
        str(report_path),
    ]
    process = subprocess.run(
        command,
        cwd=str(root),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if process.stdout:
        print(process.stdout.rstrip())
    if process.returncode != 0:
        return process.returncode

    doc = Document(str(out_path))
    print(f"[OK] Validated with python-docx: paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
