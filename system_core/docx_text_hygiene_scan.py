#!/usr/bin/env python3
"""
DOCX Text Hygiene (Scan)

Detects basic mechanical issues in DOCX text:
- multiple spaces
- spaces before punctuation
- missing space after punctuation (, ; : ! ?)
- soft/optional hyphens that split words only when formatting marks are shown
- optional: missing space after dot (.) (disabled by default)

Outputs a Markdown report with counts and sample contexts.

Usage:
  python docx_text_hygiene_scan.py --input input --out report/docx_text_hygiene_report.md
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import hashlib
import re
from pathlib import Path

from _office_common import safe_mkdir, truncate, md_escape, find_docx_files, rel_posix, write_json_file
from docx_xml_tools import read_zip_map, list_xml_parts, NS, _etree_from_bytes
from docx import Document

PAT_DOUBLE_SPACE = re.compile(r" {2,}")
PAT_SPACE_BEFORE_PUNCT = re.compile(r"\s+([,.;:!?])")
PAT_MISSING_AFTER_PUNCT = re.compile(r"([,;:!?])(?=[^\s\]\)\}\>\"'”’])")
PAT_MISSING_AFTER_DOT = re.compile(r"(?<!\d)(\.)(?=[A-Za-zА-Яа-яЁё])")  # simple heuristic
PAT_SOFT_HYPHEN = re.compile("\u00AD")
DOT_PREFIX_TOKEN_RE = re.compile(r"([A-Za-zА-Яа-яЁё]{1,16})$")
DOT_SPACING_SKIP_TOKENS = {
    "г", "д", "п", "с", "ул", "пер", "пр", "просп", "ш", "наб", "пл", "бул",
    "пос", "дер", "рп", "корп", "стр", "лит", "оф", "кв", "куб", "руб",
    "коп", "тыс", "млн", "млрд", "табл", "рис", "см", "смт", "обл", "рн",
    "им",
}

HYGIENE_LABELS_RU = {
    "double_space": "Двойные пробелы",
    "space_before_punct": "Пробел перед пунктуацией",
    "missing_after_punct": "Нет пробела после пунктуации",
    "missing_after_dot": "Нет пробела после точки",
    "soft_hyphen": "Мягкие переносы",
}

def iter_text_nodes(files: dict[str, bytes]):
    for part, _index, text in iter_text_nodes_with_index(files):
        yield part, text


def iter_text_nodes_with_index(files: dict[str, bytes]):
    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        root = tree.getroot()
        for index, t in enumerate(root.xpath(".//w:t", namespaces=NS)):
            if t.text:
                yield part, index, t.text


def is_missing_after_dot_candidate(text: str, match: re.Match[str]) -> bool:
    token_match = DOT_PREFIX_TOKEN_RE.search(text[: match.start()])
    if not token_match:
        return True
    token = token_match.group(1).casefold().replace("ё", "е")
    return token not in DOT_SPACING_SKIP_TOKENS


def count_missing_after_dot(text: str) -> int:
    return sum(1 for match in PAT_MISSING_AFTER_DOT.finditer(text) if is_missing_after_dot_candidate(text, match))


def scan_text(text: str, check_dot: bool) -> dict[str, int]:
    counts = {key: 0 for key in HYGIENE_LABELS_RU}
    for finding in find_text_issues("", 0, text, check_dot):
        counts[str(finding["class_id"])] += 1
    return counts


def count_soft_hyphen_elements(files: dict[str, bytes]) -> tuple[int, list[str], list[dict[str, object]]]:
    count = 0
    samples: list[str] = []
    findings: list[dict[str, object]] = []
    for part in list_xml_parts(files):
        if part not in files:
            continue
        tree = _etree_from_bytes(files[part])
        nodes = tree.getroot().xpath(".//w:softHyphen", namespaces=NS)
        if not nodes:
            continue
        count += len(nodes)
        if len(samples) < 5:
            samples.append(f"{part}: w:softHyphen x{len(nodes)}")
        for index, _node in enumerate(nodes):
            context = f"{part}: w:softHyphen #{index + 1}"
            findings.append(
                {
                    "class_id": "soft_hyphen",
                    "label": HYGIENE_LABELS_RU["soft_hyphen"],
                    "owner": "docx_text_hygiene",
                    "part": part,
                    "target_precision": "xml-element",
                    "element_index": index,
                    "text_node_index": -1,
                    "start": 0,
                    "end": 0,
                    "before": "w:softHyphen",
                    "after": "",
                    "context": context,
                    "context_hash": stable_hash(context),
                    "risk": "safe",
                    "autofix": "available",
                    "action": "PLAN",
                }
            )
    return count, samples, findings


def display_context(text: str) -> str:
    return text.replace("\u00AD", "¬")


def stable_hash(value: object, length: int = 12) -> str:
    return hashlib.sha1(str(value).encode("utf-8", errors="replace")).hexdigest()[:length]


def finding_context(text: str, start: int, end: int, window: int = 50) -> str:
    left = max(0, start - window)
    right = min(len(text), end + window)
    return display_context(text[left:right])


def text_finding(
    class_id: str,
    part: str,
    text_node_index: int,
    text: str,
    start: int,
    end: int,
    before: str,
    after: str,
) -> dict[str, object]:
    context = finding_context(text, start, end)
    return {
        "class_id": class_id,
        "label": HYGIENE_LABELS_RU[class_id],
        "owner": "docx_text_hygiene",
        "part": part,
        "text_node_index": text_node_index,
        "start": start,
        "end": end,
        "before": display_context(before),
        "after": display_context(after),
        "context": context,
        "context_hash": stable_hash(context),
        "risk": "safe",
        "autofix": "available",
        "action": "PLAN",
    }


def find_text_issues(part: str, text_node_index: int, text: str, check_dot: bool) -> list[dict[str, object]]:
    findings: list[dict[str, object]] = []
    for match in PAT_DOUBLE_SPACE.finditer(text):
        findings.append(text_finding("double_space", part, text_node_index, text, match.start(), match.end(), match.group(0), " "))
    for match in PAT_SPACE_BEFORE_PUNCT.finditer(text):
        findings.append(
            text_finding("space_before_punct", part, text_node_index, text, match.start(), match.end(), match.group(0), match.group(1))
        )
    for match in PAT_MISSING_AFTER_PUNCT.finditer(text):
        findings.append(
            text_finding("missing_after_punct", part, text_node_index, text, match.start(), match.end(), match.group(0), f"{match.group(1)} ")
        )
    if check_dot:
        for match in PAT_MISSING_AFTER_DOT.finditer(text):
            if not is_missing_after_dot_candidate(text, match):
                continue
            findings.append(text_finding("missing_after_dot", part, text_node_index, text, match.start(), match.end(), match.group(0), ". "))
    for match in PAT_SOFT_HYPHEN.finditer(text):
        findings.append(text_finding("soft_hyphen", part, text_node_index, text, match.start(), match.end(), match.group(0), ""))
    return sorted(findings, key=lambda item: (str(item["part"]), int(item["text_node_index"]), int(item["start"]), str(item["class_id"])))


def sample_context(text: str, pat: re.Pattern, max_samples: int = 5, window: int = 40) -> list[str]:
    out = []
    for m in pat.finditer(text):
        start = max(0, m.start() - window)
        end = min(len(text), m.end() + window)
        out.append(display_context(text[start:end]))
        if len(out) >= max_samples:
            break
    return out


def sample_missing_after_dot(text: str, max_samples: int = 5, window: int = 40) -> list[str]:
    out = []
    for m in PAT_MISSING_AFTER_DOT.finditer(text):
        if not is_missing_after_dot_candidate(text, m):
            continue
        start = max(0, m.start() - window)
        end = min(len(text), m.end() + window)
        out.append(display_context(text[start:end]))
        if len(out) >= max_samples:
            break
    return out

def write_docx_report(out_path: Path, in_dir: Path, rows: list[tuple], check_dot: bool) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Отчёт текстовой гигиены DOCX", level=1)
    doc.add_paragraph(f"Входная папка: {in_dir}")
    doc.add_paragraph(f"Файлов проверено: {len(rows)}")
    doc.add_paragraph(f"Проверять пробел после точки: {'ДА' if check_dot else 'НЕТ'}")

    doc.add_heading("Сводка", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [
        "Файл",
        HYGIENE_LABELS_RU["double_space"],
        HYGIENE_LABELS_RU["space_before_punct"],
        HYGIENE_LABELS_RU["missing_after_punct"],
        HYGIENE_LABELS_RU["missing_after_dot"],
        HYGIENE_LABELS_RU["soft_hyphen"],
    ]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_item in rows:
        fn, agg, _samples = row_item[:3]
        cells = table.add_row().cells
        cells[0].text = fn
        cells[1].text = str(agg["double_space"])
        cells[2].text = str(agg["space_before_punct"])
        cells[3].text = str(agg["missing_after_punct"])
        cells[4].text = str(agg["missing_after_dot"])
        cells[5].text = str(agg["soft_hyphen"])

    doc.add_heading("Примеры", level=2)
    for row_item in rows:
        fn, agg, samples = row_item[:3]
        if sum(agg.values()) == 0:
            continue
        doc.add_heading(fn, level=3)
        for key, title in [
            ("double_space", HYGIENE_LABELS_RU["double_space"]),
            ("space_before_punct", HYGIENE_LABELS_RU["space_before_punct"]),
            ("missing_after_punct", HYGIENE_LABELS_RU["missing_after_punct"]),
            ("missing_after_dot", HYGIENE_LABELS_RU["missing_after_dot"]),
            ("soft_hyphen", HYGIENE_LABELS_RU["soft_hyphen"]),
        ]:
            if key == "missing_after_dot" and not check_dot:
                continue
            if not samples[key]:
                continue
            doc.add_paragraph(title)
            for s in samples[key][:5]:
                doc.add_paragraph(truncate(s, 180), style=None)
    doc.save(str(out_path))

def hygiene_json_payload(in_dir: Path, rows: list[tuple], check_dot: bool) -> dict[str, object]:
    summary = {
        "total_files": len(rows),
        "pass_files": 0,
        "fail_files": 0,
        "check_dot": bool(check_dot),
        "double_space": 0,
        "space_before_punct": 0,
        "missing_after_punct": 0,
        "missing_after_dot": 0,
        "soft_hyphen": 0,
    }
    files = []
    for row in rows:
        rel_name, agg, _samples = row[:3]
        findings = row[3] if len(row) > 3 else []
        metrics = {key: int(agg.get(key, 0)) for key in HYGIENE_LABELS_RU}
        status = "FAIL" if sum(metrics.values()) else "PASS"
        summary["pass_files" if status == "PASS" else "fail_files"] += 1
        for key, value in metrics.items():
            summary[key] += value
        files.append(
            {
                "path": rel_name,
                "status": status,
                "metrics": metrics,
                "findings": findings,
            }
        )
    return {
        "tool": "docx_text_hygiene_scan",
        "version": 1,
        "input_dir": str(in_dir),
        "files": files,
        "summary": summary,
    }

def main() -> int:
    ap = argparse.ArgumentParser(description="Scan DOCX files for basic text hygiene issues.")
    ap.add_argument("--input", default="input", help="Input folder with .docx files (default: input, recursive)")
    ap.add_argument("--out", default="report/docx_text_hygiene_report.md", help="Output Markdown report path")
    ap.add_argument("--check-dot", action="store_true", help="Also check missing space after '.' (more false positives)")
    ap.add_argument("--docx-out", default="", help="Optional DOCX report path")
    ap.add_argument("--no-docx-report", action="store_true", help="Do not write a DOCX report")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    args = ap.parse_args()

    in_dir = Path(args.input).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)

    if not in_dir.exists():
        print(f"[ERROR] Input folder does not exist: {in_dir}")
        return 2

    docx_files = find_docx_files(in_dir)
    if not docx_files:
        print(f"[WARN] No .docx files found in: {in_dir}")
        out_path.write_text("# Отчёт текстовой гигиены DOCX\n\nФайлы DOCX не найдены.\n", encoding="utf-8")
        if not args.no_docx_report:
            docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
            write_docx_report(docx_out, in_dir, [], args.check_dot)
            print(f"[OK] Wrote DOCX report: {docx_out}")
        if args.json_out:
            json_out = Path(args.json_out).resolve()
            write_json_file(json_out, hygiene_json_payload(in_dir, [], args.check_dot))
            print(f"[OK] Wrote JSON report: {json_out}")
        print(f"[OK] Wrote report: {out_path}")
        return 0

    lines = []
    lines.append("# Отчёт текстовой гигиены DOCX\n")
    lines.append(f"- Входная папка: `{in_dir}`")
    lines.append(f"- Файлов проверено: **{len(docx_files)}**")
    lines.append(f"- Проверять пробел после точки: **{'ДА' if args.check_dot else 'НЕТ'}**\n")

    lines.append("## Сводка\n")
    lines.append("| Файл | Двойные пробелы | Пробел перед пунктуацией | Нет пробела после пунктуации | Нет пробела после точки | Мягкие переносы |")
    lines.append("|---|---:|---:|---:|---:|---:|")

    per_file_details = []

    for p in docx_files:
        z = read_zip_map(p)
        soft_elements, soft_element_samples, soft_element_findings = count_soft_hyphen_elements(z)
        agg = {"double_space": 0, "space_before_punct": 0, "missing_after_punct": 0, "missing_after_dot": 0, "soft_hyphen": soft_elements}

        # Build small samples per file
        samples = {
            "double_space": [],
            "space_before_punct": [],
            "missing_after_punct": [],
            "missing_after_dot": [],
            "soft_hyphen": soft_element_samples,
        }
        findings: list[dict[str, object]] = [*soft_element_findings]

        for part, text_node_index, txt in iter_text_nodes_with_index(z):
            c = scan_text(txt, args.check_dot)
            for k in agg:
                agg[k] += c[k]
            findings.extend(find_text_issues(part, text_node_index, txt, args.check_dot))

            if len(samples["double_space"]) < 5:
                samples["double_space"].extend(sample_context(txt, PAT_DOUBLE_SPACE, max_samples=5-len(samples["double_space"])))
            if len(samples["space_before_punct"]) < 5:
                samples["space_before_punct"].extend(sample_context(txt, PAT_SPACE_BEFORE_PUNCT, max_samples=5-len(samples["space_before_punct"])))
            if len(samples["missing_after_punct"]) < 5:
                samples["missing_after_punct"].extend(sample_context(txt, PAT_MISSING_AFTER_PUNCT, max_samples=5-len(samples["missing_after_punct"])))
            if args.check_dot and len(samples["missing_after_dot"]) < 5:
                samples["missing_after_dot"].extend(sample_missing_after_dot(txt, max_samples=5-len(samples["missing_after_dot"])))
            if len(samples["soft_hyphen"]) < 5:
                samples["soft_hyphen"].extend(sample_context(txt, PAT_SOFT_HYPHEN, max_samples=5-len(samples["soft_hyphen"])))

        rel_name = rel_posix(p, in_dir)
        lines.append(f"| `{md_escape(rel_name)}` | {agg['double_space']} | {agg['space_before_punct']} | {agg['missing_after_punct']} | {agg['missing_after_dot']} | {agg['soft_hyphen']} |")
        per_file_details.append((rel_name, agg, samples, findings))

    lines.append("")
    lines.append("## Примеры по файлам\n")

    for row_item in per_file_details:
        fn, agg, samples = row_item[:3]
        if sum(agg.values()) == 0:
            continue
        lines.append(f"### {fn}\n")
        for key, title in [
            ("double_space", HYGIENE_LABELS_RU["double_space"]),
            ("space_before_punct", HYGIENE_LABELS_RU["space_before_punct"]),
            ("missing_after_punct", HYGIENE_LABELS_RU["missing_after_punct"]),
            ("missing_after_dot", HYGIENE_LABELS_RU["missing_after_dot"]),
            ("soft_hyphen", HYGIENE_LABELS_RU["soft_hyphen"]),
        ]:
            if key == "missing_after_dot" and not args.check_dot:
                continue
            if not samples[key]:
                continue
            lines.append(f"**{title}**")
            lines.append("")
            for s in samples[key][:5]:
                lines.append(f"- `{truncate(s, 180)}`")
            lines.append("")
        lines.append("")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    if not args.no_docx_report:
        docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
        write_docx_report(docx_out, in_dir, per_file_details, args.check_dot)
        print(f"[OK] Wrote DOCX report: {docx_out}")
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(json_out, hygiene_json_payload(in_dir, per_file_details, args.check_dot))
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Wrote report: {out_path}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
