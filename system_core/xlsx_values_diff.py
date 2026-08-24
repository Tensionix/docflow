#!/usr/bin/env python3
"""
XLSX Values Diff (data_only)

Compares two XLSX files by computed cell values (no formulas comparison).
- Uses openpyxl with data_only=True
- Compares sheets and cell values within used bounds

Outputs Markdown and DOCX reports with:
- sheets only in A / only in B
- list of changed cells (sheet, address, A value, B value)

Usage:
  python xlsx_values_diff.py --a A.xlsx --b B.xlsx --out report/xlsx_values_diff_report.md
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
from pathlib import Path
from typing import Any, List, Tuple

from docx import Document
from openpyxl import load_workbook
from _office_common import safe_mkdir, md_escape, truncate, write_json_file

def norm_cell(v: Any) -> Any:
    if isinstance(v, str):
        return v.replace("\r\n", "\n").replace("\r", "\n").strip()
    return v


def json_cell_value(v: Any) -> Any:
    if v is None or isinstance(v, (str, int, float, bool)):
        return v
    return str(v)


def write_docx_report(lines: list[str], out_path: Path) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].replace("**", "").replace("`", ""), style="List Bullet")
        elif line.startswith("|") or line.startswith("_") or not line.strip():
            continue
        else:
            doc.add_paragraph(line.replace("**", "").replace("`", ""))
    doc.save(str(out_path))

def main() -> int:
    ap = argparse.ArgumentParser(description="Compare two XLSX files by values only (data_only=True).")
    ap.add_argument("--a", required=True, help="XLSX file A")
    ap.add_argument("--b", required=True, help="XLSX file B")
    ap.add_argument("--out", default="report/xlsx_values_diff_report.md", help="Output Markdown report path")
    ap.add_argument("--docx-out", default="", help="Optional DOCX report path")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    ap.add_argument("--max-cells", type=int, default=20000, help="Max diff rows written (default: 20000)")
    args = ap.parse_args()

    a_path = Path(args.a).resolve()
    b_path = Path(args.b).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_path.parent)

    if not a_path.exists() or not b_path.exists():
        print("[ERROR] One of the XLSX files does not exist.")
        print(f"A: {a_path}")
        print(f"B: {b_path}")
        return 2

    wb_a = load_workbook(str(a_path), data_only=True, read_only=True)
    wb_b = load_workbook(str(b_path), data_only=True, read_only=True)

    sheets_a = set(wb_a.sheetnames)
    sheets_b = set(wb_b.sheetnames)
    only_a = sorted(sheets_a - sheets_b)
    only_b = sorted(sheets_b - sheets_a)
    common = [s for s in wb_a.sheetnames if s in sheets_b]

    diffs: List[Tuple[str, str, Any, Any]] = []

    for s in common:
        ws_a = wb_a[s]
        ws_b = wb_b[s]
        max_row = max(ws_a.max_row, ws_b.max_row)
        max_col = max(ws_a.max_column, ws_b.max_column)

        for r in range(1, max_row + 1):
            for c in range(1, max_col + 1):
                va = norm_cell(ws_a.cell(row=r, column=c).value)
                vb = norm_cell(ws_b.cell(row=r, column=c).value)
                if va != vb:
                    addr = ws_a.cell(row=r, column=c).coordinate
                    diffs.append((s, addr, va, vb))

    lines: List[str] = []
    lines.append("# Отчёт сравнения значений XLSX\n")
    lines.append(f"- A: `{a_path}`")
    lines.append(f"- B: `{b_path}`\n")

    lines.append("## Различия листов\n")
    lines.append(f"- Листы только в A: {', '.join(only_a) if only_a else '(нет)'}")
    lines.append(f"- Листы только в B: {', '.join(only_b) if only_b else '(нет)'}\n")

    lines.append("## Различия ячеек по значениям\n")
    if not diffs:
        lines.append("_Различия значений не найдены._\n")
    else:
        lines.append(f"- Ячеек с отличиями: **{len(diffs)}**\n")
        lines.append("| Лист | Ячейка | Значение A | Значение B |")
        lines.append("|---|---:|---|---|")
        for (sheet, addr, va, vb) in diffs[: args.max_cells]:
            lines.append(f"| {md_escape(sheet)} | {addr} | {truncate(md_escape(va), 200)} | {truncate(md_escape(vb), 200)} |")
        if len(diffs) > args.max_cells:
            lines.append(f"\n_... показаны первые {args.max_cells} ячеек._\n")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    docx_path = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
    write_docx_report(lines, docx_path)
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        status = "FAIL" if only_a or only_b or diffs else "PASS"
        write_json_file(
            json_out,
            {
                "tool": "xlsx_values_diff",
                "version": 1,
                "input_dir": "",
                "inputs": {"a": str(a_path), "b": str(b_path)},
                "files": [
                    {
                        "path": a_path.name,
                        "status": status,
                        "metrics": {"sheets": len(sheets_a)},
                    },
                    {
                        "path": b_path.name,
                        "status": status,
                        "metrics": {"sheets": len(sheets_b)},
                    },
                ],
                "summary": {
                    "status": status,
                    "sheets_only_a": len(only_a),
                    "sheets_only_b": len(only_b),
                    "changed_cells": len(diffs),
                },
                "sheets_only_a": only_a,
                "sheets_only_b": only_b,
                "diffs": [
                    {
                        "sheet": sheet,
                        "cell": addr,
                        "a": json_cell_value(va),
                        "b": json_cell_value(vb),
                    }
                    for sheet, addr, va, vb in diffs[: args.max_cells]
                ],
            },
        )
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Wrote DOCX report: {docx_path}")
    print(f"[OK] Wrote: {out_path}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
