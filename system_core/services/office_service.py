from __future__ import annotations

from pathlib import Path
from typing import Any
import hashlib
import json
import os
import shutil
import sys

from docx import Document

from system_core.core.jobs import JobContext, iter_subprocess_lines, popen_gui_command, unbuffer_python_command


SCRIPT_RUNNER = (
    "import runpy, sys; "
    "from pathlib import Path; "
    "target = Path(sys.argv[1]).resolve(); "
    "sys.path.insert(0, str(target.parent)); "
    "sys.argv = [str(target), *sys.argv[2:]]; "
    "runpy.run_path(str(target), run_name='__main__')"
)


def _inventory(folder: Path) -> dict[str, int]:
    if folder.is_file():
        return {"files": 1, "dirs": 0}
    files = 0
    dirs = 0
    if folder.exists():
        for item in folder.rglob("*"):
            if item.name == ".gitkeep":
                continue
            if item.is_file():
                files += 1
            elif item.is_dir():
                dirs += 1
    return {"files": files, "dirs": dirs}


def _is_inside(child: Path, parent: Path) -> bool:
    try:
        child_resolved = str(child.resolve())
        parent_resolved = str(parent.resolve())
        return os.path.commonpath([child_resolved, parent_resolved]) == parent_resolved
    except (OSError, ValueError):
        return False


def _clean_managed_folder(context: JobContext, folder: Path, label: str) -> dict[str, object]:
    root = context.paths.root.resolve()
    if folder.is_symlink():
        raise RuntimeError(f"{label} является ссылкой. Очистка остановлена.")
    folder.mkdir(parents=True, exist_ok=True)
    folder_resolved = folder.resolve()
    if not _is_inside(folder_resolved, root):
        raise RuntimeError(f"{label} находится вне корня проекта. Очистка остановлена.")

    removed = 0
    skipped: list[str] = []
    for item in folder.iterdir():
        if item.name == ".gitkeep":
            continue
        try:
            if item.is_symlink():
                item.unlink()
            elif item.is_dir():
                if not _is_inside(item, folder_resolved):
                    skipped.append(f"{item.name} (выходит за пределы {label})")
                    continue
                shutil.rmtree(item)
            else:
                item.unlink()
            removed += 1
            context.log(f"Удалено из {label}: {item.name}")
        except OSError as exc:
            skipped.append(f"{item.name} ({exc})")
    return {"folder": label, "removed_items": removed, "skipped_items": skipped}


def _python_command(context: JobContext, script_name: str, *args: str | Path) -> list[str]:
    script = context.paths.system_core / script_name
    if not script.exists():
        raise FileNotFoundError(f"Не найден скрипт: {script}")
    return [sys.executable, "-u", "-c", SCRIPT_RUNNER, str(script), *[str(arg) for arg in args]]


def _main_command(_context: JobContext, *args: str | Path) -> list[str]:
    return [sys.executable, "-u", "-m", "system_core.main", *[str(arg) for arg in args]]


def _run_command(
    context: JobContext,
    command: list[str],
    *,
    allow_exit_codes: set[int] | None = None,
) -> dict[str, object]:
    allow_exit_codes = allow_exit_codes or {0}
    command = unbuffer_python_command(command)
    context.log("> " + " ".join(f'"{part}"' if " " in part else part for part in command))
    process = popen_gui_command(command, cwd=context.paths.root)
    try:
        for line in iter_subprocess_lines(process, command):
            context.log(line)
            if context.cancelled():
                process.terminate()
                break
    finally:
        if process.stdout:
            process.stdout.close()
    return_code = process.wait()
    context.progress(1.0)
    if context.cancelled() and return_code not in allow_exit_codes:
        raise RuntimeError("Команда отменена")
    if return_code not in allow_exit_codes:
        raise RuntimeError(f"Команда завершилась с кодом {return_code}")
    return {"exit_code": return_code}


def _bool_param(context: JobContext, name: str, default: bool = False) -> bool:
    value = context.operation.parameters.get(name, default)
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "y", "on", "да"}


def _str_param(context: JobContext, name: str, default: str = "") -> str:
    value = context.operation.parameters.get(name, default)
    return str(value).strip() if value is not None else default


def _optional_str_param(context: JobContext, name: str) -> str | None:
    value = context.operation.parameters.get(name)
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _extend_if_value(args: list[str | Path], option: str, value: object | None) -> None:
    if value is None:
        return
    text = str(value).strip()
    if text:
        args.extend([option, text])


def _table_preheader_args(context: JobContext) -> list[str]:
    return ["--preheader-mode", _str_param(context, "preheader_mode", "separate")]


def _table_page_args(context: JobContext, default_orientation: str = "landscape") -> list[str | Path]:
    args: list[str | Path] = [
        "--page-size",
        _str_param(context, "page_size", "A4").upper(),
        "--page-orientation",
        _str_param(context, "page_orientation", default_orientation),
    ]
    for name, option in (
        ("margin_top_mm", "--margin-top-mm"),
        ("margin_right_mm", "--margin-right-mm"),
        ("margin_bottom_mm", "--margin-bottom-mm"),
        ("margin_left_mm", "--margin-left-mm"),
    ):
        _extend_if_value(args, option, _optional_str_param(context, name))
    return args


def _int_param(context: JobContext, name: str, default: int) -> int:
    value = context.operation.parameters.get(name, default)
    try:
        return int(float(str(value).strip()))
    except (TypeError, ValueError):
        return default


def _read_json_payload(path: Path | str | None) -> dict[str, Any] | None:
    if not path:
        return None
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def _write_json_payload(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8", newline="\n")


def _report_path(context: JobContext, path: object) -> str:
    text = str(path or "")
    if not text:
        return ""
    candidate = Path(text)
    try:
        return candidate.resolve().relative_to(context.paths.root.resolve()).as_posix()
    except (OSError, ValueError):
        return text


def _md_cell(value: object) -> str:
    return str(value or "").replace("|", "\\|").replace("\n", " ")


def _deep_short(value: object, limit: int = 160) -> str:
    text = " ".join(str(value or "").split())
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def _stable_hash(value: object, length: int = 12) -> str:
    if isinstance(value, str):
        text = value
    else:
        text = json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
    return hashlib.sha1(text.encode("utf-8", errors="replace")).hexdigest()[:length]


def _deep_owner_registry() -> list[dict[str, object]]:
    return [
        {
            "owner": "docx_text_hygiene",
            "target_type": "text-range",
            "classes": ["double_space", "space_before_punct", "missing_after_punct", "generic_dot_spacing", "soft_hyphen"],
            "note": "Механическая гигиена видимого текста. Общее правило пробела после точки пропускает сокращения и единицы, которые принадлежат правилам аудита или морфологии.",
        },
        {
            "owner": "docx_audit_processor",
            "target_type": "text-range",
            "classes": ["unit_sqm", "unit_cubic", "unit_space", "percent", "degree", "number_sign", "date_suffix"],
            "note": "Детерминированные предметные правила; в глубокой гигиене работает как проход 1B.",
        },
        {
            "owner": "docx_anomaly_inspector",
            "target_type": "document-object",
            "classes": ["tables", "captions", "layout", "headers", "fields", "lists", "spacing"],
            "note": "Поиск форматных и структурных аномалий без изменения документов.",
        },
        {
            "owner": "docx_anomaly_corrector",
            "target_type": "document-object",
            "classes": ["empty_paragraphs", "row_heights", "nowrap", "table_borders", "cell_margins"],
            "note": "Безопасная коррекция форматных аномалий в копиях DOCX.",
        },
        {
            "owner": "future_morphology_address_layer",
            "target_type": "text-range",
            "classes": ["address_city", "address_street", "address_house", "settlements"],
            "note": "Будущий слой для г., ул., д., корп. и топонимов с контекстом и статусом review_required.",
        },
    ]


def _deep_domain_boundaries() -> list[dict[str, object]]:
    return [
        {
            "domain": "text_hygiene",
            "title": "Гигиена текста",
            "primary_module": "docx_text_hygiene_*",
            "owns": [
                "двойные пробелы",
                "пробелы перед пунктуацией",
                "пробел после , ; : ! ?",
                "generic-пробел после точки, если это не сокращение/единица",
                "мягкие переносы",
            ],
            "does_not_own": [
                "единицы измерения и предметные сокращения",
                "адресные обозначения и топонимы",
                "поля, секции, таблицы, колонтитулы, подписи",
                "назначение Word-стилей",
            ],
            "scan_mode": "считает и показывает механический текстовый мусор",
            "fix_mode": "исправляет только видимый текст в копии DOCX",
            "autofix_policy": "safe, кроме явно включаемого удаления зачёркнутого текста",
        },
        {
            "domain": "audit_rules",
            "title": "Правила аудита",
            "primary_module": "docx_audit_processor",
            "owns": [
                "кв. м / куб. м",
                "пробел между числом и единицей",
                "№",
                "%",
                "градусы",
                "суффикс г. после даты",
                "типовые опечатки в подписях таблиц",
            ],
            "does_not_own": [
                "общую пунктуационную гигиену",
                "адреса и топонимы без морфологического слоя",
                "форматирование таблиц и секций",
                "стили документа",
            ],
            "scan_mode": "находит предметные текстовые правила и даёт точные диапазоны",
            "fix_mode": "применяет только детерминированные безопасные правки; неоднозначное остаётся review_required",
            "autofix_policy": "безопасно для строгих правил, review_required для падежа/смысла",
        },
        {
            "domain": "document_anomalies",
            "title": "Аномалии документа",
            "primary_module": "docx_anomaly_inspector / docx_anomaly_corrector",
            "owns": [
                "секции, поля и ориентация",
                "таблицы как объекты: ширина, обрезка, границы, поля ячеек, запрет переноса текста, высоты строк",
                "подписи таблиц/рисунков и близость к объекту",
                "колонтитулы и нумерация страниц",
                "поля Word, оглавление, REF/PAGE/PAGEREF, ссылки и закладки",
                "списки, пустоты, разрывы, висячие заголовки",
            ],
            "does_not_own": [
                "обычные пробелы и пунктуацию внутри текста",
                "единицы измерения как текстовую норму",
                "широкую намеренную унификацию всех таблиц",
                "назначение Word-стилей",
            ],
            "scan_mode": "строит отчёт по структуре и форматированию без изменения документов",
            "fix_mode": "исправляет только узкий набор безопасных форматных аномалий в копии DOCX",
            "autofix_policy": "безопасно только для пустых абзацев, точной высоты строк, запрета переноса текста и выбранной унификации границ/полей",
        },
        {
            "domain": "document_styles",
            "title": "Стили документа",
            "primary_module": "docx_restyle_by_template",
            "owns": [
                "перенос стилевой базы эталона: styles.xml, нумерация, тема, шрифты",
                "разметку заголовков, подписей, перечней и текста по карте заголовков",
                "сквозную нумерацию подписей и обновление ссылок на них",
                "отсев неиспользуемых определений стилей",
                "чистку XML: атрибуты ревизий, дробленые раны, пустые свойства",
            ],
            "does_not_own": [
                "переписывание текста",
                "исправление пунктуации и единиц",
                "геометрию таблиц, секций и колонтитулов",
                "составление самой карты заголовков - её пишет человек по разведке",
            ],
            "scan_mode": "разведка: плоский дамп документа, стили эталона по частоте, сводка по мусору и структуре",
            "fix_mode": "переносит стилевую базу эталона и размечает документ по карте заголовков",
            "autofix_policy": "без карты заголовков разметка не выполняется: переносится база и идёт чистка",
        },
        {
            "domain": "table_unifier",
            "title": "Унификация таблиц",
            "primary_module": "document_tables_unifier",
            "owns": [
                "намеренную унификацию таблиц по команде пользователя",
                "шрифты таблиц по плотности",
                "границы, поля ячеек, баланс ширин",
                "вжимание таблиц в поля/ориентацию",
            ],
            "does_not_own": [
                "поиск случайных текстовых ошибок",
                "глобальные стили документа",
                "подписи и нумерацию как логические аномалии",
            ],
            "scan_mode": "анализирует таблицы для планируемой унификации",
            "fix_mode": "массово нормализует таблицы только по явной команде",
            "autofix_policy": "intentional command: это не скрытая гигиена, а отдельная операция",
        },
        {
            "domain": "future_morphology",
            "title": "Будущая морфология",
            "primary_module": "future_morphology_address_layer",
            "owns": [
                "г., ул., д., корп. и похожие адресные обозначения",
                "населённые пункты и топонимы",
                "контекстные единицы, где нужен падеж или смысл",
            ],
            "does_not_own": [
                "простые механические пробелы",
                "детерминированные правила аудита, уже закрытые строгими регулярными выражениями",
                "форматирование документа",
            ],
            "scan_mode": "будет искать контекстные текстовые случаи с морфологией",
            "fix_mode": "по умолчанию review_required для неоднозначных замен",
            "autofix_policy": "review_required, пока нет закрытого словаря и уверенного контекста",
        },
    ]


def _deep_overlap_notes() -> list[str]:
    return [
        "Общее правило пробела после точки не исправляет сокращения и единицы: их должен исправлять слой правил аудита или морфологический слой.",
        "Текст в таблицах принадлежит текстовой гигиене; ширины, границы, поля, запрет переноса текста и высоты строк принадлежат аномалиям/табличным инструментам.",
        "Стили владеют назначением стилей Word; аномалии владеют логикой подписи, номера, близости объекта и локации.",
        "Удаление зачёркнутого текста считается изменением содержания и включается только явным параметром.",
    ]


def _deep_pass_summary(pass_id: str, payload: dict[str, Any] | None) -> dict[str, object]:
    if not payload:
        return {}
    summary = payload.get("summary", {})
    if not isinstance(summary, dict):
        return {}
    if pass_id == "text_hygiene_scan":
        issue_keys = ["double_space", "space_before_punct", "missing_after_punct", "missing_after_dot", "soft_hyphen"]
        return {
            "files": summary.get("total_files", 0),
            "findings": sum(int(summary.get(key, 0) or 0) for key in issue_keys),
            "failed_files": summary.get("fail_files", 0),
            "metrics": {key: summary.get(key, 0) for key in issue_keys},
        }
    if pass_id == "text_hygiene_fix":
        return {
            "files": summary.get("files", 0),
            "changed_nodes": summary.get("changed_nodes", 0),
            "changed_chars": summary.get("changed_chars", 0),
            "removed_temp_office_files": summary.get("removed_temp_office_files", 0),
            "errors": summary.get("errors", 0),
        }
    if pass_id in {"audit_scan", "audit_fix"}:
        return {
            "files": summary.get("total_files", 0),
            "findings": summary.get("issues", 0),
            "fixes": summary.get("fixes", 0),
            "anchors": summary.get("anchors", 0),
            "changed_nodes": summary.get("changed_nodes", 0),
            "failed_files": summary.get("fail_files", 0),
        }
    if pass_id == "anomaly_scan":
        return {
            "files": summary.get("files", 0),
            "findings": summary.get("findings", 0),
            "severity": summary.get("severity", {}),
            "classes": summary.get("classes", {}),
        }
    if pass_id == "anomaly_correct":
        change_keys = ["removed_empty_paragraphs", "fixed_exact_row_heights", "removed_nowrap", "normalized_table_borders", "normalized_cell_margins"]
        return {
            "files": summary.get("files", 0),
            "changed_files": summary.get("changed_files", 0),
            "changes": sum(int(summary.get(key, 0) or 0) for key in change_keys),
            "metrics": {key: summary.get(key, 0) for key in change_keys},
            "errors": summary.get("errors", 0),
        }
    return dict(summary)


def _deep_pass_record(
    context: JobContext,
    *,
    pass_id: str,
    title: str,
    owner: str,
    mode: str,
    report: Path | None = None,
    docx_report: Path | None = None,
    json_report: Path | None = None,
    outdir: Path | None = None,
    status: str = "ok",
    enabled: bool = True,
) -> dict[str, object]:
    payload = _read_json_payload(json_report)
    return {
        "id": pass_id,
        "title": title,
        "owner": owner,
        "mode": mode,
        "enabled": enabled,
        "status": status,
        "report": _report_path(context, report),
        "docx_report": _report_path(context, docx_report),
        "json_report": _report_path(context, json_report),
        "outdir": _report_path(context, outdir),
        "summary": _deep_pass_summary(pass_id, payload),
    }


_DEEP_TEXT_LABELS = {
    "double_space": "Двойные пробелы",
    "space_before_punct": "Пробел перед пунктуацией",
    "missing_after_punct": "Нет пробела после пунктуации",
    "missing_after_dot": "Нет пробела после точки",
    "soft_hyphen": "Мягкие переносы",
}

_DEEP_CORRECTION_LABELS = {
    "removed_empty_paragraphs": "Лишние пустые абзацы",
    "fixed_exact_row_heights": "Точная высота строк таблиц",
    "removed_nowrap": "Запрет переноса текста в ячейках",
    "normalized_table_borders": "Унификация границ таблиц",
    "normalized_cell_margins": "Унификация полей ячеек",
}


def _deep_json_report_path(context: JobContext, pass_record: dict[str, object]) -> Path | None:
    value = str(pass_record.get("json_report") or "").strip()
    if not value:
        return None
    path = Path(value)
    if path.is_absolute():
        return path
    return context.paths.root / path


def _deep_location_text(location: object) -> str:
    if isinstance(location, dict):
        return str(location.get("display") or location.get("detail") or location.get("search_text") or "")
    return str(location or "")


def _deep_entry(
    *,
    pass_id: str,
    source_module: str,
    owner: str,
    class_id: str,
    file: str,
    target_fingerprint: str,
    issue_fingerprint: str,
    target_precision: str,
    location: dict[str, object],
    evidence: str,
    proposed_fix: str,
    risk: str,
    autofix: str,
    count: int = 1,
) -> dict[str, object]:
    return {
        "pass_id": pass_id,
        "source_module": source_module,
        "owner": owner,
        "ownership_key": f"{owner}:{class_id}",
        "class_id": class_id,
        "file": file,
        "target_fingerprint": target_fingerprint,
        "issue_fingerprint": issue_fingerprint,
        "target_precision": target_precision,
        "location": location,
        "location_display": _deep_location_text(location),
        "evidence": _deep_short(evidence, 260),
        "proposed_fix": _deep_short(proposed_fix, 260),
        "risk": risk,
        "autofix": autofix,
        "count": count,
        "status": "active",
        "duplicate_of": "",
        "overlap_group": "",
    }


def _deep_text_scan_entries(pass_record: dict[str, object], payload: dict[str, Any]) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    pass_id = str(pass_record.get("id") or "")
    for file_item in payload.get("files", []):
        if not isinstance(file_item, dict):
            continue
        file_name = str(file_item.get("path") or "")
        findings = file_item.get("findings")
        if isinstance(findings, list) and findings:
            for finding in findings:
                if isinstance(finding, dict):
                    entries.append(_deep_text_range_entry(pass_id, "docx_text_hygiene_scan", file_name, finding, default_autofix="available"))
            continue
        metrics = file_item.get("metrics")
        if not isinstance(metrics, dict):
            continue
        for key, value in metrics.items():
            count = int(value or 0)
            if count <= 0 or key not in _DEEP_TEXT_LABELS:
                continue
            target = f"file:{file_name}"
            issue = f"issue:docx_text_hygiene:{key}:{target}"
            entries.append(
                _deep_entry(
                    pass_id=pass_id,
                    source_module="docx_text_hygiene_scan",
                    owner="docx_text_hygiene",
                    class_id=str(key),
                    file=file_name,
                    target_fingerprint=target,
                    issue_fingerprint=issue,
                    target_precision="file",
                    location={"file": file_name, "object_type": "file", "display": f"файл: {file_name}"},
                    evidence=f"{_DEEP_TEXT_LABELS[str(key)]}: {count}",
                    proposed_fix="Исправить механической текстовой гигиеной.",
                    risk="safe",
                    autofix="safe" if pass_id.endswith("_fix") else "available",
                    count=count,
                )
            )
    return entries


def _deep_text_range_entry(
    pass_id: str,
    source_module: str,
    file_name: str,
    finding: dict[str, object],
    *,
    default_autofix: str,
) -> dict[str, object]:
    class_id = str(finding.get("class_id") or "text_hygiene")
    part = str(finding.get("part") or "")
    node_index = int(finding.get("text_node_index") or 0)
    start = int(finding.get("start") or 0)
    end = int(finding.get("end") or start)
    target_precision = str(finding.get("target_precision") or "text-range")
    element_index = int(finding.get("element_index") or 0)
    before = str(finding.get("before") or "")
    after = str(finding.get("after") or "")
    context_text = str(finding.get("context") or "")
    if target_precision == "xml-element":
        target = f"xml:{file_name}:{part}:e{element_index}:{_stable_hash(before or context_text)}"
        detail = f"{part}; XML element {element_index}"
    else:
        target = f"text:{file_name}:{part}:n{node_index}:{start}-{end}:{_stable_hash(before or context_text)}"
        detail = f"{part}; text node {node_index}; chars {start}-{end}"
    issue = f"issue:docx_text_hygiene:{class_id}:{target}"
    return _deep_entry(
        pass_id=pass_id,
        source_module=source_module,
        owner="docx_text_hygiene",
        class_id=class_id,
        file=file_name,
        target_fingerprint=target,
        issue_fingerprint=issue,
        target_precision=target_precision,
        location={
            "file": file_name,
            "object_type": "text" if target_precision == "text-range" else "xml",
            "detail": detail,
            "search_text": before or context_text,
            "display": f"{file_name}; {detail}",
        },
        evidence=f"{finding.get('label') or class_id}: {before}",
        proposed_fix=after,
        risk=str(finding.get("risk") or "safe"),
        autofix=str(finding.get("autofix") or default_autofix),
    )


def _deep_text_fix_entries(pass_record: dict[str, object], payload: dict[str, Any]) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    pass_id = str(pass_record.get("id") or "")
    for file_item in payload.get("files", []):
        if not isinstance(file_item, dict):
            continue
        file_name = str(file_item.get("file") or file_item.get("path") or "")
        findings = file_item.get("findings")
        if isinstance(findings, list) and findings:
            for finding in findings:
                if isinstance(finding, dict):
                    entries.append(_deep_text_range_entry(pass_id, "docx_text_hygiene_fix", file_name, finding, default_autofix="applied"))
            continue
        changed_nodes = int(file_item.get("changed_nodes") or 0)
        changed_chars = int(file_item.get("changed_chars") or 0)
        if changed_nodes <= 0 and changed_chars <= 0:
            continue
        target = f"file:{file_name}"
        issue = f"issue:docx_text_hygiene:text_hygiene_changed:{target}"
        entries.append(
            _deep_entry(
                pass_id=pass_id,
                source_module="docx_text_hygiene_fix",
                owner="docx_text_hygiene",
                class_id="text_hygiene_changed",
                file=file_name,
                target_fingerprint=target,
                issue_fingerprint=issue,
                target_precision="file",
                location={"file": file_name, "object_type": "file", "display": f"файл: {file_name}"},
                evidence=f"Изменённых текстовых узлов: {changed_nodes}; примерное изменение символов: {changed_chars}",
                proposed_fix="Безопасные правки текстовой гигиены применены к копии DOCX.",
                risk="safe",
                autofix="applied",
                count=max(1, changed_nodes),
            )
        )
    return entries


def _deep_audit_entries(pass_record: dict[str, object], payload: dict[str, Any]) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    pass_id = str(pass_record.get("id") or "")
    for file_item in payload.get("files", []):
        if not isinstance(file_item, dict):
            continue
        file_name = str(file_item.get("path") or "")
        for finding in file_item.get("findings", []):
            if not isinstance(finding, dict):
                continue
            class_id = str(finding.get("rule_code") or "audit_rule")
            part = str(finding.get("part") or "")
            node_index = int(finding.get("text_node_index") or 0)
            start = int(finding.get("start") or 0)
            end = int(finding.get("end") or start)
            before = str(finding.get("before") or "")
            after = str(finding.get("after") or "")
            context_text = str(finding.get("context") or "")
            target = f"text:{file_name}:{part}:n{node_index}:{start}-{end}:{_stable_hash(before or context_text)}"
            issue = f"issue:docx_audit_processor:{class_id}:{target}"
            action = str(finding.get("action") or "").upper()
            risk = "review_required" if action == "SUGGEST" else "safe"
            autofix = "applied" if action == "FIX" else ("review_required" if action == "SUGGEST" else "available")
            entries.append(
                _deep_entry(
                    pass_id=pass_id,
                    source_module="docx_audit_processor",
                    owner="docx_audit_processor",
                    class_id=class_id,
                    file=file_name,
                    target_fingerprint=target,
                    issue_fingerprint=issue,
                    target_precision="text-range",
                    location={
                        "file": file_name,
                        "object_type": "text",
                        "detail": f"{part}; text node {node_index}; chars {start}-{end}",
                        "search_text": before or context_text,
                        "display": f"{file_name}; {part}; text node {node_index}; chars {start}-{end}",
                    },
                    evidence=f"{finding.get('rule_title') or class_id}: {before}",
                    proposed_fix=str(after),
                    risk=risk,
                    autofix=autofix,
                )
            )
    return entries


def _deep_anomaly_scan_entries(pass_record: dict[str, object], payload: dict[str, Any]) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    pass_id = str(pass_record.get("id") or "")
    for file_item in payload.get("files", []):
        if not isinstance(file_item, dict):
            continue
        file_name = str(file_item.get("path") or "")
        for finding in file_item.get("findings", []):
            if not isinstance(finding, dict):
                continue
            location = finding.get("location") if isinstance(finding.get("location"), dict) else {}
            if not isinstance(location, dict):
                location = {}
            class_id = str(finding.get("code") or finding.get("class") or "anomaly")
            object_type = str(location.get("object_type") or finding.get("object") or "")
            object_number = str(location.get("object_number") or "")
            debug_path = str(location.get("debug_path") or "")
            detail = str(location.get("detail") or "")
            search_text = str(location.get("search_text") or "")
            location_key = debug_path or "|".join([object_type, object_number, detail, search_text])
            target = f"object:{file_name}:{object_type}:{object_number}:{_stable_hash(location_key)}"
            issue = f"issue:docx_anomaly_inspector:{class_id}:{target}"
            autofix = str(finding.get("autofix") or "none")
            risk = "safe" if autofix == "safe" else ("review_required" if autofix == "review_required" else "report_only")
            entries.append(
                _deep_entry(
                    pass_id=pass_id,
                    source_module="docx_anomaly_inspector",
                    owner="docx_anomaly_inspector",
                    class_id=class_id,
                    file=file_name,
                    target_fingerprint=target,
                    issue_fingerprint=issue,
                    target_precision="document-object",
                    location=location,
                    evidence=str(finding.get("evidence") or finding.get("why_it_matters") or ""),
                    proposed_fix=str(finding.get("suggested_fix") or ""),
                    risk=risk,
                    autofix=autofix,
                )
            )
    return entries


def _deep_anomaly_correct_entries(pass_record: dict[str, object], payload: dict[str, Any]) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    pass_id = str(pass_record.get("id") or "")
    for file_item in payload.get("files", []):
        if not isinstance(file_item, dict):
            continue
        file_name = str(file_item.get("file") or file_item.get("path") or "")
        for key, label in _DEEP_CORRECTION_LABELS.items():
            count = int(file_item.get(key) or 0)
            if count <= 0:
                continue
            target = f"file:{file_name}"
            issue = f"issue:docx_anomaly_corrector:{key}:{target}"
            entries.append(
                _deep_entry(
                    pass_id=pass_id,
                    source_module="docx_anomaly_corrector",
                    owner="docx_anomaly_corrector",
                    class_id=key,
                    file=file_name,
                    target_fingerprint=target,
                    issue_fingerprint=issue,
                    target_precision="file",
                    location={"file": file_name, "object_type": "file", "display": f"файл: {file_name}"},
                    evidence=f"{label}: {count}",
                    proposed_fix="Безопасная коррекция форматирования применена к копии DOCX.",
                    risk="safe",
                    autofix="applied",
                    count=count,
                )
            )
    return entries


def _deep_entries_for_pass(context: JobContext, pass_record: dict[str, object]) -> list[dict[str, object]]:
    json_path = _deep_json_report_path(context, pass_record)
    payload = _read_json_payload(json_path)
    if not payload:
        return []
    tool = str(payload.get("tool") or "")
    pass_id = str(pass_record.get("id") or "")
    if tool == "docx_text_hygiene_scan":
        return _deep_text_scan_entries(pass_record, payload)
    if tool == "docx_text_hygiene_fix":
        return _deep_text_fix_entries(pass_record, payload)
    if tool == "docx_audit_processor":
        return _deep_audit_entries(pass_record, payload)
    if tool == "docx_anomaly_inspector":
        return _deep_anomaly_scan_entries(pass_record, payload)
    if tool == "docx_anomaly_corrector":
        return _deep_anomaly_correct_entries(pass_record, payload)
    summary = payload.get("summary")
    if isinstance(summary, dict):
        file_name = str(payload.get("input") or pass_record.get("json_report") or "")
        target = f"payload:{tool or pass_id}:{_stable_hash(file_name)}"
        return [
            _deep_entry(
                pass_id=pass_id,
                source_module=tool or pass_id,
                owner=str(pass_record.get("owner") or tool or pass_id),
                class_id="summary",
                file=file_name,
                target_fingerprint=target,
                issue_fingerprint=f"issue:{tool or pass_id}:summary:{target}",
                target_precision="summary",
                location={"display": file_name},
                evidence=json.dumps(summary, ensure_ascii=False, sort_keys=True),
                proposed_fix="См. исходный отчёт прохода.",
                risk="report_only",
                autofix="none",
            )
        ]
    return []


def _deep_unified_findings(context: JobContext, passes: list[dict[str, object]]) -> tuple[list[dict[str, object]], dict[str, object]]:
    entries: list[dict[str, object]] = []
    for pass_record in passes:
        if isinstance(pass_record, dict):
            entries.extend(_deep_entries_for_pass(context, pass_record))

    seen_issues: dict[str, str] = {}
    for index, entry in enumerate(entries, 1):
        entry["id"] = f"U{index:05d}"
        issue = str(entry.get("issue_fingerprint") or "")
        if issue and issue in seen_issues:
            entry["status"] = "suppressed_duplicate"
            entry["duplicate_of"] = seen_issues[issue]
        elif issue:
            seen_issues[issue] = str(entry["id"])

    target_groups: dict[str, list[dict[str, object]]] = {}
    for entry in entries:
        if entry.get("status") != "active":
            continue
        if entry.get("target_precision") in {"file", "summary"}:
            continue
        target = str(entry.get("target_fingerprint") or "")
        if not target:
            continue
        target_groups.setdefault(target, []).append(entry)

    overlap_groups: list[dict[str, object]] = []
    for target, group_entries in target_groups.items():
        classes = sorted({str(item.get("ownership_key") or "") for item in group_entries})
        if len(group_entries) < 2 or len(classes) < 2:
            continue
        group_id = f"T{len(overlap_groups) + 1:04d}"
        ids = [str(item.get("id") or "") for item in group_entries]
        for item in group_entries:
            item["overlap_group"] = group_id
        overlap_groups.append(
            {
                "id": group_id,
                "target_fingerprint": target,
                "entries": ids,
                "classes": classes,
                "location": _deep_location_text(group_entries[0].get("location")),
            }
        )

    active = sum(1 for entry in entries if entry.get("status") == "active")
    suppressed = sum(1 for entry in entries if entry.get("status") == "suppressed_duplicate")
    coarse = sum(1 for entry in entries if entry.get("target_precision") in {"file", "summary"})
    by_owner: dict[str, int] = {}
    by_risk: dict[str, int] = {}
    for entry in entries:
        if entry.get("status") != "active":
            continue
        owner = str(entry.get("owner") or "")
        risk = str(entry.get("risk") or "")
        by_owner[owner] = by_owner.get(owner, 0) + 1
        by_risk[risk] = by_risk.get(risk, 0) + 1
    index = {
        "total": len(entries),
        "active": active,
        "suppressed_duplicates": suppressed,
        "coarse_entries": coarse,
        "overlap_groups": len(overlap_groups),
        "by_owner": by_owner,
        "by_risk": by_risk,
        "overlaps": overlap_groups,
    }
    return entries, index


def _deep_hygiene_payload(
    context: JobContext,
    mode: str,
    include_audit: bool,
    include_anomalies: bool,
    passes: list[dict[str, object]],
    outdir: Path | None,
) -> dict[str, object]:
    unified_findings, target_index = _deep_unified_findings(context, passes)
    return {
        "tool": "docx_deep_hygiene",
        "version": 2,
        "mode": mode,
        "input": str(context.paths.input),
        "outdir": str(outdir) if outdir is not None else "",
        "include_audit": include_audit,
        "include_anomalies": include_anomalies,
        "passes": passes,
        "domain_boundaries": _deep_domain_boundaries(),
        "target_index": target_index,
        "unified_findings": unified_findings,
        "ownership": _deep_owner_registry(),
        "overlap_notes": _deep_overlap_notes(),
    }


def _deep_hygiene_markdown(context: JobContext, payload: dict[str, object]) -> str:
    mode = str(payload.get("mode") or "scan")
    lines: list[str] = []
    lines.append("# Глубокая гигиена DOCX\n")
    lines.append(f"- Режим: **{'Корректировка' if mode == 'fix' else 'Проверка'}**")
    lines.append(f"- Вход: `{_md_cell(_report_path(context, payload.get('input')) or payload.get('input'))}`")
    if payload.get("outdir"):
        lines.append(f"- Выход: `{_md_cell(_report_path(context, payload.get('outdir')) or payload.get('outdir'))}`")
    lines.append(f"- Правила аудита подключены: **{'да' if payload.get('include_audit') else 'нет'}**")
    lines.append(f"- Аномалии документа подключены: **{'да' if payload.get('include_anomalies') else 'нет'}**")
    lines.append("")
    lines.append("## Проходы\n")
    lines.append("| Проход | Владелец | Режим | Статус | Сводка | Отчёты |")
    lines.append("|---|---|---|---|---|---|")
    for item in payload.get("passes", []):
        if not isinstance(item, dict):
            continue
        summary = item.get("summary") if isinstance(item.get("summary"), dict) else {}
        summary_text = ", ".join(f"{key}: {value}" for key, value in summary.items() if not isinstance(value, dict)) or "нет данных"
        report_bits = [item.get("report"), item.get("docx_report"), item.get("json_report"), item.get("outdir")]
        reports = "<br>".join(f"`{_md_cell(bit)}`" for bit in report_bits if bit)
        lines.append(
            f"| {_md_cell(item.get('title'))} | `{_md_cell(item.get('owner'))}` | {_md_cell(item.get('mode'))} | "
            f"{_md_cell(item.get('status'))} | {_md_cell(summary_text)} | {reports} |"
        )
    lines.append("")
    lines.append("## Границы зон ответственности\n")
    lines.append("| Зона | Модуль | Отвечает за | Не отвечает за | Корректировка |")
    lines.append("|---|---|---|---|---|")
    for item in payload.get("domain_boundaries", []):
        if not isinstance(item, dict):
            continue
        owns = "; ".join(str(value) for value in item.get("owns", []) if value)
        does_not_own = "; ".join(str(value) for value in item.get("does_not_own", []) if value)
        lines.append(
            f"| {_md_cell(item.get('title'))} | `{_md_cell(item.get('primary_module'))}` | "
            f"{_md_cell(owns)} | {_md_cell(does_not_own)} | {_md_cell(item.get('autofix_policy'))} |"
        )
    lines.append("")
    target_index = payload.get("target_index") if isinstance(payload.get("target_index"), dict) else {}
    lines.append("## Карта целей\n")
    lines.append(f"- Всего нормализованных записей: **{target_index.get('total', 0)}**")
    lines.append(f"- Активных записей: **{target_index.get('active', 0)}**")
    lines.append(f"- Подавленных точных дублей: **{target_index.get('suppressed_duplicates', 0)}**")
    lines.append(f"- Грубых записей уровня файла: **{target_index.get('coarse_entries', 0)}**")
    lines.append(f"- Групп пересечений по одной цели: **{target_index.get('overlap_groups', 0)}**")
    lines.append("")
    unified_findings = [item for item in payload.get("unified_findings", []) if isinstance(item, dict)]
    if unified_findings:
        lines.append("### Нормализованные находки\n")
        lines.append("| ID | Статус | Владелец | Класс | Риск | Локация | Доказательство | Правка |")
        lines.append("|---|---|---|---|---|---|---|---|")
        for item in unified_findings[:120]:
            lines.append(
                f"| `{_md_cell(item.get('id'))}` | {_md_cell(item.get('status'))} | `{_md_cell(item.get('owner'))}` | "
                f"`{_md_cell(item.get('class_id'))}` | {_md_cell(item.get('risk'))} | {_md_cell(item.get('location_display'))} | "
                f"{_md_cell(item.get('evidence'))} | {_md_cell(item.get('proposed_fix'))} |"
            )
        if len(unified_findings) > 120:
            lines.append(f"\n_Показаны первые 120 записей из {len(unified_findings)}. Полный список находится в JSON._")
        lines.append("")
    overlaps = target_index.get("overlaps") if isinstance(target_index.get("overlaps"), list) else []
    if overlaps:
        lines.append("### Пересечения\n")
        lines.append("| Группа | Записи | Классы | Локация |")
        lines.append("|---|---|---|---|")
        for item in overlaps[:80]:
            if not isinstance(item, dict):
                continue
            lines.append(
                f"| `{_md_cell(item.get('id'))}` | {_md_cell(', '.join(str(value) for value in item.get('entries', [])))} | "
                f"{_md_cell(', '.join(str(value) for value in item.get('classes', [])))} | {_md_cell(item.get('location'))} |"
            )
        lines.append("")
    lines.append("## Ownership\n")
    lines.append("| Владелец | Объект | Классы | Примечание |")
    lines.append("|---|---|---|---|")
    for item in payload.get("ownership", []):
        if not isinstance(item, dict):
            continue
        classes = ", ".join(str(value) for value in item.get("classes", []) if value)
        lines.append(f"| `{_md_cell(item.get('owner'))}` | {_md_cell(item.get('target_type'))} | {_md_cell(classes)} | {_md_cell(item.get('note'))} |")
    lines.append("")
    lines.append("## Пересечения\n")
    for note in payload.get("overlap_notes", []):
        lines.append(f"- {note}")
    lines.append("")
    return "\n".join(lines)


def _write_deep_hygiene_docx(path: Path, payload: dict[str, object]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    mode = str(payload.get("mode") or "scan")
    doc.add_heading("Глубокая гигиена DOCX", level=1)
    doc.add_paragraph(f"Режим: {'Корректировка' if mode == 'fix' else 'Проверка'}")
    doc.add_paragraph(f"Вход: {payload.get('input')}")
    if payload.get("outdir"):
        doc.add_paragraph(f"Выход: {payload.get('outdir')}")
    doc.add_paragraph(f"Правила аудита подключены: {'да' if payload.get('include_audit') else 'нет'}")
    doc.add_paragraph(f"Аномалии документа подключены: {'да' if payload.get('include_anomalies') else 'нет'}")
    doc.add_heading("Проходы", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for index, header in enumerate(["Проход", "Владелец", "Режим", "Статус", "Сводка"]):
        table.rows[0].cells[index].text = header
    for item in payload.get("passes", []):
        if not isinstance(item, dict):
            continue
        row = table.add_row().cells
        summary = item.get("summary") if isinstance(item.get("summary"), dict) else {}
        summary_text = "; ".join(f"{key}: {value}" for key, value in summary.items() if not isinstance(value, dict))
        for index, value in enumerate([item.get("title"), item.get("owner"), item.get("mode"), item.get("status"), summary_text]):
            row[index].text = str(value or "")
    doc.add_heading("Границы зон ответственности", level=2)
    boundary_table = doc.add_table(rows=1, cols=5)
    boundary_table.style = "Table Grid"
    for index, header in enumerate(["Зона", "Модуль", "Отвечает за", "Не отвечает за", "Корректировка"]):
        boundary_table.rows[0].cells[index].text = header
    for item in payload.get("domain_boundaries", []):
        if not isinstance(item, dict):
            continue
        row = boundary_table.add_row().cells
        row[0].text = str(item.get("title") or "")
        row[1].text = str(item.get("primary_module") or "")
        row[2].text = "; ".join(str(value) for value in item.get("owns", []) if value)
        row[3].text = "; ".join(str(value) for value in item.get("does_not_own", []) if value)
        row[4].text = str(item.get("autofix_policy") or "")
    target_index = payload.get("target_index") if isinstance(payload.get("target_index"), dict) else {}
    doc.add_heading("Карта целей", level=2)
    doc.add_paragraph(f"Всего нормализованных записей: {target_index.get('total', 0)}")
    doc.add_paragraph(f"Активных записей: {target_index.get('active', 0)}")
    doc.add_paragraph(f"Подавленных точных дублей: {target_index.get('suppressed_duplicates', 0)}")
    doc.add_paragraph(f"Грубых записей уровня файла: {target_index.get('coarse_entries', 0)}")
    doc.add_paragraph(f"Групп пересечений по одной цели: {target_index.get('overlap_groups', 0)}")
    unified_findings = [item for item in payload.get("unified_findings", []) if isinstance(item, dict)]
    if unified_findings:
        findings_table = doc.add_table(rows=1, cols=7)
        findings_table.style = "Table Grid"
        for index, header in enumerate(["ID", "Статус", "Владелец", "Класс", "Риск", "Локация", "Доказательство"]):
            findings_table.rows[0].cells[index].text = header
        for item in unified_findings[:80]:
            row = findings_table.add_row().cells
            values = [
                item.get("id"),
                item.get("status"),
                item.get("owner"),
                item.get("class_id"),
                item.get("risk"),
                item.get("location_display"),
                item.get("evidence"),
            ]
            for index, value in enumerate(values):
                row[index].text = str(value or "")
    doc.add_heading("Ownership", level=2)
    owner_table = doc.add_table(rows=1, cols=4)
    owner_table.style = "Table Grid"
    for index, header in enumerate(["Владелец", "Объект", "Классы", "Примечание"]):
        owner_table.rows[0].cells[index].text = header
    for item in payload.get("ownership", []):
        if not isinstance(item, dict):
            continue
        row = owner_table.add_row().cells
        row[0].text = str(item.get("owner") or "")
        row[1].text = str(item.get("target_type") or "")
        row[2].text = ", ".join(str(value) for value in item.get("classes", []) if value)
        row[3].text = str(item.get("note") or "")
    doc.add_heading("Пересечения", level=2)
    for note in payload.get("overlap_notes", []):
        doc.add_paragraph(str(note), style="List Bullet")
    doc.save(str(path))


def _write_deep_hygiene_reports(context: JobContext, payload: dict[str, object]) -> dict[str, str]:
    md_report = context.report_dir / "docx_deep_hygiene.md"
    docx_report = context.report_dir / "docx_deep_hygiene.docx"
    json_report = context.report_dir / "docx_deep_hygiene.json"
    md_report.write_text(_deep_hygiene_markdown(context, payload), encoding="utf-8", newline="\n")
    _write_deep_hygiene_docx(docx_report, payload)
    _write_json_payload(json_report, payload)
    return {"deep_report": str(md_report), "deep_docx_report": str(docx_report), "deep_json_report": str(json_report)}


def _table_fit_column_skip_args(context: JobContext) -> list[str]:
    if not _bool_param(context, "skip_fit_by_columns", True):
        return ["--skip-fit-max-columns", "0"]
    threshold = max(0, _int_param(context, "skip_fit_max_columns", 3))
    return ["--skip-fit-max-columns", str(threshold)]


def _input_file(context: JobContext, value: str, suffixes: tuple[str, ...], label: str) -> Path:
    if value:
        raw = Path(value).expanduser()
        candidate = raw.resolve() if raw.is_absolute() else (context.paths.input / raw).resolve()
        if not candidate.is_file():
            raise FileNotFoundError(f"{label}: файл не найден: {candidate}")
        if candidate.suffix.lower() not in suffixes:
            raise RuntimeError(f"{label}: неподдерживаемый формат {candidate.suffix}")
        return candidate

    source = context.paths.input.resolve()
    if source.is_file():
        files = [source] if source.suffix.lower() in suffixes else []
    else:
        files = sorted(
            path for path in source.rglob("*")
            if path.is_file() and path.suffix.lower() in suffixes
        )
    if not files:
        raise FileNotFoundError(f"{label}: в input нет файлов {', '.join(suffixes)}")
    return files[0]


def _two_input_files(
    context: JobContext,
    first: str,
    second: str,
    suffixes: tuple[str, ...],
    label: str,
) -> tuple[Path, Path]:
    if first or second:
        if not first or not second:
            raise RuntimeError(f"{label}: нужно указать оба файла")
        return (
            _input_file(context, first, suffixes, f"{label} A"),
            _input_file(context, second, suffixes, f"{label} B"),
        )

    source = context.paths.input.resolve()
    if source.is_file():
        files = [source] if source.suffix.lower() in suffixes else []
    else:
        files = sorted(
            path for path in source.rglob("*")
            if path.is_file() and path.suffix.lower() in suffixes
        )
    if len(files) < 2:
        raise FileNotFoundError(f"{label}: нужно минимум два файла в input")
    return files[0], files[1]


def validate_input(context: JobContext) -> dict[str, object]:
    if not context.paths.input.exists():
        raise FileNotFoundError(f"Источник не найден: {context.paths.input}")
    inventory = _inventory(context.paths.input)
    context.log(f"Input: {context.paths.input}")
    context.log(f"Файлов: {inventory['files']}; папок: {inventory['dirs']}")
    context.progress(1.0)
    return {"input": str(context.paths.input), "inventory": inventory}


def cleanup_input_output(context: JobContext) -> dict[str, object]:
    context.log("Очищаю управляемые папки input и output.")
    input_result = _clean_managed_folder(context, context.paths.input, "input")
    context.progress(0.5)
    output_result = _clean_managed_folder(context, context.paths.output, "output")
    removed = int(input_result["removed_items"]) + int(output_result["removed_items"])
    context.log(f"Очистка завершена. Удалено элементов: {removed}")
    return {"input": input_result, "output": output_result, "removed_items": removed}


def docx_quality_gate(context: JobContext) -> dict[str, object]:
    report = context.report_dir / "docx_quality_gate_report.md"
    docx_report = context.report_dir / "docx_quality_gate_report.docx"
    return _run_command(
        context,
        _python_command(context, "docx_quality_gate.py", "--input", context.paths.input, "--out", report, "--docx-out", docx_report),
    ) | {"report": str(report), "docx_report": str(docx_report)}


def docx_hard_gate(context: JobContext) -> dict[str, object]:
    report = context.report_dir / "docx_gate_hard_fail.md"
    docx_report = context.report_dir / "docx_gate_hard_fail.docx"
    return _run_command(
        context,
        _python_command(context, "docx_quality_gate_hard_fail.py", "--input", context.paths.input, "--out", report, "--docx-out", docx_report),
    ) | {"report": str(report), "docx_report": str(docx_report)}


def docx_text_hygiene_scan(context: JobContext) -> dict[str, object]:
    report = context.report_dir / "docx_text_hygiene_report.md"
    args: list[str | Path] = ["--input", context.paths.input, "--out", report, "--docx-out", context.report_dir / "docx_text_hygiene_report.docx"]
    if _bool_param(context, "check_dot"):
        args.append("--check-dot")
    return _run_command(context, _python_command(context, "docx_text_hygiene_scan.py", *args)) | {"report": str(report)}


def docx_text_hygiene_fix(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "hygiene_fixed"
    report = context.report_dir / "docx_text_hygiene_changes.md"
    args: list[str | Path] = ["--input", context.paths.input, "--outdir", outdir, "--report", report]
    if _bool_param(context, "fix_dot"):
        args.append("--fix-dot")
    if _bool_param(context, "remove_strikethrough", False):
        args.append("--remove-strikethrough")
    return _run_command(context, _python_command(context, "docx_text_hygiene_fix.py", *args)) | {"outdir": str(outdir), "report": str(report)}


def docx_deep_hygiene(context: JobContext) -> dict[str, object]:
    mode = _str_param(context, "deep_hygiene_mode", "scan").lower()
    fix_mode = mode in {"fix", "correct", "correction", "corrective", "корректировка", "исправить"}
    include_audit = _bool_param(context, "deep_hygiene_include_audit", True)
    include_anomalies = _bool_param(context, "deep_hygiene_include_anomalies", False)
    dot_spacing = _bool_param(context, "deep_hygiene_dot_spacing", False)
    remove_strikethrough = _bool_param(context, "deep_hygiene_remove_strikethrough", False)
    result: dict[str, object] = {
        "mode": "fix" if fix_mode else "scan",
        "include_audit": include_audit,
        "include_anomalies": include_anomalies,
        "report_dir": str(context.report_dir),
    }
    passes: list[dict[str, object]] = []

    if fix_mode:
        context.log("Проход 1: текстовая гигиена, корректировка видимого текста.")
        hygiene_outdir = context.paths.output / "hygiene_fixed"
        hygiene_report = context.report_dir / "docx_text_hygiene_changes.md"
        hygiene_json_report = context.report_dir / "docx_text_hygiene_changes.json"
        hygiene_args: list[str | Path] = [
            "--input",
            context.paths.input,
            "--outdir",
            hygiene_outdir,
            "--report",
            hygiene_report,
            "--json-out",
            hygiene_json_report,
        ]
        if dot_spacing:
            hygiene_args.append("--fix-dot")
        if remove_strikethrough:
            hygiene_args.append("--remove-strikethrough")
        _run_command(context, _python_command(context, "docx_text_hygiene_fix.py", *hygiene_args))
        result["text_hygiene_report"] = str(hygiene_report)
        result["text_hygiene_json_report"] = str(hygiene_json_report)
        result["text_hygiene_outdir"] = str(hygiene_outdir)
        passes.append(
            _deep_pass_record(
                context,
                pass_id="text_hygiene_fix",
                title="Проход 1: текстовая гигиена",
                owner="docx_text_hygiene",
                mode="fix",
                report=hygiene_report,
                docx_report=hygiene_report.with_suffix(".docx"),
                json_report=hygiene_json_report,
                outdir=hygiene_outdir,
            )
        )
        current_fix_input = hygiene_outdir

        if include_audit:
            context.log("Проход 1B: правила аудита, корректировка детерминированных текстовых правил.")
            audit_outdir = context.paths.output / "docx_deep_hygiene_audit_fixed"
            audit_report = context.report_dir / "docx_deep_hygiene_audit_changes.md"
            audit_docx_report = context.report_dir / "docx_deep_hygiene_audit_changes.docx"
            audit_json_report = context.report_dir / "docx_deep_hygiene_audit_changes.json"
            audit_args: list[str | Path] = [
                "--input",
                current_fix_input,
                "--outdir",
                audit_outdir,
                "--report",
                audit_report,
                "--docx-out",
                audit_docx_report,
                "--json-out",
                audit_json_report,
                "--fix",
            ]
            _run_command(context, _python_command(context, "docx_audit_processor.py", *audit_args))
            result["audit_report"] = str(audit_report)
            result["audit_docx_report"] = str(audit_docx_report)
            result["audit_json_report"] = str(audit_json_report)
            result["audit_outdir"] = str(audit_outdir)
            current_fix_input = audit_outdir
            passes.append(
                _deep_pass_record(
                    context,
                    pass_id="audit_fix",
                    title="Проход 1B: правила аудита",
                    owner="docx_audit_processor",
                    mode="fix",
                    report=audit_report,
                    docx_report=audit_docx_report,
                    json_report=audit_json_report,
                    outdir=audit_outdir,
                )
            )

        if include_anomalies:
            context.log("Проход 2: аномалии документа, корректировка форматирования поверх текстовых копий.")
            anomaly_outdir = context.paths.output / "docx_deep_hygiene_fixed"
            anomaly_report = context.report_dir / "docx_deep_hygiene_anomaly_corrections.md"
            anomaly_docx_report = context.report_dir / "docx_deep_hygiene_anomaly_corrections.docx"
            anomaly_json_report = context.report_dir / "docx_deep_hygiene_anomaly_corrections.json"
            anomaly_args: list[str | Path] = [
                "--input",
                current_fix_input,
                "--outdir",
                anomaly_outdir,
                "--report",
                anomaly_report,
                "--docx-report",
                anomaly_docx_report,
                "--json-out",
                anomaly_json_report,
                "--border-size",
                str(max(1, _int_param(context, "border_size", 4))),
                "--border-color",
                _str_param(context, "border_color", "000000") or "000000",
                "--cell-margin-cm",
                _str_param(context, "cell_margin_cm", "0.2") or "0.2",
            ]
            if not _bool_param(context, "fix_empty_paragraphs", True):
                anomaly_args.append("--no-fix-empty-paragraphs")
            if not _bool_param(context, "fix_row_heights", True):
                anomaly_args.append("--no-fix-row-heights")
            if not _bool_param(context, "fix_nowrap", True):
                anomaly_args.append("--no-fix-nowrap")
            if _bool_param(context, "normalize_table_borders", False):
                anomaly_args.append("--normalize-table-borders")
            if _bool_param(context, "normalize_cell_margins", False):
                anomaly_args.append("--normalize-cell-margins")
            _run_command(context, _python_command(context, "docx_anomaly_corrector.py", *anomaly_args))
            result["anomaly_report"] = str(anomaly_report)
            result["anomaly_docx_report"] = str(anomaly_docx_report)
            result["anomaly_json_report"] = str(anomaly_json_report)
            result["outdir"] = str(anomaly_outdir)
            passes.append(
                _deep_pass_record(
                    context,
                    pass_id="anomaly_correct",
                    title="Проход 2: аномалии документа",
                    owner="docx_anomaly_corrector",
                    mode="fix",
                    report=anomaly_report,
                    docx_report=anomaly_docx_report,
                    json_report=anomaly_json_report,
                    outdir=anomaly_outdir,
                )
            )
        else:
            result["outdir"] = str(current_fix_input)
        payload = _deep_hygiene_payload(context, "fix", include_audit, include_anomalies, passes, Path(str(result.get("outdir"))))
        result.update(_write_deep_hygiene_reports(context, payload))
        return result

    context.log("Проход 1: текстовая гигиена, проверка без изменения документов.")
    hygiene_report = context.report_dir / "docx_text_hygiene_report.md"
    hygiene_docx_report = context.report_dir / "docx_text_hygiene_report.docx"
    hygiene_json_report = context.report_dir / "docx_text_hygiene_report.json"
    hygiene_args = [
        "--input",
        context.paths.input,
        "--out",
        hygiene_report,
        "--docx-out",
        hygiene_docx_report,
        "--json-out",
        hygiene_json_report,
    ]
    if dot_spacing:
        hygiene_args.append("--check-dot")
    _run_command(context, _python_command(context, "docx_text_hygiene_scan.py", *hygiene_args))
    result["text_hygiene_report"] = str(hygiene_report)
    result["text_hygiene_docx_report"] = str(hygiene_docx_report)
    result["text_hygiene_json_report"] = str(hygiene_json_report)
    passes.append(
        _deep_pass_record(
            context,
            pass_id="text_hygiene_scan",
            title="Проход 1: текстовая гигиена",
            owner="docx_text_hygiene",
            mode="scan",
            report=hygiene_report,
            docx_report=hygiene_docx_report,
            json_report=hygiene_json_report,
        )
    )

    if include_audit:
        context.log("Проход 1B: правила аудита, проверка детерминированных текстовых правил без изменения документов.")
        audit_report = context.report_dir / "docx_deep_hygiene_audit.md"
        audit_docx_report = context.report_dir / "docx_deep_hygiene_audit.docx"
        audit_json_report = context.report_dir / "docx_deep_hygiene_audit.json"
        audit_args: list[str | Path] = [
            "--input",
            context.paths.input,
            "--report",
            audit_report,
            "--docx-out",
            audit_docx_report,
            "--json-out",
            audit_json_report,
        ]
        _run_command(context, _python_command(context, "docx_audit_processor.py", *audit_args))
        result["audit_report"] = str(audit_report)
        result["audit_docx_report"] = str(audit_docx_report)
        result["audit_json_report"] = str(audit_json_report)
        passes.append(
            _deep_pass_record(
                context,
                pass_id="audit_scan",
                title="Проход 1B: правила аудита",
                owner="docx_audit_processor",
                mode="scan",
                report=audit_report,
                docx_report=audit_docx_report,
                json_report=audit_json_report,
            )
        )

    if include_anomalies:
        context.log("Проход 2: аномалии документа, проверка форматирования без изменения документов.")
        anomaly_report = context.report_dir / "docx_deep_hygiene_anomalies.md"
        anomaly_docx_report = context.report_dir / "docx_deep_hygiene_anomalies.docx"
        anomaly_json_report = context.report_dir / "docx_deep_hygiene_anomalies.json"
        anomaly_args = [
            "--input",
            context.paths.input,
            "--out",
            anomaly_report,
            "--docx-out",
            anomaly_docx_report,
            "--json-out",
            anomaly_json_report,
            "--max-findings-per-file",
            str(max(1, _int_param(context, "max_findings_per_file", 300))),
        ]
        if not _bool_param(context, "check_tables", True):
            anomaly_args.append("--no-tables")
        if not _bool_param(context, "check_captions", True):
            anomaly_args.append("--no-captions")
        if not _bool_param(context, "check_layout", True):
            anomaly_args.append("--no-layout")
        if not _bool_param(context, "check_headers", True):
            anomaly_args.append("--no-headers")
        if not _bool_param(context, "check_fields", True):
            anomaly_args.append("--no-fields")
        if not _bool_param(context, "check_lists", True):
            anomaly_args.append("--no-lists")
        if not _bool_param(context, "check_spacing", True):
            anomaly_args.append("--no-spacing")
        _run_command(context, _python_command(context, "docx_anomaly_inspector.py", *anomaly_args))
        result["anomaly_report"] = str(anomaly_report)
        result["anomaly_docx_report"] = str(anomaly_docx_report)
        result["anomaly_json_report"] = str(anomaly_json_report)
        passes.append(
            _deep_pass_record(
                context,
                pass_id="anomaly_scan",
                title="Проход 2: аномалии документа",
                owner="docx_anomaly_inspector",
                mode="scan",
                report=anomaly_report,
                docx_report=anomaly_docx_report,
                json_report=anomaly_json_report,
            )
        )

    payload = _deep_hygiene_payload(context, "scan", include_audit, include_anomalies, passes, None)
    result.update(_write_deep_hygiene_reports(context, payload))
    return result


def docx_anomaly_scan(context: JobContext) -> dict[str, object]:
    report = context.report_dir / "docx_anomalies.md"
    docx_report = context.report_dir / "docx_anomalies.docx"
    json_report = context.report_dir / "docx_anomalies.json"
    args: list[str | Path] = [
        "--input",
        context.paths.input,
        "--out",
        report,
        "--docx-out",
        docx_report,
        "--json-out",
        json_report,
        "--max-findings-per-file",
        str(max(1, _int_param(context, "max_findings_per_file", 300))),
    ]
    if not _bool_param(context, "check_tables", True):
        args.append("--no-tables")
    if not _bool_param(context, "check_captions", True):
        args.append("--no-captions")
    if not _bool_param(context, "check_layout", True):
        args.append("--no-layout")
    if not _bool_param(context, "check_headers", True):
        args.append("--no-headers")
    if not _bool_param(context, "check_fields", True):
        args.append("--no-fields")
    if not _bool_param(context, "check_lists", True):
        args.append("--no-lists")
    if not _bool_param(context, "check_spacing", True):
        args.append("--no-spacing")
    return _run_command(context, _python_command(context, "docx_anomaly_inspector.py", *args)) | {
        "report": str(report),
        "docx_report": str(docx_report),
        "json_report": str(json_report),
    }


def docx_anomaly_correct(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "docx_anomaly_fixed"
    report = context.report_dir / "docx_anomaly_corrections.md"
    docx_report = context.report_dir / "docx_anomaly_corrections.docx"
    json_report = context.report_dir / "docx_anomaly_corrections.json"
    args: list[str | Path] = [
        "--input",
        context.paths.input,
        "--outdir",
        outdir,
        "--report",
        report,
        "--docx-report",
        docx_report,
        "--json-out",
        json_report,
        "--border-size",
        str(max(1, _int_param(context, "border_size", 4))),
        "--border-color",
        _str_param(context, "border_color", "000000") or "000000",
        "--cell-margin-cm",
        _str_param(context, "cell_margin_cm", "0.2") or "0.2",
    ]
    if not _bool_param(context, "fix_empty_paragraphs", True):
        args.append("--no-fix-empty-paragraphs")
    if not _bool_param(context, "fix_row_heights", True):
        args.append("--no-fix-row-heights")
    if not _bool_param(context, "fix_nowrap", True):
        args.append("--no-fix-nowrap")
    if _bool_param(context, "normalize_table_borders", False):
        args.append("--normalize-table-borders")
    if _bool_param(context, "normalize_cell_margins", False):
        args.append("--normalize-cell-margins")
    return _run_command(context, _python_command(context, "docx_anomaly_corrector.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "docx_report": str(docx_report),
        "json_report": str(json_report),
    }


def docx_nonprinting_clean(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "nonprinting_cleaned"
    report = context.report_dir / "docx_nonprinting_clean.md"
    json_report = context.report_dir / "docx_nonprinting_clean.json"
    return _run_command(
        context,
        _python_command(
            context,
            "docx_nonprinting_clean.py",
            "--input",
            context.paths.input,
            "--outdir",
            outdir,
            "--report",
            report,
            "--json-out",
            json_report,
        ),
    ) | {"outdir": str(outdir), "report": str(report), "json_report": str(json_report)}


def docx_comma_lowercase(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "comma_lowercase"
    report = context.report_dir / "comma_lowercase.json"
    md_report = context.report_dir / "comma_lowercase.md"
    per_document_report_dir = context.report_dir / "comma_lowercase"
    args: list[str | Path] = [
        "comma-lowercase-docx",
        "--input",
        context.paths.input,
        "--output",
        outdir,
        "--recursive",
        "--report",
        report,
        "--md-report",
        md_report,
        "--scope",
        "table-cells",
    ]
    keep_words = _str_param(context, "keep_words")
    if keep_words:
        args.extend(["--keep-words", keep_words])
    if _bool_param(context, "dry_run", True):
        args.append("--dry-run")
    if _bool_param(context, "overwrite", False):
        args.append("--overwrite")
    result = _run_command(context, _main_command(context, *args))
    return result | {
        "outdir": str(outdir),
        "report": str(report),
        "md_report": str(md_report),
        "per_document_report_dir": str(per_document_report_dir),
    }


def morph_replace(context: JobContext) -> dict[str, object]:
    find = _str_param(context, "find", "гаражи")
    replace = _str_param(context, "replace", "подземные гаражи")
    mode = _str_param(context, "mode", "replace") or "replace"
    if mode not in {"replace", "append"}:
        raise RuntimeError("Режим морфологической замены должен быть replace или append.")
    if not find:
        raise RuntimeError("Поле поиска не заполнено.")
    if not replace:
        raise RuntimeError("Поле замены/добавки не заполнено.")
    outdir = context.paths.output / ("morph_appended" if mode == "append" else "morph_replaced")
    report = context.report_dir / "morph_replace.json"
    md_report = context.report_dir / "morph_replace.md"
    per_document_report_dir = context.report_dir / "morph_replace"
    args: list[str | Path] = [
        "morph-replace",
        "--input",
        context.paths.input,
        "--output",
        outdir,
        "--find",
        find,
        "--replace",
        replace,
        "--mode",
        mode,
        "--recursive",
        "--extensions",
        _str_param(context, "extensions", ".docx,.xlsx") or ".docx,.xlsx",
        "--report",
        report,
        "--md-report",
        md_report,
    ]
    if _bool_param(context, "inflect_replacement", True):
        args.append("--inflect-replacement")
    if _bool_param(context, "dry_run", True):
        args.append("--dry-run")
    if _bool_param(context, "overwrite", False):
        args.append("--overwrite")
    result = _run_command(context, _main_command(context, *args))
    return result | {
        "outdir": str(outdir),
        "report": str(report),
        "md_report": str(md_report),
        "per_document_report_dir": str(per_document_report_dir),
    }


def docx_audit_processor(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "audit_processed"
    annotated_outdir = context.paths.output / "audit_annotated"
    report = context.report_dir / "docx_audit_processor.md"
    args: list[str | Path] = [
        "--input",
        context.paths.input,
        "--outdir",
        outdir,
        "--report",
        report,
        "--docx-out",
        context.report_dir / "docx_audit_processor.docx",
        "--json-out",
        context.report_dir / "docx_audit_processor.json",
    ]
    if _bool_param(context, "write_audit_anchors", True):
        args.extend(["--annotate", "--annotated-outdir", annotated_outdir])
    if _bool_param(context, "apply_fixes"):
        args.append("--fix")
    if _bool_param(context, "dry_run"):
        args.append("--dry-run")
    result = _run_command(context, _python_command(context, "docx_audit_processor.py", *args))
    return result | {"outdir": str(outdir), "annotated_outdir": str(annotated_outdir), "report": str(report)}


def docx_audit_strip_anchors(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "audit_unanchored"
    report = context.report_dir / "docx_audit_anchors.md"
    json_report = context.report_dir / "docx_audit_anchors.json"
    source_param = _str_param(context, "audit_anchor_docx")
    source: Path = _input_file(context, source_param, (".docx",), "DOCX с якорями") if source_param else context.paths.input
    args: list[str | Path] = [
        "--input",
        source,
        "--outdir",
        outdir,
        "--report",
        report,
        "--json-out",
        json_report,
    ]
    if _bool_param(context, "strip_all_docx", False):
        args.append("--all-docx")
    return _run_command(context, _python_command(context, "docx_audit_anchors.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "json_report": str(json_report),
    }


def docx_strip_comments(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "no_comments"
    return _run_command(context, _python_command(context, "docx_strip_comments.py", "--input", context.paths.input, "--outdir", outdir)) | {"outdir": str(outdir)}


def docx_accept_changes(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "accepted_changes"
    return _run_command(context, _python_command(context, "docx_accept_changes_simple.py", "--input", context.paths.input, "--outdir", outdir)) | {"outdir": str(outdir)}


def docx_finalize_black(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "final_black"
    report = context.report_dir / "docx_finalize_changes.md"
    return _run_command(
        context,
        _python_command(context, "docx_finalize_black_clean.py", "--input", context.paths.input, "--outdir", outdir, "--report", report),
    ) | {"outdir": str(outdir), "report": str(report)}


def docx_oneclick(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "final"
    summary = context.report_dir / "summary_pass_fail.md"
    args: list[str | Path] = ["--input", context.paths.input, "--outdir", outdir, "--summary", summary]
    if not _bool_param(context, "strip_comments", True):
        args.append("--keep-comments")
    if not _bool_param(context, "accept_changes", True):
        args.append("--keep-changes")
    if _bool_param(context, "fix_dot"):
        args.append("--fix-dot")
    return _run_command(context, _python_command(context, "docx_oneclick_finalize_gate.py", *args)) | {"outdir": str(outdir), "summary": str(summary)}


def docx_near_dupes(context: JobContext) -> dict[str, object]:
    report = context.report_dir / "docx_similarity_map.md"
    docx_report = context.report_dir / "docx_similarity_map.docx"
    json_report = context.report_dir / "docx_similarity_map.json"
    args: list[str | Path] = [
        "--input",
        context.paths.input,
        "--threshold",
        _str_param(context, "threshold", "0.30"),
        "--out",
        report,
        "--docx-out",
        docx_report,
        "--json-out",
        json_report,
    ]
    if _bool_param(context, "write_diff"):
        args.append("--diff")
    return _run_command(context, _python_command(context, "docx_near_dupes.py", *args)) | {
        "report": str(report),
        "docx_report": str(docx_report),
        "json_report": str(json_report),
    }


def docx_pair_diff(context: JobContext) -> dict[str, object]:
    a, b = _two_input_files(context, _str_param(context, "docx_a"), _str_param(context, "docx_b"), (".docx",), "Diff DOCX")
    report = context.report_dir / "docx_pair_diff.md"
    docx_report = context.report_dir / "docx_pair_diff.docx"
    diff_report = context.report_dir / "docx_pair_diff.diff"
    json_report = context.report_dir / "docx_pair_diff.json"
    stats_report = context.report_dir / "docx_pair_diff.stats.json"
    alignment_docx = context.report_dir / "docx_pair_alignment_map.docx"
    alignment_md = context.report_dir / "docx_pair_alignment_map.md"
    alignment_json = context.report_dir / "docx_pair_alignment_map.json"
    redline_out = context.paths.output / "docx_compare" / "docx_pair_redline.docx"
    output_mode = _str_param(context, "docflow_output_mode", "redline_report")
    command = _python_command(
        context,
        "docx_pair_diff.py",
        "--a",
        a,
        "--b",
        b,
        "--out",
        report,
        "--docx-out",
        docx_report,
        "--diff-out",
        diff_report,
        "--json-out",
        json_report,
        "--stats-out",
        stats_report,
        "--alignment-map-out",
        alignment_docx,
        "--alignment-map-md",
        alignment_md,
        "--alignment-map-json",
        alignment_json,
        "--compare-mode",
        _str_param(context, "compare_mode", "sections"),
    )
    result: dict[str, object] = {
        "report": str(report),
        "docx_report": str(docx_report),
        "diff_report": str(diff_report),
        "json_report": str(json_report),
        "stats": str(stats_report),
        "alignment_docx": str(alignment_docx),
        "alignment_md": str(alignment_md),
        "alignment_json": str(alignment_json),
    }
    if output_mode == "redline_report":
        command.extend(["--redline-out", str(redline_out)])
        result["out"] = str(redline_out)
    return _run_command(context, command) | result


def docx_word_compare(context: JobContext) -> dict[str, object]:
    a_param = _str_param(context, "word_compare_a")
    b_param = _str_param(context, "word_compare_b")
    if not a_param or not b_param:
        raise RuntimeError("Сравнение через Word COM: укажите оба DOCX файла.")
    a = _input_file(context, a_param, (".docx",), "Word Compare A")
    b = _input_file(context, b_param, (".docx",), "Word Compare B")
    out = context.paths.output / "word_compare" / "word_compare_result.docx"
    report = context.report_dir / "docx_word_compare.md"
    args: list[str | Path] = [
        "--a",
        a,
        "--b",
        b,
        "--out",
        out,
        "--report",
        report,
    ]
    split_mode = _str_param(context, "word_compare_split_mode", "none") or "none"
    if _bool_param(context, "word_compare_close_word_before_compare", False):
        args.append("--close-word-before-compare")
    if split_mode != "none":
        args.extend(["--chunked", "--split-mode", split_mode])
    elif _bool_param(context, "word_compare_auto_smart_on_problem", True):
        args.append("--auto-smart-on-problem")
    if split_mode == "chunks":
        chunks = _str_param(context, "word_compare_chunks")
        if chunks:
            args.extend(["--chunks", chunks])
    alignment_docx = context.report_dir / "word_compare_alignment_plan.docx"
    alignment_md = context.report_dir / "word_compare_alignment_plan.md"
    alignment_json = context.report_dir / "word_compare_alignment_plan.json"
    return _run_command(
        context,
        _python_command(
            context,
            "docx_word_compare.py",
            *args,
        ),
    ) | {
        "out": str(out),
        "report": str(report),
        "alignment_docx": str(alignment_docx),
        "alignment_md": str(alignment_md),
        "alignment_json": str(alignment_json),
    }


def docx_merge(context: JobContext) -> dict[str, object]:
    output_name = _str_param(context, "output_name", "merged.docx") or "merged.docx"
    if Path(output_name).is_absolute() or ".." in Path(output_name).parts:
        raise RuntimeError("Имя выходного DOCX должно быть простым относительным путём.")
    if not output_name.lower().endswith(".docx"):
        output_name += ".docx"
    out = context.paths.output / output_name
    report = context.report_dir / "docx_merge.md"
    args: list[str | Path] = ["--input", context.paths.input, "--out", out, "--report", report]
    files = _str_param(context, "files")
    if files:
        args.extend(["--files", files])
    return _run_command(context, _python_command(context, "docx_merge.py", *args)) | {"out": str(out), "report": str(report)}


def _restyle_template_arg(context: JobContext) -> list[str | Path]:
    """The reference is opt-in: the checkbox decides, the field only names it."""
    if not _bool_param(context, "use_reference", False):
        return []
    reference = _str_param(context, "reference_docx")
    path = _input_file(context, reference, (".docx",), "Эталон DOCX")
    return ["--template", path]


def _restyle_config_arg(context: JobContext) -> list[str | Path]:
    """The heading map is written by hand per pair of documents; without it the
    transfer still moves the style base and cleans, it just marks nothing."""
    value = _str_param(context, "style_config")
    if not value:
        return []
    raw = Path(value).expanduser()
    path = raw.resolve() if raw.is_absolute() else (context.paths.input / raw).resolve()
    if not path.is_file():
        raise FileNotFoundError(f"Карта заголовков: файл не найден: {path}")
    return ["--config", path]


def docx_style_scan(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "style_probe"
    report = context.report_dir / "docx_style_probe.md"
    json_report = context.report_dir / "docx_style_probe.json"
    args: list[str | Path] = [
        "--input", context.paths.input,
        "--outdir", outdir,
        "--report", report,
        "--json-out", json_report,
    ]
    if _bool_param(context, "use_reference", False):
        reference = _str_param(context, "reference_docx")
        args += ["--template", _input_file(context, reference, (".docx",), "Эталон DOCX")]
    return _run_command(context, _python_command(context, "docx_style_probe.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "json_report": str(json_report),
    }


def docx_style_fix(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "style_processed"
    report = context.report_dir / "docx_restyle.md"
    json_report = context.report_dir / "docx_restyle.json"
    args: list[str | Path] = [
        "--input", context.paths.input,
        "--outdir", outdir,
        "--report", report,
        "--json-out", json_report,
        *_restyle_template_arg(context),
        *_restyle_config_arg(context),
    ]
    return _run_command(context, _python_command(context, "docx_restyle_by_template.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "json_report": str(json_report),
    }


def docx_xml_cleanup(context: JobContext) -> dict[str, object]:
    """Cleanup alone: revision attributes, split runs, empty properties, broken
    paragraphs, fixed table layout and thousands of unused style clones. Styles
    stay as they are, run properties are untouched, so colours survive."""
    outdir = context.paths.output / "xml_cleaned"
    report = context.report_dir / "docx_xml_cleanup.md"
    json_report = context.report_dir / "docx_xml_cleanup.json"
    args: list[str | Path] = [
        "--input", context.paths.input,
        "--outdir", outdir,
        "--report", report,
        "--json-out", json_report,
        "--clean-only",
    ]
    return _run_command(context, _python_command(context, "docx_restyle_by_template.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "json_report": str(json_report),
    }


def docx_extract_tables(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "tables"
    report = context.report_dir / "docx_tables_index.md"
    return _run_command(context, _python_command(context, "docx_extract_tables.py", "--input", context.paths.input, "--outdir", outdir, "--out", report)) | {"outdir": str(outdir), "report": str(report)}


def docx_extract_media(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "media"
    report = context.report_dir / "office_media_index.md"
    return _run_command(context, _python_command(context, "docx_extract_media.py", "--input", context.paths.input, "--outdir", outdir, "--out", report)) | {"outdir": str(outdir), "report": str(report)}


def docx_clean(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "cleaned_docx"
    report = context.report_dir / "docx_clean.md"
    json_report = context.report_dir / "docx_clean.json"
    return _run_command(
        context,
        _python_command(
            context,
            "docx_clean.py",
            "--input",
            context.paths.input,
            "--outdir",
            outdir,
            "--report",
            report,
            "--json-out",
            json_report,
        ),
    ) | {"outdir": str(outdir), "report": str(report), "json_report": str(json_report)}


def docx_split(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "docx_split"
    report = context.report_dir / "docx_split.md"
    json_report = context.report_dir / "docx_split.json"
    args: list[str | Path] = [
        "--input",
        context.paths.input,
        "--outdir",
        outdir,
        "--report",
        report,
        "--json-out",
        json_report,
    ]
    chunks = _str_param(context, "docx_split_chunks")
    if chunks:
        args.extend(["--chunks", chunks])
    if _bool_param(context, "docx_split_top_level_sections", True):
        args.append("--top-level-sections")
    return _run_command(context, _python_command(context, "docx_split.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "json_report": str(json_report),
    }


def table_cell_margins(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "table_cell_margins"
    report = context.report_dir / "table_cell_margins.md"
    json_report = context.report_dir / "table_cell_margins.json"
    return _run_command(
        context,
        _python_command(
            context,
            "docx_xlsx_table_cell_margins.py",
            "--input",
            context.paths.input,
            "--outdir",
            outdir,
            "--report",
            report,
            "--json-out",
            json_report,
            "--margin-cm",
            "0.1",
        ),
    ) | {"outdir": str(outdir), "report": str(report), "json_report": str(json_report)}


def docx_table_stitcher(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "stitched"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_stitcher.py",
            "--input",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--recursive",
            *_table_preheader_args(context),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_stitcher_running_header(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "stitched_running_header"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_stitcher_reconstruct_running_header.py",
            "--input",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--recursive",
            *_table_preheader_args(context),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_unify_safe(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "unified_safe"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_unifier.py",
            "--input-dir",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--all",
            "--recursive",
            "--mode",
            "safe",
            "--layout",
            "standard",
            *_table_preheader_args(context),
            *_table_page_args(context, default_orientation="document"),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_unify_width_only(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "unified_width_only"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_unifier.py",
            "--input-dir",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--all",
            "--recursive",
            "--mode",
            "width-only",
            "--layout",
            "standard",
            *_table_preheader_args(context),
            *_table_page_args(context, default_orientation="document"),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_unify_merged_sections(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "unified_merged_sections"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_unifier.py",
            "--input-dir",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--all",
            "--recursive",
            "--mode",
            "safe",
            "--layout",
            "merged-sections",
            *_table_preheader_args(context),
            *_table_page_args(context, default_orientation="document"),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_optimize_widths(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "optimized_widths"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_width_optimizer.py",
            "--input-dir",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--all",
            "--recursive",
            "--mode",
            "preserve-width",
            *_table_preheader_args(context),
            "--fit-target",
            _str_param(context, "fit_target", "current-section"),
            *_table_page_args(context),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_adapt_orientation(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "adapted_orientation"
    width_action = _str_param(context, "width_action", "fit-to-margins")
    mode = {
        "none": "page-setup-only",
        "balance": "preserve-width",
        "fit-to-margins": "fit-to-margins",
    }.get(width_action, "fit-to-margins")
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_width_optimizer.py",
            "--input-dir",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--all",
            "--recursive",
            "--mode",
            mode,
            *_table_preheader_args(context),
            "--fit-target",
            "page-setup",
            "--page-setup-margin-fallback",
            "keep",
            *_table_page_args(context),
            *_table_fit_column_skip_args(context),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_table_fit_to_margins(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "fit_to_margins"
    result = _run_command(
        context,
        _python_command(
            context,
            "docx_table_width_optimizer.py",
            "--input-dir",
            context.paths.input,
            "--outdir",
            outdir,
            "--report-dir",
            context.report_dir,
            "--all",
            "--recursive",
            "--mode",
            "fit-to-margins",
            *_table_preheader_args(context),
            "--fit-target",
            _str_param(context, "fit_target", "current-section"),
            *_table_page_args(context),
            *_table_fit_column_skip_args(context),
        ),
    )
    return result | {"outdir": str(outdir), "report_dir": str(context.report_dir)}


def docx_document_tables_unifier(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "word_excel_tables" / "document_tables_unified"
    report_dir = context.report_dir / "word_excel_tables" / "document_tables_unified"
    args: list[str | Path] = [
        "--input-dir",
        context.paths.input,
        "--outdir",
        outdir,
        "--report-dir",
        report_dir,
        "--all",
        "--recursive",
        "--preheader-mode",
        _str_param(context, "preheader_mode", "separate"),
        "--font",
        _str_param(context, "document_table_font", "auto") or "auto",
        "--fallback-font",
        _str_param(context, "fallback_font", "Tahoma") or "Tahoma",
        "--normal-size",
        _str_param(context, "normal_size", "10.0") or "10.0",
        "--dense-size",
        _str_param(context, "dense_size", "8.0") or "8.0",
        "--cell-margin-cm",
        _str_param(context, "cell_margin_cm", "0.2") or "0.2",
        "--border-size",
        str(max(1, _int_param(context, "border_size", 4))),
        "--border-color",
        _str_param(context, "border_color", "000000") or "000000",
        "--line-twips",
        str(max(180, _int_param(context, "line_twips", 240))),
    ]
    skip_first_tables = max(0, _int_param(context, "skip_first_tables", 0))
    if skip_first_tables:
        args.extend(["--skip-first-tables", str(skip_first_tables)])
    skip_first_pages = max(0, _int_param(context, "skip_first_pages", 2))
    if skip_first_pages:
        args.extend(["--skip-first-pages", str(skip_first_pages)])
    result = _run_command(context, _python_command(context, "docx_document_tables_unifier.py", *args))
    return result | {"outdir": str(outdir), "report_dir": str(report_dir)}


def xlsx_values_diff(context: JobContext) -> dict[str, object]:
    a, b = _two_input_files(context, _str_param(context, "xlsx_a"), _str_param(context, "xlsx_b"), (".xlsx", ".xlsm"), "XLSX diff")
    report = context.report_dir / "xlsx_values_diff_report.md"
    return _run_command(context, _python_command(context, "xlsx_values_diff.py", "--a", a, "--b", b, "--out", report, "--docx-out", context.report_dir / "xlsx_values_diff_report.docx")) | {"report": str(report)}


def tabular_reconcile(context: JobContext) -> dict[str, object]:
    a, b = _two_input_files(
        context,
        _str_param(context, "table_a"),
        _str_param(context, "table_b"),
        (".docx", ".xlsx", ".xlsm", ".csv"),
        "Сверка таблиц",
    )
    report = context.report_dir / "reconcile_4lists.md"
    args: list[str | Path] = ["--a", a, "--b", b, "--out", report, "--docx-out", context.report_dir / "reconcile_4lists.docx"]
    key_columns = _str_param(context, "key_columns")
    compare_fields = _str_param(context, "compare_fields")
    if key_columns:
        args.extend(["--key-csv", key_columns])
    if compare_fields:
        args.extend(["--fields-csv", compare_fields])
    return _run_command(context, _python_command(context, "tabular_reconcile_4lists.py", *args)) | {"report": str(report)}


def md_tables_to_xlsx(context: JobContext) -> dict[str, object]:
    md_file = _input_file(context, _str_param(context, "md_file"), (".md", ".markdown"), "Markdown")
    output_name = _str_param(context, "output_name", "markdown_tables.xlsx") or "markdown_tables.xlsx"
    if Path(output_name).is_absolute() or ".." in Path(output_name).parts:
        raise RuntimeError("Имя выходного XLSX должно быть простым относительным путём.")
    if not output_name.lower().endswith(".xlsx"):
        output_name += ".xlsx"
    out = context.paths.output / output_name
    return _run_command(context, _python_command(context, "md_tables_to_xlsx.py", "--md", md_file, "--out", out)) | {"out": str(out)}


def xlsx_remove_empty_rows(context: JobContext) -> dict[str, object]:
    outdir = context.paths.output / "xlsx_no_empty_rows"
    report = context.report_dir / "xlsx_remove_empty_rows.md"
    json_report = context.report_dir / "xlsx_remove_empty_rows.json"
    args: list[str | Path] = [
        "--input",
        context.paths.input,
        "--outdir",
        outdir,
        "--report",
        report,
        "--json-out",
        json_report,
    ]
    if _bool_param(context, "xlsx_clean_all_sheets", True):
        args.append("--all-sheets")
    return _run_command(context, _python_command(context, "xlsx_remove_empty_rows.py", *args)) | {
        "outdir": str(outdir),
        "report": str(report),
        "json_report": str(json_report),
    }
