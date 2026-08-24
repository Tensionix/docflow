#!/usr/bin/env python3
"""
DOCX anomaly inspector.

Read-only scan for visible formatting and structure anomalies. The report uses
human-oriented locations: page when available, nearest numbered section, table
or figure number, object detail, and a short search fragment.
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import re
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from _office_common import find_docx_files, md_escape, norm_space, rel_posix, safe_mkdir, truncate, write_json_file
from docx_table_unifier_model import TableModel, build_table_model, normalize_text
from docx_table_width_optimizer import (
    _body_table_section_map,
    _fallback_total_width_cm,
    _read_table_total_width_cm,
    _section_available_width_cm,
)


DOTTED_HEADING_RE = re.compile(r"^\s*(?P<num>\d+\.\d+(?:\.\d+){0,4})\.?\s*(?P<title>\S.+)$")
SINGLE_HEADING_RE = re.compile(r"^\s*(?P<num>\d+)\.\s+(?P<title>\S.+)$")
APPENDIX_HEADING_RE = re.compile(r"^\s*(?P<num>Приложение\s+[А-ЯA-Z0-9]+)\.?\s*(?P<title>.*)$", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(
    r"^\s*(?:таблица|табл\.)\s*\.?\s*(?P<num>\d+(?:\.\d+)*)(?P<dots>\.{0,2})(?=\s|$|[-–—:])",
    re.IGNORECASE,
)
FIGURE_CAPTION_RE = re.compile(
    r"^\s*(?:рисунок|рис\.)\s*\.?\s*(?P<num>\d+(?:\.\d+)*)(?P<dots>\.{0,2})(?=\s|$|[-–—:])",
    re.IGNORECASE,
)
LIST_NUMBER_RE = re.compile(r"^\s*(?:\d+|[IVXLCDM]+)[\.)]\s+\S", re.IGNORECASE)
MANUAL_NUMBER_RE = re.compile(r"^\s*(?P<num>\d+(?:\.\d+){0,5})[\.)]\s+\S")
MANUAL_BULLET_RE = re.compile(r"^\s*[-–—•·●]\s+\S")
FIELD_REF_RE = re.compile(r"\b(?:REF|PAGEREF|NOTEREF)\s+(?P<name>[A-Za-zА-Яа-яЁё_][\wА-Яа-яЁё.-]*)", re.IGNORECASE)
FIELD_SEQ_RE = re.compile(r"\bSEQ\s+(?P<name>[A-Za-zА-Яа-яЁё_][\wА-Яа-яЁё.-]*)", re.IGNORECASE)
BROKEN_FIELD_RE = re.compile(
    r"(?:Error!\s*(?:Reference source not found|Bookmark not defined)|Ошибка!\s*(?:Источник ссылки не найден|Закладка не определена))",
    re.IGNORECASE,
)
TOC_STYLE_RE = re.compile(r"\b(?:toc|оглавлен|содержан)", re.IGNORECASE)
PAGE_FIELD_RE = re.compile(r"\b(?:PAGE|NUMPAGES|SECTIONPAGES)\b", re.IGNORECASE)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
PAGE_UNKNOWN = "страница: не определена"
MAX_SEARCH_TEXT = 120


@dataclass
class Location:
    file: str
    page: int | None = None
    section_number: str = ""
    section_title: str = ""
    object_type: str = ""
    object_number: str = ""
    detail: str = ""
    search_text: str = ""
    debug_path: str = ""

    def display(self) -> str:
        parts = [f"стр. {self.page}" if self.page else PAGE_UNKNOWN]
        if self.section_number or self.section_title:
            section = "раздел"
            if self.section_number:
                section += f" {self.section_number}"
            if self.section_title:
                section += f' "{truncate(self.section_title, 80)}"'
            parts.append(section)
        if self.object_type:
            label = self.object_type
            if self.object_number:
                label += f" {self.object_number}"
            parts.append(label)
        if self.detail:
            parts.append(self.detail)
        if self.search_text:
            parts.append(f"поиск: {truncate(self.search_text, 80)}")
        return "; ".join(parts)

    def to_json(self) -> dict[str, Any]:
        return {
            "file": self.file,
            "page": self.page,
            "section_number": self.section_number,
            "section_title": self.section_title,
            "object_type": self.object_type,
            "object_number": self.object_number,
            "detail": self.detail,
            "search_text": self.search_text,
            "debug_path": self.debug_path,
            "display": self.display(),
        }


@dataclass
class Finding:
    severity: str
    code: str
    cls: str
    object: str
    location: Location
    evidence: str
    why_it_matters: str
    suggested_fix: str
    autofix: str = "none"

    def to_json(self) -> dict[str, Any]:
        return {
            "severity": self.severity,
            "code": self.code,
            "class": self.cls,
            "object": self.object,
            "location": self.location.to_json(),
            "evidence": self.evidence,
            "why_it_matters": self.why_it_matters,
            "suggested_fix": self.suggested_fix,
            "autofix": self.autofix,
        }


@dataclass
class BodyItem:
    index: int
    kind: str
    element: Any
    text: str
    page: int | None
    section_index: int
    section_number: str = ""
    section_title: str = ""
    paragraph_no: int = 0
    table_no: int = 0
    figure_no: int = 0
    drawing_count: int = 0
    style_id: str = ""
    style_name: str = ""
    num_key: tuple[str, str] | None = None


@dataclass
class Caption:
    kind: str
    number: str
    raw_dots: str
    item: BodyItem


@dataclass
class FileScan:
    path: Path
    rel_path: str
    status: str = "OK"
    error: str = ""
    paragraphs: int = 0
    tables: int = 0
    figures: int = 0
    sections: int = 0
    findings: list[Finding] = field(default_factory=list)


def _local_name(element: Any) -> str:
    tag = getattr(element, "tag", "")
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _w_attr(name: str) -> str:
    return f"{{{W_NS}}}{name}"


def _w_val(element: Any, name: str = "val") -> str:
    if element is None:
        return ""
    return element.get(_w_attr(name)) or element.get(name) or ""


def _toggle_on(element: Any) -> bool:
    value = _w_val(element).strip().lower()
    return value not in {"0", "false", "off", "none"}


def _text_from_element(element: Any) -> str:
    return norm_space("".join(node.text or "" for node in element.iter(qn("w:t"))))


def _short_search(text: str) -> str:
    return truncate(norm_space(text), MAX_SEARCH_TEXT)


def _style_name_by_id(document: Document) -> dict[str, str]:
    names: dict[str, str] = {}
    for style in document.styles:
        style_id = getattr(style, "style_id", "") or ""
        if style_id:
            names[style_id] = getattr(style, "name", "") or style_id
    return names


def _paragraph_style_id(p: Any) -> str:
    p_pr = p.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    return _w_val(p_style) if p_style is not None else ""


def _paragraph_num_key(p: Any) -> tuple[str, str] | None:
    p_pr = p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    ilvl = num_pr.find(qn("w:ilvl"))
    num_id = num_pr.find(qn("w:numId"))
    return (_w_val(ilvl), _w_val(num_id))


def _paragraph_section_break(p: Any) -> bool:
    p_pr = p.find(qn("w:pPr"))
    return bool(p_pr is not None and p_pr.find(qn("w:sectPr")) is not None)


def _page_break_count(p: Any) -> int:
    explicit = 0
    for br in p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            explicit += 1
    rendered = sum(1 for _ in p.iter(qn("w:lastRenderedPageBreak")))
    return explicit + rendered


def _drawing_count(p: Any) -> int:
    return sum(1 for _ in p.iter(qn("w:drawing"))) + sum(1 for _ in p.iter(qn("w:pict")))


def _heading_from_paragraph(text: str, style_name: str) -> tuple[str, str] | None:
    if not text:
        return None
    lower_style = style_name.lower()
    is_heading_style = "heading" in lower_style or "заголовок" in lower_style
    dotted = DOTTED_HEADING_RE.match(text)
    single = SINGLE_HEADING_RE.match(text)
    match = dotted or single
    if match:
        title = match.group("title").strip()
        single_heading_like = bool(single and title[:1].isupper() and len(title) >= 12)
        if len(title) <= 180 and (bool(dotted) or is_heading_style or single_heading_like):
            return match.group("num").rstrip("."), title
    appendix = APPENDIX_HEADING_RE.match(text)
    if appendix:
        return appendix.group("num").strip(), appendix.group("title").strip()
    if is_heading_style and 0 < len(text) <= 180:
        return "", text
    return None


def _iter_body_items(document: Document) -> list[BodyItem]:
    style_names = _style_name_by_id(document)
    items: list[BodyItem] = []
    current_page = 1
    current_section = 0
    max_section = max(0, len(document.sections) - 1)
    current_section_number = ""
    current_section_title = ""
    paragraph_no = 0
    table_no = 0
    figure_no = 0

    for child in document._body._element.iterchildren():
        tag = _local_name(child)
        if tag == "p":
            paragraph_no += 1
            text = _text_from_element(child)
            style_id = _paragraph_style_id(child)
            style_name = style_names.get(style_id, style_id)
            num_key = _paragraph_num_key(child)
            heading = _heading_from_paragraph(text, style_name)
            if heading:
                current_section_number, current_section_title = heading
            drawings = _drawing_count(child)
            if drawings:
                figure_no += drawings
            items.append(
                BodyItem(
                    index=len(items),
                    kind="paragraph",
                    element=child,
                    text=text,
                    page=current_page,
                    section_index=current_section,
                    section_number=current_section_number,
                    section_title=current_section_title,
                    paragraph_no=paragraph_no,
                    figure_no=figure_no if drawings else 0,
                    drawing_count=drawings,
                    style_id=style_id,
                    style_name=style_name,
                    num_key=num_key,
                )
            )
            current_page += _page_break_count(child)
            if _paragraph_section_break(child) and current_section < max_section:
                current_section += 1
            continue
        if tag == "tbl":
            table_no += 1
            items.append(
                BodyItem(
                    index=len(items),
                    kind="table",
                    element=child,
                    text="",
                    page=current_page,
                    section_index=current_section,
                    section_number=current_section_number,
                    section_title=current_section_title,
                    table_no=table_no,
                )
            )
    return items


def _location_from_item(
    rel_path: str,
    item: BodyItem | None,
    *,
    object_type: str = "",
    object_number: str = "",
    detail: str = "",
    search_text: str = "",
    debug_path: str = "",
) -> Location:
    return Location(
        file=rel_path,
        page=item.page if item else None,
        section_number=item.section_number if item else "",
        section_title=item.section_title if item else "",
        object_type=object_type,
        object_number=object_number,
        detail=detail,
        search_text=_short_search(search_text or (item.text if item else "")),
        debug_path=debug_path,
    )


def _location_for_section(rel_path: str, section_index: int, items: list[BodyItem], detail: str) -> Location:
    item = next((candidate for candidate in items if candidate.section_index == section_index), None)
    return _location_from_item(
        rel_path,
        item,
        object_type="секция",
        object_number=str(section_index + 1),
        detail=detail,
        debug_path=f"section[{section_index}]",
    )


def _caption_from_item(item: BodyItem) -> Caption | None:
    table = TABLE_CAPTION_RE.match(item.text)
    if table:
        return Caption("table", table.group("num"), table.group("dots") or "", item)
    figure = FIGURE_CAPTION_RE.match(item.text)
    if figure:
        return Caption("figure", figure.group("num"), figure.group("dots") or "", item)
    return None


def _caption_display(kind: str, number: str) -> str:
    return ("Таблица" if kind == "table" else "Рисунок") + f" {number}"


def _number_tuple(value: str) -> tuple[int, ...] | None:
    try:
        parts = tuple(int(part) for part in value.strip(".").split(".") if part != "")
    except ValueError:
        return None
    return parts or None


def _nearest_caption(table_item: BodyItem, captions: list[Caption], max_distance: int = 4) -> tuple[Caption, str, int] | None:
    candidates: list[tuple[int, Caption, str]] = []
    for caption in captions:
        if caption.kind != "table":
            continue
        distance = abs(caption.item.index - table_item.index)
        if distance <= max_distance:
            side = "before" if caption.item.index < table_item.index else "after"
            candidates.append((distance, caption, side))
    if not candidates:
        return None
    distance, caption, side = sorted(candidates, key=lambda item: item[0])[0]
    return caption, side, distance


def _nearest_caption_for_kind(
    object_item: BodyItem,
    captions: list[Caption],
    kind: str,
    max_distance: int = 4,
) -> tuple[Caption, str, int] | None:
    candidates: list[tuple[int, Caption, str]] = []
    for caption in captions:
        if caption.kind != kind:
            continue
        distance = abs(caption.item.index - object_item.index)
        if distance <= max_distance:
            side = "before" if caption.item.index < object_item.index else "after"
            candidates.append((distance, caption, side))
    if not candidates:
        return None
    distance, caption, side = sorted(candidates, key=lambda item: item[0])[0]
    return caption, side, distance


def _scan_caption_numbers(rel_path: str, captions: list[Caption]) -> list[Finding]:
    findings: list[Finding] = []
    for kind in ("table", "figure"):
        kind_captions = [caption for caption in captions if caption.kind == kind]
        label = "таблиц" if kind == "table" else "рисунков"
        display_label = "Таблица" if kind == "table" else "Рисунок"
        counts = Counter(caption.number for caption in kind_captions)
        for number, count in counts.items():
            if count <= 1:
                continue
            first = next(caption for caption in kind_captions if caption.number == number)
            findings.append(
                Finding(
                    severity="warning",
                    code=f"caption-{kind}-duplicate-number",
                    cls="Подписи таблиц и рисунков",
                    object=f"{display_label} {number}",
                    location=_location_from_item(
                        rel_path,
                        first.item,
                        object_type="подпись",
                        object_number=number,
                        detail=f"дубль номера: {count} раза",
                        search_text=first.item.text,
                        debug_path=f"body[{first.item.index}]",
                    ),
                    evidence=f"Номер {display_label.lower()} {number} встречается {count} раза.",
                    why_it_matters="Дубли номеров ломают ссылки и ручную навигацию по документу.",
                    suggested_fix=f"Проверить последовательность подписей {label}.",
                    autofix="review_required",
                )
            )

        previous_tuple: tuple[int, ...] | None = None
        previous_caption: Caption | None = None
        for caption in kind_captions:
            number = _number_tuple(caption.number)
            if caption.raw_dots and len(caption.raw_dots) > 1:
                findings.append(
                    Finding(
                        severity="warning",
                        code=f"caption-{kind}-dot-format",
                        cls="Подписи таблиц и рисунков",
                        object=f"{display_label} {caption.number}",
                        location=_location_from_item(
                            rel_path,
                            caption.item,
                            object_type="подпись",
                            object_number=caption.number,
                            detail="подозрительные точки в номере",
                            search_text=caption.item.text,
                            debug_path=f"body[{caption.item.index}]",
                        ),
                        evidence=f"После номера найдено несколько точек: `{caption.number}{caption.raw_dots}`.",
                        why_it_matters="Нестабильные точки в номерах часто указывают на ручную правку.",
                        suggested_fix="Привести формат номера к принятому в документе.",
                        autofix="review_required",
                    )
                )
            if number is None:
                continue
            if previous_tuple is not None and previous_caption is not None:
                if number < previous_tuple:
                    findings.append(
                        Finding(
                            severity="warning",
                            code=f"caption-{kind}-reverse-number",
                            cls="Подписи таблиц и рисунков",
                            object=f"{display_label} {caption.number}",
                            location=_location_from_item(
                                rel_path,
                                caption.item,
                                object_type="подпись",
                                object_number=caption.number,
                                detail=f"номер меньше предыдущего {previous_caption.number}",
                                search_text=caption.item.text,
                                debug_path=f"body[{caption.item.index}]",
                            ),
                            evidence=f"{display_label} {caption.number} идёт после {display_label} {previous_caption.number}.",
                            why_it_matters="Обратный ход номера почти всегда выглядит как сбой ручной нумерации.",
                            suggested_fix="Проверить последовательность номеров и связанные поля SEQ/REF.",
                            autofix="review_required",
                        )
                    )
                elif (
                    len(number) == len(previous_tuple)
                    and len(number) >= 1
                    and number[:-1] == previous_tuple[:-1]
                    and number[-1] - previous_tuple[-1] > 1
                ):
                    findings.append(
                        Finding(
                            severity="warning",
                            code=f"caption-{kind}-number-gap",
                            cls="Подписи таблиц и рисунков",
                            object=f"{display_label} {caption.number}",
                            location=_location_from_item(
                                rel_path,
                                caption.item,
                                object_type="подпись",
                                object_number=caption.number,
                                detail=f"пропуск после {previous_caption.number}",
                                search_text=caption.item.text,
                                debug_path=f"body[{caption.item.index}]",
                            ),
                            evidence=f"После {display_label.lower()} {previous_caption.number} идёт {display_label.lower()} {caption.number}.",
                            why_it_matters="Пропуск номера может означать удалённый объект или ручную ошибку.",
                            suggested_fix="Проверить, есть ли отсутствующий объект или нужно перенумеровать подписи.",
                            autofix="review_required",
                        )
                    )
            previous_tuple = number
            previous_caption = caption
    return findings


def _scan_caption_proximity(rel_path: str, items: list[BodyItem], captions: list[Caption]) -> list[Finding]:
    findings: list[Finding] = []
    table_items = [item for item in items if item.kind == "table"]
    if not captions or not any(caption.kind == "table" for caption in captions):
        return findings

    paired_sides: list[str] = []
    table_pairs: dict[int, tuple[Caption, str, int]] = {}
    for table_item in table_items:
        nearest = _nearest_caption(table_item, captions)
        if nearest is None:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-caption-missing",
                    cls="Подписи таблиц и рисунков",
                    object=f"Таблица {table_item.table_no}",
                    location=_location_from_item(
                        rel_path,
                        table_item,
                        object_type="Таблица",
                        object_number=str(table_item.table_no),
                        detail="нет близкой подписи",
                        debug_path=f"body[{table_item.index}]",
                    ),
                    evidence="В документе есть подписанные таблицы, но рядом с этой таблицей подпись не найдена.",
                    why_it_matters="Таблица без подписи среди подписанных таблиц выглядит как пропуск.",
                    suggested_fix="Добавить подпись или проверить, не оторвалась ли она на другую страницу.",
                    autofix="review_required",
                )
            )
            continue
        caption, side, distance = nearest
        paired_sides.append(side)
        table_pairs[table_item.table_no] = nearest
        if distance > 2:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-caption-far",
                    cls="Подписи таблиц и рисунков",
                    object=f"Таблица {table_item.table_no}",
                    location=_location_from_item(
                        rel_path,
                        table_item,
                        object_type="Таблица",
                        object_number=str(table_item.table_no),
                        detail=f"подпись {caption.number} далеко: {distance} блок(а)",
                        search_text=caption.item.text,
                        debug_path=f"body[{table_item.index}]",
                    ),
                    evidence=f"Ближайшая подпись `{truncate(caption.item.text, 120)}` находится не вплотную к таблице.",
                    why_it_matters="Оторванная подпись затрудняет проверку и может уехать на другую страницу.",
                    suggested_fix="Поставить подпись непосредственно рядом с таблицей.",
                    autofix="review_required",
                )
            )

    if not paired_sides:
        return findings
    expected_side, expected_count = Counter(paired_sides).most_common(1)[0]
    if expected_count < 2:
        return findings
    for table_no, (caption, side, _distance) in table_pairs.items():
        if side == expected_side:
            continue
        table_item = next((item for item in table_items if item.table_no == table_no), None)
        findings.append(
            Finding(
                severity="warning",
                code="table-caption-side-inconsistent",
                cls="Подписи таблиц и рисунков",
                object=f"Таблица {table_no}",
                location=_location_from_item(
                    rel_path,
                    table_item,
                    object_type="Таблица",
                    object_number=str(table_no),
                    detail="подпись с другой стороны",
                    search_text=caption.item.text,
                    debug_path=f"body[{table_item.index}]" if table_item else "",
                ),
                evidence=f"В документе подписи обычно {'до' if expected_side == 'before' else 'после'} таблицы, а здесь иначе.",
                why_it_matters="Смена положения подписи часто означает ручной сбой в оформлении.",
                suggested_fix="Привести положение подписи к принятому в документе.",
                autofix="review_required",
            )
        )
    return findings


def _scan_figure_caption_proximity(rel_path: str, items: list[BodyItem], captions: list[Caption]) -> list[Finding]:
    findings: list[Finding] = []
    figure_items = [item for item in items if item.kind == "paragraph" and item.drawing_count > 0]
    if not figure_items or not any(caption.kind == "figure" for caption in captions):
        return findings

    paired_sides: list[str] = []
    figure_pairs: dict[int, tuple[Caption, str, int]] = {}
    for figure_item in figure_items:
        nearest = _nearest_caption_for_kind(figure_item, captions, "figure")
        figure_number = str(figure_item.figure_no or len(figure_pairs) + 1)
        if nearest is None:
            findings.append(
                Finding(
                    severity="warning",
                    code="figure-caption-missing",
                    cls="Подписи таблиц и рисунков",
                    object=f"Рисунок {figure_number}",
                    location=_location_from_item(
                        rel_path,
                        figure_item,
                        object_type="Рисунок",
                        object_number=figure_number,
                        detail="нет близкой подписи",
                        debug_path=f"paragraph[{figure_item.paragraph_no}]/drawing",
                    ),
                    evidence="В документе есть подписанные рисунки, но рядом с этим рисунком подпись не найдена.",
                    why_it_matters="Рисунок без подписи среди подписанных рисунков выглядит как пропуск.",
                    suggested_fix="Добавить подпись или проверить, не оторвалась ли она на другую страницу.",
                    autofix="review_required",
                )
            )
            continue
        caption, side, distance = nearest
        paired_sides.append(side)
        figure_pairs[figure_item.figure_no or len(figure_pairs) + 1] = nearest
        if distance > 2:
            findings.append(
                Finding(
                    severity="warning",
                    code="figure-caption-far",
                    cls="Подписи таблиц и рисунков",
                    object=f"Рисунок {figure_number}",
                    location=_location_from_item(
                        rel_path,
                        figure_item,
                        object_type="Рисунок",
                        object_number=figure_number,
                        detail=f"подпись {caption.number} далеко: {distance} блок(а)",
                        search_text=caption.item.text,
                        debug_path=f"paragraph[{figure_item.paragraph_no}]/drawing",
                    ),
                    evidence=f"Ближайшая подпись `{truncate(caption.item.text, 120)}` находится не вплотную к рисунку.",
                    why_it_matters="Оторванная подпись затрудняет проверку и может уехать на другую страницу.",
                    suggested_fix="Поставить подпись непосредственно рядом с рисунком.",
                    autofix="review_required",
                )
            )

    if not paired_sides:
        return findings
    expected_side, expected_count = Counter(paired_sides).most_common(1)[0]
    if expected_count < 2:
        return findings
    for figure_no, (caption, side, _distance) in figure_pairs.items():
        if side == expected_side:
            continue
        figure_item = next((item for item in figure_items if item.figure_no == figure_no), None)
        findings.append(
            Finding(
                severity="warning",
                code="figure-caption-side-inconsistent",
                cls="Подписи таблиц и рисунков",
                object=f"Рисунок {figure_no}",
                location=_location_from_item(
                    rel_path,
                    figure_item,
                    object_type="Рисунок",
                    object_number=str(figure_no),
                    detail="подпись с другой стороны",
                    search_text=caption.item.text,
                    debug_path=f"paragraph[{figure_item.paragraph_no}]/drawing" if figure_item else "",
                ),
                evidence=f"В документе подписи обычно {'до' if expected_side == 'before' else 'после'} рисунка, а здесь иначе.",
                why_it_matters="Смена положения подписи часто означает ручной сбой в оформлении.",
                suggested_fix="Привести положение подписи к принятому в документе.",
                autofix="review_required",
            )
        )
    return findings


def _row_elements(table: Any) -> list[Any]:
    return [child for child in table._tbl.iterchildren() if _local_name(child) == "tr"]


def _has_repeating_header(table: Any, header_rows: int) -> bool:
    for row in _row_elements(table)[: max(0, header_rows)]:
        tr_pr = row.find(qn("w:trPr"))
        if tr_pr is None:
            continue
        for header in tr_pr.findall(qn("w:tblHeader")):
            if _toggle_on(header):
                return True
    return False


def _fixed_height_rows(table: Any) -> tuple[int, int]:
    exact_rows = 0
    at_least_rows = 0
    for row in _row_elements(table):
        tr_pr = row.find(qn("w:trPr"))
        if tr_pr is None:
            continue
        for height in tr_pr.findall(qn("w:trHeight")):
            rule = (_w_val(height, "hRule") or "").lower()
            if rule == "exact":
                exact_rows += 1
            elif rule == "atleast":
                at_least_rows += 1
    return exact_rows, at_least_rows


def _nowrap_count(table: Any) -> int:
    return sum(1 for _ in table._tbl.iter(qn("w:noWrap")))


def _border_signature_counts(table: Any) -> Counter[tuple[str, str, str]]:
    counts: Counter[tuple[str, str, str]] = Counter()
    for borders_name in ("tblBorders", "tcBorders"):
        for borders in table._tbl.iter(qn(f"w:{borders_name}")):
            for border in borders:
                local = _local_name(border)
                if local not in {"top", "left", "bottom", "right", "insideH", "insideV", "start", "end"}:
                    continue
                value = _w_val(border) or "none"
                size = _w_val(border, "sz") or ""
                color = (_w_val(border, "color") or "").lower()
                counts[(value, size, color)] += 1
    return counts


def _cell_margin_signatures(table: Any) -> Counter[tuple[tuple[str, str], ...]]:
    counts: Counter[tuple[tuple[str, str], ...]] = Counter()
    for margin_name in ("tblCellMar", "tcMar"):
        for margin in table._tbl.iter(qn(f"w:{margin_name}")):
            sides: list[tuple[str, str]] = []
            for side in margin:
                local = _local_name(side)
                if local not in {"top", "bottom", "left", "right", "start", "end"}:
                    continue
                sides.append((local, _w_val(side, "w") or ""))
            if sides:
                counts[tuple(sorted(sides))] += 1
    return counts


def _table_font_stats(table: Any) -> tuple[Counter[str], Counter[float]]:
    fonts: Counter[str] = Counter()
    sizes: Counter[float] = Counter()
    for r_pr in table._tbl.iter(qn("w:rPr")):
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                value = r_fonts.get(_w_attr(attr))
                if value:
                    fonts[value] += 1
        size = r_pr.find(qn("w:sz"))
        if size is None:
            size = r_pr.find(qn("w:szCs"))
        raw_size = _w_val(size) if size is not None else ""
        if raw_size:
            try:
                sizes[float(raw_size) / 2.0] += 1
            except ValueError:
                pass
    return fonts, sizes


def _empty_rows_cols(model: TableModel) -> tuple[list[int], list[int]]:
    empty_rows: list[int] = []
    empty_cols: list[int] = []
    if not model.display_grid or model.width <= 0:
        return empty_rows, empty_cols
    for row_idx, row in enumerate(model.display_grid, start=1):
        if all(not normalize_text(cell) for cell in row):
            empty_rows.append(row_idx)
    for col_idx in range(model.width):
        values = [row[col_idx] if col_idx < len(row) else "" for row in model.display_grid]
        if values and all(not normalize_text(value) for value in values):
            empty_cols.append(col_idx + 1)
    return empty_rows, empty_cols


def _body_merged_cells(model: TableModel) -> list[Any]:
    return [origin for origin in model.origins if origin.is_merged and origin.row >= max(1, model.header_row_count)]


def _density_metrics(model: TableModel) -> tuple[float, int]:
    texts = [normalize_text(cell) for row in model.display_grid for cell in row if normalize_text(cell)]
    if not texts:
        return 0.0, 0
    avg = sum(len(text) for text in texts) / max(len(texts), 1)
    return avg, max(len(text) for text in texts)


def _scan_tables(
    rel_path: str,
    document: Document,
    items: list[BodyItem],
    *,
    max_findings: int,
) -> tuple[list[Finding], dict[int, dict[str, float]]]:
    findings: list[Finding] = []
    table_items = {item.table_no: item for item in items if item.kind == "table"}
    table_section_map = _body_table_section_map(document)
    fallback_width = _fallback_total_width_cm(document)
    table_widths_by_section: dict[int, dict[str, float]] = defaultdict(lambda: {"count": 0.0, "max_width": 0.0})

    for table_no, table in enumerate(document.tables, start=1):
        if len(findings) >= max_findings:
            break
        item = table_items.get(table_no)
        section_index = table_section_map.get(id(table._tbl), item.section_index if item else 0)
        available_width = _section_available_width_cm(document, table, table_section_map)
        total_width = _read_table_total_width_cm(table, fallback_width)
        stats = table_widths_by_section[section_index]
        stats["count"] += 1
        stats["max_width"] = max(stats["max_width"], total_width)
        loc = _location_from_item(
            rel_path,
            item,
            object_type="Таблица",
            object_number=str(table_no),
            debug_path=f"table[{table_no}]",
        )

        overflow_width = total_width - available_width
        if overflow_width > 0.3:
            findings.append(
                Finding(
                    severity="critical" if overflow_width >= 0.8 else "warning",
                    code="table-width-over-margins",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="шире полезной области",
                        debug_path=f"table[{table_no}]/width",
                    ),
                    evidence=f"Ширина таблицы примерно {total_width:.1f} см, полезная ширина секции {available_width:.1f} см, превышение {overflow_width:.1f} см.",
                    why_it_matters="Такая таблица может быть обрезана или съехать за поля страницы.",
                    suggested_fix="Вписать таблицу в поля текущей секции или изменить ориентацию страницы.",
                    autofix="safe",
                )
            )

        try:
            model = build_table_model(table, table_no)
        except Exception as exc:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-model-failed",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=loc,
                    evidence=f"Не удалось построить модель таблицы: {exc}",
                    why_it_matters="Без модели нельзя проверить пустые строки, объединения и плотность.",
                    suggested_fix="Проверить таблицу вручную.",
                    autofix="none",
                )
            )
            continue

        empty_rows, empty_cols = _empty_rows_cols(model)
        if empty_rows:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-empty-rows",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="полностью пустые строки",
                        debug_path=f"table[{table_no}]",
                    ),
                    evidence=f"Полностью пустые строки: {', '.join(map(str, empty_rows[:12]))}.",
                    why_it_matters="Пустые строки обычно являются следом ручной верстки или разрыва таблицы.",
                    suggested_fix="Удалить пустые строки, если они не несут смысловой роли.",
                    autofix="review_required",
                )
            )
        if empty_cols:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-empty-columns",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="полностью пустые колонки",
                        debug_path=f"table[{table_no}]",
                    ),
                    evidence=f"Полностью пустые колонки: {', '.join(map(str, empty_cols[:12]))}.",
                    why_it_matters="Пустые колонки часто появляются после конвертации или ручной правки.",
                    suggested_fix="Удалить пустые колонки после проверки структуры шапки.",
                    autofix="review_required",
                )
            )

        body_merges = _body_merged_cells(model)
        if body_merges and len(body_merges) <= 4:
            sample = body_merges[0]
            findings.append(
                Finding(
                    severity="warning",
                    code="table-suspicious-body-merge",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail=f"объединение в теле: R{sample.row + 1}C{sample.col + 1}",
                        search_text=sample.text,
                        debug_path=f"table[{table_no}]/merge[{sample.row},{sample.col}]",
                    ),
                    evidence=f"Найдено {len(body_merges)} объединение(й) вне шапки таблицы.",
                    why_it_matters="Единичные объединения в теле часто ломают регулярную сетку данных.",
                    suggested_fix="Проверить, действительно ли объединение нужно в теле таблицы.",
                    autofix="review_required",
                )
            )

        exact_rows, _at_least_rows = _fixed_height_rows(table)
        no_wrap = _nowrap_count(table)
        if exact_rows:
            findings.append(
                Finding(
                    severity="critical" if no_wrap else "warning",
                    code="table-fixed-row-height",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="фиксированная высота строк",
                        debug_path=f"table[{table_no}]/rows",
                    ),
                    evidence=f"Строк с точной высотой: {exact_rows}; ячеек с запретом переноса текста: {no_wrap}.",
                    why_it_matters="Фиксированная высота может визуально резать текст, особенно без переноса.",
                    suggested_fix="Снять точную высоту строк и включить перенос текста.",
                    autofix="safe",
                )
            )

        border_counts = _border_signature_counts(table)
        if len(border_counts) > 2:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-border-mixed",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="разные границы",
                        debug_path=f"table[{table_no}]/borders",
                    ),
                    evidence=f"Найдено {len(border_counts)} вариантов границ: {', '.join('/'.join(sig) for sig, _ in border_counts.most_common(3))}.",
                    why_it_matters="Разные толщины и цвета границ внутри одной таблицы выглядят как сбой оформления.",
                    suggested_fix="Унифицировать границы таблицы.",
                    autofix="safe",
                )
            )

        margin_counts = _cell_margin_signatures(table)
        if len(margin_counts) > 2:
            findings.append(
                Finding(
                    severity="warning",
                    code="table-cell-margins-mixed",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="разные поля ячеек",
                        debug_path=f"table[{table_no}]/cellMargins",
                    ),
                    evidence=f"Найдено {len(margin_counts)} вариантов внутренних полей ячеек.",
                    why_it_matters="Разные поля ячеек меняют плотность и визуальный ритм таблицы.",
                    suggested_fix="Унифицировать поля ячеек.",
                    autofix="safe",
                )
            )

        fonts, sizes = _table_font_stats(table)
        if len(fonts) > 3 or (sizes and max(sizes) - min(sizes) >= 3.0):
            evidence_parts = []
            if fonts:
                evidence_parts.append("шрифты: " + ", ".join(f"{name} x{count}" for name, count in fonts.most_common(4)))
            if sizes:
                evidence_parts.append(f"кегли: {min(sizes):g}-{max(sizes):g} pt")
            findings.append(
                Finding(
                    severity="warning",
                    code="table-fonts-mixed",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="смесь шрифтов/кеглей",
                        debug_path=f"table[{table_no}]/fonts",
                    ),
                    evidence="; ".join(evidence_parts) or "Найдена неоднородная прямая разметка шрифтов.",
                    why_it_matters="Разнобой шрифтов внутри таблицы обычно является следом ручного форматирования.",
                    suggested_fix="Прогнать унификацию таблиц или проверить исключения вручную.",
                    autofix="safe",
                )
            )

        avg_density, max_cell = _density_metrics(model)
        if model.width >= 8 and (avg_density >= 28 or max_cell >= 140):
            findings.append(
                Finding(
                    severity="info",
                    code="table-dense-cells",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="плотные ячейки",
                        debug_path=f"table[{table_no}]/density",
                    ),
                    evidence=f"Колонок: {model.width}; средняя длина непустой ячейки: {avg_density:.1f}; максимум: {max_cell}.",
                    why_it_matters="Плотные таблицы часто требуют меньший кегль, баланс колонок или альбомную ориентацию.",
                    suggested_fix="Проверить баланс колонок и режим унификации таблиц по плотности.",
                    autofix="safe",
                )
            )

        if model.height >= 25 and model.header_row_count > 0 and not _has_repeating_header(table, model.header_row_count):
            findings.append(
                Finding(
                    severity="warning",
                    code="table-header-not-repeated",
                    cls="Таблицы",
                    object=f"Таблица {table_no}",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="Таблица",
                        object_number=str(table_no),
                        detail="шапка не повторяется",
                        debug_path=f"table[{table_no}]/header",
                    ),
                    evidence=f"Строк: {model.height}; распознанных строк шапки: {model.header_row_count}; повторяемая шапка не включена.",
                    why_it_matters="У многостраничной таблицы без повторяемой шапки теряется контекст на следующих страницах.",
                    suggested_fix="Включить повтор шапки по строгой эвристике строк перед шапкой.",
                    autofix="review_required",
                )
            )
    return findings[:max_findings], table_widths_by_section


def _collect_table_widths_by_section(document: Document) -> dict[int, dict[str, float]]:
    table_section_map = _body_table_section_map(document)
    fallback_width = _fallback_total_width_cm(document)
    table_widths_by_section: dict[int, dict[str, float]] = defaultdict(lambda: {"count": 0.0, "max_width": 0.0})
    for table in document.tables:
        section_index = table_section_map.get(id(table._tbl), 0)
        total_width = _read_table_total_width_cm(table, fallback_width)
        stats = table_widths_by_section[section_index]
        stats["count"] += 1
        stats["max_width"] = max(stats["max_width"], total_width)
    return table_widths_by_section


def _section_margin_profile(section: Any) -> tuple[int, int, int, int]:
    return (
        round(section.top_margin.cm * 10),
        round(section.right_margin.cm * 10),
        round(section.bottom_margin.cm * 10),
        round(section.left_margin.cm * 10),
    )


def _scan_layout(
    rel_path: str,
    document: Document,
    items: list[BodyItem],
    table_widths_by_section: dict[int, dict[str, float]],
) -> list[Finding]:
    findings: list[Finding] = []
    sections = list(document.sections)
    if not sections:
        return findings

    profiles = [_section_margin_profile(section) for section in sections]
    common_profile = Counter(profiles).most_common(1)[0][0] if profiles else None
    for section_index, section in enumerate(sections):
        is_landscape = section.page_width > section.page_height
        stats = table_widths_by_section.get(section_index, {"count": 0.0, "max_width": 0.0})
        max_table_width = stats.get("max_width", 0.0)
        if is_landscape and max_table_width <= 20.0:
            findings.append(
                Finding(
                    severity="warning",
                    code="landscape-section-without-wide-table",
                    cls="Секции, страницы и ориентация",
                    object=f"Секция {section_index + 1}",
                    location=_location_for_section(
                        rel_path,
                        section_index,
                        items,
                        "альбомная ориентация без широкой таблицы",
                    ),
                    evidence=f"Секция альбомная; максимальная ширина таблицы примерно {max_table_width:.1f} см.",
                    why_it_matters="Альбомная страница без широкой таблицы или рисунка часто остаётся после ручной правки.",
                    suggested_fix="Проверить, нужна ли альбомная ориентация этой секции.",
                    autofix="review_required",
                )
            )
        if common_profile is not None and profiles[section_index] != common_profile and max_table_width <= 20.0:
            profile_cm = tuple(value / 10 for value in profiles[section_index])
            findings.append(
                Finding(
                    severity="warning",
                    code="section-margin-profile-differs",
                    cls="Секции, страницы и ориентация",
                    object=f"Секция {section_index + 1}",
                    location=_location_for_section(
                        rel_path,
                        section_index,
                        items,
                        "поля отличаются от основного профиля",
                    ),
                    evidence=f"Поля секции (верх/право/низ/лево): {profile_cm}; широкой таблицы как причины не найдено.",
                    why_it_matters="Разные поля без видимой причины могут сдвинуть колонтитулы, номера страниц и таблицы.",
                    suggested_fix="Сравнить поля секции с соседними секциями.",
                    autofix="review_required",
                )
            )
    return findings


def _header_footer_text(section: Any, attr: str) -> str:
    try:
        part = getattr(section, attr)
        return norm_space(" ".join(paragraph.text for paragraph in part.paragraphs if paragraph.text))
    except Exception:
        return ""


def _section_page_start(section: Any) -> str:
    try:
        pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    except Exception:
        return ""
    return _w_val(pg_num_type, "start") if pg_num_type is not None else ""


def _scan_headers_footers(
    rel_path: str,
    document: Document,
    items: list[BodyItem],
    table_widths_by_section: dict[int, dict[str, float]],
) -> list[Finding]:
    findings: list[Finding] = []
    sections = list(document.sections)
    if not sections:
        return findings

    page_starts = [_section_page_start(section) for section in sections]
    for section_index, start in enumerate(page_starts):
        if not start:
            continue
        if section_index == 0 and start != "1":
            findings.append(
                Finding(
                    severity="warning",
                    code="page-number-first-section-start",
                    cls="Секции, страницы и ориентация",
                    object="Нумерация страниц",
                    location=_location_for_section(
                        rel_path,
                        section_index,
                        items,
                        f"нумерация начинается с {start}",
                    ),
                    evidence=f"В первой секции задан старт номера страницы: {start}.",
                    why_it_matters="Номер страницы на первом листе часто уезжает из-за ручной настройки секции.",
                    suggested_fix="Проверить параметры нумерации страниц в первой секции.",
                    autofix="review_required",
                )
            )
        elif section_index > 0:
            findings.append(
                Finding(
                    severity="warning",
                    code="page-number-section-restart",
                    cls="Секции, страницы и ориентация",
                    object="Нумерация страниц",
                    location=_location_for_section(
                        rel_path,
                        section_index,
                        items,
                        f"перезапуск нумерации с {start}",
                    ),
                    evidence=f"В секции {section_index + 1} явно задан старт номера страницы: {start}.",
                    why_it_matters="Перезапуск нумерации может быть намеренным, но часто выглядит как сбой после разрыва секции.",
                    suggested_fix="Проверить, должен ли номер страницы продолжаться от предыдущей секции.",
                    autofix="review_required",
                )
            )

    if len(sections) < 2:
        return findings

    for attr, label in (
        ("header", "верхний колонтитул"),
        ("footer", "нижний колонтитул"),
        ("first_page_header", "верхний колонтитул первой страницы"),
        ("first_page_footer", "нижний колонтитул первой страницы"),
    ):
        profiles = [_header_footer_text(section, attr) for section in sections]
        non_empty = [profile for profile in profiles if profile]
        if len(set(non_empty)) <= 1:
            continue
        common = Counter(non_empty).most_common(1)[0][0] if non_empty else ""
        for section_index, profile in enumerate(profiles):
            if not profile or profile == common:
                continue
            stats = table_widths_by_section.get(section_index, {"max_width": 0.0})
            findings.append(
                Finding(
                    severity="warning",
                    code="header-footer-profile-differs",
                    cls="Секции, страницы и ориентация",
                    object="Колонтитул",
                    location=_location_for_section(
                        rel_path,
                        section_index,
                        items,
                        f"{label} отличается",
                    ),
                    evidence=f"{label.capitalize()} секции отличается от основного профиля: `{truncate(profile, 140)}`.",
                    why_it_matters="Съехавший или отвязанный колонтитул часто появляется после разрыва секции.",
                    suggested_fix="Проверить связь колонтитула с предыдущей секцией и параметры страницы.",
                    autofix="review_required",
                )
            )
            if stats.get("max_width", 0.0) <= 20.0:
                break
    return findings


def _scan_spacing(rel_path: str, items: list[BodyItem]) -> list[Finding]:
    findings: list[Finding] = []
    run_start: BodyItem | None = None
    run_count = 0
    for item in items:
        if item.kind != "paragraph":
            if run_count >= 2 and run_start is not None:
                findings.append(_blank_run_finding(rel_path, run_start, run_count))
            run_start = None
            run_count = 0
            continue
        if item.text:
            if run_count >= 2 and run_start is not None:
                findings.append(_blank_run_finding(rel_path, run_start, run_count))
            run_start = None
            run_count = 0
            continue
        if run_start is None:
            run_start = item
        run_count += 1
    if run_count >= 2 and run_start is not None:
        findings.append(_blank_run_finding(rel_path, run_start, run_count))
    return findings


def _blank_run_finding(rel_path: str, item: BodyItem, count: int) -> Finding:
    return Finding(
        severity="warning",
        code="duplicated-empty-paragraphs",
        cls="Разрывы, пустоты и привязка объектов",
        object="Пустые абзацы",
        location=_location_from_item(
            rel_path,
            item,
            object_type="абзац",
            object_number=str(item.paragraph_no),
            detail=f"{count} пустых абзаца подряд",
            debug_path=f"paragraph[{item.paragraph_no}]",
        ),
        evidence=f"Найдено {count} пустых абзаца подряд.",
        why_it_matters="Две и более пустые строки подряд обычно являются маркером ручного сбоя форматирования.",
        suggested_fix="Удалить лишние пустые абзацы или заменить их настройками интервала.",
        autofix="safe",
    )


def _scan_headings(rel_path: str, items: list[BodyItem]) -> list[Finding]:
    findings: list[Finding] = []
    for index, item in enumerate(items):
        if item.kind != "paragraph" or not item.text:
            continue
        heading = DOTTED_HEADING_RE.match(item.text) or SINGLE_HEADING_RE.match(item.text) or APPENDIX_HEADING_RE.match(item.text)
        if not heading:
            continue
        next_content = next((candidate for candidate in items[index + 1 :] if candidate.kind == "table" or candidate.text), None)
        if next_content is None:
            findings.append(
                Finding(
                    severity="warning",
                    code="heading-without-following-content",
                    cls="Разрывы, пустоты и привязка объектов",
                    object="Заголовок",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="заголовок",
                        detail="нет последующего содержимого",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]",
                    ),
                    evidence=f"После заголовка `{truncate(item.text, 120)}` нет видимого содержимого.",
                    why_it_matters="Заголовок без текста после него часто остаётся в конце страницы или документа.",
                    suggested_fix="Проверить разрывы после заголовка и связать его со следующим блоком.",
                    autofix="review_required",
                )
            )
    return findings


def _field_instruction_text(element: Any) -> str:
    return norm_space(" ".join(node.text or "" for node in element.iter(qn("w:instrText"))))


def _bookmark_maps(document: Document) -> tuple[dict[str, list[Any]], dict[str, list[Any]], Counter[str]]:
    starts_by_id: dict[str, list[Any]] = defaultdict(list)
    ends_by_id: dict[str, list[Any]] = defaultdict(list)
    names: Counter[str] = Counter()
    for start in document._element.iter(qn("w:bookmarkStart")):
        bookmark_id = _w_val(start, "id")
        name = _w_val(start, "name")
        if bookmark_id:
            starts_by_id[bookmark_id].append(start)
        if name and not name.startswith("_"):
            names[name] += 1
    for end in document._element.iter(qn("w:bookmarkEnd")):
        bookmark_id = _w_val(end, "id")
        if bookmark_id:
            ends_by_id[bookmark_id].append(end)
    return starts_by_id, ends_by_id, names


def _bookmark_names(document: Document) -> set[str]:
    return {
        _w_val(start, "name")
        for start in document._element.iter(qn("w:bookmarkStart"))
        if _w_val(start, "name")
    }


def _item_containing(items: list[BodyItem], target: Any) -> BodyItem | None:
    for item in items:
        if item.element is target:
            return item
        if any(node is target for node in item.element.iter()):
            return item
    return None


def _scan_fields_and_links(rel_path: str, document: Document, items: list[BodyItem]) -> list[Finding]:
    findings: list[Finding] = []
    bookmark_names = _bookmark_names(document)
    starts_by_id, ends_by_id, name_counts = _bookmark_maps(document)

    for name, count in name_counts.items():
        if count <= 1:
            continue
        first_start = next(
            (start for start in document._element.iter(qn("w:bookmarkStart")) if _w_val(start, "name") == name),
            None,
        )
        item = _item_containing(items, first_start) if first_start is not None else None
        findings.append(
            Finding(
                severity="warning",
                code="bookmark-duplicate-name",
                cls="Оглавление, поля и ссылки",
                object="Закладка",
                location=_location_from_item(
                    rel_path,
                    item,
                    object_type="закладка",
                    object_number=name,
                    detail=f"имя повторяется {count} раза",
                    debug_path=f"bookmarkStart[name={name}]",
                ),
                evidence=f"Закладка `{name}` встречается {count} раза.",
                why_it_matters="Дубли имён закладок могут ломать REF/PAGEREF и внутренние ссылки.",
                suggested_fix="Проверить закладки в Word и оставить уникальное имя.",
                autofix="review_required",
            )
        )

    for bookmark_id, starts in starts_by_id.items():
        if bookmark_id not in ends_by_id:
            item = _item_containing(items, starts[0]) if starts else None
            name = _w_val(starts[0], "name") if starts else bookmark_id
            findings.append(
                Finding(
                    severity="warning",
                    code="bookmark-start-without-end",
                    cls="Оглавление, поля и ссылки",
                    object="Закладка",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="закладка",
                        object_number=name,
                        detail="нет конца закладки",
                        debug_path=f"bookmarkStart[id={bookmark_id}]",
                    ),
                    evidence=f"У закладки `{name}` есть start без matching end.",
                    why_it_matters="Повреждённая закладка может давать битые ссылки.",
                    suggested_fix="Проверить структуру закладки в Word.",
                    autofix="review_required",
                )
            )
    for bookmark_id, ends in ends_by_id.items():
        if bookmark_id not in starts_by_id:
            item = _item_containing(items, ends[0]) if ends else None
            findings.append(
                Finding(
                    severity="warning",
                    code="bookmark-end-without-start",
                    cls="Оглавление, поля и ссылки",
                    object="Закладка",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="закладка",
                        object_number=bookmark_id,
                        detail="нет начала закладки",
                        debug_path=f"bookmarkEnd[id={bookmark_id}]",
                    ),
                    evidence=f"У bookmarkEnd id `{bookmark_id}` нет matching start.",
                    why_it_matters="Осиротевшая закладка указывает на повреждение структуры ссылок.",
                    suggested_fix="Проверить закладки и поля REF/PAGEREF.",
                    autofix="review_required",
                )
            )

    for hyperlink in document._element.iter(qn("w:hyperlink")):
        anchor = hyperlink.get(_w_attr("anchor"))
        if not anchor or anchor in bookmark_names:
            continue
        item = _item_containing(items, hyperlink)
        findings.append(
            Finding(
                severity="warning",
                code="hyperlink-missing-anchor",
                cls="Оглавление, поля и ссылки",
                object="Ссылка",
                location=_location_from_item(
                    rel_path,
                    item,
                    object_type="ссылка",
                    object_number=anchor,
                    detail="целевая закладка не найдена",
                    search_text=_text_from_element(hyperlink),
                    debug_path=f"hyperlink[anchor={anchor}]",
                ),
                evidence=f"Внутренняя ссылка ведёт на закладку `{anchor}`, но такой закладки нет.",
                why_it_matters="Битая внутренняя ссылка не приведёт пользователя к нужному месту.",
                suggested_fix="Восстановить целевую закладку или удалить ссылку.",
                autofix="review_required",
            )
        )

    for item in items:
        if item.kind != "paragraph":
            continue
        instr = _field_instruction_text(item.element)
        if BROKEN_FIELD_RE.search(item.text):
            findings.append(
                Finding(
                    severity="critical",
                    code="field-broken-visible-result",
                    cls="Оглавление, поля и ссылки",
                    object="Поле Word",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="поле",
                        detail="видимая ошибка поля",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]/field",
                    ),
                    evidence=f"Видимый текст поля содержит ошибку: `{truncate(item.text, 160)}`.",
                    why_it_matters="Word уже показывает, что ссылка, закладка или поле повреждены.",
                    suggested_fix="Обновить поля и восстановить источник ссылки/закладку.",
                    autofix="review_required",
                )
            )
        if not instr:
            continue
        for match in FIELD_REF_RE.finditer(instr):
            name = match.group("name")
            if name in bookmark_names or name.startswith("_"):
                continue
            findings.append(
                Finding(
                    severity="warning",
                    code="field-ref-missing-bookmark",
                    cls="Оглавление, поля и ссылки",
                    object="Поле Word",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="поле",
                        object_number=name,
                        detail="REF/PAGEREF без закладки",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]/instrText",
                    ),
                    evidence=f"Инструкция поля `{truncate(instr, 140)}` ссылается на отсутствующую закладку `{name}`.",
                    why_it_matters="При обновлении поля ссылка может превратиться в ошибку.",
                    suggested_fix="Восстановить закладку или исправить REF/PAGEREF.",
                    autofix="review_required",
                )
            )
    return findings


def _toc_items(items: list[BodyItem]) -> list[BodyItem]:
    return [item for item in items if item.kind == "paragraph" and TOC_STYLE_RE.search(item.style_name or item.style_id)]


def _normalize_toc_text(text: str) -> str:
    text = re.sub(r"\s+\d+\s*$", "", norm_space(text))
    text = text.replace("\t", " ")
    return norm_space(text).casefold().replace("ё", "е")


def _scan_toc(rel_path: str, items: list[BodyItem]) -> list[Finding]:
    findings: list[Finding] = []
    toc_field_items = [item for item in items if item.kind == "paragraph" and "TOC" in _field_instruction_text(item.element).upper()]
    toc_items = _toc_items(items)
    if not toc_field_items and not toc_items:
        return findings

    headings = [
        (item.section_number, item.section_title, item)
        for item in items
        if item.kind == "paragraph" and (item.section_number or item.section_title) and _heading_from_paragraph(item.text, item.style_name)
    ]
    toc_blob = "\n".join(_normalize_toc_text(item.text) for item in toc_items if item.text)
    toc_anchor = toc_field_items[0] if toc_field_items else (toc_items[0] if toc_items else None)
    if toc_field_items and not toc_items:
        findings.append(
            Finding(
                severity="info",
                code="toc-field-without-parsed-entries",
                cls="Оглавление, поля и ссылки",
                object="Оглавление",
                location=_location_from_item(
                    rel_path,
                    toc_anchor,
                    object_type="оглавление",
                    detail="поле оглавления найдено, строки не распознаны",
                    debug_path=f"paragraph[{toc_anchor.paragraph_no}]/toc" if toc_anchor else "",
                ),
                evidence="В документе есть поле оглавления, но строки оглавления не распознаны по стилям оглавления.",
                why_it_matters="Проверка полноты оглавления ограничена без распознанных строк оглавления.",
                suggested_fix="Открыть документ в Word и обновить оглавление.",
                autofix="none",
            )
        )
        return findings

    missing: list[tuple[str, str, BodyItem]] = []
    for number, title, item in headings:
        if not number and not title:
            continue
        needle = _normalize_toc_text(f"{number} {title}")
        title_needle = _normalize_toc_text(title)
        if needle and needle in toc_blob:
            continue
        if title_needle and title_needle in toc_blob:
            continue
        missing.append((number, title, item))
    if missing:
        number, title, item = missing[0]
        findings.append(
            Finding(
                severity="warning",
                code="toc-heading-missing",
                cls="Оглавление, поля и ссылки",
                object="Оглавление",
                location=_location_from_item(
                    rel_path,
                    item,
                    object_type="заголовок",
                    object_number=number,
                    detail="не найден в оглавлении",
                    search_text=item.text,
                    debug_path=f"paragraph[{item.paragraph_no}]/heading",
                ),
                evidence=f"Заголовков, не найденных в распознанном оглавлении: {len(missing)}; первый: `{truncate(item.text, 140)}`.",
                why_it_matters="Заголовок документа может отсутствовать в оглавлении.",
                suggested_fix="Проверить стиль заголовка и обновить оглавление в Word.",
                autofix="review_required",
            )
        )

    heading_keys = {_normalize_toc_text(f"{number} {title}") for number, title, _item in headings}
    heading_titles = {_normalize_toc_text(title) for _number, title, _item in headings if title}
    stale_items: list[BodyItem] = []
    for toc_item in toc_items:
        text = _normalize_toc_text(toc_item.text)
        if not text or text in heading_keys or text in heading_titles:
            continue
        if DOTTED_HEADING_RE.match(toc_item.text) or SINGLE_HEADING_RE.match(toc_item.text):
            stale_items.append(toc_item)
    if stale_items:
        item = stale_items[0]
        findings.append(
            Finding(
                severity="warning",
                code="toc-entry-stale",
                cls="Оглавление, поля и ссылки",
                object="Оглавление",
                location=_location_from_item(
                    rel_path,
                    item,
                    object_type="оглавление",
                    detail="строка не похожа на текущий заголовок",
                    search_text=item.text,
                    debug_path=f"paragraph[{item.paragraph_no}]/toc",
                ),
                evidence=f"Строк оглавления без совпадения с текущими заголовками: {len(stale_items)}; первая: `{truncate(item.text, 140)}`.",
                why_it_matters="В оглавлении может быть заголовок, которого уже нет в документе.",
                suggested_fix="Обновить оглавление и проверить стили заголовков.",
                autofix="review_required",
            )
        )
    return findings


def _manual_number_tuple(text: str) -> tuple[int, ...] | None:
    match = MANUAL_NUMBER_RE.match(text)
    if not match:
        return None
    try:
        return tuple(int(part) for part in match.group("num").split("."))
    except ValueError:
        return None


def _scan_lists(rel_path: str, items: list[BodyItem]) -> list[Finding]:
    findings: list[Finding] = []
    manual_numbered: list[tuple[BodyItem, tuple[int, ...]]] = []
    manual_bullets: list[BodyItem] = []
    for item in items:
        if item.kind != "paragraph" or not item.text:
            continue
        if item.num_key is not None:
            continue
        if _heading_from_paragraph(item.text, item.style_name):
            continue
        number = _manual_number_tuple(item.text)
        if number:
            manual_numbered.append((item, number))
        elif MANUAL_BULLET_RE.match(item.text):
            manual_bullets.append(item)

    if len(manual_numbered) >= 2:
        first_item, first_number = manual_numbered[0]
        findings.append(
            Finding(
                severity="warning",
                code="manual-numbered-list",
                cls="Нумерация и списки",
                object="Ручная нумерация",
                location=_location_from_item(
                    rel_path,
                    first_item,
                    object_type="список",
                    object_number=".".join(map(str, first_number)),
                    detail=f"ручных номеров: {len(manual_numbered)}",
                    search_text=first_item.text,
                    debug_path=f"paragraph[{first_item.paragraph_no}]/manualNumber",
                ),
                evidence=f"Найдено {len(manual_numbered)} абзацев, похожих на ручную нумерацию без встроенной нумерации Word.",
                why_it_matters="Ручные номера легко дают пропуски, дубли и обратный ход при правке документа.",
                suggested_fix="Проверить блоки списков и при необходимости заменить ручные номера на встроенную нумерацию Word.",
                autofix="review_required",
            )
        )

    if len(manual_bullets) >= 3:
        first_item = manual_bullets[0]
        findings.append(
            Finding(
                severity="info",
                code="manual-bulleted-list",
                cls="Нумерация и списки",
                object="Ручные маркеры",
                location=_location_from_item(
                    rel_path,
                    first_item,
                    object_type="список",
                    detail=f"ручных маркеров: {len(manual_bullets)}",
                    search_text=first_item.text,
                    debug_path=f"paragraph[{first_item.paragraph_no}]/manualBullet",
                ),
                evidence=f"Найдено {len(manual_bullets)} абзацев с ручными маркерами без встроенной нумерации Word.",
                why_it_matters="Ручные маркеры могут отличаться от настоящего списка по отступам и стилям.",
                suggested_fix="Проверить оформление списков в этом блоке.",
                autofix="review_required",
            )
        )

    seen: set[tuple[int, ...]] = set()
    previous: tuple[int, ...] | None = None
    previous_item: BodyItem | None = None
    for item, number in manual_numbered:
        if len(number) > 1 and number[:-1] not in seen:
            findings.append(
                Finding(
                    severity="warning",
                    code="manual-list-orphan-subpoint",
                    cls="Нумерация и списки",
                    object="Подпункт",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="список",
                        object_number=".".join(map(str, number)),
                        detail="нет родительского пункта выше",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]/manualNumber",
                    ),
                    evidence=f"Подпункт {'.'.join(map(str, number))} встретился без родителя {'.'.join(map(str, number[:-1]))}.",
                    why_it_matters="Висячий подпункт обычно означает пропуск или сбой нумерации.",
                    suggested_fix="Проверить родительский пункт и структуру списка.",
                    autofix="review_required",
                )
            )
            break
        if number in seen:
            findings.append(
                Finding(
                    severity="warning",
                    code="manual-list-duplicate-number",
                    cls="Нумерация и списки",
                    object="Дубль номера",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="список",
                        object_number=".".join(map(str, number)),
                        detail="дубль ручного номера",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]/manualNumber",
                    ),
                    evidence=f"Номер {'.'.join(map(str, number))} уже встречался выше.",
                    why_it_matters="Дубль ручного номера ломает последовательность списка.",
                    suggested_fix="Проверить нумерацию блока.",
                    autofix="review_required",
                )
            )
            break
        if (
            previous is not None
            and previous_item is not None
            and len(number) == len(previous)
            and number[:-1] == previous[:-1]
            and number[-1] - previous[-1] > 1
        ):
            findings.append(
                Finding(
                    severity="warning",
                    code="manual-list-number-gap",
                    cls="Нумерация и списки",
                    object="Пропуск номера",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="список",
                        object_number=".".join(map(str, number)),
                        detail=f"пропуск после {'.'.join(map(str, previous))}",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]/manualNumber",
                    ),
                    evidence=f"После {'.'.join(map(str, previous))} идёт {'.'.join(map(str, number))}.",
                    why_it_matters="Пропуск номера может указывать на удалённый пункт или ручную ошибку.",
                    suggested_fix="Проверить соседние пункты списка.",
                    autofix="review_required",
                )
            )
            break
        if previous is not None and previous_item is not None and len(number) == len(previous) and number < previous:
            findings.append(
                Finding(
                    severity="warning",
                    code="manual-list-number-reverse",
                    cls="Нумерация и списки",
                    object="Обратный ход номера",
                    location=_location_from_item(
                        rel_path,
                        item,
                        object_type="список",
                        object_number=".".join(map(str, number)),
                        detail=f"после {'.'.join(map(str, previous))}",
                        search_text=item.text,
                        debug_path=f"paragraph[{item.paragraph_no}]/manualNumber",
                    ),
                    evidence=f"Номер {'.'.join(map(str, number))} идёт после {'.'.join(map(str, previous))}.",
                    why_it_matters="Обратный ход номера почти всегда похож на сбой списка.",
                    suggested_fix="Проверить порядок пунктов и формат нумерации.",
                    autofix="review_required",
                )
            )
            break
        seen.add(number)
        previous = number
        previous_item = item
    return findings


def _scan_file(
    path: Path,
    input_root: Path,
    *,
    include_tables: bool,
    include_captions: bool,
    include_layout: bool,
    include_spacing: bool,
    include_fields: bool,
    include_lists: bool,
    include_headers: bool,
    max_findings: int,
) -> FileScan:
    rel_path = path.name if path.parent == input_root and input_root.is_file() else rel_posix(path, input_root if input_root.is_dir() else path.parent)
    report = FileScan(path=path, rel_path=rel_path)
    try:
        document = Document(str(path))
    except Exception as exc:
        report.status = "ERROR"
        report.error = str(exc)
        report.findings.append(
            Finding(
                severity="critical",
                code="docx-open-failed",
                cls="Файл",
                object="DOCX",
                location=Location(file=rel_path, detail="файл не открыт"),
                evidence=f"python-docx не смог открыть файл: {exc}",
                why_it_matters="Документ нельзя проверить без чтения структуры DOCX.",
                suggested_fix="Проверить, что файл не повреждён и не является временным Office-файлом.",
                autofix="none",
            )
        )
        return report

    items = _iter_body_items(document)
    captions = [caption for item in items if (caption := _caption_from_item(item)) is not None]
    report.paragraphs = sum(1 for item in items if item.kind == "paragraph")
    report.tables = len(document.tables)
    report.figures = sum(item.drawing_count for item in items)
    report.sections = len(document.sections)

    table_widths_by_section: dict[int, dict[str, float]] = {}
    if include_tables:
        table_findings, table_widths_by_section = _scan_tables(
            rel_path,
            document,
            items,
            max_findings=max_findings,
        )
        report.findings.extend(table_findings)
    elif include_layout:
        table_widths_by_section = _collect_table_widths_by_section(document)

    if include_captions:
        report.findings.extend(_scan_caption_numbers(rel_path, captions))
        report.findings.extend(_scan_caption_proximity(rel_path, items, captions))
        report.findings.extend(_scan_figure_caption_proximity(rel_path, items, captions))

    if include_layout:
        report.findings.extend(_scan_layout(rel_path, document, items, table_widths_by_section))

    if include_headers:
        if not table_widths_by_section:
            table_widths_by_section = _collect_table_widths_by_section(document)
        report.findings.extend(_scan_headers_footers(rel_path, document, items, table_widths_by_section))

    if include_fields:
        report.findings.extend(_scan_fields_and_links(rel_path, document, items))
        report.findings.extend(_scan_toc(rel_path, items))

    if include_lists:
        report.findings.extend(_scan_lists(rel_path, items))

    if include_spacing:
        report.findings.extend(_scan_spacing(rel_path, items))
        report.findings.extend(_scan_headings(rel_path, items))

    if len(report.findings) > max_findings:
        report.findings = report.findings[:max_findings]
        report.findings.append(
            Finding(
                severity="info",
                code="findings-truncated",
                cls="Отчёт",
                object="Лимит",
                location=Location(file=rel_path),
                evidence=f"Показаны первые {max_findings} находок по файлу.",
                why_it_matters="Лимит защищает отчёт от разрастания на сильно повреждённых документах.",
                suggested_fix="Увеличить лимит находок, если нужен полный список.",
                autofix="none",
            )
        )
    return report


def _resolve_docx_files(input_path: Path) -> tuple[list[Path], Path]:
    resolved = input_path.resolve()
    if resolved.is_file():
        if resolved.suffix.lower() != ".docx" or resolved.name.startswith("~$"):
            return [], resolved.parent
        return [resolved], resolved.parent
    if resolved.is_dir():
        return find_docx_files(resolved), resolved
    return [], resolved


def _summary_payload(scans: list[FileScan]) -> dict[str, Any]:
    severity = Counter()
    classes = Counter()
    for scan in scans:
        for finding in scan.findings:
            severity[finding.severity] += 1
            classes[finding.cls] += 1
    return {
        "files": len(scans),
        "findings": sum(severity.values()),
        "severity": dict(severity),
        "classes": dict(classes),
    }


def _json_payload(input_path: Path, scans: list[FileScan], options: dict[str, Any]) -> dict[str, Any]:
    return {
        "tool": "docx_anomaly_inspector",
        "version": 1,
        "input": str(input_path),
        "options": options,
        "summary": _summary_payload(scans),
        "files": [
            {
                "path": scan.rel_path,
                "status": scan.status,
                "error": scan.error,
                "metrics": {
                    "paragraphs": scan.paragraphs,
                    "tables": scan.tables,
                    "figures": scan.figures,
                    "sections": scan.sections,
                    "findings": len(scan.findings),
                },
                "findings": [finding.to_json() for finding in scan.findings],
            }
            for scan in scans
        ],
    }


def _write_markdown_report(out_path: Path, input_path: Path, scans: list[FileScan], options: dict[str, Any]) -> None:
    safe_mkdir(out_path.parent)
    summary = _summary_payload(scans)
    lines: list[str] = []
    lines.append("# Проверка аномалий DOCX\n")
    lines.append(f"- Вход: `{input_path}`")
    lines.append(f"- Файлов проверено: **{summary['files']}**")
    lines.append(f"- Находок: **{summary['findings']}**")
    lines.append(
        "- Классы проверок: "
        + ", ".join(
            label
            for enabled, label in [
                (options.get("include_tables"), "таблицы"),
                (options.get("include_captions"), "подписи"),
                (options.get("include_layout"), "секции/ориентация"),
                (options.get("include_headers"), "колонтитулы/нумерация страниц"),
                (options.get("include_fields"), "поля/оглавление/ссылки"),
                (options.get("include_lists"), "списки/нумерация"),
                (options.get("include_spacing"), "пустоты/заголовки"),
            ]
            if enabled
        )
    )
    lines.append("")

    lines.append("## Сводка по файлам\n")
    lines.append("| Файл | Статус | Абзацы | Таблицы | Рисунки | Секции | Находки |")
    lines.append("|---|---|---:|---:|---:|---:|---:|")
    for scan in scans:
        lines.append(
            f"| `{md_escape(scan.rel_path)}` | {scan.status} | {scan.paragraphs} | {scan.tables} | "
            f"{scan.figures} | {scan.sections} | {len(scan.findings)} |"
        )
    lines.append("")

    if summary["severity"]:
        lines.append("## Сводка по серьёзности\n")
        lines.append("| Уровень | Количество |")
        lines.append("|---|---:|")
        for key in ("critical", "warning", "info"):
            if summary["severity"].get(key, 0):
                lines.append(f"| {key} | {summary['severity'][key]} |")
        lines.append("")

    lines.append("## Находки\n")
    any_findings = False
    for scan in scans:
        if not scan.findings:
            continue
        any_findings = True
        lines.append(f"### {scan.rel_path}\n")
        lines.append("| Уровень | Класс | Локация | Доказательство | Предложение |")
        lines.append("|---|---|---|---|---|")
        for finding in scan.findings:
            lines.append(
                f"| {finding.severity} | {md_escape(finding.cls)} | {md_escape(finding.location.display())} | "
                f"{md_escape(truncate(finding.evidence, 180))} | {md_escape(truncate(finding.suggested_fix, 180))} |"
            )
        lines.append("")
    if not any_findings:
        lines.append("Аномалии по выбранным классам не найдены.\n")

    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _write_docx_report(out_path: Path, input_path: Path, scans: list[FileScan], options: dict[str, Any]) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    summary = _summary_payload(scans)
    doc.add_heading("Проверка аномалий DOCX", level=1)
    doc.add_paragraph(f"Вход: {input_path}")
    doc.add_paragraph(f"Файлов проверено: {summary['files']}")
    doc.add_paragraph(f"Находок: {summary['findings']}")
    doc.add_paragraph(
        "Классы проверок: "
        + ", ".join(
            label
            for enabled, label in [
                (options.get("include_tables"), "таблицы"),
                (options.get("include_captions"), "подписи"),
                (options.get("include_layout"), "секции/ориентация"),
                (options.get("include_headers"), "колонтитулы/нумерация страниц"),
                (options.get("include_fields"), "поля/оглавление/ссылки"),
                (options.get("include_lists"), "списки/нумерация"),
                (options.get("include_spacing"), "пустоты/заголовки"),
            ]
            if enabled
        )
    )

    doc.add_heading("Сводка по файлам", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Файл", "Статус", "Абзацы", "Таблицы", "Рисунки", "Секции", "Находки"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for scan in scans:
        row = table.add_row().cells
        values = [scan.rel_path, scan.status, scan.paragraphs, scan.tables, scan.figures, scan.sections, len(scan.findings)]
        for index, value in enumerate(values):
            row[index].text = str(value)

    doc.add_heading("Находки", level=2)
    any_findings = False
    for scan in scans:
        if not scan.findings:
            continue
        any_findings = True
        doc.add_heading(scan.rel_path, level=3)
        findings_table = doc.add_table(rows=1, cols=5)
        findings_table.style = "Table Grid"
        for index, header in enumerate(["Уровень", "Класс", "Локация", "Доказательство", "Предложение"]):
            findings_table.rows[0].cells[index].text = header
        for finding in scan.findings[:120]:
            row = findings_table.add_row().cells
            values = [
                finding.severity,
                finding.cls,
                finding.location.display(),
                truncate(finding.evidence, 220),
                truncate(finding.suggested_fix, 220),
            ]
            for index, value in enumerate(values):
                row[index].text = str(value)
    if not any_findings:
        doc.add_paragraph("Аномалии по выбранным классам не найдены.")
    doc.save(str(out_path))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Read-only DOCX anomaly inspector.")
    parser.add_argument("--input", default="input", help="DOCX file or folder with DOCX files.")
    parser.add_argument("--out", default="report/docx_anomalies.md", help="Markdown report path.")
    parser.add_argument("--docx-out", default="", help="Optional DOCX report path.")
    parser.add_argument("--json-out", default="", help="Optional JSON report path.")
    parser.add_argument("--no-docx-report", action="store_true", help="Do not write a DOCX report.")
    parser.add_argument("--no-tables", action="store_true", help="Skip table anomaly checks.")
    parser.add_argument("--no-captions", action="store_true", help="Skip table/figure caption checks.")
    parser.add_argument("--no-layout", action="store_true", help="Skip section/page setup checks.")
    parser.add_argument("--no-headers", action="store_true", help="Skip header/footer and page numbering checks.")
    parser.add_argument("--no-fields", action="store_true", help="Skip TOC, field, hyperlink, and bookmark checks.")
    parser.add_argument("--no-lists", action="store_true", help="Skip list and numbering checks.")
    parser.add_argument("--no-spacing", action="store_true", help="Skip duplicated blank paragraphs and dangling heading checks.")
    parser.add_argument("--max-findings-per-file", type=int, default=300, help="Limit findings per file.")
    args = parser.parse_args(argv)

    input_path = Path(args.input).resolve()
    out_path = Path(args.out).resolve()
    docx_files, input_root = _resolve_docx_files(input_path)
    options = {
        "include_tables": not args.no_tables,
        "include_captions": not args.no_captions,
        "include_layout": not args.no_layout,
        "include_headers": not args.no_headers,
        "include_fields": not args.no_fields,
        "include_lists": not args.no_lists,
        "include_spacing": not args.no_spacing,
        "max_findings_per_file": max(1, int(args.max_findings_per_file or 300)),
    }

    if not input_path.exists():
        safe_mkdir(out_path.parent)
        out_path.write_text(f"# Проверка аномалий DOCX\n\nВход не найден: `{input_path}`\n", encoding="utf-8")
        if args.json_out:
            write_json_file(Path(args.json_out).resolve(), _json_payload(input_path, [], options))
        print(f"[ERROR] Input does not exist: {input_path}")
        return 2

    scans = [
        _scan_file(
            path,
            input_root,
            include_tables=options["include_tables"],
            include_captions=options["include_captions"],
            include_layout=options["include_layout"],
            include_spacing=options["include_spacing"],
            include_fields=options["include_fields"],
            include_lists=options["include_lists"],
            include_headers=options["include_headers"],
            max_findings=options["max_findings_per_file"],
        )
        for path in docx_files
    ]

    _write_markdown_report(out_path, input_path, scans, options)
    if not args.no_docx_report:
        docx_out = Path(args.docx_out).resolve() if args.docx_out else out_path.with_suffix(".docx")
        _write_docx_report(docx_out, input_path, scans, options)
        print(f"[OK] Wrote DOCX report: {docx_out}")
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(json_out, _json_payload(input_path, scans, options))
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Wrote report: {out_path}")
    print(f"[OK] Files checked: {len(scans)}; findings: {sum(len(scan.findings) for scan in scans)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
