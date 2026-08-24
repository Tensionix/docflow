#!/usr/bin/env python3
r"""
DOCX/XLSX Table Cell Margins

Batch tool for making table cells more compact.

DOCX:
- sets every table cell margin to the requested value;
- the default is 0.1 cm, stored as 57 twips in WordprocessingML.

XLSX:
- Excel workbooks do not have a real per-cell margin/padding property;
- the tool resets the available text indent in existing worksheet cells.

Usage:
  python docx_xlsx_table_cell_margins.py --input input --outdir output/table_cell_margins --report report/table_cell_margins.md
  python docx_xlsx_table_cell_margins.py --file input/a.docx --out output/a.docx
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree
from openpyxl import load_workbook
from openpyxl.styles import Alignment

from _office_common import md_escape, mirrored_output_path, rel_posix, safe_mkdir, write_json_file
from docx_xml_tools import NS, _etree_from_bytes, _etree_to_bytes, list_xml_parts, read_zip_map, write_zip_map


SUPPORTED_SUFFIXES = {".docx", ".xlsx", ".xlsm"}
DOCX_MARGIN_SIDES = ("top", "start", "bottom", "end", "left", "right")
W_NS = NS["w"]


def w_tag(name: str) -> str:
    return f"{{{W_NS}}}{name}"


def parse_cm(value: str) -> float:
    parsed = float(str(value).strip().replace(",", "."))
    if parsed < 0:
        raise argparse.ArgumentTypeError("margin must be non-negative")
    return parsed


def cm_to_twips(value_cm: float) -> int:
    return max(0, int(round(value_cm * 1440 / 2.54)))


def find_office_files(root: Path) -> list[Path]:
    if root.is_file():
        eligible = root.suffix.lower() in SUPPORTED_SUFFIXES and not root.name.startswith("~$")
        return [root] if eligible else []
    return sorted(
        path for path in root.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_SUFFIXES
        and not path.name.startswith("~$")
    )


def ensure_child(parent: etree._Element, name: str) -> tuple[etree._Element, bool]:
    child = parent.find(f"w:{name}", namespaces=NS)
    if child is not None:
        return child, False
    child = etree.Element(w_tag(name))
    parent.append(child)
    return child, True


def set_cell_margin(cell: etree._Element, margin_twips: int) -> bool:
    modified = False
    tc_pr = cell.find("w:tcPr", namespaces=NS)
    if tc_pr is None:
        tc_pr = etree.Element(w_tag("tcPr"))
        cell.insert(0, tc_pr)
        modified = True

    tc_mar, created = ensure_child(tc_pr, "tcMar")
    modified = modified or created

    value = str(margin_twips)
    for side in DOCX_MARGIN_SIDES:
        side_el, created = ensure_child(tc_mar, side)
        if created:
            modified = True
        if side_el.get(w_tag("w")) != value:
            side_el.set(w_tag("w"), value)
            modified = True
        if side_el.get(w_tag("type")) != "dxa":
            side_el.set(w_tag("type"), "dxa")
            modified = True
    return modified


def set_docx_part_margins(xml_bytes: bytes, margin_twips: int) -> tuple[bytes, dict[str, int], bool]:
    tree = _etree_from_bytes(xml_bytes)
    root = tree.getroot()
    tables = root.xpath(".//w:tbl", namespaces=NS)
    cells = root.xpath(".//w:tc", namespaces=NS)
    changed_cells = 0
    modified = False

    for cell in cells:
        if set_cell_margin(cell, margin_twips):
            changed_cells += 1
            modified = True

    stats = {
        "tables": len(tables),
        "cells": len(cells),
        "changed_cells": changed_cells,
    }
    if not modified:
        return xml_bytes, stats, False
    return _etree_to_bytes(tree), stats, True


def process_docx(source: Path, out_path: Path, margin_twips: int, dry_run: bool) -> dict[str, Any]:
    files = read_zip_map(source)
    total_tables = 0
    total_cells = 0
    changed_cells = 0
    modified = False

    for part in list_xml_parts(files):
        new_xml, stats, part_modified = set_docx_part_margins(files[part], margin_twips)
        files[part] = new_xml
        total_tables += stats["tables"]
        total_cells += stats["cells"]
        changed_cells += stats["changed_cells"]
        modified = modified or part_modified

    if not dry_run:
        safe_mkdir(out_path.parent)
        if modified:
            write_zip_map(out_path, files)
        else:
            shutil.copy2(source, out_path)

    note = "Поля ячеек DOCX установлены."
    if total_tables == 0:
        note = "Таблицы DOCX не найдены; файл скопирован без изменений."
    elif changed_cells == 0:
        note = "Все ячейки DOCX уже имели нужные поля; файл скопирован без изменений."

    return {
        "type": "docx",
        "tables": total_tables,
        "worksheets": 0,
        "cells": total_cells,
        "changed_cells": changed_cells,
        "status": "OK",
        "note": note,
    }


def zero_indent_alignment(alignment: Alignment) -> Alignment:
    return Alignment(
        horizontal=alignment.horizontal,
        vertical=alignment.vertical,
        textRotation=alignment.textRotation,
        wrapText=alignment.wrapText,
        shrinkToFit=alignment.shrinkToFit,
        indent=0,
        relativeIndent=0,
        justifyLastLine=alignment.justifyLastLine,
        readingOrder=alignment.readingOrder,
    )


def nonzero(value: Any) -> bool:
    try:
        return float(value or 0) != 0
    except (TypeError, ValueError):
        return bool(value)


def process_xlsx(source: Path, out_path: Path, dry_run: bool) -> dict[str, Any]:
    keep_vba = source.suffix.lower() == ".xlsm"
    wb = load_workbook(source, keep_vba=keep_vba)
    worksheets = len(wb.worksheets)
    excel_tables = 0
    cells_seen = 0
    changed_cells = 0

    for ws in wb.worksheets:
        try:
            excel_tables += len(ws.tables)
        except TypeError:
            excel_tables += 0
        for cell in list(ws._cells.values()):
            cells_seen += 1
            alignment = cell.alignment
            if nonzero(alignment.indent) or nonzero(alignment.relativeIndent):
                cell.alignment = zero_indent_alignment(alignment)
                changed_cells += 1

    if not dry_run:
        safe_mkdir(out_path.parent)
        if changed_cells:
            wb.save(out_path)
        else:
            shutil.copy2(source, out_path)

    note = (
        "XLSX не хранит настоящие поля ячеек как Word; доступный text indent "
        "сброшен до 0 там, где он был задан."
    )
    if changed_cells == 0:
        note = (
            "XLSX не хранит настоящие поля ячеек как Word; text indent уже был 0, "
            "файл скопирован без изменений."
        )
    return {
        "type": "xlsx",
        "tables": excel_tables,
        "worksheets": worksheets,
        "cells": cells_seen,
        "changed_cells": changed_cells,
        "status": "OK",
        "note": note,
    }


def process_file(source: Path, out_path: Path, margin_twips: int, dry_run: bool) -> dict[str, Any]:
    if source.suffix.lower() == ".docx":
        return process_docx(source, out_path, margin_twips, dry_run)
    if source.suffix.lower() in {".xlsx", ".xlsm"}:
        return process_xlsx(source, out_path, dry_run)
    raise RuntimeError(f"Unsupported file type: {source.suffix}")


def render_markdown(rows: list[dict[str, Any]], input_dir: Path, out_dir: Path, margin_cm: float, margin_twips: int, dry_run: bool) -> str:
    ok_rows = [row for row in rows if row["status"] == "OK"]
    error_rows = [row for row in rows if row["status"] != "OK"]
    docx_rows = [row for row in ok_rows if row["type"] == "docx"]
    xlsx_rows = [row for row in ok_rows if row["type"] == "xlsx"]

    lines: list[str] = []
    lines.append("# Поля ячеек таблиц DOCX/XLSX\n")
    lines.append(f"- Входная папка: `{input_dir}`")
    lines.append(f"- Выходная папка: `{out_dir}`")
    lines.append(f"- Режим dry-run: **{'да' if dry_run else 'нет'}**")
    lines.append(f"- Целевое значение DOCX: **{margin_cm:g} см** (`{margin_twips}` twips)")
    lines.append(f"- Файлов обработано: **{len(rows)}**")
    lines.append(f"- DOCX: **{len(docx_rows)}**; XLSX/XLSM: **{len(xlsx_rows)}**; ошибок: **{len(error_rows)}**")
    lines.append("")
    lines.append("> Важно: XLSX не имеет настоящего внутреннего поля ячейки в формате файла. Для XLSX команда сбрасывает только доступный текстовый отступ (`indent`) у существующих ячеек.")
    lines.append("")
    if not rows:
        lines.append("Файлы DOCX/XLSX не найдены.\n")
        return "\n".join(lines)

    lines.append("| Файл | Тип | Выходной файл | Таблицы | Листы | Ячейки | Изменено ячеек | Статус | Примечание |")
    lines.append("|---|---|---|---:|---:|---:|---:|---|---|")
    for row in rows:
        lines.append(
            f"| `{md_escape(row['file'])}` | {md_escape(row['type'])} | `{md_escape(row['output'])}` | "
            f"{row['tables']} | {row['worksheets']} | {row['cells']} | {row['changed_cells']} | "
            f"{md_escape(row['status'])} | {md_escape(row['note'])} |"
        )
    return "\n".join(lines) + "\n"


def write_docx_report(out_path: Path, rows: list[dict[str, Any]], margin_cm: float, margin_twips: int, dry_run: bool) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Поля ячеек таблиц DOCX/XLSX", level=1)
    doc.add_paragraph(f"Целевое значение DOCX: {margin_cm:g} см ({margin_twips} twips).")
    doc.add_paragraph(f"Dry-run: {'да' if dry_run else 'нет'}.")
    doc.add_paragraph("XLSX не имеет настоящего внутреннего поля ячейки; для XLSX сбрасывается только text indent.")
    if not rows:
        doc.add_paragraph("Файлы DOCX/XLSX не найдены.")
        doc.save(str(out_path))
        return

    headers = ("Файл", "Тип", "Выходной файл", "Таблицы", "Листы", "Ячейки", "Изменено", "Статус", "Примечание")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title
    for row in rows:
        cells = table.add_row().cells
        values = (
            row["file"],
            row["type"],
            row["output"],
            row["tables"],
            row["worksheets"],
            row["cells"],
            row["changed_cells"],
            row["status"],
            row["note"],
        )
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
    doc.save(str(out_path))


def build_json_payload(rows: list[dict[str, Any]], input_dir: Path, out_dir: Path, margin_cm: float, margin_twips: int, dry_run: bool) -> dict[str, Any]:
    return {
        "tool": "docx_xlsx_table_cell_margins",
        "version": 1,
        "input_dir": str(input_dir),
        "output_dir": str(out_dir),
        "margin_cm": margin_cm,
        "margin_twips": margin_twips,
        "dry_run": dry_run,
        "files": rows,
        "summary": {
            "total_files": len(rows),
            "pass_files": sum(1 for row in rows if row["status"] == "OK"),
            "fail_files": sum(1 for row in rows if row["status"] != "OK"),
            "docx_files": sum(1 for row in rows if row["type"] == "docx"),
            "xlsx_files": sum(1 for row in rows if row["type"] == "xlsx"),
            "cells_seen": sum(int(row["cells"]) for row in rows if row["status"] == "OK"),
            "changed_cells": sum(int(row["changed_cells"]) for row in rows if row["status"] == "OK"),
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Set DOCX table cell margins to 0.1 cm and compact XLSX text indents.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--input", help="Input folder with DOCX/XLSX files (recursive)")
    group.add_argument("--file", help="Single DOCX/XLSX file")
    parser.add_argument("--outdir", default="output/table_cell_margins", help="Output folder for batch mode")
    parser.add_argument("--out", default="", help="Output file path for single-file mode")
    parser.add_argument("--report", default="report/table_cell_margins.md", help="Output Markdown report path")
    parser.add_argument("--json-out", default="", help="Optional JSON report path")
    parser.add_argument("--margin-cm", type=parse_cm, default=0.1, help="DOCX cell margin in centimeters. Default: 0.1")
    parser.add_argument("--dry-run", action="store_true", help="Report planned outputs without writing files")
    args = parser.parse_args()

    margin_twips = cm_to_twips(args.margin_cm)
    out_dir = Path(args.outdir).resolve()
    report_path = Path(args.report).resolve()
    rows: list[dict[str, Any]] = []

    if args.input:
        input_dir = Path(args.input).resolve()
        sources = find_office_files(input_dir)
        for source in sources:
            out_path = mirrored_output_path(source, input_dir, out_dir)
            rel_source = rel_posix(source, input_dir)
            rel_output = rel_posix(out_path, out_dir)
            try:
                result = process_file(source, out_path, margin_twips, args.dry_run)
            except Exception as exc:
                result = {
                    "type": source.suffix.lower().lstrip(".") or "unknown",
                    "tables": 0,
                    "worksheets": 0,
                    "cells": 0,
                    "changed_cells": 0,
                    "status": "ERROR",
                    "note": str(exc),
                }
            rows.append({"file": rel_source, "output": rel_output, **result})
    else:
        source = Path(args.file).resolve()
        if not source.exists():
            print(f"[ERROR] File not found: {source}")
            return 2
        if source.suffix.lower() not in SUPPORTED_SUFFIXES:
            print(f"[ERROR] Unsupported file type: {source.suffix}")
            return 2
        if not args.out:
            print("[ERROR] --out is required when using --file")
            return 2
        out_path = Path(args.out).resolve()
        try:
            result = process_file(source, out_path, margin_twips, args.dry_run)
        except Exception as exc:
            result = {
                "type": source.suffix.lower().lstrip(".") or "unknown",
                "tables": 0,
                "worksheets": 0,
                "cells": 0,
                "changed_cells": 0,
                "status": "ERROR",
                "note": str(exc),
            }
        rows.append({"file": source.name, "output": out_path.name, **result})
        input_dir = source.parent
        out_dir = out_path.parent

    safe_mkdir(report_path.parent)
    report_path.write_text(render_markdown(rows, input_dir, out_dir, args.margin_cm, margin_twips, args.dry_run), encoding="utf-8")
    docx_report = report_path.with_suffix(".docx")
    write_docx_report(docx_report, rows, args.margin_cm, margin_twips, args.dry_run)
    if args.json_out:
        write_json_file(Path(args.json_out).resolve(), build_json_payload(rows, input_dir, out_dir, args.margin_cm, margin_twips, args.dry_run))

    errors = sum(1 for row in rows if row["status"] != "OK")
    changed = sum(int(row["changed_cells"]) for row in rows if row["status"] == "OK")
    print(f"[OK] Files checked: {len(rows)}")
    print(f"[OK] Changed cells: {changed}")
    if not args.dry_run:
        print(f"[OK] Output folder: {out_dir}")
    print(f"[OK] Wrote DOCX report: {docx_report}")
    print(f"[OK] Wrote report: {report_path}")
    return 1 if errors else 0


if __name__ == "__main__":
    raise SystemExit(main())
