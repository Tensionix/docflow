#!/usr/bin/env python3
r"""
DOCX Table Extractor

Recursively extracts all tables from DOCX files under input into XLSX files.
Also writes DOCX/MD index reports in report.

Usage:
  python docx_extract_tables.py --input input --outdir output/tables --out report/docx_tables_index.md
"""

from __future__ import annotations

# >>> audion CLI bootstrap >>>
import os as _os, sys as _sys
_HERE = _os.path.dirname(_os.path.abspath(__file__))
if _HERE not in _sys.path:
    _sys.path.insert(0, _HERE)
try:
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
# <<< audion CLI bootstrap <<<

import argparse
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import Workbook

from _office_common import safe_mkdir, find_docx_files, rel_posix, md_escape, truncate, write_json_file


def iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def clean_filename_part(value: str, max_len: int = 80) -> str:
    value = re.sub(r"\s+", " ", value or "").strip()
    value = re.sub(r'[<>:"/\\\\|?*\x00-\x1f]+', "_", value)
    value = value.strip(" ._")
    if not value:
        value = "table"
    return value[:max_len].rstrip(" ._")


def title_slug(title: str, fallback: str) -> str:
    title = re.sub(r"\s+", " ", title or "").strip()
    if not title:
        return fallback
    words = title.split()
    if words and words[0].lower().startswith(("таблица", "table")):
        picked = words[:8]
    else:
        picked = words[:6]
    return clean_filename_part(" ".join(picked), max_len=90)


def table_values(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def write_table_xlsx(rows: list[list[str]], out_path: Path) -> None:
    safe_mkdir(out_path.parent)
    wb = Workbook()
    ws = wb.active
    ws.title = "Table"
    for row in rows:
        ws.append(row)
    wb.save(out_path)


def write_docx_index(out_path: Path, rows: list[dict[str, Any]]) -> None:
    safe_mkdir(out_path.parent)
    doc = Document()
    doc.add_heading("Индекс таблиц DOCX", level=1)
    if not rows:
        doc.add_paragraph("Таблицы не найдены.")
        doc.save(str(out_path))
        return
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, title in enumerate(("Документ", "Таблица", "Название", "Строки", "Столбцы", "Выходной файл", "Статус")):
        hdr[idx].text = title
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row["document"])
        cells[1].text = str(row["table_no"])
        cells[2].text = str(row["title"])
        cells[3].text = str(row["rows"])
        cells[4].text = str(row["columns"])
        cells[5].text = str(row["output"])
        cells[6].text = str(row["status"])
    doc.save(str(out_path))


def build_json_payload(input_dir: Path, out_dir: Path, docx_files: list[Path], rows: list[dict[str, Any]]) -> dict[str, object]:
    rows_by_doc: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        rows_by_doc.setdefault(str(row["document"]), []).append(row)
    files_payload = []
    for docx_path in docx_files:
        rel_doc = rel_posix(docx_path, input_dir)
        doc_rows = rows_by_doc.get(rel_doc, [])
        error_count = sum(1 for row in doc_rows if str(row["status"]).startswith("ERROR"))
        ok_rows = [row for row in doc_rows if row["status"] == "OK"]
        files_payload.append(
            {
                "path": rel_doc,
                "status": "FAIL" if error_count else "PASS",
                "metrics": {
                    "tables_extracted": len(ok_rows),
                    "errors": error_count,
                    "rows_total": sum(int(row["rows"]) for row in ok_rows),
                    "max_columns": max((int(row["columns"]) for row in ok_rows), default=0),
                },
            }
        )
    tables = [
        {
            "document": row["document"],
            "table_no": row["table_no"],
            "title": row["title"],
            "rows": row["rows"],
            "columns": row["columns"],
            "output": row["output"],
            "status": row["status"],
        }
        for row in rows
    ]
    return {
        "tool": "docx_extract_tables",
        "version": 1,
        "input_dir": str(input_dir),
        "output_dir": str(out_dir),
        "files": files_payload,
        "tables": tables,
        "summary": {
            "total_files": len(docx_files),
            "pass_files": sum(1 for item in files_payload if item["status"] == "PASS"),
            "fail_files": sum(1 for item in files_payload if item["status"] == "FAIL"),
            "tables_extracted": sum(1 for row in rows if row["status"] == "OK"),
            "errors": sum(1 for row in rows if str(row["status"]).startswith("ERROR")),
        },
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract DOCX tables recursively into XLSX files.")
    ap.add_argument("--input", default="input", help="Input folder with DOCX files (recursive)")
    ap.add_argument("--outdir", default="output/tables", help="Output folder for extracted XLSX tables")
    ap.add_argument("--out", default="report/docx_tables_index.md", help="Output Markdown index path")
    ap.add_argument("--json-out", default="", help="Optional JSON report path")
    args = ap.parse_args()

    in_dir = Path(args.input).resolve()
    out_dir = Path(args.outdir).resolve()
    out_path = Path(args.out).resolve()
    safe_mkdir(out_dir)
    safe_mkdir(out_path.parent)

    docx_files = find_docx_files(in_dir)
    index_rows: list[dict[str, Any]] = []

    for docx_path in docx_files:
        rel_doc = rel_posix(docx_path, in_dir)
        try:
            doc = Document(str(docx_path))
        except Exception as exc:
            index_rows.append(
                {
                    "document": rel_doc,
                    "table_no": "",
                    "title": "",
                    "rows": 0,
                    "columns": 0,
                    "output": "",
                    "status": f"ERROR: {exc}",
                }
            )
            continue

        table_no = 0
        last_paragraph = ""
        doc_out_dir = out_dir / docx_path.relative_to(in_dir).with_suffix("")
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    last_paragraph = text
                continue

            table_no += 1
            values = table_values(block)
            row_count = len(values)
            col_count = max((len(row) for row in values), default=0)
            slug = title_slug(last_paragraph, f"table_{table_no:03d}")
            xlsx_name = f"{table_no:03d}_{slug}.xlsx"
            xlsx_path = doc_out_dir / xlsx_name
            write_table_xlsx(values, xlsx_path)
            index_rows.append(
                {
                    "document": rel_doc,
                    "table_no": table_no,
                    "title": last_paragraph,
                    "rows": row_count,
                    "columns": col_count,
                    "output": rel_posix(xlsx_path, out_dir),
                    "status": "OK",
                }
            )
            last_paragraph = ""

    lines: list[str] = []
    lines.append("# Индекс таблиц DOCX\n")
    lines.append(f"- Входная папка: `{in_dir}`")
    lines.append(f"- Файлов DOCX проверено: **{len(docx_files)}**")
    lines.append(f"- Таблиц извлечено: **{sum(1 for r in index_rows if r['status'] == 'OK')}**")
    lines.append(f"- Выходная папка: `{out_dir}`\n")
    if not index_rows:
        lines.append("Таблицы не найдены.\n")
    else:
        lines.append("| Документ | Таблица | Название перед таблицей | Строки | Столбцы | Выходной файл | Статус |")
        lines.append("|---|---:|---|---:|---:|---|---|")
        for row in index_rows:
            lines.append(
                f"| `{md_escape(row['document'])}` | {row['table_no']} | {md_escape(truncate(row['title'], 120))} | "
                f"{row['rows']} | {row['columns']} | `{md_escape(row['output'])}` | {md_escape(row['status'])} |"
            )
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    docx_out = out_path.with_suffix(".docx")
    write_docx_index(docx_out, index_rows)
    if args.json_out:
        json_out = Path(args.json_out).resolve()
        write_json_file(json_out, build_json_payload(in_dir, out_dir, docx_files, index_rows))
        print(f"[OK] Wrote JSON report: {json_out}")
    print(f"[OK] Extracted tables: {sum(1 for r in index_rows if r['status'] == 'OK')}")
    print(f"[OK] Output folder: {out_dir}")
    print(f"[OK] Wrote DOCX report: {docx_out}")
    print(f"[OK] Wrote report: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
