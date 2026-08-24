from __future__ import annotations

import argparse
import json
import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Cm

TWIPS_PER_CM = 567.0
INDEX_RE = re.compile(r"^\d+(?:[.\-]\d+)*\.?$")


def _cm_to_twips(cm_value: float) -> int:
    return int(round(cm_value * TWIPS_PER_CM))


CURRENT_DIR = Path(__file__).resolve().parent
if str(CURRENT_DIR) not in sys.path:
    sys.path.insert(0, str(CURRENT_DIR))

from docx_table_unifier_model import build_table_model, normalize_text
from docx_table_unifier_writers import KNOWN_WIDTHS_CM

SUPPORTED_EXTENSIONS = {'.docx'}
PAGE_SIZES_CM = {
    'A4': (21.0, 29.7),
    'A3': (29.7, 42.0),
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description='Optimize column widths in existing DOCX tables while preserving page layout and table structure.'
    )
    parser.add_argument('--input', help='Path to a single DOCX file')
    parser.add_argument('--input-dir', default='input', help='Input folder with DOCX files')
    parser.add_argument('--outdir', default='output/word_excel_tables/widths', help='Output folder for generated DOCX files')
    parser.add_argument('--report-dir', default='report/word_excel_tables/widths', help='Output folder for JSON reports')
    parser.add_argument('--all', action='store_true', help='Process all supported files from input directory')
    parser.add_argument('--recursive', action='store_true', help='Process DOCX files recursively under input directory')
    parser.add_argument(
        '--mode',
        choices=('preserve-width', 'fit-to-margins', 'page-setup-only'),
        default='preserve-width',
        help='Preserve current table width, fit table width to current document margins before rebalancing, or only apply page setup.',
    )
    parser.add_argument('--preheader-mode', choices=('include', 'separate'), default='separate', help='Include pre-header rows in header weighting or keep them out of repeat/header weighting.')
    parser.add_argument('--fit-target', choices=('current-section', 'page-setup'), default='current-section', help='For fit-to-margins: use current section geometry or apply explicit page setup first.')
    parser.add_argument('--page-size', choices=('A4', 'A3'), default='A4', help='Page size when --fit-target page-setup is used.')
    parser.add_argument('--page-orientation', choices=('portrait', 'landscape'), default='landscape', help='Page orientation when --fit-target page-setup is used.')
    parser.add_argument('--page-setup-margin-fallback', choices=('one-cm', 'keep'), default='one-cm', help='When applying page setup and an explicit margin is omitted, use 1 cm or keep the source section margin.')
    parser.add_argument('--margin-top-mm', type=float, help='Top margin in millimeters for explicit page/margin setup.')
    parser.add_argument('--margin-right-mm', type=float, help='Right margin in millimeters for explicit page/margin setup.')
    parser.add_argument('--margin-bottom-mm', type=float, help='Bottom margin in millimeters for explicit page/margin setup.')
    parser.add_argument('--margin-left-mm', type=float, help='Left margin in millimeters for explicit page/margin setup.')
    parser.add_argument('--skip-fit-max-columns', type=int, default=3, help='Do not rebalance or fit tables with this many columns or fewer. Use 0 to disable.')
    return parser.parse_args()


def _resolve_project_path(root: Path, value: str) -> Path:
    path = Path(value)
    if not path.is_absolute():
        path = root / path
    return path.resolve()


def _margin_cm(value_mm: float | None, fallback_cm: float) -> float:
    if value_mm is None:
        return fallback_cm
    return max(0.0, float(value_mm) / 10.0)


def _page_dimensions_cm(page_size: str, orientation: str) -> tuple[float, float]:
    width, height = PAGE_SIZES_CM.get(page_size.upper(), PAGE_SIZES_CM['A4'])
    if orientation == 'landscape':
        return max(width, height), min(width, height)
    return min(width, height), max(width, height)


def _apply_page_or_margin_setup(document: Document, args: argparse.Namespace) -> dict[str, object]:
    apply_page = args.fit_target == 'page-setup'
    apply_margins = any(
        value is not None
        for value in (args.margin_top_mm, args.margin_right_mm, args.margin_bottom_mm, args.margin_left_mm)
    )
    page_width_cm, page_height_cm = _page_dimensions_cm(args.page_size, args.page_orientation)
    for section in document.sections:
        if apply_page:
            section.orientation = WD_ORIENT.LANDSCAPE if args.page_orientation == 'landscape' else WD_ORIENT.PORTRAIT
            section.page_width = Cm(page_width_cm)
            section.page_height = Cm(page_height_cm)
        if apply_page or apply_margins:
            fallback_top = section.top_margin.cm
            fallback_right = section.right_margin.cm
            fallback_bottom = section.bottom_margin.cm
            fallback_left = section.left_margin.cm
            if apply_page and args.page_setup_margin_fallback == 'one-cm':
                fallback_top = fallback_right = fallback_bottom = fallback_left = 1.0
            section.top_margin = Cm(_margin_cm(args.margin_top_mm, fallback_top))
            section.right_margin = Cm(_margin_cm(args.margin_right_mm, fallback_right))
            section.bottom_margin = Cm(_margin_cm(args.margin_bottom_mm, fallback_bottom))
            section.left_margin = Cm(_margin_cm(args.margin_left_mm, fallback_left))
    return {
        'fit_target': args.fit_target,
        'page_setup_applied': apply_page,
        'margins_applied': apply_page or apply_margins,
        'page_setup_margin_fallback': args.page_setup_margin_fallback,
        'page_size': args.page_size,
        'page_orientation': args.page_orientation,
    }


def iter_sources(input_dir: Path, args: argparse.Namespace) -> list[Path]:
    if args.input:
        return [Path(args.input).resolve()]
    if input_dir.is_file():
        eligible = input_dir.suffix.lower() == '.docx' and not input_dir.name.startswith('~$')
        return [input_dir] if eligible else []
    globber = input_dir.rglob if args.recursive else input_dir.glob
    files = sorted(path for path in globber('*.docx') if path.is_file() and not path.name.startswith('~$'))
    if args.all:
        return files
    return files[:1]


def _relative_stem(source: Path, input_dir: Path) -> Path:
    try:
        return source.relative_to(input_dir).with_suffix('')
    except ValueError:
        return Path(source.stem)


def _percentile(values: list[int], p: float) -> float:
    if not values:
        return 0.0
    items = sorted(values)
    if len(items) == 1:
        return float(items[0])
    pos = max(0.0, min(1.0, p)) * (len(items) - 1)
    lo = math.floor(pos)
    hi = math.ceil(pos)
    if lo == hi:
        return float(items[lo])
    frac = pos - lo
    return items[lo] * (1.0 - frac) + items[hi] * frac


def _is_numeric(text: str) -> bool:
    text = normalize_text(text).replace(' ', '')
    if not text:
        return False
    try:
        float(text.replace(',', '.'))
        return True
    except Exception:
        return False


def _looks_index(text: str) -> bool:
    return bool(INDEX_RE.fullmatch(normalize_text(text).replace(' ', '')))


def _is_section_row(row: list[str], header_rows: int, row_index: int) -> bool:
    if row_index < header_rows:
        return False
    non_empty = [idx for idx, value in enumerate(row) if normalize_text(value)]
    return len(non_empty) == 1 and non_empty[0] == 0


def _read_table_total_width_cm(table, fallback_total_width_cm: float) -> float:
    try:
        grid_cols = getattr(getattr(table._tbl, 'tblGrid', None), 'gridCol_lst', [])
        widths_twips = []
        for grid_col in grid_cols:
            value = getattr(grid_col, 'w', None)
            if value is None:
                continue
            try:
                widths_twips.append(int(value))
            except Exception:
                pass
        if widths_twips and sum(widths_twips) > 0:
            raw_total = sum(widths_twips)
            # Some source documents store tblGrid widths in EMU-like units instead of twips.
            if max(widths_twips) > 10000:
                return raw_total / 360000.0
            return raw_total / 567.0
    except Exception:
        pass

    try:
        col_widths = []
        for column in table.columns:
            width = getattr(column, 'width', None)
            if width:
                col_widths.append(width.cm)
        if col_widths and sum(col_widths) > 0:
            return sum(col_widths)
    except Exception:
        pass

    return fallback_total_width_cm


def _fallback_total_width_cm(document: Document) -> float:
    if not document.sections:
        return 17.0
    section = document.sections[0]
    return max(8.0, section.page_width.cm - section.left_margin.cm - section.right_margin.cm)


def _body_table_section_map(document: Document) -> dict[int, int]:
    mapping: dict[int, int] = {}
    current_section = 0
    sections = list(document.sections)
    max_section_index = max(0, len(sections) - 1)

    body = document._body._element
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'tbl':
            mapping[id(child)] = min(current_section, max_section_index)
            continue
        if tag != 'p':
            continue
        p_pr = getattr(child, 'pPr', None)
        if p_pr is None:
            continue
        sect_pr = getattr(p_pr, 'sectPr', None)
        if sect_pr is not None and current_section < max_section_index:
            current_section += 1
    return mapping


def _body_table_page_map(document: Document) -> dict[int, int]:
    mapping: dict[int, int] = {}
    current_page = 1
    table_index = 0
    body = document._body._element
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'tbl':
            table_index += 1
            mapping[table_index] = current_page
            continue
        if tag != 'p':
            continue
        explicit_breaks = sum(1 for br in child.iter(qn('w:br')) if br.get(qn('w:type')) == 'page')
        rendered_breaks = sum(1 for _ in child.iter(qn('w:lastRenderedPageBreak')))
        current_page += explicit_breaks + rendered_breaks
    return mapping


def _section_available_width_cm(document: Document, table, table_section_map: dict[int, int]) -> float:
    sections = list(document.sections)
    if not sections:
        return 17.0
    section_index = table_section_map.get(id(table._tbl), 0)
    section_index = max(0, min(section_index, len(sections) - 1))
    section = sections[section_index]
    return max(8.0, section.page_width.cm - section.left_margin.cm - section.right_margin.cm)


def _scale_widths(widths: list[float], target_total_width_cm: float) -> list[float]:
    total = sum(widths) or 1.0
    factor = target_total_width_cm / total
    scaled = [round(width * factor, 4) for width in widths]
    if scaled:
        scaled[-1] = round(scaled[-1] + (target_total_width_cm - sum(scaled)), 4)
    return scaled


def _sample_display_rows(model, preheader_mode: str = 'include', data_row_limit: int = 56) -> list[list[str]]:
    start_row = model.preheader_row_count if preheader_mode == 'separate' else 0
    header_rows = max(0, model.header_row_count - start_row)
    sampled: list[list[str]] = []
    data_seen = 0
    for row_index, row in enumerate(model.display_grid[start_row:]):
        if row_index < header_rows:
            sampled.append(row)
            continue
        if _is_section_row(row, header_rows, row_index):
            continue
        sampled.append(row)
        data_seen += 1
        if data_seen >= data_row_limit:
            break
    return sampled


def _count_data_rows(model, preheader_mode: str = 'include') -> int:
    start_row = model.preheader_row_count if preheader_mode == 'separate' else 0
    header_rows = max(0, model.header_row_count - start_row)
    count = 0
    for row_index, row in enumerate(model.display_grid[start_row:]):
        if row_index < header_rows:
            continue
        if _is_section_row(row, header_rows, row_index):
            continue
        if any(normalize_text(cell) for cell in row):
            count += 1
    return count


def _skip_fit_max_columns(args: argparse.Namespace) -> int:
    return max(0, int(args.skip_fit_max_columns or 0))


def _should_optimize_table(model, page_estimate: int, preheader_mode: str, skip_fit_max_columns: int) -> tuple[bool, str]:
    if page_estimate <= 2:
        return False, 'skip_first_two_pages'
    if skip_fit_max_columns > 0 and model.width <= skip_fit_max_columns:
        return False, 'too_few_columns'
    if model.height <= max(model.header_row_count, 0) + 1:
        return False, 'too_few_rows'
    if _count_data_rows(model, preheader_mode=preheader_mode) < 2:
        return False, 'too_few_data_rows'
    return True, 'ok'


def _first_column_is_index(sampled_rows: list[list[str]], header_rows: int) -> bool:
    values: list[str] = []
    for row_index, row in enumerate(sampled_rows):
        if row_index < header_rows or not row:
            continue
        text = normalize_text(row[0])
        if text:
            values.append(text)
    if not values:
        return False
    hits = sum(1 for value in values if _looks_index(value) or _is_numeric(value))
    return hits / len(values) >= 0.70


def _rebalance_overloaded_columns(
    widths: list[float],
    scores: list[float],
    numeric_flags: list[bool],
    min_width: float,
) -> list[float]:
    if len(widths) < 4:
        return widths

    densities = [score / max(width, 0.45) for score, width in zip(scores, widths)]
    candidate_indexes = [idx for idx, is_numeric in enumerate(numeric_flags) if not is_numeric]
    if not candidate_indexes:
        return widths

    reference = sorted(densities[idx] for idx in candidate_indexes)
    median_density = reference[len(reference) // 2]
    if median_density <= 0:
        return widths

    overloaded = [idx for idx in candidate_indexes if densities[idx] >= median_density * 1.78]
    donors = [idx for idx in candidate_indexes if densities[idx] <= median_density * 0.72 and widths[idx] > min_width + 0.18]
    if not overloaded or not donors:
        return widths

    donor_room = {idx: max(0.0, widths[idx] - min_width) for idx in donors}
    total_room = sum(donor_room.values())
    if total_room <= 0:
        return widths

    overload_weights = {idx: max(0.0, densities[idx] / median_density - 1.0) for idx in overloaded}
    total_overload = sum(overload_weights.values()) or 1.0
    max_shift = min(total_room, max(0.35, sum(widths) * 0.12))

    for idx in donors:
        shift = max_shift * (donor_room[idx] / total_room)
        widths[idx] -= shift

    for idx in overloaded:
        gain = max_shift * (overload_weights[idx] / total_overload)
        widths[idx] += gain

    return widths


def _content_balanced_widths(model, total_width_cm: float, preheader_mode: str = 'include') -> list[float]:
    col_count = model.width
    if col_count <= 0:
        return []

    sampled_rows = _sample_display_rows(model, preheader_mode=preheader_mode)
    if not sampled_rows:
        sampled_rows = model.display_grid

    base_widths = None
    if col_count in KNOWN_WIDTHS_CM:
        base_widths = _scale_widths(KNOWN_WIDTHS_CM[col_count], total_width_cm)

    min_width = 0.95 if col_count <= 6 else 0.85 if col_count <= 9 else 0.75
    header_rows = max(0, model.header_row_count - (model.preheader_row_count if preheader_mode == 'separate' else 0))

    content_scores: list[float] = []
    numeric_flags: list[bool] = []
    header_lengths_all: list[int] = []
    header_max_by_col: list[int] = []

    # First pass for header crowding awareness.
    for col_idx in range(col_count):
        header_lengths = []
        for row_index, row in enumerate(sampled_rows[:header_rows]):
            if col_idx >= len(row):
                continue
            text = normalize_text(row[col_idx])
            if text:
                header_lengths.append(len(text.replace('\n', ' ')))
        header_max = max(header_lengths) if header_lengths else 0
        header_max_by_col.append(header_max)
        if header_max:
            header_lengths_all.append(header_max)
    global_header_p75 = _percentile(header_lengths_all, 0.75)

    for col_idx in range(col_count):
        header_lengths: list[int] = []
        data_lengths: list[int] = []
        line_break_counts: list[int] = []
        numeric_hits = 0
        filled_rows = 0
        considered_rows = 0
        long_hits = 0

        for row_index, row in enumerate(sampled_rows):
            if col_idx >= len(row):
                continue
            text = normalize_text(row[col_idx])
            if row_index < header_rows:
                if text:
                    header_lengths.append(len(text.replace('\n', ' ')))
                continue
            considered_rows += 1
            if not text:
                continue
            filled_rows += 1
            compact = text.replace('\n', ' ')
            text_len = len(compact)
            data_lengths.append(text_len)
            line_break_counts.append(text.count('\n'))
            if text_len >= 44:
                long_hits += 1
            if _is_numeric(compact) or _looks_index(compact):
                numeric_hits += 1

        header_max = max(header_lengths) if header_lengths else 0
        header_avg = sum(header_lengths) / len(header_lengths) if header_lengths else 0.0
        p85 = _percentile(data_lengths, 0.85)
        p60 = _percentile(data_lengths, 0.60)
        avg = sum(data_lengths) / len(data_lengths) if data_lengths else 0.0
        fill_ratio = filled_rows / considered_rows if considered_rows else 0.0
        numeric_ratio = numeric_hits / len(data_lengths) if data_lengths else 0.0
        long_ratio = long_hits / len(data_lengths) if data_lengths else 0.0
        line_break_avg = sum(line_break_counts) / len(line_break_counts) if line_break_counts else 0.0

        score = 1.0
        score += header_max * 0.85
        score += header_avg * 0.25
        score += p85 * 0.80
        score += p60 * 0.25
        score += avg * 0.15
        score += fill_ratio * 5.5
        score += long_ratio * 12.0
        score += line_break_avg * 4.0
        if header_max and header_max >= max(global_header_p75 * 1.25, 22):
            score *= 1.10
        if numeric_ratio >= 0.80:
            score *= 0.50
        elif numeric_ratio >= 0.45:
            score *= 0.78
        content_scores.append(max(score, 0.5))
        numeric_flags.append(numeric_ratio >= 0.70)

    content_total = sum(content_scores) or 1.0
    adaptive = [max(min_width, total_width_cm * (score / content_total)) for score in content_scores]

    total = sum(adaptive)
    if total > total_width_cm:
        shrinkable = [max(0.0, width - min_width) for width in adaptive]
        shrink_total = sum(shrinkable)
        overflow = total - total_width_cm
        if shrink_total > 0:
            adaptive = [width - overflow * (room / shrink_total) for width, room in zip(adaptive, shrinkable)]
    total = sum(adaptive)
    if total < total_width_cm and adaptive:
        adaptive[-1] += total_width_cm - total

    if base_widths:
        blended = [round(base * 0.45 + auto * 0.55, 4) for base, auto in zip(base_widths, adaptive)]
    else:
        blended = [round(width, 4) for width in adaptive]

    if _first_column_is_index(sampled_rows, header_rows) and blended:
        max_index_width = max(1.15, min(1.65, total_width_cm * 0.085))
        if blended[0] > max_index_width:
            freed = blended[0] - max_index_width
            blended[0] = max_index_width
            recipients = [idx for idx in range(1, len(blended)) if not numeric_flags[idx]] or list(range(1, len(blended)))
            if recipients and freed > 0:
                total_scores = sum(content_scores[idx] for idx in recipients) or 1.0
                for idx in recipients:
                    blended[idx] += freed * (content_scores[idx] / total_scores)

    blended = _rebalance_overloaded_columns(blended, content_scores, numeric_flags, min_width)

    # Final normalization.
    blended = [max(min_width, round(width, 4)) for width in blended]
    total = sum(blended)
    if blended and total != total_width_cm:
        blended[-1] = round(blended[-1] + (total_width_cm - total), 4)
    return blended


def _set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        from docx.oxml import OxmlElement
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    table.autofit = False


def _apply_widths(table, widths_cm: list[float]) -> None:
    _set_table_layout_fixed(table)
    grid_cols = getattr(getattr(table._tbl, 'tblGrid', None), 'gridCol_lst', [])
    for idx, width in enumerate(widths_cm):
        try:
            if idx < len(grid_cols):
                grid_cols[idx].set(qn('w:w'), str(_cm_to_twips(width)))
        except Exception:
            pass
    for row in table.rows:
        if len(row.cells) < len(widths_cm):
            continue
        for idx, width in enumerate(widths_cm):
            try:
                row.cells[idx].width = Cm(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tcw = tc_pr.find(qn('w:tcW'))
                if tcw is None:
                    from docx.oxml import OxmlElement
                    tcw = OxmlElement('w:tcW')
                    tc_pr.append(tcw)
                tcw.set(qn('w:w'), str(_cm_to_twips(width)))
                tcw.set(qn('w:type'), 'dxa')
            except Exception:
                pass


def process_one(input_dir: Path, out_dir: Path, report_dir: Path, source: Path, args: argparse.Namespace) -> int:
    if not source.exists():
        print(f'[ERROR] Input file was not found: {source}')
        return 1
    if source.suffix.lower() not in SUPPORTED_EXTENSIONS:
        print(f'[ERROR] Unsupported input type: {source.suffix}. Supported: .docx')
        return 1

    document = Document(str(source))
    page_setup = _apply_page_or_margin_setup(document, args)
    fallback_total_width_cm = _fallback_total_width_cm(document)
    table_section_map = _body_table_section_map(document)
    table_page_map = _body_table_page_map(document)
    report = {
        'source_file': source.name,
        'width_strategy': 'existing-docx-smart-balance',
        'mode': args.mode,
        'preheader_mode': args.preheader_mode,
        'skip_fit_max_columns': _skip_fit_max_columns(args),
        'page_setup': page_setup,
        'tables': [],
        'skipped_tables': [],
    }

    optimized_tables = 0
    if args.mode == 'page-setup-only':
        report['tables_seen'] = len(document.tables)
    else:
        for idx, table in enumerate(document.tables, start=1):
            model = build_table_model(table, idx)
            page_estimate = table_page_map.get(idx, 1)
            if model.width <= 0 or model.height <= 0:
                report['skipped_tables'].append({
                    'table_index': idx,
                    'page_estimate': page_estimate,
                    'reason': 'empty_table_model',
                })
                continue
            should_optimize, reason = _should_optimize_table(
                model,
                page_estimate,
                args.preheader_mode,
                _skip_fit_max_columns(args),
            )
            if not should_optimize:
                report['skipped_tables'].append({
                    'table_index': idx,
                    'page_estimate': page_estimate,
                    'column_count': model.width,
                    'reason': reason,
                })
                continue
            if args.mode == 'fit-to-margins':
                total_width_cm = _section_available_width_cm(document, table, table_section_map)
            else:
                total_width_cm = _read_table_total_width_cm(table, fallback_total_width_cm)
            widths_cm = _content_balanced_widths(model, total_width_cm, preheader_mode=args.preheader_mode)
            if not widths_cm:
                report['skipped_tables'].append({
                    'table_index': idx,
                    'page_estimate': page_estimate,
                    'column_count': model.width,
                    'reason': 'no_widths_computed',
                })
                continue
            _apply_widths(table, widths_cm)
            report['tables'].append({
                'table_index': idx,
                'page_estimate': page_estimate,
                'column_count': model.width,
                'target_total_width_cm': round(total_width_cm, 4),
                'optimized_widths_cm': [round(w, 4) for w in widths_cm],
            })
            optimized_tables += 1

    out_dir.mkdir(parents=True, exist_ok=True)
    report_dir.mkdir(parents=True, exist_ok=True)
    rel_stem = _relative_stem(source, input_dir)
    if args.fit_target == 'page-setup':
        if args.mode == 'page-setup-only':
            suffix = '__adapted_orientation'
        elif args.mode == 'preserve-width':
            suffix = '__adapted_orientation_balanced_widths'
        else:
            suffix = '__adapted_orientation_fit_to_margins'
    else:
        suffix = '__optimized_widths' if args.mode == 'preserve-width' else '__fit_to_margins_optimized_widths'
    out_docx = (out_dir / rel_stem).with_name(f'{rel_stem.name}{suffix}.docx')
    out_json = (report_dir / rel_stem).with_name(f'{rel_stem.name}{suffix}.json')
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    out_json.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_docx))
    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')

    print(f'[INFO] Optimized tables: {optimized_tables}')
    print(f'[OK] Output DOCX : {out_docx}')
    print(f'[OK] Output JSON : {out_json}')
    return 0


def main() -> int:
    args = parse_args()
    root = Path(__file__).resolve().parents[1]
    input_dir = _resolve_project_path(root, args.input_dir)
    out_dir = _resolve_project_path(root, args.outdir)
    report_dir = _resolve_project_path(root, args.report_dir)
    sources = iter_sources(input_dir, args)
    if not sources:
        print(f'[ERROR] No DOCX files were found in: {input_dir}')
        print('[INFO] Put one or more .docx files into the input folder and run again.')
        return 1

    overall_rc = 0
    for source in sources:
        rc = process_one(input_dir, out_dir, report_dir, source, args)
        if rc != 0:
            overall_rc = rc
    return overall_rc


if __name__ == '__main__':
    raise SystemExit(main())
