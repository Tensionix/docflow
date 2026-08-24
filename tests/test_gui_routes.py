from __future__ import annotations

from dataclasses import replace
from pathlib import Path
import sys
import tempfile
import unittest

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from system_core.core.jobs import execute_operation
from system_core.core.manifest import Operation
from system_core.ui_nicegui import app


class GuiRouteTests(unittest.TestCase):
    def test_active_paths_follow_workbench_source_and_target(self) -> None:
        original_source = app.state.get("source_path")
        original_target = app.state.get("destination_path")
        try:
            with tempfile.TemporaryDirectory() as tmp:
                root = Path(tmp)
                source = root / "selected.docx"
                target = root / "result"
                source.touch()
                app.state["source_path"] = str(source)
                app.state["destination_path"] = str(target)
                active = app.active_project_paths()
                self.assertEqual(active.input, source.resolve())
                self.assertEqual(active.output, target.resolve())
        finally:
            app.state["source_path"] = original_source
            app.state["destination_path"] = original_target

    def test_delete_guard_refuses_roots_and_project(self) -> None:
        with self.assertRaises(RuntimeError):
            app.validate_workspace_delete_target(app.ROOT)
        with self.assertRaises(RuntimeError):
            app.validate_workspace_delete_target(Path(app.ROOT.anchor))

    def test_delete_contents_keeps_container(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            folder = Path(tmp) / "target"
            folder.mkdir()
            (folder / "one.txt").write_text("one", encoding="utf-8")
            (folder / "nested").mkdir()
            (folder / "nested" / "two.txt").write_text("two", encoding="utf-8")
            result = app.delete_workspace_path_contents(folder)
            self.assertEqual(result["removed"], 2)
            self.assertTrue(folder.is_dir())
            self.assertEqual(list(folder.iterdir()), [])

    def test_selected_single_file_reaches_office_backend_and_target(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            temp = Path(tmp)
            source = temp / "selected.docx"
            output = temp / "target"
            logs = temp / "logs"
            report = temp / "report"
            workspace = temp / "workspace"
            for folder in (output, logs, report, workspace):
                folder.mkdir(parents=True, exist_ok=True)
            document = Document()
            document.add_paragraph("Workbench route probe")
            document.save(source)

            test_paths = replace(
                app.paths,
                input=source,
                output=output,
                logs=logs,
                report=report,
                workspace=workspace,
            )
            operation = Operation(
                id="docx_clean_route_test",
                title="DOCX clean route test",
                description="",
                service="system_core.services.office_service:docx_clean",
            )
            result = execute_operation(test_paths, operation)
            self.assertTrue(result.ok, result.message)
            produced = output / "cleaned_docx" / source.name
            self.assertTrue(produced.is_file(), f"Missing routed output: {produced}")
            self.assertEqual(Document(produced).paragraphs[0].text, "Workbench route probe")

    def test_restored_table_backends_accept_one_workbench_source_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            temp = Path(tmp)
            source = temp / "tables.docx"
            output = temp / "target"
            logs = temp / "logs"
            report = temp / "report"
            workspace = temp / "workspace"
            for folder in (output, logs, report, workspace):
                folder.mkdir(parents=True, exist_ok=True)
            document = Document()
            for prefix in ("A", "B"):
                table = document.add_table(rows=2, cols=2)
                table.cell(0, 0).text = "Код"
                table.cell(0, 1).text = "Значение"
                table.cell(1, 0).text = prefix
                table.cell(1, 1).text = f"{prefix}-1"
            document.save(source)
            test_paths = replace(
                app.paths,
                input=source,
                output=output,
                logs=logs,
                report=report,
                workspace=workspace,
            )

            for operation_id, service, output_folder in (
                (
                    "running_header_route_test",
                    "system_core.services.office_service:docx_table_stitcher_running_header",
                    output / "word_excel_tables" / "stitched_running_header",
                ),
                (
                    "width_only_route_test",
                    "system_core.services.office_service:docx_table_unify_width_only",
                    output / "word_excel_tables" / "unified_width_only",
                ),
            ):
                operation = Operation(operation_id, operation_id, "", service)
                result = execute_operation(test_paths, operation)
                self.assertTrue(result.ok, result.message)
                self.assertTrue(list(output_folder.rglob("*.docx")), f"No DOCX written by {service}")

    def test_gui_host_is_loopback_only_by_default(self) -> None:
        app.assert_gui_host_allowed("127.0.0.1")
        app.assert_gui_host_allowed("::1")
        with self.assertRaises(SystemExit):
            app.assert_gui_host_allowed("0.0.0.0")

    def test_only_control_free_leaf_commands_become_direct_actions(self) -> None:
        plain = Operation("plain", "Plain", "", "system_core.services.office_service:validate_input")
        controlled = Operation(
            "controlled",
            "Controlled",
            "",
            "system_core.services.office_service:validate_input",
            fields=({"id": "mode", "type": "radio", "options": ["one", "two"]},),
        )
        route_only = Operation(
            "route",
            "Route",
            "",
            "system_core.services.office_service:validate_input",
            fields=({"id": "input_path", "type": "path", "route": "source"},),
        )
        self.assertTrue(app.is_direct_action_node(app.operation_to_command_node(plain)))
        self.assertFalse(app.is_direct_action_node(app.operation_to_command_node(controlled)))
        self.assertTrue(app.is_direct_action_node(app.operation_to_command_node(route_only)))

    def test_rules_audit_gui_group_keeps_both_backend_actions(self) -> None:
        audit = app.command_node_by_id("audit_rules")
        self.assertIsNotNone(audit)
        assert audit is not None
        self.assertEqual(
            [child.id for child in audit.children],
            ["docx_audit_processor", "docx_audit_strip_anchors"],
        )
        self.assertTrue(callable(app.folder_button))


if __name__ == "__main__":
    unittest.main()
